"""
Etapa 6 (19 apr 2026) — smoke tests E2E pentru TOATE template-ele MDLPA.

Verifică că pentru fiecare template din public/templates/ procesarea
critică (compute_checkboxes, build_checkbox_index, append_legal_supplement,
insert_apartment_table, etc.) NU crashează și produce output valid.

Acoperire categorii Mc 001-2022 + tipologii MDLPA:
  RI / RC / RA / BI / ED / SA / HC / CO / SP / AL

Rulare:
  cd energy-app
  python -m pytest api/_tests/test_smoke_all_templates.py -v
"""
import base64
import importlib.util
import io
import sys
import unittest
from pathlib import Path

from docx import Document

API_DIR = Path(__file__).resolve().parent.parent
MODULE_PATH = API_DIR / "generate-document.py"

_spec = importlib.util.spec_from_file_location("generate_document_smoke", MODULE_PATH)
gd = importlib.util.module_from_spec(_spec)
sys.modules["generate_document_smoke"] = gd
_spec.loader.exec_module(gd)

TEMPLATES_DIR = API_DIR.parent / "public" / "templates"


# ── Payload sintetic minim viabil pentru orice categorie ──
def _baseline_data(category="RC"):
    """Date sintetice care exersează toate ramurile din process_document."""
    return {
        # Date generale clădire
        "year": "1985", "year_built": "1985",
        "address": "Str. Test nr. 1, București",
        "city": "București", "county": "B",
        "gps": "44,4300 x 26,1000",
        "regime": "P+10E", "scope": "Vânzare",
        "software": "ZEPHREN test",
        "area_ref": "1200", "area_gross": "1380",
        "volume": "3500", "nr_units": "20",
        "category_label": "Test categorie",
        "climate_zone": "zona II", "climate_zone_num": "2",
        # EP/CO2
        "ep_total_real": "210000", "ep_total_ref": "150000",
        "qf_thermal": "120", "qf_electric": "30",
        "qf_thermal_ref": "85", "qf_electric_ref": "22",
        "ep_specific": "175", "ep_ref": "125",
        "co2_val": "32.4",
        "ep_incalzire": "120", "ep_racire": "0",
        "ep_acm": "30", "ep_ventilare": "10", "ep_iluminat": "15",
        "sre_st": "0", "sre_pv": "0", "sre_pc": "0",
        "sre_bio": "0", "sre_other": "0", "sre_total": "0",
        "energy_class": "C",
        "ep_class_real": "C", "ep_class_ref": "B",
        "co2_class_real": "D",
        "rer": "12,5", "nzeb": "NU", "gwp": "32,1",
        # Auditor
        "auditor_name": "ing. Test Auditor",
        "auditor_atestat": "RO/12345",
        "auditor_grade": "I",
        "auditor_company": "Test SRL",
        "auditor_phone": "0700-000-000",
        "auditor_email": "test@example.com",
        "auditor_date": "19.04.2026",
        "auditor_mdlpa": "12345",
        # Sprint 14/15/17 — supliment legal
        "cpe_code": "12345_2026-04-19_Test_Auditor_RO_12345_1_CPE_a3f7b9c2",
        "registry_index": "1",
        "expiry": "19.04.2031",
        "validity_years": 5,
        "validity_label": "valabil 5 ani (clasa C, EPBD 2024)",
        "cadastral_number": "123456-C1-U5",
        "land_book": "98765/București",
        "area_built": "1380",
        "n_apartments": "20",
        "passport_uuid": "550e8400-e29b-41d4-a716-446655440000",
        "qr_verify_url": "https://zephren.ro/verify/test",
        "passport_url": "https://zephren.ro/passport/test",
        "passport_qr_url": "https://zephren.ro/passport/test",
        "signature_png_b64": "",
        "stamp_png_b64": "",
        "co2_max_ppm": "850", "pm25_avg": "12.4",
        "ev_charging_points": "2", "ev_charging_prepared": "5",
        # Etapa 2 — BACS/SRI/n50/penalties
        "bacs_class": "C",
        "sri_total": "55",
        "sri_grade": "B",
        "n50": "3.5",
        "penalties_summary": '{"summary":{"count_applied":2,"total_delta_pct":18,"ep_multiplier":1.18},"applied":[{"id":"p0","reason":"test","delta_EP_pct":15},{"id":"p2","reason":"punte","delta_EP_pct":3}]}',
        # Building category-specific
        "heating_source": "gaz_conv",
        "heating_fuel": "gaz natural",
        "heating_control": "manual",
        "heating_power": "30",
        "acm_source": "ct_prop",
        "acm_storage_volume": "200",
        "cooling_source": "",
        "cooling_has": "false",
        "ventilation_type": "natural_org",
        "lighting_type": "fluorescent",
        "lighting_control": "",
        "solar_thermal_enabled": "false",
        "pv_enabled": "false",
        "heat_pump_enabled": "false",
        "biomass_enabled": "false",
        "wind_enabled": "false",
        "structure": "Cadre beton armat",
        # Anvelopă (pentru penalități + checkboxes)
        "opaque_u_values": '[{"type":"PE","u":0.50,"area":300},{"type":"PB","u":0.40,"area":120}]',
        "glazing_max_u": "2.5",
        "glazing_g_value": "0.65",
        "overheating_hours": "0",
        # Scale (pentru CPE)
        "s_ap": "73", "s_a": "101", "s_b": "198", "s_c": "297",
        "s_d": "396", "s_e": "495", "s_f": "595",
        "co2_ap": "12,7", "co2_a": "17,6", "co2_b": "34,6", "co2_c": "52,2",
        "co2_d": "69,9", "co2_e": "87,4", "co2_f": "104,9",
    }


# ── Lista template-elor disponibile ──
def _list_templates(prefix=""):
    if not TEMPLATES_DIR.exists():
        return []
    return sorted(p for p in TEMPLATES_DIR.glob("*.docx")
                  if p.name.startswith(prefix) or not prefix)


def _open_template(name):
    path = TEMPLATES_DIR / name
    if not path.exists():
        raise unittest.SkipTest(f"Template lipsă: {name}")
    return Document(str(path))


class TestSmokeAllCpeTemplates(unittest.TestCase):
    """Smoke test: append_legal_supplement nu crashează pe niciun template CPE."""

    CPE_TEMPLATES = [
        "2-CPE-forma-generala-apartament.docx",
        "3-CPE-forma-generala-cladire.docx",
        "4-CPE-apartament-bloc-INC-ACC-RAC-VENT-IL.docx",
        "5-CPE-cladire-locuit-individuala-INC-ACC-RAC-VENT-IL.docx",
        "6-CPE-cladire-locuit-colectiva-INC-ACC-RAC-VENT-IL.docx",
        "7-CPE-cladire-birouri-INC-ACC-RAC-VENT-IL.docx",
        "8-CPE-cladire-invatamant-INC-ACC-RAC-VENT-IL.docx",
        "9-CPE-cladire-sanitar-INC-ACC-RAC-VENT-IL.docx",
        "10-CPE-cladire-comert-INC-ACC-RAC-VENT-IL.docx",
        "11-CPE-cladire-turism-INC-ACC-RAC-VENT-IL.docx",
        "12-CPE-cladire-sport-INC-ACC-RAC-VENT-IL.docx",
    ]

    def test_supplement_appends_to_all_cpe_templates(self):
        """append_legal_supplement nu crashează pe nicio variantă MDLPA."""
        data = _baseline_data()
        ok_count = 0
        errors = []
        for tpl_name in self.CPE_TEMPLATES:
            try:
                doc = _open_template(tpl_name)
                result = gd.append_legal_supplement(doc, data)
                self.assertTrue(result, f"{tpl_name}: append a returnat False")
                # Verifică că documentul e valid după modificare
                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)
                Document(buf)  # Re-open → throws dacă invalid
                ok_count += 1
            except unittest.SkipTest:
                continue
            except Exception as e:
                errors.append(f"{tpl_name}: {e}")
        if errors:
            self.fail("Erori smoke test:\n" + "\n".join(errors))
        self.assertGreater(ok_count, 0, "Niciun template CPE găsit pentru smoke test")


class TestSmokeAllAnexaTemplates(unittest.TestCase):
    """Smoke test: compute_checkboxes + build_checkbox_index pe Anexa templates."""

    ANEXA_TEMPLATES = [
        "ANEXA-1-si-ANEXA-2-la-CPE-cladire.docx",
        "ANEXA-1-si-ANEXA-2-la-CPE-apartament.docx",
    ]

    CATEGORIES = ["RI", "RC", "RA", "BI", "ED", "SA", "HC", "CO", "SP", "AL"]

    def test_compute_checkboxes_for_all_categories(self):
        """compute_checkboxes returnează listă validă pentru orice categorie Mc 001-2022."""
        data = _baseline_data()
        for cat in self.CATEGORIES:
            with self.subTest(category=cat):
                indices = gd.compute_checkboxes(data, cat)
                self.assertIsInstance(indices, list)
                # Toate indicii sunt int pozitivi
                for idx in indices:
                    self.assertIsInstance(idx, int)
                    self.assertGreaterEqual(idx, 0)

    def test_compute_checkbox_keys_for_all_categories(self):
        """compute_checkbox_keys returnează chei semantice pentru orice categorie."""
        data = _baseline_data()
        for cat in self.CATEGORIES:
            with self.subTest(category=cat):
                keys = gd.compute_checkbox_keys(data, cat)
                self.assertIsInstance(keys, list)
                self.assertGreater(len(keys), 0, f"{cat}: lista chei goală")
                # Există cel puțin o cheie de categorie
                self.assertTrue(
                    any(k.startswith("CAT_") for k in keys),
                    f"{cat}: nicio cheie CAT_ generată"
                )

    def test_build_checkbox_index_on_anexa_templates(self):
        """build_checkbox_index găsește chei semantice pe ambele template-uri Anexa."""
        for tpl_name in self.ANEXA_TEMPLATES:
            with self.subTest(template=tpl_name):
                doc = _open_template(tpl_name)
                idx = gd.build_checkbox_index(doc)
                self.assertGreater(len(idx), 10,
                    f"{tpl_name}: doar {len(idx)} chei găsite — prea puține")
                # Cheile esențiale trebuie să fie găsite pe orice template Anexa
                self.assertIn("REC_PE_INSULATE", idx, f"{tpl_name}: REC_PE_INSULATE lipsă")

    def test_toggle_checkboxes_by_keys_no_crash(self):
        """Bifare cu chei semantice nu crashează pe niciun template Anexa."""
        for tpl_name in self.ANEXA_TEMPLATES:
            with self.subTest(template=tpl_name):
                doc = _open_template(tpl_name)
                keys = gd.compute_checkbox_keys(_baseline_data(), "RC")
                result = gd.toggle_checkboxes_by_keys(doc, keys)
                self.assertIsInstance(result, dict)
                self.assertIn("found", result)
                self.assertIn("missing", result)
                # Documentul trebuie să rămână valid după bifare
                buf = io.BytesIO()
                doc.save(buf)
                buf.seek(0)
                Document(buf)


class TestSmokeAnexaBlocAllSizes(unittest.TestCase):
    """Smoke test: insert_apartment_table cu liste de diferite mărimi."""

    def _gen_apartments(self, n):
        """Generează N apartamente sintetice variate."""
        results = []
        for i in range(n):
            floor = 0 if i == 0 else (10 if i == n - 1 else (i % 5))
            results.append({
                "number": f"{i+1}A",
                "staircase": "A" if i < n // 2 else "B",
                "floor": floor,
                "areaUseful": 50 + (i * 3) % 30,
                "orientation": ["N", "E"] if i % 2 == 0 else ["S", "V"],
                "occupants": 2 + (i % 3),
                "corner": (i % 4 == 0),
                "topFloor": (i == n - 1),
                "groundFloor": (i == 0),
                "posKey": "ground_corner" if i == 0 else (
                    "top_corner" if i == n - 1 else "mid_interior"
                ),
                "posFactor": 1.18 if i == 0 else (1.15 if i == n - 1 else 1.0),
                "epAptM2": 150 + (i * 5) % 80,
                "co2AptM2": 25 + (i * 1.5) % 15,
                "enClass": ["A", "B", "C", "D", "E"][i % 5],
                "co2Class": ["B", "C", "D", "E"][i % 4],
                "allocatedPct": (100 / n),
            })
        return results

    def test_apartment_table_for_small_block(self):
        """Bloc mic: 4 apartamente."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("Test")
        apts = self._gen_apartments(4)
        result = gd.insert_apartment_table(doc, apts, {
            "totalAu": 220, "epAvgWeighted": 165,
            "co2AvgWeighted": 28, "avgEnergyClass": "C",
            "avgCo2Class": "D", "classDistribution": {"B": 1, "C": 2, "D": 1},
            "count": 4,
        })
        self.assertTrue(result)

    def test_apartment_table_for_medium_block(self):
        """Bloc mediu: 24 apartamente."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("Test")
        apts = self._gen_apartments(24)
        result = gd.insert_apartment_table(doc, apts, {
            "totalAu": 1500, "epAvgWeighted": 175,
            "co2AvgWeighted": 30, "avgEnergyClass": "C",
            "avgCo2Class": "D",
            "classDistribution": {"A": 2, "B": 5, "C": 10, "D": 5, "E": 2},
            "count": 24,
        })
        self.assertTrue(result)
        # Header + 24 apartamente + 1 sumar
        self.assertEqual(len(doc.tables[-1].rows), 26)

    def test_apartment_table_for_large_block(self):
        """Bloc mare: 80 apartamente (sanity check performanță)."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("Test")
        apts = self._gen_apartments(80)
        result = gd.insert_apartment_table(doc, apts, {
            "totalAu": 5000, "epAvgWeighted": 180,
            "co2AvgWeighted": 32, "avgEnergyClass": "C",
            "avgCo2Class": "D",
            "classDistribution": {"B": 10, "C": 30, "D": 25, "E": 15},
            "count": 80,
        })
        self.assertTrue(result)
        # Document trebuie să rămână valid
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        Document(buf)


class TestSmokePayloadIntegrity(unittest.TestCase):
    """Smoke: payload sintetic exersează toate ramurile fără crash."""

    CATEGORIES = ["RI", "RC", "RA", "BI", "ED", "SA", "HC", "CO", "SP", "AL"]

    def test_supplement_handles_all_categories_data(self):
        """append_legal_supplement nu crashează indiferent de categorie."""
        from docx import Document
        for cat in self.CATEGORIES:
            with self.subTest(category=cat):
                doc = Document()
                doc.add_paragraph("Test")
                data = _baseline_data(cat)
                result = gd.append_legal_supplement(doc, data)
                self.assertTrue(result, f"Categoria {cat}: supplement returns False")

    def test_supplement_with_minimal_payload_no_crash(self):
        """Date minimale (doar cpe_code) nu produc crash."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("Test")
        result = gd.append_legal_supplement(doc, {"cpe_code": "MIN_TEST"})
        self.assertTrue(result)

    def test_supplement_with_extreme_values_no_crash(self):
        """Valori extreme (string-uri foarte lungi, caractere speciale) nu strică DOCX."""
        from docx import Document
        doc = Document()
        data = {
            "cpe_code": "TEST_" + "X" * 200,  # cod foarte lung
            "auditor_date": "32.13.9999",  # dată invalidă
            "expiry": "—",
            "cadastral_number": "<XML>&special;chars\"'",
            "land_book": "Mai mult text Cu Diacritice ăâîșț ŞȚ și unicode 中文",
            "area_built": "abc def",  # non-numeric
            "co2_max_ppm": "",
            "pm25_avg": "0",
            "penalties_summary": "INVALID_JSON{",  # JSON invalid
        }
        result = gd.append_legal_supplement(doc, data)
        self.assertTrue(result)
        # Document trebuie să rămână valid
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        Document(buf)


if __name__ == "__main__":
    unittest.main()
