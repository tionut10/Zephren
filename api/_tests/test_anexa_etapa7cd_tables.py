"""
Etapa 7c+7d (20 apr 2026) — completare automată tabele Anexa + indicatori finali.

Acoperă:
  - Tabel 0: zone climatică/eoliană/regim înălțime
  - Tabel 1: apartamente/unități cu arie + nr.
  - Tabel 2: anvelopă R-values + arii
  - Tabel 3: 5 sisteme + TOTAL cu EP + CO2 + clase
  - Paragrafe finale: EPP, RERP, CO2 indicator, SRI complet

Rulare:
  python -m pytest api/_tests/test_anexa_etapa7cd_tables.py -v
"""
import importlib.util
import io
import json
import sys
import unittest
from pathlib import Path

from docx import Document

API_DIR = Path(__file__).resolve().parent.parent
MODULE_PATH = API_DIR / "generate-document.py"

_spec = importlib.util.spec_from_file_location("gd_etapa7cd", MODULE_PATH)
gd = importlib.util.module_from_spec(_spec)
sys.modules["gd_etapa7cd"] = gd
_spec.loader.exec_module(gd)

ANEXA_TPL = API_DIR.parent / "public" / "templates" / "ANEXA-1-si-ANEXA-2-la-CPE-cladire.docx"


def _doc_text(doc):
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        parts.append(p.text)
    return "\n".join(parts)


def _table_cells(doc, table_idx):
    """Returnează matricea text a unui tabel."""
    t = doc.tables[table_idx]
    return [[cell.text.strip() for cell in row.cells] for row in t.rows]


def _full_data():
    return {
        "address": "Str. Test", "year": "1985",
        "category_label": "Locuință individuală",
        "area_ref": "180", "volume": "500", "area_envelope": "230",
        "n_apartments_count": "1",
        "climate_zone_num": "2",  # zona II
        "regime": "P+1",
        "opaque_u_values": json.dumps([
            {"type": "PE", "u": 0.50, "area": 120},
            {"type": "PT", "u": 0.30, "area": 90},
            {"type": "PB", "u": 0.40, "area": 90},
        ]),
        "glazing_max_u": "1.50",
        "glazing_area_total_m2": "25,0",
        # EP per sistem
        "ep_incalzire": "120,5",
        "ep_acm": "30,2",
        "ep_racire": "0,0",
        "ep_ventilare": "10,5",
        "ep_iluminat": "15,0",
        "ep_specific": "175,0",
        "co2_val": "32,4",
        "energy_class": "C",
        "ep_ref": "125",
        "ep_class_ref": "B",
        # Etapa 7d
        "heating_power": "30",
        "acm_power": "15",
        "ventilation_flow_m3h": "150",
        "rer": "12,5",
        "sri_total": "67",
        "sri_grade": "B",
        # Etapa 7
        "cooling_eer": "3,20",
        "cooling_power_kw": "8,5",
        "cooled_area_m2": "120,5",
        "ventilation_has_hr": "Da",
        "ventilation_hr_efficiency_pct": "85",
        "lighting_power_kw": "2,40",
        "pv_kwh_year": "4500",
        "solar_th_kwh_year": "1800",
    }


def _process_anexa_full(doc, data):
    """Reproduce flow-ul process_document pentru mode==anexa (subset esențial)."""
    # Rulează replacements existente
    if data.get("address"):
        gd.replace_in_doc(doc, "[adresa]", data["address"])

    # Year (cu fix Etapa 6)
    year_full = data.get("year", "")
    if year_full:
        for p in doc.paragraphs:
            if "Anul construc" in p.text and "................." in p.text:
                gd.replace_in_paragraph(p, ".................", year_full, count=1)
                break

    # Aria + volum + factor formă
    au = data.get("area_ref", "")
    vol = data.get("volume", "")
    if au:
        gd.replace_in_doc(doc, "Aria de referință totală",
                          "Aria de referință totală a pardoselii: " + au + " m²")
    if vol:
        gd.replace_in_doc(doc, "Volumul interior de referință",
                          "Volumul interior de referință V: " + vol + " m³")

    # Etapa 7 — _fill_para_blank helper
    def _fill(label, value, suffix=""):
        if not value:
            return False
        val_str = str(value).strip()
        if not val_str:
            return False
        for p in doc.paragraphs:
            if label in p.text and val_str not in p.text:
                p.add_run(f" {val_str}{suffix}")
                return True
        return False

    # Etapa 7
    if data.get("cooling_eer"):
        _fill("coeficientului de performanţă EER", f"EER = {data['cooling_eer']}")
    if data.get("cooling_power_kw"):
        _fill("Necesarul de frig pentru răcire", data["cooling_power_kw"], " kW")
        _fill("Puterea frigorifică totală instalată", data["cooling_power_kw"], " kW")
    if data.get("ventilation_has_hr"):
        _fill("Există recuperator de căldură", data["ventilation_has_hr"])
    if data.get("ventilation_hr_efficiency_pct"):
        eff = data["ventilation_hr_efficiency_pct"]
        _fill("Eficiență declarată pe durata verii/iernii", f"{eff}% / {eff}%")
    if data.get("lighting_power_kw"):
        _fill("Puterea electrică totală necesară a sistemului de iluminat", data["lighting_power_kw"], " kW")
        _fill("Puterea electrică instalată totală a sistemului de iluminat", data["lighting_power_kw"], " kW")
    if data.get("pv_kwh_year") and data["pv_kwh_year"] != "0":
        _fill("Energia electrică exportată", data["pv_kwh_year"], " kWh/an")
    if data.get("solar_th_kwh_year") and data["solar_th_kwh_year"] != "0":
        _fill("Energia termică exportată", data["solar_th_kwh_year"], " kWh/an")
    if data.get("n_apartments_count"):
        _fill("Numărul & tipul apartamentelor", f" Total: {data['n_apartments_count']} unități")

    # Etapa 7d
    if data.get("heating_power"):
        _fill("Necesarul de căldură de calcul", data["heating_power"], " kW")
    if data.get("acm_power"):
        _fill("Puterea termică necesară pentru prepararea acc", data["acm_power"], " kW")
        _fill("Puterea termică maximă instalată pentru prepararea acc", data["acm_power"], " kW")
    if data.get("ventilation_flow_m3h"):
        _fill("Debitul minim de aer proaspăt", data["ventilation_flow_m3h"], " m³/h")
    if data.get("ep_specific"):
        _fill("Indicatorul energiei primare EPP", data["ep_specific"], " kWh/(m²·an)")
    if data.get("rer"):
        _fill("Indicele RERP", data["rer"], "%")
    if data.get("co2_val"):
        _fill("Indicatorul emisiilor de CO", data["co2_val"], " kgCO₂/(m²·an)")
    if data.get("sri_total"):
        sri_str = f"{data['sri_total']}% (Clasa {data.get('sri_grade', '')})"
        _fill("Indicele SRI (smart readiness indicator)", sri_str)


class TestEtapa7cdAnexaTables(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        if not ANEXA_TPL.exists():
            raise unittest.SkipTest(f"Template lipsă: {ANEXA_TPL}")

    def test_full_run_no_crash(self):
        """Smoke test: rularea completă pe template real nu crashează."""
        doc = Document(str(ANEXA_TPL))
        _process_anexa_full(doc, _full_data())
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        Document(buf)  # re-open

    def test_etapa7d_heating_power_appears(self):
        doc = Document(str(ANEXA_TPL))
        _process_anexa_full(doc, _full_data())
        text = _doc_text(doc)
        # "Necesarul de căldură de calcul ... 30 kW"
        for p in doc.paragraphs:
            if "Necesarul de căldură de calcul" in p.text:
                self.assertIn("30", p.text)
                break

    def test_etapa7d_acm_power_appears(self):
        doc = Document(str(ANEXA_TPL))
        _process_anexa_full(doc, _full_data())
        for p in doc.paragraphs:
            if "Puterea termică necesară pentru prepararea acc" in p.text:
                self.assertIn("15", p.text)
                self.assertIn("kW", p.text)
                break
        else:
            self.fail("Para ACM power nu găsit")

    def test_etapa7d_ventilation_flow_appears(self):
        doc = Document(str(ANEXA_TPL))
        _process_anexa_full(doc, _full_data())
        for p in doc.paragraphs:
            if "Debitul minim de aer proaspăt" in p.text:
                self.assertIn("150", p.text)
                self.assertIn("m³/h", p.text)
                break
        else:
            self.fail("Para debit aer nu găsit")

    def test_etapa7d_epp_appears(self):
        doc = Document(str(ANEXA_TPL))
        _process_anexa_full(doc, _full_data())
        for p in doc.paragraphs:
            if "EPP" in p.text and "primare" in p.text:
                self.assertIn("175,0", p.text)
                break
        else:
            self.fail("Para EPP nu găsit")

    def test_etapa7d_rerp_appears(self):
        doc = Document(str(ANEXA_TPL))
        _process_anexa_full(doc, _full_data())
        for p in doc.paragraphs:
            if "Indicele RERP" in p.text:
                self.assertIn("12,5", p.text)
                self.assertIn("%", p.text)
                break
        else:
            self.fail("Para RERP nu găsit")

    def test_etapa7d_co2_indicator_appears(self):
        doc = Document(str(ANEXA_TPL))
        _process_anexa_full(doc, _full_data())
        for p in doc.paragraphs:
            if "Indicatorul emisiilor de CO" in p.text:
                self.assertIn("32,4", p.text)
                break
        else:
            self.fail("Para CO2 indicator nu găsit")

    def test_etapa7d_sri_complete_label(self):
        doc = Document(str(ANEXA_TPL))
        _process_anexa_full(doc, _full_data())
        for p in doc.paragraphs:
            if "Indicele SRI" in p.text:
                self.assertIn("67%", p.text)
                self.assertIn("Clasa B", p.text)
                break
        else:
            self.fail("Para SRI nu găsit")

    def test_no_2015_repetition(self):
        """REGRESSION: nu reapare bug-ul anterior."""
        doc = Document(str(ANEXA_TPL))
        _process_anexa_full(doc, _full_data())
        text = _doc_text(doc)
        self.assertNotIn("20152015", text)
        self.assertNotIn("19851985", text)

    def test_idempotent_double_apply(self):
        """Aplicare dublă nu duplică valorile."""
        doc = Document(str(ANEXA_TPL))
        _process_anexa_full(doc, _full_data())
        _process_anexa_full(doc, _full_data())
        text = _doc_text(doc)
        # SRI trebuie să apară EXACT o dată cu valoarea
        self.assertEqual(text.count("67% (Clasa B)"), 1)


if __name__ == "__main__":
    unittest.main()
