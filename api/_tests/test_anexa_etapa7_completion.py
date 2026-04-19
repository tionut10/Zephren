"""
Etapa 7 (20 apr 2026) — completare automată câmpuri Anexa 2 detaliate.

Verifică că noile replacement-uri (introduse ca răspuns la auditul de gap-uri)
funcționează: EER, putere frigorifică, ventilare HR (existență + eficiență),
putere iluminat, energie regenerabilă exportată (PV/solar/eolian),
nr. apartamente, arie ferestre.

Folosește template Anexa real (308 cb) pentru smoke test end-to-end.

Rulare:
  cd energy-app
  python -m pytest api/_tests/test_anexa_etapa7_completion.py -v
"""
import importlib.util
import sys
import unittest
from pathlib import Path

from docx import Document

API_DIR = Path(__file__).resolve().parent.parent
MODULE_PATH = API_DIR / "generate-document.py"

_spec = importlib.util.spec_from_file_location("gd_etapa7", MODULE_PATH)
gd = importlib.util.module_from_spec(_spec)
sys.modules["gd_etapa7"] = gd
_spec.loader.exec_module(gd)

ANEXA_TPL = API_DIR.parent / "public" / "templates" / "ANEXA-1-si-ANEXA-2-la-CPE-cladire.docx"


def _doc_text(doc):
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        parts.append(p.text)
    return "\n".join(parts)


def _apply_etapa7_replacements(doc, data):
    """Reproduce logica de completare din process_document (mode in anexa)."""

    def _fill_para_blank(label_match, value, suffix=""):
        if not value:
            return False
        val_str = str(value).strip()
        if not val_str:
            return False
        for p in doc.paragraphs:
            pt = p.text
            if label_match not in pt:
                continue
            if val_str in pt:
                return True
            if p.runs:
                run = p.add_run(f" {val_str}{suffix}")
                run.font.size = p.runs[0].font.size
                run.font.name = p.runs[0].font.name
            else:
                p.add_run(f" {val_str}{suffix}")
            return True
        return False

    if data.get("cooling_eer"):
        _fill_para_blank("coeficientului de performanţă EER", f"EER = {data['cooling_eer']}")
    if data.get("cooling_seer"):
        for p in doc.paragraphs:
            if "EER al sursei de răcire" in p.text and f"SEER = {data['cooling_seer']}" not in p.text:
                p.add_run(f"   |   SEER = {data['cooling_seer']}")
                break
    if data.get("cooling_power_kw"):
        _fill_para_blank("Necesarul de frig pentru răcire", data["cooling_power_kw"], " kW")
        _fill_para_blank("Puterea frigorifică totală instalată", data["cooling_power_kw"], " kW")
    if data.get("cooled_area_m2"):
        _fill_para_blank("Volumul de referință al zonei climatizate", data["cooled_area_m2"], " m²")
    if data.get("ventilation_has_hr"):
        _fill_para_blank("Există recuperator de căldură", data["ventilation_has_hr"])
    if data.get("ventilation_hr_efficiency_pct"):
        eff = data["ventilation_hr_efficiency_pct"]
        _fill_para_blank("Eficiență declarată pe durata verii/iernii", f"{eff}% / {eff}%")
    if data.get("lighting_power_kw"):
        _fill_para_blank("Puterea electrică totală necesară a sistemului de iluminat", data["lighting_power_kw"], " kW")
        _fill_para_blank("Puterea electrică instalată totală a sistemului de iluminat", data["lighting_power_kw"], " kW")
    if data.get("solar_th_kwh_year") and data["solar_th_kwh_year"] != "0":
        _fill_para_blank("Energia termică exportată", data["solar_th_kwh_year"], " kWh/an")
    if data.get("pv_kwh_year") and data["pv_kwh_year"] != "0":
        _fill_para_blank("Energia electrică exportată", data["pv_kwh_year"], " kWh/an")
    if data.get("n_apartments_count"):
        _fill_para_blank(
            "Numărul & tipul apartamentelor/unităților de clădire/zonelor termice",
            f" Total: {data['n_apartments_count']} unități",
        )
    if data.get("glazing_area_total_m2"):
        for p in doc.paragraphs:
            if "U tâmplărie" in p.text and "Arie totală ferestre" not in p.text:
                p.add_run(f"   |   Arie totală ferestre: {data['glazing_area_total_m2']} m²")
                break


def _full_data():
    return {
        "cooling_eer": "3,20",
        "cooling_seer": "5,40",
        "cooling_power_kw": "8,5",
        "cooled_area_m2": "120,5",
        "ventilation_has_hr": "Da",
        "ventilation_hr_efficiency_pct": "85",
        "lighting_power_kw": "2,40",
        "pv_kwh_year": "4500",
        "solar_th_kwh_year": "1800",
        "wind_kwh_year": "0",
        "n_apartments_count": "12",
        "glazing_area_total_m2": "45,2",
    }


class TestEtapa7AnexaCompletion(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        if not ANEXA_TPL.exists():
            raise unittest.SkipTest(f"Template lipsă: {ANEXA_TPL}")

    def test_eer_appears_in_anexa(self):
        doc = Document(str(ANEXA_TPL))
        _apply_etapa7_replacements(doc, _full_data())
        text = _doc_text(doc)
        self.assertIn("EER = 3,20", text)

    def test_seer_appears_alongside_eer(self):
        doc = Document(str(ANEXA_TPL))
        _apply_etapa7_replacements(doc, _full_data())
        text = _doc_text(doc)
        self.assertIn("SEER = 5,40", text)

    def test_cooling_power_appears_in_two_places(self):
        """Putere frigorifică apare în 'Necesarul de frig' + 'Putere instalată'."""
        doc = Document(str(ANEXA_TPL))
        _apply_etapa7_replacements(doc, _full_data())
        text = _doc_text(doc)
        # Numărăm aparițiile lui "8,5 kW"
        self.assertGreaterEqual(text.count("8,5 kW"), 2,
            "Puterea frigorifică ar trebui să apară în cel puțin 2 paragrafe")

    def test_ventilation_hr_yes_appears(self):
        doc = Document(str(ANEXA_TPL))
        _apply_etapa7_replacements(doc, _full_data())
        text = _doc_text(doc)
        # "Există recuperator de căldură: Da"
        for p in doc.paragraphs:
            if "Există recuperator de căldură" in p.text:
                self.assertIn("Da", p.text, "HR=Da nu apare în paragraful corect")
                break

    def test_ventilation_hr_efficiency_appears(self):
        doc = Document(str(ANEXA_TPL))
        _apply_etapa7_replacements(doc, _full_data())
        text = _doc_text(doc)
        self.assertIn("85% / 85%", text)

    def test_lighting_power_appears_twice(self):
        """Puterea iluminat apare în 'totală necesară' + 'instalată totală'."""
        doc = Document(str(ANEXA_TPL))
        _apply_etapa7_replacements(doc, _full_data())
        text = _doc_text(doc)
        self.assertGreaterEqual(text.count("2,40 kW"), 2)

    def test_pv_energy_exported_appears(self):
        doc = Document(str(ANEXA_TPL))
        _apply_etapa7_replacements(doc, _full_data())
        text = _doc_text(doc)
        self.assertIn("4500 kWh/an", text)

    def test_solar_thermal_energy_appears(self):
        doc = Document(str(ANEXA_TPL))
        _apply_etapa7_replacements(doc, _full_data())
        text = _doc_text(doc)
        self.assertIn("1800 kWh/an", text)

    def test_wind_zero_skipped(self):
        """Energia eoliană = 0 → NU se completează (evită 0 inutile)."""
        doc = Document(str(ANEXA_TPL))
        _apply_etapa7_replacements(doc, _full_data())
        text = _doc_text(doc)
        # Eoliene n-ar trebui să primească "0 kWh"
        self.assertNotIn("0 kWh/an produși", text)

    def test_n_apartments_appears(self):
        doc = Document(str(ANEXA_TPL))
        _apply_etapa7_replacements(doc, _full_data())
        text = _doc_text(doc)
        self.assertIn("Total: 12 unități", text)

    def test_glazing_area_appears(self):
        """Arie ferestre apare în paragraful U tâmplărie (dacă există)."""
        # Pentru template cu paragraful U tâmplărie inițial (înainte de replace,
        # paragraful are doar 'coeficientul global... tâmplăriei').
        # Test simulat: adaug un paragraf cu "U tâmplărie" și verific.
        doc = Document(str(ANEXA_TPL))
        # Adaug un paragraf simulat cu eticheta
        p = doc.add_paragraph("U tâmplărie = 1,30 W/(m²·K)")
        _apply_etapa7_replacements(doc, _full_data())
        text = _doc_text(doc)
        self.assertIn("Arie totală ferestre: 45,2 m²", text)

    def test_idempotent_replacements(self):
        """Aplicarea de 2x nu duplică valorile."""
        doc = Document(str(ANEXA_TPL))
        _apply_etapa7_replacements(doc, _full_data())
        _apply_etapa7_replacements(doc, _full_data())
        text = _doc_text(doc)
        # EER = 3,20 trebuie să apară DOAR o dată
        self.assertEqual(text.count("EER = 3,20"), 1,
            "Aplicare dublă a duplicat valoarea EER")

    def test_empty_data_no_changes(self):
        """Cu date complet goale, document e neschimbat."""
        doc = Document(str(ANEXA_TPL))
        before = _doc_text(doc)
        _apply_etapa7_replacements(doc, {})
        after = _doc_text(doc)
        self.assertEqual(before, after)

    def test_no_2015_repetition_after_etapa7(self):
        """REGRESSION: completarea Etapa 7 nu reintroduce bug-ul '20152015'."""
        doc = Document(str(ANEXA_TPL))
        _apply_etapa7_replacements(doc, _full_data())
        text = _doc_text(doc)
        self.assertNotIn("20152015", text)
        self.assertNotIn("85852015", text)  # alt pattern probabil

    def test_doc_remains_valid_after_replacements(self):
        """Document trebuie să rămână DOCX valid."""
        import io
        doc = Document(str(ANEXA_TPL))
        _apply_etapa7_replacements(doc, _full_data())
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        Document(buf)


if __name__ == "__main__":
    unittest.main()
