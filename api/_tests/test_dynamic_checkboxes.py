"""
Etapa 3 (19 apr 2026) — teste pentru sistemul de mapping dinamic checkbox.

Verifică:
  1. _normalize_text() normalizează corect (lowercase + diacritice + spații)
  2. _get_checkbox_context() extrage contextul corect din template real
  3. build_checkbox_index() găsește indici pentru cheile semantice principale
     pe AMBELE template-uri (clădire 308 cb + apartament 244 cb)
  4. Match-urile sunt corecte: cheia REC_PE_INSULATE → CB[0] etc.
  5. compute_checkbox_keys() generează chei active conform datelor de input
  6. toggle_checkboxes_by_keys() bifează corect și raportează missing keys

Rulare:
  cd energy-app
  python -m pytest api/_tests/test_dynamic_checkboxes.py -v
"""
import importlib.util
import sys
import unittest
from pathlib import Path

from docx import Document

# ── Încărcare dinamică generate-document.py ──
API_DIR = Path(__file__).resolve().parent.parent
MODULE_PATH = API_DIR / "generate-document.py"

_spec = importlib.util.spec_from_file_location("generate_document_dyn", MODULE_PATH)
gd = importlib.util.module_from_spec(_spec)
sys.modules["generate_document_dyn"] = gd
_spec.loader.exec_module(gd)

# ── Path-uri template ──
TEMPLATES_DIR = API_DIR.parent / "public" / "templates"
ANEXA_CLADIRE = TEMPLATES_DIR / "ANEXA-1-si-ANEXA-2-la-CPE-cladire.docx"
ANEXA_APARTAMENT = TEMPLATES_DIR / "ANEXA-1-si-ANEXA-2-la-CPE-apartament.docx"


class TestNormalizeText(unittest.TestCase):
    def test_lowercase(self):
        self.assertEqual(gd._normalize_text("HELLO World"), "hello world")

    def test_strip_diacritics(self):
        self.assertEqual(gd._normalize_text("Țepeș Ștefan"), "tepes stefan")
        self.assertEqual(gd._normalize_text("ăâîșțĂÂÎȘȚ"), "aaistaaist")

    def test_collapse_whitespace(self):
        self.assertEqual(gd._normalize_text("  multi   spaces  here  "), "multi spaces here")

    def test_none_and_empty(self):
        self.assertEqual(gd._normalize_text(None), "")
        self.assertEqual(gd._normalize_text(""), "")

    def test_combined(self):
        self.assertEqual(
            gd._normalize_text("  ÎNCĂLZIRE   cu  CORPURI Statice  "),
            "incalzire cu corpuri statice"
        )


class TestCheckboxContext(unittest.TestCase):
    """Verifică extracția contextuală pe template-ul real."""

    @classmethod
    def setUpClass(cls):
        if not ANEXA_CLADIRE.exists():
            raise unittest.SkipTest(f"Template lipsă: {ANEXA_CLADIRE}")
        cls.doc = Document(str(ANEXA_CLADIRE))
        cls.checkboxes = cls.doc.element.findall(".//w:checkBox", gd.NSMAP)

    def test_template_has_308_checkboxes(self):
        self.assertEqual(len(self.checkboxes), 308)

    def test_cb0_context_contains_perete_exterior(self):
        ctx = gd._get_checkbox_context(self.checkboxes[0], self.doc)
        self.assertIn("pereti", ctx)
        self.assertIn("exteriori", ctx)

    def test_cb135_context_contains_existenta_incalzire(self):
        ctx = gd._get_checkbox_context(self.checkboxes[135], self.doc)
        self.assertIn("existenta", ctx)
        self.assertIn("incalzire", ctx)

    def test_cb288_context_contains_panouri_termosolare(self):
        ctx = gd._get_checkbox_context(self.checkboxes[288], self.doc)
        self.assertIn("termosolare", ctx)


class TestBuildCheckboxIndexCladire(unittest.TestCase):
    """Index dinamic pe template clădire (308 cb)."""

    @classmethod
    def setUpClass(cls):
        if not ANEXA_CLADIRE.exists():
            raise unittest.SkipTest(f"Template lipsă: {ANEXA_CLADIRE}")
        cls.doc = Document(str(ANEXA_CLADIRE))
        cls.idx = gd.build_checkbox_index(cls.doc)

    def test_index_finds_majority_of_keys(self):
        """Index-ul trebuie să găsească ≥80% din cele ~60 chei semantice."""
        total = len(gd.CHECKBOX_KEYWORD_MAP)
        found = len(self.idx)
        coverage = found / total
        self.assertGreaterEqual(coverage, 0.80,
            f"Acoperire keyword-uri prea mică: {found}/{total} ({coverage:.0%}). "
            f"Lipsă: {set(gd.CHECKBOX_KEYWORD_MAP) - set(self.idx)}")

    def test_anvelopa_keys_match_expected_indices(self):
        """Cheile principale anvelopă trebuie să se mapeze pe indicii cunoscuți."""
        # Conform _get_checkbox_context, CB[0] = perete exterior
        self.assertEqual(self.idx.get("REC_PE_INSULATE"), 0)
        self.assertEqual(self.idx.get("REC_PB_INSULATE"), 1)
        self.assertEqual(self.idx.get("REC_PT_INSULATE"), 2)
        self.assertEqual(self.idx.get("REC_PL_INSULATE"), 3)
        self.assertEqual(self.idx.get("REC_GLAZING"), 5)
        self.assertEqual(self.idx.get("REC_GRILES_VENT"), 6)

    def test_instalatii_keys_match_expected(self):
        """Chei instalații (CB ~13, 21, 25, 26, 27, 28)."""
        self.assertEqual(self.idx.get("REC_THERM_VALVES"), 13)
        self.assertEqual(self.idx.get("REC_AUTOMATION"), 21)
        self.assertEqual(self.idx.get("REC_LIGHT_LED"), 25)
        self.assertEqual(self.idx.get("REC_PRESENCE_SENS"), 26)
        self.assertEqual(self.idx.get("REC_RENEWABLES"), 27)
        self.assertEqual(self.idx.get("REC_HEAT_RECOVERY"), 28)

    def test_existenta_instalatii_keys(self):
        """Chei existență instalații (CB ~135, 176, 202, 256, 272)."""
        self.assertEqual(self.idx.get("HEAT_EXISTS_OK"), 135)
        self.assertEqual(self.idx.get("DHW_EXISTS_OK"), 176)
        self.assertEqual(self.idx.get("COOL_EXISTS_OK"), 202)
        self.assertEqual(self.idx.get("VENT_EXISTS_OK"), 256)
        self.assertEqual(self.idx.get("LIGHT_EXISTS_OK"), 272)

    def test_regenerabile_keys(self):
        """Chei regenerabile pereche YES/NO (CB 288-307)."""
        self.assertEqual(self.idx.get("RENEW_SOLAR_TH_YES"), 288)
        self.assertEqual(self.idx.get("RENEW_SOLAR_TH_NO"), 289)
        self.assertEqual(self.idx.get("RENEW_PV_YES"), 290)
        self.assertEqual(self.idx.get("RENEW_PV_NO"), 291)
        self.assertEqual(self.idx.get("RENEW_HP_YES"), 292)

    def test_index_is_cached(self):
        """A doua apelare returnează același obiect (cache)."""
        idx2 = gd.build_checkbox_index(self.doc)
        self.assertIs(idx2, self.idx)


class TestBuildCheckboxIndexApartament(unittest.TestCase):
    """Index dinamic pe template apartament (244 cb) — diferit de clădire."""

    @classmethod
    def setUpClass(cls):
        if not ANEXA_APARTAMENT.exists():
            raise unittest.SkipTest(f"Template lipsă: {ANEXA_APARTAMENT}")
        cls.doc = Document(str(ANEXA_APARTAMENT))
        cls.idx = gd.build_checkbox_index(cls.doc)

    def test_template_has_244_checkboxes(self):
        d2 = Document(str(ANEXA_APARTAMENT))
        cbs2 = d2.element.findall(".//w:checkBox", gd.NSMAP)
        self.assertEqual(len(cbs2), 244)

    def test_anvelopa_keys_still_match_on_apartament(self):
        """Cheia REC_PE_INSULATE trebuie să se găsească indiferent de template
        (deși indexul XML va fi DIFERIT față de template clădire)."""
        # Pe apartament, semnificația e poate la alt index — dar trebuie găsită
        self.assertIn("REC_PE_INSULATE", self.idx)
        self.assertIn("REC_GLAZING", self.idx)

    def test_existenta_incalzire_key_matches_on_apartament(self):
        """Cheia HEAT_EXISTS_OK trebuie găsită pe orice template Anexa."""
        self.assertIn("HEAT_EXISTS_OK", self.idx)


class TestComputeCheckboxKeys(unittest.TestCase):
    """Verifică logica de generare a cheilor semantice din date de audit."""

    def _base_data(self):
        return {
            "opaque_u_values": "[]",
            "glazing_max_u": "0",
            "ventilation_type": "",
            "heating_source": "gaz_conv",
            "heating_control": "manual",
            "lighting_type": "fluorescent",
            "lighting_control": "",
            "solar_thermal_enabled": "false",
            "pv_enabled": "false",
            "year_built": "1985",
            "acm_source": "ct_prop",
            "cooling_has": "false",
            "heat_pump_enabled": "false",
            "biomass_enabled": "false",
            "wind_enabled": "false",
        }

    def test_returns_list_of_strings(self):
        keys = gd.compute_checkbox_keys(self._base_data(), "RC")
        self.assertIsInstance(keys, list)
        for k in keys:
            self.assertIsInstance(k, str)

    def test_residential_block_has_correct_category_key(self):
        keys = gd.compute_checkbox_keys(self._base_data(), "RC")
        self.assertIn("CAT_RES_BLOC", keys)
        self.assertIn("BLDG_EXISTING", keys)

    def test_office_building_uses_office_category(self):
        keys = gd.compute_checkbox_keys(self._base_data(), "BI")
        self.assertIn("CAT_OFFICE", keys)
        self.assertNotIn("CAT_RES_BLOC", keys)

    def test_no_renewables_marks_REC_RENEWABLES(self):
        keys = gd.compute_checkbox_keys(self._base_data(), "RC")
        self.assertIn("REC_RENEWABLES", keys)
        self.assertIn("RENEW_SOLAR_TH_NO", keys)
        self.assertIn("RENEW_PV_NO", keys)

    def test_active_renewables_marks_yes(self):
        d = self._base_data()
        d["solar_thermal_enabled"] = "true"
        d["pv_enabled"] = "true"
        keys = gd.compute_checkbox_keys(d, "RC")
        self.assertIn("RENEW_SOLAR_TH_YES", keys)
        self.assertIn("RENEW_PV_YES", keys)
        self.assertNotIn("REC_RENEWABLES", keys)

    def test_no_heating_source_marks_NONE(self):
        d = self._base_data()
        d["heating_source"] = ""
        keys = gd.compute_checkbox_keys(d, "RC")
        self.assertIn("HEAT_EXISTS_NONE", keys)
        self.assertNotIn("HEAT_EXISTS_OK", keys)

    def test_walls_above_threshold_mark_PE_INSULATE(self):
        d = self._base_data()
        d["opaque_u_values"] = '[{"type":"PE","u":0.50}]'  # > 0.25 (RC res)
        keys = gd.compute_checkbox_keys(d, "RC")
        self.assertIn("REC_PE_INSULATE", keys)

    def test_walls_below_threshold_skip_PE_INSULATE(self):
        d = self._base_data()
        d["opaque_u_values"] = '[{"type":"PE","u":0.20}]'  # < 0.25
        keys = gd.compute_checkbox_keys(d, "RC")
        self.assertNotIn("REC_PE_INSULATE", keys)


class TestToggleCheckboxesByKeys(unittest.TestCase):
    """Verifică bifarea efectivă pe template real."""

    @classmethod
    def setUpClass(cls):
        if not ANEXA_CLADIRE.exists():
            raise unittest.SkipTest(f"Template lipsă: {ANEXA_CLADIRE}")

    def test_toggle_with_known_keys_succeeds(self):
        doc = Document(str(ANEXA_CLADIRE))
        result = gd.toggle_checkboxes_by_keys(doc, ["REC_PE_INSULATE", "HEAT_EXISTS_OK"])
        self.assertEqual(set(result["found"]), {"REC_PE_INSULATE", "HEAT_EXISTS_OK"})
        self.assertEqual(result["missing"], [])

    def test_toggle_with_unknown_key_reports_missing(self):
        doc = Document(str(ANEXA_CLADIRE))
        result = gd.toggle_checkboxes_by_keys(doc, ["NONEXISTENT_KEY_XYZ"])
        self.assertEqual(result["found"], [])
        self.assertEqual(result["missing"], ["NONEXISTENT_KEY_XYZ"])

    def test_toggle_empty_list_returns_empty(self):
        doc = Document(str(ANEXA_CLADIRE))
        result = gd.toggle_checkboxes_by_keys(doc, [])
        self.assertEqual(result, {"found": [], "missing": []})

    def test_toggle_actually_marks_checkbox_in_xml(self):
        """Verificare end-to-end: după toggle, w:default w:val='1' e setat."""
        doc = Document(str(ANEXA_CLADIRE))
        # Bifez REC_PE_INSULATE → CB[0]
        gd.toggle_checkboxes_by_keys(doc, ["REC_PE_INSULATE"])
        cb0 = doc.element.findall(".//w:checkBox", gd.NSMAP)[0]
        default = cb0.find("w:default", gd.NSMAP)
        self.assertIsNotNone(default)
        self.assertEqual(default.get(gd.qn("w:val")), "1")


class TestBackwardCompat(unittest.TestCase):
    """compute_checkboxes() vechi rămâne funcțional pentru backward compat."""

    def test_compute_checkboxes_returns_int_list(self):
        data = {"opaque_u_values": "[]", "ventilation_type": "natural_neorg",
                "heating_source": "gaz_conv", "heating_control": "manual"}
        result = gd.compute_checkboxes(data, "RC")
        self.assertIsInstance(result, list)
        for v in result:
            self.assertIsInstance(v, int)

    def test_toggle_checkboxes_legacy_int_indices_still_work(self):
        """Vechiul flux cu indici hardcodate trebuie să funcționeze ca înainte."""
        if not ANEXA_CLADIRE.exists():
            self.skipTest(f"Template lipsă: {ANEXA_CLADIRE}")
        doc = Document(str(ANEXA_CLADIRE))
        gd.toggle_checkboxes(doc, [0, 5, 13])
        # Verifică toate 3 sunt bifate
        cbs = doc.element.findall(".//w:checkBox", gd.NSMAP)
        for idx in [0, 5, 13]:
            d = cbs[idx].find("w:default", gd.NSMAP)
            self.assertIsNotNone(d)
            self.assertEqual(d.get(gd.qn("w:val")), "1")


if __name__ == "__main__":
    unittest.main()
