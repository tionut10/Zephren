"""
Sprint 14 / Etapa 1 (19 apr 2026) — teste pentru pagina supliment legală CPE.

Verifică:
  1. append_legal_supplement() returnează False pentru cpe_code gol
  2. append_legal_supplement() adaugă exact o pagină nouă
  3. Tabelul conține toate etichetele cerute de Ord. MDLPA 16/2023 + L.238/2024 + EPBD 2024
  4. Valorile lipsă afișează "—" (transparență — nu omite rândul)
  5. Imaginile PNG b64 sunt inserate ca <w:drawing>
  6. Data URL prefix `data:image/png;base64,` este tolerat la decodare
  7. QR PNG generat e valid (header \\x89PNG) când segno e disponibil
  8. Indicatorii EPBD 2024 Art. 14 apar în document (CO₂, PM2.5, EV)
  9. Nota finală e prezentă

Directorul `_tests` cu prefix underscore — Vercel ignoră fișierele/dir-urile cu underscore
la deploy, deci testele nu măresc bundle-ul de funcții (limit Hobby 12 funcții).

Rulare:
  cd energy-app
  python -m pytest api/_tests/ -v
"""
import base64
import importlib.util
import io
import sys
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

# ── Încărcare dinamică generate-document.py (dash în nume → nu se importă direct) ──
API_DIR = Path(__file__).resolve().parent.parent
MODULE_PATH = API_DIR / "generate-document.py"

if not MODULE_PATH.exists():
    raise RuntimeError(f"Nu găsesc {MODULE_PATH}")

_spec = importlib.util.spec_from_file_location("generate_document_module", MODULE_PATH)
gd = importlib.util.module_from_spec(_spec)
sys.modules["generate_document_module"] = gd
_spec.loader.exec_module(gd)


# ── PNG 1×1 transparent valid (pentru testele de imagine) ──
PNG_1X1_TRANSPARENT = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNkYPhfDwAChwGA60e6kgAAAABJRU5ErkJggg=="
)
PNG_1X1_B64 = base64.b64encode(PNG_1X1_TRANSPARENT).decode("ascii")


def _make_minimal_doc():
    """Creează un document minimal cu un singur paragraf, simulând un template MDLPA gol."""
    doc = Document()
    doc.add_paragraph("CPE Test — corpul principal")
    return doc


def _count_drawings(doc):
    """Numără elementele <w:drawing> din document — fiecare e o imagine PNG inserată."""
    return sum(1 for _ in doc.element.body.iter(qn("w:drawing")))


def _count_page_breaks(doc):
    """Numără elementele <w:br w:type='page'/> din document."""
    count = 0
    for br in doc.element.body.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            count += 1
    return count


def _doc_text(doc):
    """Concatenează tot textul vizibil din document (paragrafe + tabele)."""
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        parts.append(p.text)
    return "\n".join(parts)


# ═══════════════════════════════════════════════════════
# TESTS
# ═══════════════════════════════════════════════════════

class TestAppendLegalSupplement(unittest.TestCase):
    """Sprint 14/15/17 — pagina supliment cu metadate legale CPE."""

    def setUp(self):
        self.full_data = {
            "cpe_code": "12345_2026-04-19_Popescu_Ion_RO_4567_1_CPE_a3f7b9c2",
            "auditor_date": "19.04.2026",
            "expiry": "19.04.2031",
            "validity_years": 5,
            "validity_label": "valabil 5 ani (clasa D, EPBD 2024/1275 Art. 17)",
            "cadastral_number": "123456-C1-U5",
            "land_book": "98765/Cluj-Napoca",
            "area_built": "85.4",
            "n_apartments": "12",
            "passport_uuid": "550e8400-e29b-41d4-a716-446655440000",
            "qr_verify_url": "https://zephren.ro/verify/12345_2026-04-19_Popescu_Ion_RO_4567_1_CPE_a3f7b9c2",
            "passport_url": "https://zephren.ro/passport/550e8400-e29b-41d4-a716-446655440000",
            "passport_qr_url": "https://zephren.ro/passport/550e8400-e29b-41d4-a716-446655440000",
            "signature_png_b64": PNG_1X1_B64,
            "stamp_png_b64": PNG_1X1_B64,
            "co2_max_ppm": "850",
            "pm25_avg": "12.4",
            "ev_charging_points": "2",
            "ev_charging_prepared": "5",
        }

    def test_returns_false_when_cpe_code_missing(self):
        """Fără cpe_code, funcția returnează False și nu modifică documentul."""
        doc = _make_minimal_doc()
        before_paragraphs = len(doc.paragraphs)
        result = gd.append_legal_supplement(doc, {"cpe_code": ""})
        self.assertFalse(result)
        self.assertEqual(len(doc.paragraphs), before_paragraphs)

    def test_returns_false_when_data_empty(self):
        """Fără cheia cpe_code în data, returnează False."""
        doc = _make_minimal_doc()
        result = gd.append_legal_supplement(doc, {})
        self.assertFalse(result)

    def test_returns_false_when_cpe_code_whitespace_only(self):
        """cpe_code doar cu spații → False (strip)."""
        doc = _make_minimal_doc()
        result = gd.append_legal_supplement(doc, {"cpe_code": "   "})
        self.assertFalse(result)

    def test_adds_page_break(self):
        """Cu cpe_code valid, se adaugă exact un page break."""
        doc = _make_minimal_doc()
        before_breaks = _count_page_breaks(doc)
        result = gd.append_legal_supplement(doc, self.full_data)
        self.assertTrue(result)
        after_breaks = _count_page_breaks(doc)
        self.assertEqual(after_breaks - before_breaks, 1)

    def test_table_contains_all_required_labels(self):
        """Tabelul include toate etichetele cerute de Ord. MDLPA 16/2023 + L.238/2024."""
        doc = _make_minimal_doc()
        gd.append_legal_supplement(doc, self.full_data)
        text = _doc_text(doc)
        required_labels = [
            "Cod unic CPE",
            "Data emiterii",
            "Data expirării",
            "Valabilitate (ani)",
            "Etichetă valabilitate",
            "Nr. cadastral",
            "Carte funciară",
            "Arie construită desfășurată",
            "Nr. apartamente",
            "UUID pașaport renovare",
            "URL verificare CPE",
            "URL pașaport renovare",
            "Semnătură auditor",
            "Ștampilă auditor",
            "QR verificare CPE",
            "QR pașaport renovare",
        ]
        for lbl in required_labels:
            self.assertIn(lbl, text, f"Lipsește eticheta: {lbl}")

    def test_table_contains_actual_values(self):
        """Valorile reale (nu placeholder-e) apar în tabel."""
        doc = _make_minimal_doc()
        gd.append_legal_supplement(doc, self.full_data)
        text = _doc_text(doc)
        self.assertIn(self.full_data["cpe_code"], text)
        self.assertIn(self.full_data["cadastral_number"], text)
        self.assertIn(self.full_data["land_book"], text)
        self.assertIn(self.full_data["passport_uuid"], text)

    def test_missing_values_show_dash(self):
        """Valori lipsă (None/empty) afișează '—' — transparență, nu omitere."""
        doc = _make_minimal_doc()
        partial = {
            "cpe_code": "TEST_CODE_123",
            # Toate celelalte intenționat absente
        }
        result = gd.append_legal_supplement(doc, partial)
        self.assertTrue(result)
        text = _doc_text(doc)
        self.assertIn("TEST_CODE_123", text)
        self.assertIn("—", text)  # cel puțin un dash pentru câmpuri lipsă

    def test_signature_png_inserted_when_provided(self):
        """PNG b64 valid pentru semnătură → drawing inserat în doc."""
        doc = _make_minimal_doc()
        before = _count_drawings(doc)
        gd.append_legal_supplement(doc, self.full_data)
        after = _count_drawings(doc)
        # Avem cel puțin 2 imagini (semnătură + ștampilă), plus eventual 2 QR (dacă segno disponibil)
        self.assertGreaterEqual(after - before, 2)

    def test_data_url_prefix_tolerated(self):
        """Prefixul `data:image/png;base64,` e strip-uit înainte de decodare."""
        doc = _make_minimal_doc()
        data_with_prefix = dict(self.full_data)
        data_with_prefix["signature_png_b64"] = "data:image/png;base64," + PNG_1X1_B64
        before = _count_drawings(doc)
        gd.append_legal_supplement(doc, data_with_prefix)
        after = _count_drawings(doc)
        self.assertGreater(after, before, "Imaginea cu prefix data URL nu a fost decodată")

    def test_invalid_b64_does_not_crash(self):
        """B64 invalid → mesaj eroare în celulă, fără excepție."""
        doc = _make_minimal_doc()
        bad = dict(self.full_data)
        bad["signature_png_b64"] = "ACEAST_NU_E_BASE64_VALID!!!"
        try:
            result = gd.append_legal_supplement(doc, bad)
            self.assertTrue(result)
        except Exception as e:
            self.fail(f"append_legal_supplement a aruncat: {e}")

    def test_epbd_indicators_present(self):
        """Indicatorii EPBD 2024 Art. 14 (CO₂, PM2.5, EV) apar în document."""
        doc = _make_minimal_doc()
        gd.append_legal_supplement(doc, self.full_data)
        text = _doc_text(doc)
        self.assertIn("EPBD 2024/1275 Art. 14", text)
        self.assertIn("CO₂", text)
        self.assertIn("PM2.5", text)
        self.assertIn("Puncte încărcare EV", text)
        self.assertIn(self.full_data["co2_max_ppm"], text)

    def test_footer_note_present(self):
        """Nota finală cu referințe normative apare în document."""
        doc = _make_minimal_doc()
        gd.append_legal_supplement(doc, self.full_data)
        text = _doc_text(doc)
        self.assertIn("Ord. MDLPA 16/2023", text)
        self.assertIn("EPBD 2024/1275", text)

    def test_area_built_gets_m2_unit(self):
        """area_built fără unitate primește ' m²' automat în supliment."""
        doc = _make_minimal_doc()
        data = dict(self.full_data)
        data["area_built"] = "150"
        gd.append_legal_supplement(doc, data)
        text = _doc_text(doc)
        self.assertIn("150 m²", text)

    def test_area_built_preserves_existing_unit(self):
        """area_built care deja are 'm²' nu se duplichează."""
        doc = _make_minimal_doc()
        data = dict(self.full_data)
        data["area_built"] = "200 m²"
        gd.append_legal_supplement(doc, data)
        text = _doc_text(doc)
        self.assertIn("200 m²", text)
        self.assertNotIn("200 m² m²", text)

    def test_doc_remains_valid_docx(self):
        """Documentul rezultat poate fi salvat și re-deschis fără erori."""
        doc = _make_minimal_doc()
        gd.append_legal_supplement(doc, self.full_data)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reopened = Document(buf)
        self.assertGreater(len(reopened.paragraphs), 0)

    def test_supplement_table_has_2_columns(self):
        """Tabelul supliment are exact 2 coloane (etichetă + valoare)."""
        doc = _make_minimal_doc()
        gd.append_legal_supplement(doc, self.full_data)
        # Ultimul tabel adăugat e cel suplimentar
        self.assertGreater(len(doc.tables), 0)
        last_tbl = doc.tables[-1]
        self.assertEqual(len(last_tbl.columns), 2)


class TestQRGeneration(unittest.TestCase):
    """Verifică generate_qr_png() — folosit de append_legal_supplement pentru QR-uri."""

    def test_qr_png_header_valid(self):
        """QR generat are header PNG valid (\\x89PNG\\r\\n\\x1a\\n)."""
        if not gd._SEGNO_AVAILABLE:
            self.skipTest("segno nu e instalat — skip QR test")
        png = gd.generate_qr_png("https://zephren.ro/verify/test", scale=4, border=2)
        self.assertIsNotNone(png)
        self.assertTrue(png.startswith(b"\x89PNG"), "Headerul PNG lipsește")

    def test_qr_returns_none_for_empty_url(self):
        """URL gol → None (nu PNG vid)."""
        result = gd.generate_qr_png("", scale=4, border=2)
        self.assertIsNone(result)


class TestHelperFunctions(unittest.TestCase):
    """Funcții helper interne ale supliment-ului."""

    def test_format_area_adds_unit(self):
        self.assertEqual(gd._format_area("100"), "100 m²")
        self.assertEqual(gd._format_area("85.4"), "85.4 m²")

    def test_format_area_preserves_unit(self):
        self.assertEqual(gd._format_area("100 m²"), "100 m²")
        self.assertEqual(gd._format_area("100 m2"), "100 m2")

    def test_format_area_empty(self):
        self.assertEqual(gd._format_area(""), "")
        self.assertEqual(gd._format_area(None), "")


if __name__ == "__main__":
    unittest.main()
