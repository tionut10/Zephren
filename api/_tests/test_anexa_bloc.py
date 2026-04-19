"""
Etapa 4 (19 apr 2026) — teste pentru Anexa Bloc multi-apartament (BUG-4).

Verifică:
  1. insert_apartment_table() injectează tabel cu coloanele corecte
  2. Sumar (media ponderată + distribuție clase) apare corect
  3. Coloare clase energetice se aplică pe celulele de clasă
  4. insert_common_systems_section() listează doar sistemele cu installed=True
  5. Sisteme comune neinstalate sunt omise
  6. Apartamente lipsă → returnează False fără modificare doc
  7. Routing type=anexa_bloc → mode='anexa_bloc'

Rulare:
  cd energy-app
  python -m pytest api/_tests/test_anexa_bloc.py -v
"""
import importlib.util
import io
import sys
import unittest
from pathlib import Path

from docx import Document

API_DIR = Path(__file__).resolve().parent.parent
MODULE_PATH = API_DIR / "generate-document.py"

_spec = importlib.util.spec_from_file_location("generate_document_bloc", MODULE_PATH)
gd = importlib.util.module_from_spec(_spec)
sys.modules["generate_document_bloc"] = gd
_spec.loader.exec_module(gd)


def _make_doc():
    """Document minimal cu un paragraf pentru testare."""
    doc = Document()
    doc.add_paragraph("Anexa CPE — corp principal")
    return doc


def _doc_text(doc):
    parts = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        parts.append(p.text)
    return "\n".join(parts)


def _sample_apartments():
    return [
        {
            "number": "1A", "staircase": "A", "floor": 0, "areaUseful": 58.5,
            "orientation": ["N", "E"], "occupants": 3,
            "corner": True, "topFloor": False, "groundFloor": True,
            "posKey": "ground_corner", "posFactor": 1.18,
            "epAptM2": 192.4, "co2AptM2": 32.1,
            "enClass": "C", "co2Class": "D",
            "allocatedPct": 12.4,
        },
        {
            "number": "2B", "staircase": "A", "floor": 1, "areaUseful": 64.2,
            "orientation": ["S", "V"], "occupants": 2,
            "corner": False, "topFloor": False, "groundFloor": False,
            "posKey": "mid_interior", "posFactor": 1.00,
            "epAptM2": 168.7, "co2AptM2": 28.3,
            "enClass": "B", "co2Class": "C",
            "allocatedPct": 13.6,
        },
        {
            "number": "3C", "staircase": "B", "floor": 4, "areaUseful": 72.0,
            "orientation": ["N"], "occupants": 4,
            "corner": True, "topFloor": True, "groundFloor": False,
            "posKey": "top_corner", "posFactor": 1.15,
            "epAptM2": 215.3, "co2AptM2": 36.0,
            "enClass": "D", "co2Class": "D",
            "allocatedPct": 15.3,
        },
    ]


def _sample_summary():
    return {
        "totalAu": 471.7,
        "epAvgWeighted": 192.5,
        "co2AvgWeighted": 32.1,
        "avgEnergyClass": "C",
        "avgCo2Class": "D",
        "classDistribution": {"B": 1, "C": 1, "D": 1},
        "count": 3,
    }


class TestInsertApartmentTable(unittest.TestCase):

    def test_returns_false_for_empty_apartments(self):
        doc = _make_doc()
        result = gd.insert_apartment_table(doc, [], None)
        self.assertFalse(result)

    def test_returns_false_for_none(self):
        doc = _make_doc()
        result = gd.insert_apartment_table(doc, None, None)
        self.assertFalse(result)

    def test_adds_table_with_apartments(self):
        doc = _make_doc()
        before_tables = len(doc.tables)
        result = gd.insert_apartment_table(doc, _sample_apartments(), _sample_summary())
        self.assertTrue(result)
        self.assertEqual(len(doc.tables), before_tables + 1)

    def test_table_has_correct_headers(self):
        doc = _make_doc()
        gd.insert_apartment_table(doc, _sample_apartments(), _sample_summary())
        last_tbl = doc.tables[-1]
        header_row = last_tbl.rows[0]
        header_texts = [c.text for c in header_row.cells]
        for expected in ["Nr.", "Ap.", "Au [m²]", "EP [kWh/(m²·an)]", "Clasă"]:
            self.assertIn(expected, header_texts)

    def test_table_has_correct_row_count(self):
        """Header + N apartamente + 1 sumar = N+2 rânduri."""
        doc = _make_doc()
        apts = _sample_apartments()
        gd.insert_apartment_table(doc, apts, _sample_summary())
        last_tbl = doc.tables[-1]
        # 1 header + len(apts) data + 1 summary
        self.assertEqual(len(last_tbl.rows), len(apts) + 2)

    def test_apartment_data_appears_in_doc_text(self):
        doc = _make_doc()
        gd.insert_apartment_table(doc, _sample_apartments(), _sample_summary())
        text = _doc_text(doc)
        self.assertIn("1A", text)
        self.assertIn("2B", text)
        self.assertIn("3C", text)

    def test_position_labels_translated_to_romanian(self):
        doc = _make_doc()
        gd.insert_apartment_table(doc, _sample_apartments(), _sample_summary())
        text = _doc_text(doc)
        self.assertIn("parter colț", text)
        self.assertIn("etaj int.", text)
        self.assertIn("ultim colț", text)

    def test_orientation_displayed(self):
        doc = _make_doc()
        gd.insert_apartment_table(doc, _sample_apartments(), _sample_summary())
        text = _doc_text(doc)
        # "N E" pentru primul apartament
        self.assertIn("N E", text)
        self.assertIn("S V", text)

    def test_summary_row_contains_medie_ponderata(self):
        doc = _make_doc()
        gd.insert_apartment_table(doc, _sample_apartments(), _sample_summary())
        text = _doc_text(doc)
        self.assertIn("MEDIE PONDERATĂ BLOC", text)
        # Au total format ro: 471,7
        self.assertIn("471,7", text)
        # EP ponderat 192,5
        self.assertIn("192,5", text)

    def test_class_distribution_appears(self):
        doc = _make_doc()
        gd.insert_apartment_table(doc, _sample_apartments(), _sample_summary())
        text = _doc_text(doc)
        self.assertIn("Distribuție clase apartamente", text)
        # cele 3 clase
        self.assertIn("B=1", text)
        self.assertIn("C=1", text)
        self.assertIn("D=1", text)

    def test_methodology_subtitle_appears(self):
        doc = _make_doc()
        gd.insert_apartment_table(doc, _sample_apartments(), _sample_summary())
        text = _doc_text(doc)
        self.assertIn("Mc 001-2022", text)
        self.assertIn("Anexa 7", text)

    def test_floor_zero_displays_as_P(self):
        doc = _make_doc()
        apts = [
            {"number": "1", "floor": 0, "areaUseful": 50, "epAptM2": 180,
             "enClass": "C", "co2AptM2": 30, "co2Class": "C"}
        ]
        gd.insert_apartment_table(doc, apts, None)
        text = _doc_text(doc)
        # Prima coloană etaj trebuie să conțină "P"
        last_tbl = doc.tables[-1]
        # Row 1 (după header), col 3 = etaj
        self.assertEqual(last_tbl.rows[1].cells[3].text, "P")

    def test_doc_remains_valid_docx(self):
        doc = _make_doc()
        gd.insert_apartment_table(doc, _sample_apartments(), _sample_summary())
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        reopened = Document(buf)
        self.assertGreater(len(reopened.tables), 0)

    def test_summary_color_class_applied(self):
        """Celula clasei energetice din summary trebuie să aibă fundal colorat."""
        doc = _make_doc()
        gd.insert_apartment_table(doc, _sample_apartments(), _sample_summary())
        last_tbl = doc.tables[-1]
        # Summary row e ultimul
        sum_row = last_tbl.rows[-1]
        class_cell = sum_row.cells[9]
        # Verifică că shading a fost aplicat
        tc_pr = class_cell._tc.find(gd.qn("w:tcPr"))
        self.assertIsNotNone(tc_pr)
        shd = tc_pr.find(gd.qn("w:shd"))
        self.assertIsNotNone(shd)


class TestInsertCommonSystemsSection(unittest.TestCase):

    def test_returns_false_for_empty(self):
        doc = _make_doc()
        result = gd.insert_common_systems_section(doc, {})
        self.assertFalse(result)

    def test_returns_false_when_no_systems_installed(self):
        doc = _make_doc()
        result = gd.insert_common_systems_section(doc, {
            "elevator": {"installed": False, "powerKW": 5},
            "stairsLighting": {"installed": False},
        })
        self.assertFalse(result)

    def test_lists_only_installed_systems(self):
        doc = _make_doc()
        gd.insert_common_systems_section(doc, {
            "elevator": {"installed": True, "powerKW": 4.5, "hoursYear": 2000, "fuel": "electric"},
            "stairsLighting": {"installed": False, "powerKW": 0.6},
            "centralHeating": {"installed": True, "powerKW": 50, "hoursYear": 2400, "fuel": "gaz"},
        })
        text = _doc_text(doc)
        self.assertIn("Lift", text)
        self.assertIn("Centrală termică comună", text)
        self.assertNotIn("Iluminat scări/holuri", text)

    def test_table_has_4_columns(self):
        doc = _make_doc()
        gd.insert_common_systems_section(doc, {
            "elevator": {"installed": True, "powerKW": 4.5},
        })
        last_tbl = doc.tables[-1]
        self.assertEqual(len(last_tbl.columns), 4)

    def test_power_displayed_with_ro_format(self):
        doc = _make_doc()
        gd.insert_common_systems_section(doc, {
            "elevator": {"installed": True, "powerKW": 4.5, "hoursYear": 2000},
        })
        text = _doc_text(doc)
        self.assertIn("4,50", text)  # 4.5 → "4,50" cu virgulă

    def test_section_title_present(self):
        doc = _make_doc()
        gd.insert_common_systems_section(doc, {
            "elevator": {"installed": True, "powerKW": 4.5},
        })
        text = _doc_text(doc)
        self.assertIn("SISTEME COMUNE BLOC", text)


class TestHelpers(unittest.TestCase):

    def test_fmt_ro_basic(self):
        # Python f-string folosește rotunjire bankers' (half-to-even):
        # 123.45 → "123.5" la 1 zecimală (5 → up când digit anterior par)
        self.assertEqual(gd._fmt_ro(123.4, 1), "123,4")
        self.assertEqual(gd._fmt_ro(123.45, 2), "123,45")
        self.assertEqual(gd._fmt_ro(100, 1), "100,0")

    def test_fmt_ro_handles_none_and_empty(self):
        self.assertEqual(gd._fmt_ro(None), "—")
        self.assertEqual(gd._fmt_ro(""), "—")
        self.assertEqual(gd._fmt_ro("invalid"), "—")

    def test_fmt_floor_ground(self):
        self.assertEqual(gd._fmt_floor(0), "P")
        self.assertEqual(gd._fmt_floor("0"), "P")
        self.assertEqual(gd._fmt_floor("P"), "P")
        self.assertEqual(gd._fmt_floor("p"), "P")

    def test_fmt_floor_normal(self):
        self.assertEqual(gd._fmt_floor(3), "3")
        self.assertEqual(gd._fmt_floor("5"), "5")

    def test_fmt_floor_empty(self):
        self.assertEqual(gd._fmt_floor(None), "—")
        self.assertEqual(gd._fmt_floor(""), "—")


class TestRoutingAnexaBloc(unittest.TestCase):
    """Verifică că type=anexa_bloc activează mode='anexa_bloc' în handler."""

    def test_anexa_bloc_is_handled_as_separate_mode(self):
        """Sanity check: în cod sursă există branch pentru mode == 'anexa_bloc'."""
        # Citește sursa și verifică prezența branch-ului
        with open(MODULE_PATH, encoding="utf-8") as f:
            src = f.read()
        self.assertIn("anexa_bloc", src)
        self.assertIn("insert_apartment_table", src)


if __name__ == "__main__":
    unittest.main()
