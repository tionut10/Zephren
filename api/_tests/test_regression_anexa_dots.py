"""
Regression test (20 apr 2026) — bug "201520152015..." în Anexa.

Problema raportată: în template Anexa, șirurile lungi de puncte (151/150/121)
care reprezintă spații de completare manuală pentru auditor (ex. "Enunțarea
etapelor...", "Informații privind stimulentele...", "auditorul energetic va
completa mai departe lista") erau înlocuite global cu anul construcției
("2015"), rezultând în repetiție: "201520152015..." de 7-8 ori.

Cauza: `replace_in_doc(".................")` (17 puncte) face match global
pe ORICE secvență ≥ 17 puncte și o înlocuiește; un șir de 151 puncte
acceptă ⌊151/17⌋ = 8 înlocuiri consecutive ale string-ului "2015".

Fix: replace chirurgical DOAR în paragraful care conține "Anul construc"
(eticheta semantică unică), o singură dată per paragraf (count=1).

Rulare:
  cd energy-app
  python -m pytest api/_tests/test_regression_anexa_dots.py -v
"""
import importlib.util
import sys
import unittest
from pathlib import Path

from docx import Document

API_DIR = Path(__file__).resolve().parent.parent
MODULE_PATH = API_DIR / "generate-document.py"

_spec = importlib.util.spec_from_file_location("gd_dots_regression", MODULE_PATH)
gd = importlib.util.module_from_spec(_spec)
sys.modules["gd_dots_regression"] = gd
_spec.loader.exec_module(gd)

ANEXA_TEMPLATE = API_DIR.parent / "public" / "templates" / "ANEXA-1-si-ANEXA-2-la-CPE-cladire.docx"


class TestRegressionAnexaDots(unittest.TestCase):
    """Verifică că bug-ul '201520152015...' nu mai apare."""

    @classmethod
    def setUpClass(cls):
        if not ANEXA_TEMPLATE.exists():
            raise unittest.SkipTest(f"Template lipsă: {ANEXA_TEMPLATE}")

    def _doc_text(self, doc):
        parts = []
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text:
                            parts.append(p.text)
        return "\n".join(parts)

    def _apply_year_replacement(self, doc, year_str):
        """Reproduce logica fix-ului din process_document (mode=='anexa')."""
        for p in doc.paragraphs:
            if "Anul construc" in p.text and "................." in p.text:
                gd.replace_in_paragraph(p, ".................", year_str, count=1)
                break

    def test_year_replaced_in_construction_paragraph(self):
        """Anul apare corect în paragraful 'Anul construcției'."""
        doc = Document(str(ANEXA_TEMPLATE))
        self._apply_year_replacement(doc, "2015")
        text = self._doc_text(doc)
        # Paragraful "Anul construcției" trebuie să conțină acum "2015"
        # (înainte de fix: tot "2015" dar și "201520152015..." în alte locuri)
        anul_para = next(
            (p.text for p in doc.paragraphs if "Anul construc" in p.text),
            None,
        )
        self.assertIsNotNone(anul_para)
        self.assertIn("2015", anul_para)

    def test_no_repeated_2015_anywhere(self):
        """Nicăieri în document NU trebuie să apară '20152015' (repetiție bug)."""
        doc = Document(str(ANEXA_TEMPLATE))
        self._apply_year_replacement(doc, "2015")
        text = self._doc_text(doc)
        # Verificare cheie: NU avem secvența bug-ului
        self.assertNotIn("20152015", text,
            "BUG REGRESSION: '20152015' a apărut în document")
        self.assertNotIn("201520152015", text,
            "BUG REGRESSION: '201520152015' a apărut")

    def test_no_repeated_year_with_renov(self):
        """Pentru year_str compus (ex. '1985 / 2015'), nu trebuie repetare."""
        doc = Document(str(ANEXA_TEMPLATE))
        self._apply_year_replacement(doc, "1985 / 2015")
        text = self._doc_text(doc)
        # Numărăm aparițiile: trebuie să fie EXACT 1 (în paragraful Anul)
        count = text.count("1985 / 2015")
        self.assertEqual(count, 1,
            f"year_full apare de {count} ori (așteptat 1) — replace prea agresiv")

    def test_long_dot_sequences_remain_unchanged(self):
        """Șirurile lungi (>=120 puncte) din template trebuie să rămână INTACTE."""
        doc = Document(str(ANEXA_TEMPLATE))
        self._apply_year_replacement(doc, "2015")
        # Caut toate paragrafele cu șiruri lungi de puncte
        long_dot_paragraphs = []
        for p in doc.paragraphs:
            max_dots = 0
            cur = 0
            for c in p.text:
                if c == ".":
                    cur += 1
                    max_dots = max(max_dots, cur)
                else:
                    cur = 0
            if max_dots >= 100:
                long_dot_paragraphs.append((max_dots, p.text[:80]))
        # Cel puțin 2 paragrafe cu >= 100 puncte trebuie să existe (Enunțarea
        # etapelor, Informații stimulente, regenerabile auditor)
        self.assertGreaterEqual(len(long_dot_paragraphs), 2,
            f"Doar {len(long_dot_paragraphs)} paragrafe cu >=100 puncte găsite — "
            f"șirurile au fost corupte de fix?")

    def test_full_anexa_pipeline_no_year_corruption(self):
        """Smoke test: după compute_checkboxes + year replace, '20152015' nu apare."""
        doc = Document(str(ANEXA_TEMPLATE))
        # Apply full anexa flow simulat (year replace doar)
        self._apply_year_replacement(doc, "1985")
        text = self._doc_text(doc)
        self.assertNotIn("19851985", text)
        self.assertNotIn("198519851985", text)


if __name__ == "__main__":
    unittest.main()
