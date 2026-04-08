"""
Vercel Python Serverless Function — Generare Raport Audit Energetic DOCX
Produce un document Word complet (.docx) cu toate secțiunile auditului energetic.
Fără template necesar — document generat de la zero cu python-docx.
"""
from http.server import BaseHTTPRequestHandler
import json, io, base64
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_bg(cell, color_hex):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor.from_string(color)
    return h

def add_kv_row(table, label, value, bold_val=False):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[1].text = str(value)
    if bold_val:
        row.cells[1].paragraphs[0].runs[0].bold = True

def generate_audit_report(data):
    doc = Document()

    # ── Stiluri globale ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # Margini
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    b = data.get('building', {})
    inst = data.get('instSummary', {})
    renew = data.get('renewSummary', {})
    aud = data.get('auditor', {})
    opaque = data.get('opaqueElements', [])
    glazing = data.get('glazingElements', [])
    bridges = data.get('thermalBridges', [])
    measured = data.get('measuredConsumption', {})
    today = datetime.date.today().strftime('%d.%m.%Y')

    # ════════════════════════════════════════
    # PAGINA DE TITLU
    # ════════════════════════════════════════
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run('RAPORT DE AUDIT ENERGETIC')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # blue-800

    doc.add_paragraph()

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = sub_para.add_run(b.get('address', 'Adresă necunoscută'))
    sub.font.size = Pt(13)
    sub.bold = True

    doc.add_paragraph()

    # Tabel informații de bază
    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = 'Table Grid'
    for label, value in [
        ('Adresă imobil', b.get('address','—')),
        ('Localitate / Județ', f"{b.get('city','—')} / {b.get('county','—')}"),
        ('Categorie clădire', b.get('category','—')),
        ('Suprafață utilă (Au)', f"{b.get('areaUseful','—')} m²"),
        ('An construcție', b.get('yearBuilt','—')),
        ('Auditor energetic', aud.get('name','—')),
        ('Nr. certificat auditor', aud.get('certNr','—')),
        ('Firma / CUI', aud.get('firma','—')),
        ('Data raportului', today),
    ]:
        add_kv_row(info_table, label, value)
    info_table.columns[0].width = Cm(7)
    info_table.columns[1].width = Cm(9)

    doc.add_page_break()

    # ════════════════════════════════════════
    # 1. REZUMAT EXECUTIV
    # ════════════════════════════════════════
    add_heading(doc, '1. Rezumat Executiv', level=1)

    ep = inst.get('ep_total_m2', 0)
    co2 = inst.get('co2_total_m2', 0)
    cls = data.get('energyClass', '—')
    rer = renew.get('rer', 0)

    # Tabel KPI
    kpi_table = doc.add_table(rows=0, cols=4)
    kpi_table.style = 'Table Grid'
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = kpi_table.add_row()
    for i, h in enumerate(['Indicator', 'Valoare', 'Unitate', 'Evaluare']):
        hdr.cells[i].text = h
        hdr.cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(hdr.cells[i], '1E40AF')
        hdr.cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

    def ep_eval(ep_val):
        if ep_val < 50: return 'Excelent (A+)'
        if ep_val < 100: return 'Bun (A/B)'
        if ep_val < 200: return 'Mediu (C/D)'
        if ep_val < 300: return 'Slab (E/F)'
        return 'Foarte slab (G)'

    for label, val, unit, ev in [
        ('Energie primară (EP)', f"{ep:.1f}", 'kWh/m²·an', ep_eval(ep)),
        ('Emisii CO₂', f"{co2:.1f}", 'kg CO₂/m²·an', '—'),
        ('Clasă energetică', cls, '—', cls if cls in ['A+','A','B'] else f'Renovare recomandată'),
        ('Cotă energie regenerabilă', f"{rer:.1f}", '%', 'Conform nZEB' if rer >= 30 else 'Sub 30% (nZEB)'),
    ]:
        row = kpi_table.add_row()
        for i, txt in enumerate([label, val, unit, ev]):
            row.cells[i].text = txt

    doc.add_paragraph()
    ep_masurat = measured.get('ep_masurat', 0)
    if ep_masurat > 0:
        diff_pct = (ep_masurat - ep) / max(ep, 1) * 100
        p = doc.add_paragraph()
        run = p.add_run(f'Reconciliere consum facturat vs calculat: ')
        run.bold = True
        p.add_run(f'EP măsurat = {ep_masurat:.1f} kWh/m²·an, EP calculat = {ep:.1f} kWh/m²·an, diferență = {diff_pct:+.1f}%.')

    doc.add_page_break()

    # ════════════════════════════════════════
    # 2. DATE GENERALE CLĂDIRE
    # ════════════════════════════════════════
    add_heading(doc, '2. Date Generale ale Clădirii', level=1)

    bldg_table = doc.add_table(rows=0, cols=2)
    bldg_table.style = 'Table Grid'
    bldg_fields = [
        ('Categorie / Destinație', b.get('category','—')),
        ('Structură rezistență', b.get('structure','—')),
        ('Regim de înălțime', b.get('floors','—')),
        ('Nr. unități locative', b.get('units','—')),
        ('Suprafață utilă Au', f"{b.get('areaUseful','—')} m²"),
        ('Suprafață construită Ac', f"{b.get('areaTotal','—')} m²"),
        ('Înălțime etaj', f"{b.get('heightFloor','—')} m"),
        ('An construcție', b.get('yearBuilt','—')),
        ('Localitate', b.get('city','—')),
        ('Județ', b.get('county','—')),
        ('Zona climatică', b.get('climateZone','—')),
    ]
    for label, val in bldg_fields:
        add_kv_row(bldg_table, label, val)
    bldg_table.columns[0].width = Cm(7)
    bldg_table.columns[1].width = Cm(9)

    # ════════════════════════════════════════
    # 3. ANVELOPA CLĂDIRII
    # ════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, '3. Anvelopa Clădirii', level=1)

    add_heading(doc, '3.1 Elemente opace', level=2)
    if opaque:
        op_table = doc.add_table(rows=1, cols=4)
        op_table.style = 'Table Grid'
        hdr = op_table.rows[0]
        for i, h in enumerate(['Element', 'Suprafață (m²)', 'U [W/m²K]', 'Evaluare']):
            hdr.cells[i].text = h
            hdr.cells[i].paragraphs[0].runs[0].bold = True
        for el in opaque:
            u = float(el.get('U_calc', el.get('u', 0)) or 0)
            ev = '✓ Bun' if u < 0.5 else ('⚠ Acceptabil' if u < 0.8 else '✗ Slab')
            row = op_table.add_row()
            for i, txt in enumerate([el.get('name','—'), str(el.get('area','—')), f"{u:.3f}", ev]):
                row.cells[i].text = txt
    else:
        doc.add_paragraph('Nu sunt înregistrate elemente opace.')

    add_heading(doc, '3.2 Elemente vitrate', level=2)
    if glazing:
        gl_table = doc.add_table(rows=1, cols=5)
        gl_table.style = 'Table Grid'
        hdr = gl_table.rows[0]
        for i, h in enumerate(['Element', 'Suprafață (m²)', 'U [W/m²K]', 'g [-]', 'Orientare']):
            hdr.cells[i].text = h
            hdr.cells[i].paragraphs[0].runs[0].bold = True
        for el in glazing:
            row = gl_table.add_row()
            for i, txt in enumerate([el.get('name','—'), str(el.get('area','—')), str(el.get('u','—')), str(el.get('g','—')), el.get('orientation','—')]):
                row.cells[i].text = txt
    else:
        doc.add_paragraph('Nu sunt înregistrate elemente vitrate.')

    add_heading(doc, '3.3 Punți termice', level=2)
    if bridges:
        br_table = doc.add_table(rows=1, cols=3)
        br_table.style = 'Table Grid'
        hdr = br_table.rows[0]
        for i, h in enumerate(['Denumire', 'ψ [W/mK]', 'Lungime (m)']):
            hdr.cells[i].text = h
            hdr.cells[i].paragraphs[0].runs[0].bold = True
        for br in bridges:
            row = br_table.add_row()
            for i, txt in enumerate([br.get('name','—'), str(br.get('psi','—')), str(br.get('length','—'))]):
                row.cells[i].text = txt
    else:
        doc.add_paragraph('Nu sunt înregistrate punți termice.')

    # ════════════════════════════════════════
    # 4. SISTEME TEHNICE
    # ════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, '4. Sisteme Tehnice', level=1)

    sys_data = data.get('systems', {})
    sys_table = doc.add_table(rows=0, cols=2)
    sys_table.style = 'Table Grid'
    for label, val in [
        ('Sursă încălzire', sys_data.get('heatSource','—')),
        ('Combustibil', sys_data.get('fuel','—')),
        ('Randament generator (%)', sys_data.get('boilerEta','—')),
        ('Sistem emisie', sys_data.get('emissionSystem','—')),
        ('Control termic', sys_data.get('controlType','—')),
        ('Sursă ACM', sys_data.get('acmSource','—')),
        ('Sistem răcire', sys_data.get('coolingSystem','—')),
        ('Ventilare', sys_data.get('ventilationType','—')),
        ('Iluminat', sys_data.get('lightingType','—')),
    ]:
        add_kv_row(sys_table, label, str(val) if val else '—')
    sys_table.columns[0].width = Cm(7)
    sys_table.columns[1].width = Cm(9)

    # ════════════════════════════════════════
    # 5. BILANȚ ENERGETIC
    # ════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, '5. Bilanț Energetic', level=1)
    doc.add_paragraph('Calculat conform metodologiei EN ISO 13790 (calcul lunar static), Mc 001-2022.')

    en_table = doc.add_table(rows=1, cols=3)
    en_table.style = 'Table Grid'
    hdr = en_table.rows[0]
    for i, h in enumerate(['Categorie consum', 'EP [kWh/m²·an]', 'Pondere (%)']):
        hdr.cells[i].text = h
        hdr.cells[i].paragraphs[0].runs[0].bold = True

    ep_total = inst.get('ep_total_m2', 0) or 0
    for label, key in [
        ('Încălzire', 'ep_heating_m2'), ('Răcire', 'ep_cooling_m2'),
        ('Apă caldă menajeră', 'ep_dhw_m2'), ('Iluminat', 'ep_lighting_m2'),
        ('Ventilare mecanică', 'ep_vent_m2'),
    ]:
        val = inst.get(key, 0) or 0
        if val > 0:
            pct = val / max(ep_total, 1) * 100
            row = en_table.add_row()
            for i, txt in enumerate([label, f"{val:.1f}", f"{pct:.1f}%"]):
                row.cells[i].text = txt

    total_row = en_table.add_row()
    total_row.cells[0].text = 'TOTAL EP'
    total_row.cells[0].paragraphs[0].runs[0].bold = True
    total_row.cells[1].text = f"{ep_total:.1f}"
    total_row.cells[1].paragraphs[0].runs[0].bold = True
    total_row.cells[2].text = "100%"

    # ════════════════════════════════════════
    # 6. SURSE REGENERABILE
    # ════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, '6. Surse de Energie Regenerabilă', level=1)

    ren_table = doc.add_table(rows=0, cols=2)
    ren_table.style = 'Table Grid'
    for label, val in [
        ('Cotă energie regenerabilă (RER)', f"{renew.get('rer',0):.1f}%"),
        ('Producție solară termică', f"{renew.get('solar_thermal_kWh',0):.0f} kWh/an"),
        ('Producție fotovoltaică', f"{renew.get('pv_annual_kWh',0):.0f} kWh/an"),
        ('Putere PV instalată', f"{renew.get('pv_peak_kWp',0):.1f} kWp"),
        ('Energie pompă de căldură', f"{renew.get('hp_kWh',0):.0f} kWh/an"),
        ('Conformitate nZEB (RER ≥ 30%)', 'DA' if renew.get('rer',0) >= 30 else 'NU'),
    ]:
        add_kv_row(ren_table, label, str(val))
    ren_table.columns[0].width = Cm(9)
    ren_table.columns[1].width = Cm(7)

    # ════════════════════════════════════════
    # 7. RECOMANDĂRI
    # ════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, '7. Recomandări de Reabilitare', level=1)

    recs = []
    if ep_total > 200:
        recs.append(('Termoizolație pereți exteriori (EPS ≥ 10cm)', 'Reducere estimată EP: 15-25%', 'Înaltă'))
    if ep_total > 150:
        recs.append(('Înlocuire ferestre cu triplu vitraj Low-E', 'Reducere estimată EP: 8-15%', 'Medie'))
    if inst.get('ep_dhw_m2', 0) > 20:
        recs.append(('Panouri solare termice pentru ACM (4-6 m²)', 'Acoperire 60-70% necesar ACM', 'Medie'))
    if renew.get('rer', 0) < 30:
        recs.append(('Sistem fotovoltaic 3-5 kWp', 'Creștere RER la ≥30% (nZEB)', 'Înaltă'))
    if renew.get('rer', 0) >= 30 and ep_total < 130:
        recs.append(('Clădirea este aproape de conformanța nZEB', 'Optimizați sistemul de control (BACS cls. B)', 'Scăzută'))

    if recs:
        rec_table = doc.add_table(rows=1, cols=3)
        rec_table.style = 'Table Grid'
        hdr = rec_table.rows[0]
        for i, h in enumerate(['Măsură recomandată', 'Impact estimat', 'Prioritate']):
            hdr.cells[i].text = h
            hdr.cells[i].paragraphs[0].runs[0].bold = True
        for rec_name, impact, prio in recs:
            row = rec_table.add_row()
            for i, txt in enumerate([rec_name, impact, prio]):
                row.cells[i].text = txt
    else:
        doc.add_paragraph('✓ Clădirea prezintă o performanță energetică bună. Nu sunt recomandări majore de reabilitare.')

    # ════════════════════════════════════════
    # 8. SEMNĂTURI
    # ════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, '8. Certificare și Semnături', level=1)

    doc.add_paragraph(
        f"Prezentul raport de audit energetic a fost întocmit conform Legii 372/2005 republicată, "
        f"HG 917/2021 și normativului Mc 001-2022."
    )
    doc.add_paragraph()

    sign_table = doc.add_table(rows=3, cols=2)
    sign_table.style = 'Table Grid'
    sign_table.cell(0, 0).text = 'Auditor energetic'
    sign_table.cell(0, 1).text = 'Beneficiar / Proprietar'
    sign_table.cell(0, 0).paragraphs[0].runs[0].bold = True
    sign_table.cell(0, 1).paragraphs[0].runs[0].bold = True
    sign_table.cell(1, 0).text = f"Nume: {aud.get('name','')}"
    sign_table.cell(1, 0).add_paragraph(f"Certificat nr.: {aud.get('certNr','')}")
    sign_table.cell(1, 0).add_paragraph(f"Firma: {aud.get('firma','')}")
    sign_table.cell(1, 1).text = 'Nume: ___________________________'
    sign_table.cell(2, 0).text = f'\n\nSemnătură și ștampilă:\n\n\n'
    sign_table.cell(2, 1).text = f'\n\nSemnătură:\n\n\n'
    sign_table.cell(2, 0).add_paragraph(f'Data: {today}')
    sign_table.cell(2, 1).add_paragraph(f'Data: _______________')

    # ── Generare bytes ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(length)
            data = json.loads(body)

            docx_bytes = generate_audit_report(data)
            b64 = base64.b64encode(docx_bytes).decode('utf-8')

            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(json.dumps({'docx': b64, 'filename': 'raport_audit_energetic.docx'}).encode())

        except Exception as e:
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(json.dumps({'error': str(e)}).encode())

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type, Authorization')
        self.end_headers()
