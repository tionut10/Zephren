"""
Vercel Python Serverless Function — Generare CPE DOCX cu python-docx.
Primește template DOCX (base64) + date JSON, returnează DOCX completat.
Lucrează direct pe template-urile originale MDLPA fără pre-procesare.
"""
from http.server import BaseHTTPRequestHandler
import json, io, base64, re, copy
from docx import Document
from docx.shared import Inches
from lxml import etree

# ═══════════════════════════════════════════════════════
# CONSTANTE — scalele EP și CO2 din template-uri (Mc 001-2022)
# Template-urile rezidențiale folosesc varianta _cool
# ═══════════════════════════════════════════════════════
EP_TEMPLATE_SCALES = {
    "RI": [91, 129, 257, 390, 522, 652, 783],
    "RC": [73, 101, 198, 297, 396, 495, 595],
    "RA": [73, 101, 198, 297, 396, 495, 595],
    "BI": [68, 97, 193, 302, 410, 511, 614],
    "ED": [55, 78, 157, 248, 340, 425, 510],
    "SA": [130, 190, 380, 570, 760, 950, 1140],
    "HC": [85, 120, 240, 370, 500, 625, 750],
    "CO": [75, 107, 213, 330, 447, 558, 670],
    "SP": [70, 100, 200, 310, 420, 525, 630],
    "AL": [68, 97, 193, 302, 410, 511, 614],
}

CO2_TEMPLATE_SCALES = {
    "RI": [16.1, 22.8, 45.5, 70.1, 94.8, 118.4, 142.1],
    "RC": [12.7, 17.6, 34.6, 52.2, 69.9, 87.4, 104.9],
    "RA": [12.7, 17.6, 34.6, 52.2, 69.9, 87.4, 104.9],
    "BI": [10.4, 14.8, 29.7, 46.1, 62.4, 77.8, 93.4],
    "ED": [8.5, 12.0, 24.0, 37.0, 50.0, 62.5, 75.0],
    "SA": [19.0, 27.0, 54.0, 83.0, 112.0, 140.0, 168.0],
    "HC": [13.0, 18.5, 37.0, 57.0, 77.0, 96.0, 115.0],
    "CO": [11.5, 16.4, 32.8, 50.5, 68.5, 85.5, 102.5],
    "SP": [10.7, 15.3, 30.5, 47.5, 64.5, 80.5, 96.5],
    "AL": [10.4, 14.8, 29.7, 46.1, 62.4, 77.8, 93.4],
}

NSMAP = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def qn(tag):
    """Qualified name helper for lxml."""
    prefix, local = tag.split(":")
    return "{%s}%s" % (NSMAP[prefix], local)


# ═══════════════════════════════════════════════════════
# REPLACE TEXT ACROSS RUNS — păstrează formatarea
# ═══════════════════════════════════════════════════════
def replace_in_paragraph(para, old_text, new_text, count=0):
    """Replace old_text with new_text in a paragraph, across split runs.
    Returns number of replacements made. If count>0, limits replacements."""
    full = para.text
    if old_text not in full:
        return 0

    # Build a map of character positions to runs
    runs = para.runs
    if not runs:
        return 0

    char_map = []  # [(run_idx, char_idx_in_run)]
    for ri, run in enumerate(runs):
        for ci in range(len(run.text)):
            char_map.append((ri, ci))

    replacements = 0
    offset = 0
    while True:
        pos = full.find(old_text, offset)
        if pos == -1:
            break
        if count > 0 and replacements >= count:
            break

        # Find which runs this spans
        start_ri, start_ci = char_map[pos]
        end_pos = pos + len(old_text) - 1
        end_ri, end_ci = char_map[end_pos]

        # Put replacement text in the first run at the start position
        first_run = runs[start_ri]
        before = first_run.text[:start_ci]
        if start_ri == end_ri:
            after = first_run.text[end_ci + 1:]
            first_run.text = before + new_text + after
        else:
            first_run.text = before + new_text
            # Clear text from intermediate and last runs
            for ri in range(start_ri + 1, end_ri):
                runs[ri].text = ""
            last_run = runs[end_ri]
            last_run.text = last_run.text[end_ci + 1:]

        replacements += 1
        # Rebuild full text and char_map for next iteration
        full = para.text
        char_map = []
        for ri, run in enumerate(runs):
            for ci in range(len(run.text)):
                char_map.append((ri, ci))
        offset = pos + len(new_text)

    return replacements


def _iter_all_paragraphs(doc, include_txbx=True):
    """Generator: yields all Paragraph objects from body, tables, and optionally text boxes."""
    from docx.text.paragraph import Paragraph as DocxPara
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para
    if include_txbx:
        txbx_tag = qn("w:txbxContent")
        p_tag = qn("w:p")
        for txbx_elem in doc.element.body.iter(txbx_tag):
            for p_elem in txbx_elem.iter(p_tag):
                yield DocxPara(p_elem, None)


def replace_in_doc(doc, old_text, new_text, max_count=0):
    """Replace text across entire document (paragraphs + tables + text boxes)."""
    total = 0
    remaining = max_count
    for para in _iter_all_paragraphs(doc, include_txbx=True):
        n = replace_in_paragraph(para, old_text, new_text, remaining if max_count > 0 else 0)
        total += n
        if max_count > 0:
            remaining -= n
            if remaining <= 0:
                return total
    return total


def replace_in_txbx_only(doc, old_text, new_text):
    """Replace text ONLY in text boxes (w:txbxContent). Safe for numeric scale bars."""
    from docx.text.paragraph import Paragraph as DocxPara
    total = 0
    txbx_tag = qn("w:txbxContent")
    p_tag = qn("w:p")
    for txbx_elem in doc.element.body.iter(txbx_tag):
        for p_elem in txbx_elem.iter(p_tag):
            total += replace_in_paragraph(DocxPara(p_elem, None), old_text, new_text)
    return total


def replace_seq(doc, old_text, values):
    """Replace sequential occurrences of old_text with different values (body + tables + text boxes)."""
    idx = [0]

    def do_para(para):
        if idx[0] >= len(values):
            return
        full = para.text
        while old_text in full and idx[0] < len(values):
            replace_in_paragraph(para, old_text, values[idx[0]], count=1)
            idx[0] += 1
            full = para.text

    for para in _iter_all_paragraphs(doc, include_txbx=True):
        do_para(para)


# ═══════════════════════════════════════════════════════
# REPLACE SCALE THRESHOLDS — EP and CO2
# ═══════════════════════════════════════════════════════
def replace_scales(doc, category, new_ep_scale, new_co2_scale=None):
    """Replace EP/CO2 scale values in text boxes ONLY (safe — no risk of corrupting other numbers)."""
    old_ep = EP_TEMPLATE_SCALES.get(category, EP_TEMPLATE_SCALES["AL"])

    for i in range(7):
        if old_ep[i] != new_ep_scale[i]:
            replace_in_txbx_only(doc, str(old_ep[i]), str(new_ep_scale[i]))

    if new_co2_scale:
        old_co2 = CO2_TEMPLATE_SCALES.get(category, CO2_TEMPLATE_SCALES["AL"])
        for i in range(7):
            old_s = format_ro(old_co2[i], 1)
            new_s = format_ro(new_co2_scale[i], 1)
            if old_s != new_s:
                replace_in_txbx_only(doc, old_s, new_s)


# ═══════════════════════════════════════════════════════
# REPLACE CLASS INDICATORS — the colored arrows on scales
# Includes REPOSITIONING the shape vertically to the correct class row
# ═══════════════════════════════════════════════════════

# Vertical positions (EMU) for each class row on the scale
# Derived from template: class A posV=510540, class B posV=977265
# Spacing per class: 977265 - 510540 = 466725 EMU
_CLASS_SPACING = 466725
_CLASS_POS_V = {
    "A+": 510540 - _CLASS_SPACING,    #   43815
    "A":  510540,                      #  510540
    "B":  977265,                      #  977265
    "C":  977265 + _CLASS_SPACING,     # 1443990
    "D":  977265 + 2 * _CLASS_SPACING, # 1910715
    "E":  977265 + 3 * _CLASS_SPACING, # 2377440
    "F":  977265 + 4 * _CLASS_SPACING, # 2844165
    "G":  977265 + 5 * _CLASS_SPACING, # 3310890
}

# CO2 scale has same spacing but slightly different base
_CO2_CLASS_POS_V = {
    "A+": 963930 - _CLASS_SPACING * 2, #   30480
    "A":  963930 - _CLASS_SPACING,     #  497205
    "B":  963930,                      #  963930
    "C":  963930 + _CLASS_SPACING,     # 1430655
    "D":  963930 + 2 * _CLASS_SPACING, # 1897380
    "E":  963930 + 3 * _CLASS_SPACING, # 2364105
    "F":  963930 + 4 * _CLASS_SPACING, # 2830830
    "G":  963930 + 5 * _CLASS_SPACING, # 3297555
}


# Culorile REALE din imaginile de fundal ale template-urilor CPE MDLPA
# Extrase pixel-by-pixel din template-ul DOCX oficial (image1-16.jpeg)
_EP_CLASS_COLORS = {
    "A+": {"vml": "#009B00", "hex": "009B00"},  # verde închis
    "A":  {"vml": "#32C831", "hex": "32C831"},  # verde mediu
    "B":  {"vml": "#00FF00", "hex": "00FF00"},  # verde pur / lime
    "C":  {"vml": "#FFFF00", "hex": "FFFF00"},  # galben pur
    "D":  {"vml": "#F39C00", "hex": "F39C00"},  # chihlimbar
    "E":  {"vml": "#FF6400", "hex": "FF6400"},  # portocaliu
    "F":  {"vml": "#FE4101", "hex": "FE4101"},  # portocaliu-roșu
    "G":  {"vml": "#FE0000", "hex": "FE0000"},  # roșu pur
}

# Scala CO2 din template MDLPA — albastru (A+ → C) → gri (D → G)
_CO2_CLASS_COLORS = {
    "A+": {"vml": "#0000FE", "hex": "0000FE"},  # albastru închis
    "A":  {"vml": "#3265FF", "hex": "3265FF"},  # albastru mediu
    "B":  {"vml": "#009BFF", "hex": "009BFF"},  # albastru deschis
    "C":  {"vml": "#9CD2FF", "hex": "9CD2FF"},  # albastru foarte deschis
    "D":  {"vml": "#BEBEBE", "hex": "BEBEBE"},  # gri deschis
    "E":  {"vml": "#969696", "hex": "969696"},  # gri mediu
    "F":  {"vml": "#646464", "hex": "646464"},  # gri închis
    "G":  {"vml": "#333333", "hex": "333333"},  # gri foarte închis
}


def _text_color_for_bg(hex_color):
    """Alege alb sau negru pentru text pe baza luminanței relative WCAG a fundalului.
    Funcționează corect pentru AMBELE palete EP (verde/galben/roșu) și CO2 (albastru/gri)."""
    try:
        r = int(hex_color[0:2], 16) / 255.0
        g = int(hex_color[2:4], 16) / 255.0
        b = int(hex_color[4:6], 16) / 255.0
        # Linearizare sRGB
        def lin(c):
            return c / 12.92 if c <= 0.04045 else ((c + 0.055) / 1.055) ** 2.4
        lum = 0.2126 * lin(r) + 0.7152 * lin(g) + 0.0722 * lin(b)
        # Contrast cu alb: (lum+0.05)/0.05; contrast cu negru: 1.05/(lum+0.05)
        # Alege culoarea cu contrast mai mare
        return "FFFFFF" if (lum + 0.05) < 0.5 else "000000"
    except Exception:
        return "000000"


def _update_shape_color(shape_xml_str, color_info):
    """Update fill color in both DrawingML (a:srgbClr) and VML (fillcolor) representations."""
    result = shape_xml_str
    hex_color = color_info["hex"]
    vml_color = color_info["vml"]

    # Update DrawingML solidFill — first srgbClr occurrence (shape fill, not outline)
    result = re.sub(
        r'(<a:solidFill>\s*<a:srgbClr val=")[A-Fa-f0-9]{6}(")',
        r'\g<1>' + hex_color + r'\g<2>',
        result,
        count=1
    )

    # Update VML fillcolor attribute
    result = re.sub(
        r'fillcolor="[^"]*"',
        'fillcolor="' + vml_color + '"',
        result,
        count=1
    )

    return result


def _update_shape_pos_v(shape_xml_str, new_pos_v):
    """Update BOTH the modern (wp:anchor posOffset) and VML fallback (style top:) positions."""
    # 1. Update wp:anchor posOffset
    result = re.sub(
        r'(<wp:positionV[^>]*>[\s\S]*?<wp:posOffset>)-?\d+(</wp:posOffset>)',
        lambda m: m.group(1) + str(new_pos_v) + m.group(2),
        shape_xml_str
    )
    # 2. Update VML fallback style "top:XXpt" — convert EMU to pt (1pt = 12700 EMU)
    new_top_pt = round(new_pos_v / 12700, 1)
    result = re.sub(
        r'(style="[^"]*?)top:\s*-?[\d.]+pt',
        lambda m: m.group(1) + "top:" + str(new_top_pt) + "pt",
        result
    )
    return result


def replace_class_indicators(doc, ep_class_real, ep_class_ref, co2_class_real):
    """Replace class letter text AND reposition the arrow shapes on the scales."""
    import zipfile as _zf

    # Work directly on the XML string for precise shape manipulation
    xml_file = doc.part._element.getparent()
    # Access the raw XML via the document element
    body = doc.element.body
    # Serialize body to string for regex-based shape manipulation
    body_xml = etree.tostring(body, encoding="unicode")

    # Find mc:AlternateContent blocks with class indicators (H > 150000)
    alt_pattern = re.compile(r'(<mc:AlternateContent>[\s\S]*?</mc:AlternateContent>)')
    blocks = list(alt_pattern.finditer(body_xml))

    # Colectăm TOATE shape-urile indicator — atât textbox (cu litera) cât și pentagon (path)
    # Fiecare indicator are 2 shape-uri consecutive: textbox + pentagon (sau invers)
    text_indicators = []   # (match, letter, posH, posV) — textbox cu litera
    path_indicators = []   # (match, posH, posV) — pentagon/arrow fără text

    for m in blocks:
        content = m.group(1)
        letters = re.findall(r'<w:t[^>]*>([A-G]\+?)</w:t>', content)
        pos_h = re.search(r'positionH[^>]*>[\s\S]*?posOffset>(-?\d+)<', content)
        pos_v = re.search(r'positionV[^>]*>[\s\S]*?posOffset>(-?\d+)<', content)
        if not pos_h or not pos_v:
            continue
        h = int(pos_h.group(1))
        v = int(pos_v.group(1))
        is_path = 'coordsize' in content or '<v:path' in content

        if h > 150000 and len(letters) > 0 and len(letters[0]) <= 2 and not is_path:
            text_indicators.append((m, letters[0], h, v))
        elif is_path and -100000 <= h < 700000:
            # Pentagon/arrow shapes — range extins pentru a include EP pentagon (h≈380365)
            path_indicators.append((m, h, v))

    # Clasificare prin ORDINE în document:
    # Clădiri: 3 text indicators → EP_real(1), EP_ref(2), CO2(3)
    # Apartamente: 2 text indicators → EP_real(1), CO2(2)
    if len(text_indicators) >= 3:
        ep_indicators = text_indicators[:2]
        co2_indicators = text_indicators[2:]
    elif len(text_indicators) == 2:
        ep_indicators = text_indicators[:1]
        co2_indicators = text_indicators[1:]
    else:
        ep_indicators = text_indicators
        co2_indicators = []

    # Perechi text-pentagon: pentagonul urmează imediat după textbox
    # Construim o mapare: pentru fiecare text indicator, găsim pentagonul cel mai apropiat (următor)
    def find_next_path(text_match, all_paths):
        """Find the path shape immediately AFTER a text shape in the XML."""
        text_end = text_match.end()
        best = None
        best_dist = float('inf')
        for pm, ph, pv in all_paths:
            dist = pm.start() - text_end
            if 0 < dist < best_dist:
                best = (pm, ph, pv)
                best_dist = dist
        return best

    # Replace EP real indicators (first group — higher posV = further down = typically "B")
    # Template has: ref (A, lower V) then real (B, higher V)
    # So sorted: [0]=ref(A, V=510540), [1]=real(B, V=977265) for single indicators
    # For 4 indicators: [0,1]=ref pair, [2,3]=real pair

    new_xml = body_xml

    def move_indicator(text_match, new_class, class_pos_map, old_v, is_co2=False):
        """Move both text indicator AND its paired pentagon shape, and update colors."""
        nonlocal new_xml
        new_pos = class_pos_map.get(new_class, old_v)
        delta_v = new_pos - old_v  # displacement in EMU
        color_map = _CO2_CLASS_COLORS if is_co2 else _EP_CLASS_COLORS
        color = color_map.get(new_class, color_map.get("B"))

        # 1. Update textbox (letter + position + fill color + text color)
        new_content = text_match.group(1)
        new_content = re.sub(r'(<w:t[^>]*>)[A-G]\+?(</w:t>)', r'\g<1>' + new_class + r'\g<2>', new_content)
        new_content = _update_shape_pos_v(new_content, new_pos)
        # Actualizează și culoarea de fundal a textbox-ului (solidFill) — altfel litera
        # rămâne colorată cu culoarea implicită din template (nu cu culoarea clasei reale)
        new_content = _update_shape_color(new_content, color)
        # Culoarea textului calculată din luminanța WCAG a fundalului (valabil EP și CO2)
        text_color = _text_color_for_bg(color["hex"])
        if '<w:color' not in new_content:
            new_content = re.sub(
                r'(<w:rPr>)([\s\S]*?)(</w:rPr>[\s\S]*?<w:t)',
                r'\g<1>\g<2><w:color w:val="' + text_color + r'"/>\g<3>',
                new_content,
                count=2
            )
        else:
            new_content = re.sub(r'<w:color w:val="[^"]*"/>', '<w:color w:val="' + text_color + '"/>', new_content)
        new_xml = new_xml.replace(text_match.group(1), new_content, 1)

        # 2. Find and update paired pentagon (path shape right after textbox)
        path = find_next_path(text_match, path_indicators)
        if path:
            pm, ph, pv = path
            new_path_pos = pv + delta_v  # same displacement
            new_path_content = pm.group(1)
            new_path_content = _update_shape_pos_v(new_path_content, new_path_pos)
            new_path_content = _update_shape_color(new_path_content, color)
            new_xml = new_xml.replace(pm.group(1), new_path_content, 1)

    # EP indicators: [0]=real, [1]=ref (ordine din document)
    if len(ep_indicators) >= 1:
        m, letter, h, v = ep_indicators[0]
        move_indicator(m, ep_class_real, _CLASS_POS_V, v)

    if len(ep_indicators) >= 2:
        m, letter, h, v = ep_indicators[1]
        move_indicator(m, ep_class_ref, _CLASS_POS_V, v)

    # CO2 indicators
    for m, letter, h, v in co2_indicators:
        move_indicator(m, co2_class_real, _CO2_CLASS_POS_V, v, is_co2=True)

    # Aplică modificările IN-PLACE (nu înlocuim tot body-ul) — evită round-trip-ul XML
    # care poate schimba subtil layout-ul documentului (overflow pagina 2).
    alt_blocks_orig = [m.group(1) for m in alt_pattern.finditer(body_xml)]
    alt_blocks_new  = [m.group(1) for m in alt_pattern.finditer(new_xml)]

    MC_NS  = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    mc_tag = f"{{{MC_NS}}}AlternateContent"
    body_mc_elems = list(body.iter(mc_tag))

    # Construiește string-ul de declarații xmlns din nsmap-ul body-ului
    ns_parts = []
    for prefix, uri in body.nsmap.items():
        if prefix is None:
            ns_parts.append(f'xmlns="{uri}"')
        else:
            ns_parts.append(f'xmlns:{prefix}="{uri}"')
    ns_decls = " ".join(ns_parts)

    for i, (orig_str, new_str) in enumerate(zip(alt_blocks_orig, alt_blocks_new)):
        if orig_str != new_str and i < len(body_mc_elems):
            old_elem = body_mc_elems[i]
            parent_el = old_elem.getparent()
            idx = list(parent_el).index(old_elem)
            try:
                wrapper = etree.fromstring(
                    f'<_wr {ns_decls}>{new_str}</_wr>'.encode("utf-8")
                )
                new_mc_elem = wrapper[0]
                parent_el.remove(old_elem)
                parent_el.insert(idx, new_mc_elem)
            except Exception:
                pass  # Dacă parsing eșuează, păstrăm elementul original
    # Nu e nevoie de doc._Document__body = None — body-ul nu a fost înlocuit


# ═══════════════════════════════════════════════════════
# TOGGLE CHECKBOXES — Anexa form fields
# ═══════════════════════════════════════════════════════
def toggle_checkboxes(doc, indices):
    """Check checkboxes at given indices (0-based) in the document."""
    checkboxes = doc.element.findall(".//w:checkBox", NSMAP)
    indices_set = set(indices)
    for i, cb in enumerate(checkboxes):
        if i in indices_set:
            default_elem = cb.find("w:default", NSMAP)
            if default_elem is not None:
                default_elem.set(qn("w:val"), "1")


# ═══════════════════════════════════════════════════════
# COMPUTE CHECKBOXES — mapare date clădire → indici checkbox Anexa
# ═══════════════════════════════════════════════════════

# U referință Mc 001-2022 (nZEB)
_U_REF_RES = {"PE":0.25,"PR":0.67,"PS":0.29,"PT":0.15,"PP":0.15,"PB":0.29,"PL":0.20,"SE":0.20}
_U_REF_NRES = {"PE":0.33,"PR":0.80,"PS":0.35,"PT":0.17,"PP":0.17,"PB":0.35,"PL":0.22,"SE":0.22}


def compute_checkboxes(data, category):
    """Compute checkbox indices to toggle based on building data.
    Returns list of 0-based indices for the 308-checkbox Anexa clădire template."""
    cbs = []
    is_res = category in ("RI", "RC", "RA")
    u_ref = _U_REF_RES if is_res else _U_REF_NRES

    # Parse opaque U-values
    try:
        opaque_u = json.loads(data.get("opaque_u_values", "[]"))
    except:
        opaque_u = []
    try:
        glaz_max_u = float(data.get("glazing_max_u", "0") or "0")
        if glaz_max_u != glaz_max_u or glaz_max_u < 0:  # NaN sau negativ
            glaz_max_u = 0.0
    except (ValueError, TypeError):
        glaz_max_u = 0.0

    # ══════════════════════════════
    # ANEXA 1 — RECOMANDĂRI (CB 0-47)
    # ══════════════════════════════

    # Anvelopă — bifăm dacă U calculat > U referință
    if any(e.get("type") == "PE" and float(e.get("u", 0)) > u_ref.get("PE", 0.25) for e in opaque_u):
        cbs.append(0)  # Pereți exteriori
    if any(e.get("type") == "PB" and float(e.get("u", 0)) > u_ref.get("PB", 0.29) for e in opaque_u):
        cbs.append(1)  # Planșeu subsol
    if any(e.get("type") in ("PT", "PP") and float(e.get("u", 0)) > u_ref.get(e["type"], 0.15) for e in opaque_u):
        cbs.append(2)  # Terasă/pod
    if any(e.get("type") in ("PL", "SE") and float(e.get("u", 0)) > u_ref.get(e["type"], 0.20) for e in opaque_u):
        cbs.append(3)  # Planșee contact exterior
    # CB5: tâmplărie
    u_glaz_ref = 1.30 if is_res else 1.80
    if glaz_max_u > u_glaz_ref:
        cbs.append(5)
    # CB6: grile ventilare
    vent_type = data.get("ventilation_type", "")
    if not vent_type or vent_type == "natural_neorg":
        cbs.append(6)
    # CB13: robinete termostat
    h_src = data.get("heating_source", "")
    if h_src and h_src not in ("electric_direct", "pc_aer_aer"):
        cbs.append(13)
    # CB21: automatizare
    h_ctrl = data.get("heating_control", "")
    if not h_ctrl or h_ctrl == "manual":
        cbs.append(21)
    # CB25: iluminat LED
    l_type = data.get("lighting_type", "")
    if l_type and l_type != "led":
        cbs.append(25)
    # CB26: senzori prezență
    l_ctrl = data.get("lighting_control", "")
    if not l_ctrl or l_ctrl not in ("sensor_presence", "daylight_dimming"):
        cbs.append(26)
    # CB27: regenerabile
    st_en = data.get("solar_thermal_enabled", "") == "true"
    pv_en = data.get("pv_enabled", "") == "true"
    if not st_en and not pv_en:
        cbs.append(27)
    # CB28: recuperare căldură
    if not vent_type or "hr" not in vent_type:
        cbs.append(28)

    # Cost/Savings/Payback (CB 48-64) — valori estimate
    # Fără financialAnalysis, folosim default-uri
    cbs.append(50)  # 10k-25k EUR (default mediu)
    cbs.append(56)  # 20-30% savings
    cbs.append(62)  # 3-7 ani payback

    # ══════════════════════════════
    # ANEXA 2 — DATE CLĂDIRE (CB 65+)
    # ══════════════════════════════

    # Tip clădire (CB 65=existentă, 66=nouă, 67=existentă nefinalizată)
    year_b = int(data.get("year_built", "2000") or "2000")
    import datetime
    if year_b >= datetime.datetime.now().year - 1:
        cbs.append(66)
    else:
        cbs.append(65)

    # Categoria clădirii (CB 68-111)
    cat_map = {
        "RI": [68, 69], "RC": [68, 71], "RA": [68, 71],
        "BI": [79, 80], "ED": [74, 76], "SA": [86, 87],
        "HC": [94, 95], "CO": [103, 104], "SP": [99, 100],
        "AL": [108, 111],
    }
    for cb in cat_map.get(category, cat_map["AL"]):
        cbs.append(cb)

    # Zone climatice (CB 112-116 = zone I-V)
    try:
        zone_num = int(data.get("climate_zone_num", "3") or "3")
    except (ValueError, TypeError):
        zone_num = 3
    zone_num = max(1, min(5, zone_num))  # clamp 1-5
    cbs.append(111 + zone_num)

    # Zone eoliene (CB 117-120)
    wind_zone = 1 if zone_num <= 2 else (2 if zone_num <= 4 else 3)
    cbs.append(116 + wind_zone)

    # Structura constructivă (CB 127-134) — matching parțial (building.structure e string lung)
    struct_map = [
        ("Zidărie portantă", 127), ("Cadre beton armat", 129),
        ("Panouri prefabricate mari", 133), ("Structură metalică", 132),
        ("Structură lemn", 131), ("Mixtă", 134),
    ]
    struct_text = data.get("structure", "")
    struct_cb = None
    for key, cb in struct_map:
        if key in struct_text:
            struct_cb = cb
            break
    if struct_cb:
        cbs.append(struct_cb)

    # ══════════════════════════════
    # ANEXA 2 — INSTALAȚII
    # ══════════════════════════════

    # ÎNCĂLZIRE (CB 135+)
    if h_src:
        cbs.append(135)  # Da, funcțională
    else:
        cbs.append(137)  # Nu

    if h_src:
        heat_src_map = {
            "gaz_conv": 144, "gaz_cond": 144, "termoficare": 146,
            "electric_direct": 139, "pc_aer_apa": 149, "pc_aer_aer": 139,
            "pc_sol_apa": 149, "pc_apa_apa": 149, "centrala_gpl": 144,
            "cazan_lemn": 138, "cazan_peleti": 138, "soba_teracota": 138,
            "pompa_caldura": 149,
        }
        h_cb = heat_src_map.get(h_src)
        if h_cb:
            cbs.append(h_cb)

    # Tip sistem încălzire
    if h_src == "soba_teracota":
        cbs.append(150)
    elif h_src == "electric_direct":
        cbs.append(154)
    elif h_src:
        cbs.append(151)  # corpuri statice

    # Distribuție
    cbs.append(160)  # inferioară (default)

    # ACM (CB 176+)
    acm_src = data.get("acm_source", "")
    if acm_src:
        cbs.append(176)  # Da
    else:
        cbs.append(178)  # Nu

    if acm_src:
        acm_map = {
            "ct_prop": 181, "boiler_electric": 180, "termoficare": 183,
            "solar_termic": 179, "pc": 186,
        }
        a_cb = acm_map.get(acm_src)
        if a_cb:
            cbs.append(a_cb)

    if acm_src == "boiler_electric":
        cbs.append(187)
    elif acm_src == "ct_prop":
        cbs.append(188)

    cbs.append(195)  # Recirculare nu există (default)

    # RĂCIRE (CB 202+)
    has_cool = data.get("cooling_has", "") == "true"
    if has_cool:
        cbs.append(202)
    else:
        cbs.append(204)

    if has_cool:
        cool_src = data.get("cooling_source", "")
        cool_map = {
            "split": 214, "chiller_aer": 205, "chiller_apa": 206,
            "pc_aer_apa": 207, "pc_apa_apa": 208, "pc_aer_aer": 209,
            "monobloc": 213,
        }
        c_cb = cool_map.get(cool_src)
        if c_cb:
            cbs.append(c_cb)
        cbs.append(229)  # Complet climatizat
        cbs.append(232)  # Fără controlul umidității

    # VENTILARE (CB 256+)
    has_vent = vent_type and vent_type != "natural_neorg"
    if has_vent:
        cbs.append(256)
    else:
        cbs.append(258)

    if vent_type == "natural_neorg":
        cbs.append(259)
    elif vent_type == "natural_org":
        cbs.append(260)
    elif has_vent:
        cbs.append(261)

    if vent_type and "hr" in vent_type:
        cbs.append(270)
    else:
        cbs.append(271)

    # ILUMINAT (CB 272+)
    if l_type:
        cbs.append(272)
    else:
        cbs.append(274)

    if l_ctrl == "manual":
        cbs.append(276)
    elif l_ctrl == "daylight_dimming":
        cbs.extend([277, 278])
    elif l_ctrl == "sensor_presence":
        cbs.extend([277, 279])
    else:
        cbs.append(275)

    if l_type == "fluorescent":
        cbs.append(281)
    elif l_type == "incandescent":
        cbs.append(282)
    elif l_type == "led":
        cbs.append(283)
    else:
        cbs.append(284)

    cbs.append(285)  # Stare bună (default)

    # REGENERABILE (CB 288+)
    if st_en:
        cbs.append(288)
    else:
        cbs.append(289)
    if pv_en:
        cbs.append(290)
    else:
        cbs.append(291)

    hp_en = data.get("heat_pump_enabled", "") == "true"
    if hp_en:
        cbs.append(292)
    else:
        cbs.append(293)

    if hp_en:
        hp_type_map = {
            "sol_apa_deschisa": 294, "sol_apa_inchisa": 295,
            "aer_apa": 296, "aer_aer": 297, "apa_aer": 298, "sol_aer": 299,
        }
        hp_cb = hp_type_map.get(data.get("heat_pump_type", ""))
        if hp_cb:
            cbs.append(hp_cb)

    bio_en = data.get("biomass_enabled", "") == "true"
    if bio_en:
        cbs.append(301)
    else:
        cbs.append(302)

    if bio_en:
        bio_type = data.get("biomass_type", "")
        if bio_type == "peleti":
            cbs.append(303)
        elif bio_type == "brichete":
            cbs.append(304)
        else:
            cbs.append(305)

    wind_en = data.get("wind_enabled", "") == "true"
    if wind_en:
        cbs.append(306)
    else:
        cbs.append(307)

    # ══════════════════════════════
    # BACS (CB 308+) — EN 15232-1
    # Dacă template-ul are câmpuri BACS (Anexe cu CB 308-311)
    # ══════════════════════════════
    bacs_class = data.get("bacs_class", "C")
    bacs_cb_map = {"A": 308, "B": 309, "C": 310, "D": 311}
    bacs_cb = bacs_cb_map.get(bacs_class)
    if bacs_cb:
        cbs.append(bacs_cb)

    # SRI readiness (CB 312+)
    sri_total = int(data.get("sri_total", "0") or "0")
    if sri_total >= 70:
        cbs.append(312)  # SRI class A
    elif sri_total >= 50:
        cbs.append(313)  # SRI class B
    elif sri_total >= 30:
        cbs.append(314)  # SRI class C
    else:
        cbs.append(315)  # SRI class D

    return cbs


# ═══════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════
def format_ro(val, dec=1):
    """Format number with Romanian comma decimal separator."""
    return f"{float(val):.{dec}f}".replace(".", ",")


# ═══════════════════════════════════════════════════════
# COLORARE CLASE ENERGETICE PER UTILITATE — tabelul de jos CPE
# ═══════════════════════════════════════════════════════
def _highlight_utility_class_cells(doc, data):
    """Colorează cu culoarea clasei energetice celula din tabelul de jos CPE
    care conține intervalul corespunzător consumului per utilitate.
    Culorile corespund scalei EP: A+=verde închis → G=roșu."""
    import re as _re
    from docx.oxml import OxmlElement as _OxmlElement

    # EP valori per utilitate (kWh/m²·an) — trimise din frontend
    def _parse_ro(s):
        try:
            return float((s or "0").replace(",", "."))
        except Exception:
            return 0.0

    ep_vals = {
        "incalzire": _parse_ro(data.get("ep_incalzire", "")),
        "acm":       _parse_ro(data.get("ep_acm", "")),
        "racire":    _parse_ro(data.get("ep_racire", "")),
        "ventilare": _parse_ro(data.get("ep_ventilare", "")),
        "iluminat":  _parse_ro(data.get("ep_iluminat", "")),
    }

    # Culori per clasă (0=A+, 1=A, ..., 7=G) — identice cu _EP_CLASS_COLORS
    _COL_FILLS = ["009B00", "32C831", "00FF00", "FFFF00", "F39C00", "FF6400", "FE4101", "FE0000"]

    # Mapare cuvinte-cheie rând → cheie ep
    _ROW_MAP = [
        (["ncălzire", "ncalzire", "eating"],  "incalzire"),
        (["Ap", "caldă", "calda", "ACM"],     "acm"),
        (["Răcire", "Racire", "ooling"],       "racire"),
        (["Ventilare", "entilat"],             "ventilare"),
        (["luminat", "ghting"],                "iluminat"),
    ]

    def _apply_shading(cell, fill_hex):
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = _OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        shd = tcPr.find(qn("w:shd"))
        if shd is None:
            shd = _OxmlElement("w:shd")
            tcPr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)

    def _text_color_for_bg(fill_hex):
        """Negru sau alb — alege culoarea cu contrast WCAG mai mare față de fundal."""
        def to_lin(c8):
            v = c8 / 255.0
            return v / 12.92 if v <= 0.04045 else ((v + 0.055) / 1.055) ** 2.4
        r = int(fill_hex[0:2], 16)
        g = int(fill_hex[2:4], 16)
        b = int(fill_hex[4:6], 16)
        L = 0.2126 * to_lin(r) + 0.7152 * to_lin(g) + 0.0722 * to_lin(b)
        # Pragul de echilibru: L=0.179 → contrast alb = contrast negru
        # Sub prag → alb are contrast mai mare; peste prag → negru are contrast mai mare
        return "FFFFFF" if L < 0.179 else "000000"

    def _format_ro(val):
        """Formatează număr cu 1 zecimală, virgulă românească."""
        return "{:.1f}".format(val).replace(".", ",")

    def _set_cell_text(cell, text, fill_hex):
        """Înlocuiește textul celulei cu valoarea reală, păstrând formatarea."""
        tc = cell._tc
        text_color = _text_color_for_bg(fill_hex)
        # Colectăm proprietățile run-ului existent (dimensiune font etc.)
        existing_sz = None
        for p in tc.iter(qn("w:p")):
            for r in p.iter(qn("w:r")):
                rPr = r.find(qn("w:rPr"))
                if rPr is not None:
                    sz_el = rPr.find(qn("w:sz"))
                    if sz_el is not None:
                        existing_sz = sz_el.get(qn("w:val"))
                break
            if existing_sz:
                break
        # Ștergem toate paragrafele și recreem unul singur centrat
        for p in list(tc.findall(qn("w:p"))):
            tc.remove(p)
        # Paragraph nou
        p_new = _OxmlElement("w:p")
        pPr = _OxmlElement("w:pPr")
        jc = _OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)
        sp = _OxmlElement("w:spacing")
        sp.set(qn("w:before"), "0")
        sp.set(qn("w:after"), "0")
        pPr.append(sp)
        p_new.append(pPr)
        # Run nou
        r_new = _OxmlElement("w:r")
        rPr = _OxmlElement("w:rPr")
        # Bold
        b_el = _OxmlElement("w:b")
        rPr.append(b_el)
        # Culoare text
        color_el = _OxmlElement("w:color")
        color_el.set(qn("w:val"), text_color)
        rPr.append(color_el)
        # Dimensiune font (dacă am extras-o)
        if existing_sz:
            sz_new = _OxmlElement("w:sz")
            sz_new.set(qn("w:val"), existing_sz)
            rPr.append(sz_new)
            szCs_new = _OxmlElement("w:szCs")
            szCs_new.set(qn("w:val"), existing_sz)
            rPr.append(szCs_new)
        r_new.append(rPr)
        t_el = _OxmlElement("w:t")
        t_el.text = text
        r_new.append(t_el)
        p_new.append(r_new)
        tc.append(p_new)

    def _parse_range(text):
        """Returnează (lo, hi) din text ca '≤30', '30...42', '>484'."""
        text = text.replace("\xa0", " ").strip()
        le_m = _re.search(r"[≤<]\s*([\d,\.]+)", text)
        gt_m = _re.search(r"[>]\s*([\d,\.]+)", text)
        rng_m = _re.search(r"([\d,\.]+)\s*[.\-–…]+\s*([\d,\.]+)", text)
        if le_m:
            return (None, float(le_m.group(1).replace(",", ".")))
        if rng_m:
            return (float(rng_m.group(1).replace(",", ".")), float(rng_m.group(2).replace(",", ".")))
        if gt_m:
            return (float(gt_m.group(1).replace(",", ".")), None)
        return None

    def _in_range(v, rng):
        lo, hi = rng
        if lo is None: return v <= hi
        if hi is None: return v > lo
        return lo <= v <= hi

    # Găsește tabelul cu "Clasă energetică" și utility rows
    target = None
    for tbl in doc.tables:
        txt = " ".join(c.text for r in tbl.rows for c in r.cells)
        if ("lasă energetic" in txt or "lasa energetic" in txt) and \
           ("ncălzire" in txt or "ncalzire" in txt or "eating" in txt):
            target = tbl
            break
    if not target:
        return

    for row in target.rows:
        cells = row.cells
        if len(cells) < 3:
            continue
        first_text = cells[0].text.strip()

        ep_key = None
        for keywords, key in _ROW_MAP:
            if any(kw in first_text for kw in keywords):
                ep_key = key
                break
        if ep_key is None:
            continue

        ep_val = ep_vals.get(ep_key, 0.0)
        if ep_val <= 0:
            continue

        # Deduplifică celulele fuzionate (python-docx returnează duplicate la merge)
        seen_ids = set()
        unique = []
        for c in cells:
            cid = id(c._tc)
            if cid not in seen_ids:
                seen_ids.add(cid)
                unique.append(c)
        # unique[0] = coloana etichetă; unique[1..8] = A+ → G
        col_idx = None
        for i in range(1, len(unique)):
            rng = _parse_range(unique[i].text)
            if rng and _in_range(ep_val, rng):
                col_idx = i
                break

        if col_idx is None:
            continue

        class_idx = col_idx - 1  # 0=A+, 1=A, ..., 7=G
        if 0 <= class_idx < len(_COL_FILLS):
            fill = _COL_FILLS[class_idx]
            _apply_shading(unique[col_idx], fill)
            _set_cell_text(unique[col_idx], _format_ro(ep_val), fill)


# ═══════════════════════════════════════════════════════
# MAIN HANDLER
# ═══════════════════════════════════════════════════════
class handler(BaseHTTPRequestHandler):
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    def do_POST(self):
        try:
            content_length = int(self.headers.get("Content-Length", 0))
            raw = self.rfile.read(content_length)
            body = json.loads(raw)

            tpl_bytes = base64.b64decode(body["template"])
            data = body.get("data", {})
            mode = body.get("mode", "cpe")
            category = body.get("category", "AL")

            doc = Document(io.BytesIO(tpl_bytes))

            # ═══════════════════════════════════════
            # 1. TEXT REPLACEMENTS — mapare directă
            # ═══════════════════════════════════════
            # ═══ CRITICAL: ordinea contează! ═══
            # Înlocuim pattern-urile LUNGI înainte de cele SCURTE
            # "xxxx,x" ÎNAINTE de "xxxx" (altfel "xxxx" corupe "xxxx,x")
            # "xxx,x" ÎNAINTE de "xx,x" (altfel "xx,x" corupe "xxx,x")
            # "ZZ.LL.AAAA" ÎNAINTE de "AAAA"

            # 1. Secvențiale (cele mai lungi pattern-uri primele)
            ep_total_vals = [data.get("ep_total_real", "0,0"), data.get("ep_total_ref", "0,0")]
            replace_seq(doc, "xxxx,x", ep_total_vals)

            xxx_vals = [data.get("area_ref", "0,0"), data.get("co2_val", "0,0"),
                        data.get("sre_st", "0,0"), data.get("sre_pv", "0,0"),
                        data.get("sre_pc", "0,0"), data.get("sre_bio", "0,0"),
                        data.get("sre_other", "0,0"), data.get("sre_total", "0,0")]
            replace_seq(doc, "xxx,x", xxx_vals)

            xx_vals = [data.get("qf_thermal", "0,0"), data.get("qf_electric", "0,0"),
                       data.get("ep_specific", "0,0"), data.get("ep_ref", "0,0")]
            replace_seq(doc, "xx,x", xx_vals)

            # 2. Înlocuiri simple (de la cele mai lungi la cele mai scurte)
            ordered_replacements = [
                ("II,IIII x LL,LLLL", data.get("gps", "")),
                ("ZZ.LL.AAAA", data.get("auditor_date", "")),
                ("ZZ/LL/AAAA", data.get("auditor_date", "").replace(".", "/")),
                ("XX/XXXXX", data.get("auditor_atestat", "")),
                ("zz/ll/aa", data.get("expiry", "")),
                ("GWP,G", data.get("gwp", "0,0")),
                ("RR,R", data.get("rer", "0,0")),
                ("zzz,z", data.get("area_ref", "")),
                ("yyy,y", data.get("area_gross", "")),
                ("AAAA", data.get("year", "____")),
                ("xxxx", data.get("volume", "")),
                ("regim", data.get("regime", "")),
            ]

            for old, new in ordered_replacements:
                if new:
                    replace_in_doc(doc, old, new)

            # Adresa — nodul conține "adresa" cu puncte
            for para in doc.paragraphs:
                if "adresa" in para.text.lower() and "..." in para.text:
                    for run in para.runs:
                        if "adresa" in run.text.lower() or "..." in run.text:
                            run.text = ""
                    if para.runs:
                        para.runs[0].text = data.get("address", "")
                    break
            # Also check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if "adresa" in para.text.lower() and ("..." in para.text or "." * 5 in para.text):
                                for run in para.runs:
                                    if "adresa" in run.text.lower() or "..." in run.text or "." * 5 in run.text:
                                        run.text = ""
                                if para.runs:
                                    para.runs[0].text = data.get("address", "")

            # Scop CPE — split text "Vânzare/Închiriere/Recepție/Inf"
            scope = data.get("scope", "")
            if scope:
                for para in doc.paragraphs:
                    full = para.text
                    if "nzare" in full or "nchir" in full or "Recep" in full:
                        # Clear all runs with scope text parts and set first one
                        scope_runs = [r for r in para.runs if any(x in r.text for x in ["Vânz", "are", "Închir", "ie", "Recep", "/Inf"])]
                        for r in scope_runs:
                            r.text = ""
                        if scope_runs:
                            scope_runs[0].text = scope
                        break
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                full = para.text
                                if "nzare" in full or "nchir" in full or "Recep" in full:
                                    scope_runs = [r for r in para.runs if any(x in r.text for x in ["Vânz", "are", "Închir", "ie", "Recep", "/Inf"])]
                                    for r in scope_runs:
                                        r.text = ""
                                    if scope_runs:
                                        scope_runs[0].text = scope

            # Program calcul — înlocuiește "................versiunea" cu "ZEPHREN"
            software = data.get("software", "Zephren v2.0")
            program_name = "ZEPHREN"
            def fill_program_field(paragraphs):
                for para in paragraphs:
                    if "Program de calcul utilizat" in para.text or "versiunea" in para.text:
                        inserted = False
                        for run in para.runs:
                            if re.search(r'\.{4,}', run.text):
                                run.text = re.sub(r'\.{4,}', program_name if not inserted else '', run.text)
                                inserted = True
                            elif "versiunea" in run.text:
                                run.text = run.text.replace("versiunea", "")
            fill_program_field(doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        fill_program_field(cell.paragraphs)

            # Auditor name/company/phone/email
            auditor_replacements = {
                "Nume & prenume auditor energetic": data.get("auditor_name", ""),
                "Nume auditor": data.get("auditor_name", ""),
                "nume auditor": data.get("auditor_name", ""),
                "Firma/PFA": data.get("auditor_company", ""),
                "denumire firma": data.get("auditor_company", ""),
                "nr. telefon": data.get("auditor_phone", ""),
                "adresa email": data.get("auditor_email", ""),
                "cod unic": data.get("auditor_mdlpa", ""),
                "Cod unic": data.get("auditor_mdlpa", ""),
            }
            for old, new in auditor_replacements.items():
                if new:
                    replace_in_doc(doc, old, new)

            # nZEB status
            replace_in_doc(doc, "nZEB DA/NU", "nZEB " + data.get("nzeb", "NU"))

            # (secvențialele xxxx,x / xxx,x / xx,x au fost mutate mai sus, ordinea contează)

            # Nr camere (RA)
            if category == "RA":
                replace_in_doc(doc, " x ", " " + data.get("nr_units", "3") + " ")

            # Category label
            cat_label = data.get("category_label", "")
            if cat_label:
                replace_in_doc(doc, "categorie funcțională", cat_label)
                replace_in_doc(doc, "categorie functionala", cat_label)

            # Location
            replace_in_doc(doc, "localitatea", data.get("city", ""))
            replace_in_doc(doc, "județul", data.get("county", ""))
            replace_in_doc(doc, "judetul", data.get("county", ""))
            replace_in_doc(doc, "zona climatică", data.get("climate_zone", ""))
            replace_in_doc(doc, "zona climatica", data.get("climate_zone", ""))

            # GWP lifecycle
            gwp_text = data.get("gwp", "0,0") + " kgCO2eq/m2an"
            replace_in_doc(doc, "GWP lifecycle", gwp_text)

            # ═══════════════════════════════════════
            # 2. SCALE EP — înlocuiesc pragurile template cu EPBD
            # ═══════════════════════════════════════
            new_ep = [int(data.get(k, 0)) for k in ["s_ap", "s_a", "s_b", "s_c", "s_d", "s_e", "s_f"]]
            if any(v > 0 for v in new_ep):
                replace_scales(doc, category, new_ep)

            # ═══════════════════════════════════════
            # 3. CLASS INDICATORS — săgețile pe scale
            # ═══════════════════════════════════════
            ep_cls_real = data.get("ep_class_real", "")
            ep_cls_ref = data.get("ep_class_ref", "")
            co2_cls_real = data.get("co2_class_real", "")
            if ep_cls_real:
                replace_class_indicators(doc, ep_cls_real, ep_cls_ref, co2_cls_real)

            # ═══════════════════════════════════════
            # 4. CHECKBOXES (Anexa)
            # ═══════════════════════════════════════
            if mode == "anexa":
                # Compute checkbox indices from building data (server-side)
                cb_indices = compute_checkboxes(data, category)
                # Also accept client-side overrides if provided
                client_cbs = body.get("checkboxes", [])
                if client_cbs:
                    cb_indices = list(set(cb_indices + client_cbs))
                if cb_indices:
                    toggle_checkboxes(doc, cb_indices)

            # ═══════════════════════════════════════
            # 4b. ANEXA 2 — TEXT REPLACEMENTS
            # ═══════════════════════════════════════
            if mode == "anexa":
                # Adresa și nr certificat
                replace_in_doc(doc, "[adresa]", data.get("address", ""))
                # An construcție/renovare
                replace_in_doc(doc, ".................", data.get("year", "") + (" / " + data.get("year_renov", "") if data.get("year_renov") else ""))
                # Arie referință totală
                au = data.get("area_ref", "")
                vol = data.get("volume", "")
                if au:
                    replace_in_doc(doc, "Aria de referință totală", "Aria de referință totală a pardoselii: " + au + " m²")
                if vol:
                    replace_in_doc(doc, "Volumul interior de referință", "Volumul interior de referință V: " + vol + " m³")
                # Factor formă
                try:
                    au_f = float(au.replace(",", ".")) if au else 0
                    vol_f = float(vol.replace(",", ".")) if vol else 0
                    ae = float(data.get("area_envelope", "0").replace(",", ".")) if data.get("area_envelope") else au_f * 1.3
                    se_v = ae / vol_f if vol_f > 0 else 0
                    replace_in_doc(doc, "Factorul de formă", "Factorul de formă al clădirii, SE/V: " + format_ro(se_v, 3))
                except:
                    pass
                # Nr persoane
                try:
                    is_res = category in ("RI", "RC", "RA")
                    nr_pers = max(1, round(au_f / (30 if is_res else 15)))
                    replace_in_doc(doc, "pers.", str(nr_pers) + " pers.")
                except:
                    pass
                # Detalii instalații — combustibil
                fuel_labels = {"gaz_nat": "gaz natural", "gpl": "GPL", "motorina": "motorină",
                               "lemn": "lemne", "peleti": "peleți", "carbune": "cărbune",
                               "electric": "electricitate", "biogaz": "biogaz"}
                fuel = fuel_labels.get(data.get("heating_fuel", ""), data.get("heating_fuel", ""))
                if fuel:
                    replace_in_doc(doc, "combustibil .....................", "combustibil " + fuel)
                    replace_in_doc(doc, "combustibil ...........", "combustibil " + fuel)
                # Putere nominală încălzire
                hp = data.get("heating_power", "0")
                if hp and hp != "0":
                    replace_in_doc(doc, "Necesarul de căldură de calcul", "Necesarul de căldură de calcul: " + hp + " kW")
                    replace_in_doc(doc, "Puterea termică instalată totală pentru încălzire", "Puterea termică instalată totală: " + hp + " kW")

                # ── U-values per tip element ──────────────────────────────
                # Înlocuim texte tip "U pereți exteriori" / "U acoperiș" etc.
                try:
                    opaque_u = json.loads(data.get("opaque_u_values", "[]"))
                    raw_glaz = data.get("glazing_max_u", "0") or "0"
                    try:
                        glaz_u = float(raw_glaz)
                        if glaz_u != glaz_u or glaz_u < 0:
                            glaz_u = 0.0
                    except (ValueError, TypeError):
                        glaz_u = 0.0
                    glaz_g = float(data.get("glazing_g_value", "0") or "0")

                    pe_u = [e for e in opaque_u if e.get("type") == "PE"]
                    pt_u = [e for e in opaque_u if e.get("type") in ("PT","PP")]
                    pb_u = [e for e in opaque_u if e.get("type") == "PB"]
                    pl_u = [e for e in opaque_u if e.get("type") == "PL"]

                    if pe_u:
                        u_val = format_ro(pe_u[0].get("u", 0), 2)
                        replace_in_doc(doc, "U pereți exteriori", f"U pereți exteriori = {u_val} W/(m²·K)")
                        replace_in_doc(doc, "coeficientul global de transfer termic al pereților exteriori",
                                       f"U pereți ext. = {u_val} W/(m²·K)")
                    if pt_u:
                        u_val = format_ro(pt_u[0].get("u", 0), 2)
                        replace_in_doc(doc, "U terasă/acoperiș", f"U terasă = {u_val} W/(m²·K)")
                    if pb_u:
                        u_val = format_ro(pb_u[0].get("u", 0), 2)
                        replace_in_doc(doc, "U planșeu subsol", f"U planșeu subsol = {u_val} W/(m²·K)")
                    if glaz_u > 0:
                        replace_in_doc(doc, "U tâmplărie", f"U tâmplărie = {format_ro(glaz_u, 2)} W/(m²·K)")
                        replace_in_doc(doc, "coeficientul global de transfer termic al tâmplăriei",
                                       f"U tâmplărie = {format_ro(glaz_u, 2)} W/(m²·K)")
                    if glaz_g > 0:
                        replace_in_doc(doc, "factorul solar g", f"factorul solar g = {format_ro(glaz_g, 2)}")
                except Exception:
                    pass

                # ── EP breakdown pe destinații ────────────────────────────
                ep_breakdown = {
                    "ep_incalzire": data.get("ep_incalzire", ""),
                    "ep_racire": data.get("ep_racire", ""),
                    "ep_acm": data.get("ep_acm", ""),
                    "ep_ventilare": data.get("ep_ventilare", ""),
                    "ep_iluminat": data.get("ep_iluminat", ""),
                }
                ep_labels = {
                    "ep_incalzire": "EP încălzire",
                    "ep_racire": "EP răcire",
                    "ep_acm": "EP ACM",
                    "ep_ventilare": "EP ventilare",
                    "ep_iluminat": "EP iluminat",
                }
                for key, val in ep_breakdown.items():
                    if val:
                        replace_in_doc(doc, ep_labels[key], ep_labels[key] + " = " + val + " kWh/(m²·an)")

                # ── BACS clasă automatizare ───────────────────────────────
                bacs_class = data.get("bacs_class", "")
                bacs_labels = {
                    "A": "Clasa A — control predictiv (economie 25-40%)",
                    "B": "Clasa B — control avansat (economie 10-25%)",
                    "C": "Clasa C — automatizare de bază (referință)",
                    "D": "Clasa D — fără automatizare",
                }
                if bacs_class in bacs_labels:
                    replace_in_doc(doc, "Clasa BACS", "BACS: " + bacs_labels[bacs_class])
                    replace_in_doc(doc, "clasă automatizare", "Clasă automatizare BACS: " + bacs_class)

                # ── SRI — Smart Readiness Indicator ──────────────────────
                sri_val = data.get("sri_total", "")
                sri_grade = data.get("sri_grade", "")
                if sri_val:
                    replace_in_doc(doc, "SRI %", f"SRI = {sri_val}% (Clasa {sri_grade})")

                # ── Detalii ACM ───────────────────────────────────────────
                acm_vol = data.get("acm_storage_volume", "")
                acm_src_label = {
                    "ct_prop": "Centrală termică proprie",
                    "boiler_electric": "Boiler electric",
                    "termoficare": "Termoficare",
                    "solar_termic": "Solar termic",
                    "pc": "Pompă de căldură",
                }.get(data.get("acm_source", ""), "")
                if acm_vol:
                    replace_in_doc(doc, "Volumul vasului de acumulare", "Volum vas ACM: " + acm_vol + " L")
                if acm_src_label:
                    replace_in_doc(doc, "sursa de preparare ACM", "Sursă ACM: " + acm_src_label)

                # ── Infiltrații / Etanșeitate ─────────────────────────────
                n50 = data.get("n50", "")
                if n50 and float(n50 or 0) > 0:
                    replace_in_doc(doc, "n50 =", "n50 = " + format_ro(float(n50), 1) + " h⁻¹")
                    replace_in_doc(doc, "rata de infiltrații", "Rata infiltrații n50 = " + format_ro(float(n50), 1) + " h⁻¹")

            # ═══════════════════════════════════════
            # 5. FOTO CLĂDIRE — inserare imagine în text box
            # ═══════════════════════════════════════
            photo_b64 = body.get("photo")
            if photo_b64 and photo_b64.startswith("data:image"):
                # Suport png/jpeg/jpg/webp/gif/bmp
                match = re.match(r"data:image/(png|jpeg|jpg|webp|gif|bmp);base64,(.+)", photo_b64, re.DOTALL)
                if match:
                    img_data = base64.b64decode(match.group(2))
                    from docx.oxml.ns import qn as docx_qn
                    from docx.shared import Emu

                    # Dimensiunile frame-ului FOTO din template (EMU): 647700 x 584200
                    FOTO_W_EMU = 620000   # puțin sub lățimea frame-ului (~1,7 cm)
                    FOTO_H_EMU = 560000   # puțin sub înălțimea frame-ului (~1,55 cm)

                    # Calculează dimensiunile imaginii păstrând aspect ratio
                    try:
                        import struct
                        raw = img_data
                        img_w_px, img_h_px = None, None
                        # PNG: lățime/înălțime la bytes 16-24
                        if raw[:4] == b'\x89PNG':
                            img_w_px, img_h_px = struct.unpack('>II', raw[16:24])
                        # JPEG: cauta SOF marker
                        elif raw[:2] == b'\xff\xd8':
                            i = 2
                            while i < len(raw) - 8:
                                if raw[i] == 0xff and raw[i+1] in (0xC0, 0xC2):
                                    img_h_px = struct.unpack('>H', raw[i+5:i+7])[0]
                                    img_w_px = struct.unpack('>H', raw[i+7:i+9])[0]
                                    break
                                seg_len = struct.unpack('>H', raw[i+2:i+4])[0]
                                i += 2 + seg_len
                    except Exception:
                        img_w_px, img_h_px = None, None

                    # Alege dimensiunea finală respectând frame-ul și aspect ratio
                    if img_w_px and img_h_px and img_w_px > 0 and img_h_px > 0:
                        ratio = img_w_px / img_h_px
                        if ratio >= 1:  # landscape sau pătrat → constrâns de lățime
                            pic_w = FOTO_W_EMU
                            pic_h = int(FOTO_W_EMU / ratio)
                            if pic_h > FOTO_H_EMU:  # totuși prea înalt
                                pic_h = FOTO_H_EMU
                                pic_w = int(FOTO_H_EMU * ratio)
                        else:  # portrait → constrâns de înălțime
                            pic_h = FOTO_H_EMU
                            pic_w = int(FOTO_H_EMU * ratio)
                            if pic_w > FOTO_W_EMU:
                                pic_w = FOTO_W_EMU
                                pic_h = int(FOTO_W_EMU / ratio)
                    else:
                        # Fallback: folosim lățimea frame-ului
                        pic_w = FOTO_W_EMU
                        pic_h = None

                    # Găsește TOATE text box-urile care conțin "FOTO" și înlocuiește-le
                    txbx_tag = qn("w:txbxContent")
                    p_tag = qn("w:p")
                    from docx.text.paragraph import Paragraph as DocxPara

                    foto_txbx_list = []
                    for txbx_elem in doc.element.body.iter(txbx_tag):
                        for p_elem in txbx_elem.iter(p_tag):
                            if "FOTO" in (DocxPara(p_elem, None).text or ""):
                                foto_txbx_list.append(txbx_elem)
                                break  # un singur text box per găsire

                    for foto_txbx in foto_txbx_list:
                        # Șterge TOATE paragrafele existente din text box
                        for p_elem in list(foto_txbx.findall(p_tag)):
                            foto_txbx.remove(p_elem)

                        # Creează un singur paragraf centrat cu imaginea
                        new_p = etree.SubElement(foto_txbx, p_tag)
                        pPr = etree.SubElement(new_p, qn("w:pPr"))
                        sp = etree.SubElement(pPr, qn("w:spacing"))
                        sp.set(qn("w:after"), "0")
                        jc = etree.SubElement(pPr, qn("w:jc"))
                        jc.set(qn("w:val"), "center")
                        new_r = etree.SubElement(new_p, qn("w:r"))

                        # Adaugă imaginea via python-docx (calea standard)
                        img_stream = io.BytesIO(img_data)
                        img_stream.seek(0)
                        temp_para = doc.add_paragraph()
                        temp_run = temp_para.add_run()
                        if pic_h:
                            temp_run.add_picture(img_stream, width=Emu(pic_w), height=Emu(pic_h))
                        else:
                            temp_run.add_picture(img_stream, width=Emu(pic_w))
                        drawing = temp_run._element.find(docx_qn("w:drawing"))
                        if drawing is not None:
                            new_r.append(drawing)
                        temp_para._element.getparent().remove(temp_para._element)

            # ═══════════════════════════════════════
            # 6. ANEXĂ FOTOGRAFII CLĂDIRE (doar în modul "anexa")
            # ═══════════════════════════════════════
            building_photos = body.get("buildingPhotos", [])
            if mode == "anexa" and building_photos:
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml.ns import qn as docx_qn
                from docx.oxml import OxmlElement

                # Separator pagină nouă
                doc.add_page_break()

                # Titlu secțiune H (conform structurii oficiale Anexa CPE)
                title_p = doc.add_paragraph()
                title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_p.add_run("H. DOCUMENTARE FOTOGRAFICĂ A CLĂDIRII")
                title_run.bold = True
                title_run.font.size = Pt(11)

                cat_labels = {
                    "exterior": "Exterior",
                    "interior": "Interior",
                    "ir": "Termoviziune IR",
                    "instalatii": "Instalații",
                    "defecte": "Defecte / Degradări",
                    "altele": "Altele",
                }

                # Grupare pe categorii
                grouped = {}
                for ph in building_photos:
                    cat = ph.get("zone", "altele")
                    grouped.setdefault(cat, []).append(ph)

                for cat, photos in grouped.items():
                    # Subtitlu categorie
                    cat_p = doc.add_paragraph()
                    cat_run = cat_p.add_run(f"{cat_labels.get(cat, cat)}  ({len(photos)})")
                    cat_run.bold = True
                    cat_run.font.size = Pt(9)

                    # Tabel 2 coloane pentru poze
                    i = 0
                    while i < len(photos):
                        row_photos = photos[i:i+2]
                        tbl = doc.add_table(rows=1, cols=2)
                        # Elimină borduri tabel
                        tbl_pr = tbl._tbl.tblPr
                        if tbl_pr is None:
                            tbl_pr = OxmlElement("w:tblPr")
                            tbl._tbl.insert(0, tbl_pr)
                        tbl_borders = OxmlElement("w:tblBorders")
                        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                            border_el = OxmlElement(f"w:{border_name}")
                            border_el.set(docx_qn("w:val"), "none")
                            tbl_borders.append(border_el)
                        tbl_pr.append(tbl_borders)

                        for j in range(2):
                            cell = tbl.rows[0].cells[j]
                            if j < len(row_photos):
                                ph = row_photos[j]
                                ph_url = ph.get("url", "")
                                if ph_url and "," in ph_url:
                                    try:
                                        img_data = base64.b64decode(ph_url.split(",")[1])
                                        img_stream = io.BytesIO(img_data)
                                        # Detectează dimensiuni JPEG/PNG fără Pillow
                                        import struct as _struct
                                        raw = img_data
                                        img_w_px, img_h_px = None, None
                                        if raw[:4] == b'\x89PNG':
                                            img_w_px, img_h_px = _struct.unpack('>II', raw[16:24])
                                        elif raw[:2] == b'\xff\xd8':
                                            ii = 2
                                            while ii < len(raw) - 8:
                                                if raw[ii] == 0xff and raw[ii+1] in (0xC0, 0xC2):
                                                    img_h_px = _struct.unpack('>H', raw[ii+5:ii+7])[0]
                                                    img_w_px = _struct.unpack('>H', raw[ii+7:ii+9])[0]
                                                    break
                                                seg_len = _struct.unpack('>H', raw[ii+2:ii+4])[0]
                                                ii += 2 + seg_len
                                        # Calculează height păstrând aspect ratio
                                        pic_w = Inches(2.8)
                                        if img_w_px and img_h_px and img_w_px > 0:
                                            ratio = img_h_px / img_w_px
                                            pic_h = Inches(2.8 * ratio)
                                        else:
                                            pic_h = Inches(2.1)  # fallback 4:3
                                        img_p = cell.paragraphs[0]
                                        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        img_run = img_p.add_run()
                                        img_run.add_picture(img_stream, width=pic_w, height=pic_h)
                                    except Exception as e:
                                        import sys
                                        print(f"[FOTO ERR] {e}", file=sys.stderr)
                                label = ph.get("label", "")
                                note = ph.get("note", "")
                                if label:
                                    lp = cell.add_paragraph(label)
                                    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    if lp.runs:
                                        lp.runs[0].font.size = Pt(7)
                                if note:
                                    np_ = cell.add_paragraph(note)
                                    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    if np_.runs:
                                        np_.runs[0].font.size = Pt(6)
                        i += 2

            # ═══════════════════════════════════════
            # 6b. COLORARE CELULE CLASE ENERGETICE PER UTILITATE
            # ═══════════════════════════════════════
            # Evidențiază cu culoarea clasei celula din tabelul de jos care conține
            # intervalul în care se încadrează consumul specific al fiecărei utilități
            if mode == "cpe":
                _highlight_utility_class_cells(doc, data)

            # ═══════════════════════════════════════
            # 7. STRIP TRAILING EMPTY PARAGRAPHS
            # ═══════════════════════════════════════
            # Elimină paragrafele goale de la finalul documentului (după tabelul de bare cod)
            # Tabelul "COD UNIC DE BARE" se păstrează — face parte din modelul oficial CPE
            body_el = doc.element.body
            sect_pr = body_el.find(qn("w:sectPr"))
            if sect_pr is not None:
                def _get_texts(el):
                    return [t.text for t in el.iter(qn("w:t")) if t.text and t.text.strip()]

                # Elimină paragrafele goale de la final (după conținut/tabel bare cod)
                children = list(body_el)
                sect_idx = children.index(sect_pr)
                for child in reversed(children[:sect_idx]):
                    if child.tag != qn("w:p"):
                        break
                    if _get_texts(child):
                        break
                    body_el.remove(child)

                # ─── CPE: elimină spacing body paragraphs pentru a preveni overflow pagina 2 ──
                # Înlocuirile de text (firma, adresă, etc.) pot mări celule → tabelele cresc
                # → conținut depășește pagina 1. Reducem spacing-ul paragrafelor de la nivel body.
                # Economie: ~11mm (suficient pentru orice date reale de audit)
                if mode == "cpe":
                    for child in list(body_el):
                        if child.tag == qn("w:p"):
                            pPr = child.find(qn("w:pPr"))
                            if pPr is None:
                                continue
                            sp = pPr.find(qn("w:spacing"))
                            if sp is None:
                                continue
                            for attr in (qn("w:after"), qn("w:before")):
                                val = sp.get(attr)
                                if val and int(val) > 0:
                                    sp.set(attr, "0")
                        elif child.tag == qn("w:tbl"):
                            # Reducere spacing în tabelul de bare cod
                            all_t = " ".join(t.text for t in child.iter(qn("w:t")) if t.text)
                            if "COD" in all_t.upper() and "BARE" in all_t.upper():
                                for p in child.iter(qn("w:p")):
                                    pPr2 = p.find(qn("w:pPr"))
                                    if pPr2 is None:
                                        continue
                                    sp2 = pPr2.find(qn("w:spacing"))
                                    if sp2 is None:
                                        continue
                                    for attr in (qn("w:after"), qn("w:before")):
                                        val = sp2.get(attr)
                                        if val and int(val) > 0:
                                            sp2.set(attr, "0")

            # ═══════════════════════════════════════
            # 8. SAVE & RETURN
            # ═══════════════════════════════════════
            buf = io.BytesIO()
            doc.save(buf)
            result = buf.getvalue()

            self.send_response(200)
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            self.send_header("Content-Disposition", "attachment; filename=CPE.docx")
            self.send_header("Content-Length", str(len(result)))
            self.end_headers()
            self.wfile.write(result)

        except Exception as e:
            self.send_response(500)
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Content-Type", "application/json")
            self.end_headers()
            self.wfile.write(json.dumps({"error": str(e)}).encode())
