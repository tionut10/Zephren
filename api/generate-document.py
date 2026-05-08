"""
Vercel Python Serverless Function — Generare documente DOCX consolidat (Sprint 14).

Endpoint unic pentru toate tipurile de documente:
  ?type=cpe     → Anexa 1 simplă (blob DOCX)
  ?type=anexa   → Anexa 1+2 extins (blob DOCX)
  ?type=audit   → Raport audit energetic (JSON {docx: base64, filename})
  ?type=anexa1  → Alias cpe (variantă PDF)
  ?type=anexa2  → Alias anexa (variantă PDF)

Retro-compat: dacă nu se trimite ?type=, se folosește body.mode.

Aplică enforce_a4_portrait() pentru toate documentele generate —
Ord. MDLPA 16/2023 cere strict A4 portret + Calibri 11pt.
"""
from http.server import BaseHTTPRequestHandler
from urllib.parse import urlparse, parse_qs
import json, io, base64, re, copy, datetime
from docx import Document
from docx.shared import Inches, Emu, Pt, Cm
from docx.enum.section import WD_ORIENT
from lxml import etree

# ── Sprint 15 — segno pentru QR scanabil (pure Python, ~200 KB) ──
try:
    import segno  # type: ignore
    _SEGNO_AVAILABLE = True
except ImportError:  # pragma: no cover — dev fallback
    segno = None
    _SEGNO_AVAILABLE = False

# ── #13 (audit Pas 6+7 V8, 7 mai 2026) — Pillow pentru bar chart EP per utilitate ──
try:
    from PIL import Image as _PIL_Image, ImageDraw as _PIL_ImageDraw, ImageFont as _PIL_ImageFont  # type: ignore
    _PIL_AVAILABLE = True
except ImportError:  # pragma: no cover — dev fallback (Vercel cold start poate lipsi temporar)
    _PIL_Image = None
    _PIL_ImageDraw = None
    _PIL_ImageFont = None
    _PIL_AVAILABLE = False


# ═══════════════════════════════════════════════════════════════════════
# #13 — GENERATOR BAR CHART EP PER UTILITATE (PILLOW)
# ═══════════════════════════════════════════════════════════════════════
# Render distribuția consumului energetic per utilitate (Mc 001-2022 §5.1)
# ca PNG cu bar chart orizontal colorat după clasă energetică A+..G
# (culorile standard MDLPA verde→roșu).
#
# Variantă LIGHTWEIGHT (Pillow only, fără matplotlib) — bundle <10MB pentru
# compatibilitate Vercel Hobby. Suficient pentru o vizualizare clară a
# distribuției EP per serviciu (încălzire / ACM / răcire / ventilare / iluminat)
# în Anexa 2 DOCX. Pentru vizualizări mai elaborate (Sankey energy flow),
# necesită upgrade la Vercel Pro + matplotlib.

# Culori clase energetice — identice cu _COL_FILLS folosite la utility cells
_ENERGY_CLASS_COLORS = {
    "A+": (0, 155, 0),     # #009B00 — verde închis
    "A":  (50, 200, 49),   # #32C831
    "B":  (0, 255, 0),     # #00FF00
    "C":  (255, 255, 0),   # #FFFF00
    "D":  (243, 156, 0),   # #F39C00
    "E":  (255, 100, 0),   # #FF6400
    "F":  (254, 65, 1),    # #FE4101
    "G":  (254, 0, 0),     # #FE0000
    "—":  (180, 180, 180), # gri pentru clasă necunoscută
}


def render_ep_distribution_chart(ep_per_service, cls_per_service, width_px=1200, height_px=600):
    """Generează un bar chart orizontal cu EP per utilitate, colorat după clasă.

    Args:
        ep_per_service: dict {nume_serviciu: ep_value_kwh_m2}
            ex: {"Încălzire": 641.0, "ACM": 171.8, "Răcire": 0.0, ...}
        cls_per_service: dict {nume_serviciu: clasă_string}
            ex: {"Încălzire": "G", "ACM": "G", "Răcire": "A+", ...}
        width_px / height_px: dimensiuni imagine PNG.

    Returns:
        bytes — conținut PNG, sau None dacă Pillow indisponibil.
    """
    if not _PIL_AVAILABLE:
        return None
    try:
        # Filtrăm utilități cu valoare > 0 (păstrăm ordinea de inserare a dict-ului)
        items = [(name, val, cls_per_service.get(name, "—"))
                 for name, val in ep_per_service.items() if val > 0]
        if not items:
            return None

        # Margini și layout
        MARGIN_L = 240      # spațiu pentru etichete utilitate (font mărit 18pt)
        MARGIN_R = 200      # spațiu pentru valoare + clasă (font mărit 18pt)
        MARGIN_T = 80       # spațiu titlu mărit 26pt + subtitlu 18pt
        MARGIN_B = 50
        bar_area_w = width_px - MARGIN_L - MARGIN_R
        bar_area_h = height_px - MARGIN_T - MARGIN_B
        bar_height = max(30, int(bar_area_h / max(1, len(items)) * 0.7))
        bar_spacing = int(bar_area_h / max(1, len(items)))

        # Valoarea maximă pentru scalare proporțională
        max_val = max(v for _, v, _ in items)
        if max_val <= 0:
            return None

        # Create canvas
        img = _PIL_Image.new("RGB", (width_px, height_px), (255, 255, 255))
        draw = _PIL_ImageDraw.Draw(img)

        # Title
        try:
            title_font = _PIL_ImageFont.truetype("DejaVuSans-Bold.ttf", 26)
            label_font = _PIL_ImageFont.truetype("DejaVuSans.ttf", 18)
            value_font = _PIL_ImageFont.truetype("DejaVuSans-Bold.ttf", 18)
        except Exception:
            # Fallback la default font dacă DejaVu lipsește (Vercel Linux base)
            title_font = _PIL_ImageFont.load_default()
            label_font = _PIL_ImageFont.load_default()
            value_font = _PIL_ImageFont.load_default()

        draw.text((MARGIN_L, 12), "Distribuția EP per utilitate (kWh/m²·an)",
                  fill=(13, 71, 161), font=title_font)
        draw.text((MARGIN_L, 46),
                  "Conform Mc 001-2022 §5.1 + Tab I.1 (clase per serviciu)",
                  fill=(100, 100, 130), font=label_font)

        # Draw bars
        for i, (name, val, cls) in enumerate(items):
            y = MARGIN_T + i * bar_spacing
            bar_w = int(bar_area_w * val / max_val)
            color = _ENERGY_CLASS_COLORS.get(cls, _ENERGY_CLASS_COLORS["—"])
            # Border bar (chenar negru subțire pentru contrast)
            draw.rectangle([MARGIN_L, y, MARGIN_L + bar_w, y + bar_height],
                           fill=color, outline=(50, 50, 50), width=1)
            # Etichetă utilitate (la stânga)
            draw.text((10, y + bar_height // 2 - 10), name,
                      fill=(0, 0, 0), font=label_font)
            # Valoare + clasă (la dreapta)
            value_text = f"{val:.1f} kWh/m²·an"
            draw.text((MARGIN_L + bar_w + 10, y + bar_height // 2 - 16), value_text,
                      fill=(0, 0, 0), font=value_font)
            class_text = f"({cls})"
            draw.text((MARGIN_L + bar_w + 10, y + bar_height // 2 + 6), class_text,
                      fill=color, font=value_font)

        # Legendă culori clase (jos)
        legend_y = MARGIN_T + len(items) * bar_spacing + 15
        if legend_y < height_px - 30:
            legend_x = MARGIN_L
            draw.text((legend_x, legend_y), "Clasă:", fill=(80, 80, 100), font=label_font)
            legend_x += 80
            for cls in ["A+", "A", "B", "C", "D", "E", "F", "G"]:
                color = _ENERGY_CLASS_COLORS[cls]
                draw.rectangle([legend_x, legend_y, legend_x + 22, legend_y + 18],
                               fill=color, outline=(50, 50, 50))
                draw.text((legend_x + 26, legend_y - 1), cls,
                          fill=(0, 0, 0), font=label_font)
                legend_x += 64

        # Export PNG bytes
        buf = io.BytesIO()
        img.save(buf, format="PNG", optimize=True)
        return buf.getvalue()
    except Exception as _err:
        print(f"[render_ep_distribution_chart] eroare: {_err}", flush=True)
        return None


# ═══════════════════════════════════════════════════════
# A4 PORTRAIT ENFORCEMENT — Ord. MDLPA 16/2023 Anexa 1
# ═══════════════════════════════════════════════════════
# A4: 11906 × 16838 DXA (1 DXA = 1/1440 inch = 635 EMU)
A4_WIDTH_EMU = 11906 * 635     # 21.0 cm
A4_HEIGHT_EMU = 16838 * 635    # 29.7 cm
A4_MARGIN_EMU = 1417 * 635     # 2.5 cm


def enforce_a4_portrait(doc, preserve_margins=False, preserve_fonts=False):
    """Forțează toate secțiunile la A4 portret. Margini + fonturi opționale.

    preserve_margins=True — păstrează marginile originale ale template-ului
    (template-urile MDLPA CPE au 20mm top/bot, 25mm left/right).

    preserve_fonts=True — NU suprascrie fontul Normal style. Folosit pentru
    template-urile oficiale MDLPA care au deja fonturi corecte setate la nivel
    de run/paragraf; suprascrierea Normal.font cascadează pe TOATE paragrafele
    și modifică layout-ul față de modelul oficial.
    """
    for section in doc.sections:
        section.page_width = Emu(A4_WIDTH_EMU)
        section.page_height = Emu(A4_HEIGHT_EMU)
        section.orientation = WD_ORIENT.PORTRAIT
        if not preserve_margins:
            section.top_margin = Emu(A4_MARGIN_EMU)
            section.bottom_margin = Emu(A4_MARGIN_EMU)
            section.left_margin = Emu(A4_MARGIN_EMU)
            section.right_margin = Emu(A4_MARGIN_EMU)
    if not preserve_fonts:
        try:
            if "Normal" in doc.styles:
                normal = doc.styles["Normal"]
                if normal.font.name is None:
                    normal.font.name = "Calibri"
                if normal.font.size is None:
                    normal.font.size = Pt(11)
        except Exception:
            pass
    return doc

# ═══════════════════════════════════════════════════════
# CONSTANTE — scalele EP și CO2 din template-uri (Mc 001-2022)
# Template-urile rezidențiale folosesc varianta _cool
# ═══════════════════════════════════════════════════════
EP_TEMPLATE_SCALES = {
    "RI": [91, 129, 257, 390, 522, 652, 783],
    "RC": [73, 101, 198, 297, 396, 495, 595],
    "RA": [73, 101, 198, 297, 396, 495, 595],
    "BI": [68, 97, 193, 302, 410, 511, 614],
    "ED": [48, 68, 135, 246, 358, 447, 536],
    "SA": [117, 165, 331, 501, 671, 838, 1005],
    "HC": [67, 93, 188, 321, 452, 565, 678],
    "CO": [88, 124, 248, 320, 393, 492, 591],
    "SP": [75, 104, 206, 350, 494, 617, 741],
    "AL": [68, 97, 193, 302, 410, 511, 614],
}

CO2_TEMPLATE_SCALES = {
    "RI": [16.1, 22.8, 45.5, 70.1, 94.8, 118.4, 142.1],
    "RC": [12.7, 17.6, 34.6, 52.2, 69.9, 87.4, 104.9],
    "RA": [12.7, 17.6, 34.6, 52.2, 69.9, 87.4, 104.9],
    "BI": [10.4, 14.8, 29.7, 46.1, 62.4, 77.8, 93.4],
    "ED": [8.3, 11.6, 23.0, 42.5, 62.2, 77.6, 93.1],
    "SA": [19.7, 27.8, 55.8, 84.0, 112.3, 140.2, 168.1],
    "HC": [11.8, 16.4, 33.1, 57.0, 80.6, 100.7, 120.8],
    "CO": [15.4, 21.6, 43.4, 54.5, 65.7, 82.3, 98.9],
    "SP": [12.3, 17.0, 33.7, 57.4, 81.2, 101.4, 121.7],
    "AL": [10.4, 14.8, 29.7, 46.1, 62.4, 77.8, 93.4],
}

NSMAP = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def qn(tag):
    """Qualified name helper for lxml."""
    prefix, local = tag.split(":")
    return "{%s}%s" % (NSMAP[prefix], local)


# ═══════════════════════════════════════════════════════
# REPLACE TEXT ACROSS RUNS — păstrează formatarea
# ═══════════════════════════════════════════════════════
def replace_in_paragraph(para, old_text, new_text, count=0):
    """Replace old_text with new_text in a paragraph, across split runs.
    Returns number of replacements made. If count>0, limits replacements."""
    full = para.text
    if old_text not in full:
        return 0

    # Build a map of character positions to runs
    runs = para.runs
    if not runs:
        return 0

    char_map = []  # [(run_idx, char_idx_in_run)]
    for ri, run in enumerate(runs):
        for ci in range(len(run.text)):
            char_map.append((ri, ci))

    replacements = 0
    offset = 0
    while True:
        pos = full.find(old_text, offset)
        if pos == -1:
            break
        if count > 0 and replacements >= count:
            break

        # Find which runs this spans
        start_ri, start_ci = char_map[pos]
        end_pos = pos + len(old_text) - 1
        end_ri, end_ci = char_map[end_pos]

        # Put replacement text in the first run at the start position
        first_run = runs[start_ri]
        before = first_run.text[:start_ci]
        if start_ri == end_ri:
            after = first_run.text[end_ci + 1:]
            first_run.text = before + new_text + after
        else:
            first_run.text = before + new_text
            # Clear text from intermediate and last runs
            for ri in range(start_ri + 1, end_ri):
                runs[ri].text = ""
            last_run = runs[end_ri]
            last_run.text = last_run.text[end_ci + 1:]

        replacements += 1
        # Rebuild full text and char_map for next iteration
        full = para.text
        char_map = []
        for ri, run in enumerate(runs):
            for ci in range(len(run.text)):
                char_map.append((ri, ci))
        offset = pos + len(new_text)

    return replacements


def _iter_all_paragraphs(doc, include_txbx=True):
    """Generator: yields all Paragraph objects from body, tables, and optionally text boxes."""
    from docx.text.paragraph import Paragraph as DocxPara
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para
    if include_txbx:
        txbx_tag = qn("w:txbxContent")
        p_tag = qn("w:p")
        for txbx_elem in doc.element.body.iter(txbx_tag):
            for p_elem in txbx_elem.iter(p_tag):
                yield DocxPara(p_elem, None)


# ═══════════════════════════════════════════════════════
# Sprint 15 — Semnătură + Ștampilă + QR helper (injecție imagini)
# ═══════════════════════════════════════════════════════

def _iter_paragraphs_for_placeholder(doc):
    """Iterator peste paragrafe + celule de tabel — pentru căutarea placeholder-urilor
    în întregul document (inclusiv tabele)."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _replace_placeholder_with_image(doc, placeholder, img_bytes, width_cm=4.0):
    """Găsește placeholder în text și îl înlocuiește cu o imagine PNG embed.

    Întoarce numărul de înlocuiri efectuate. Șterge textul placeholder-ului și
    adaugă imaginea în același paragraf. Ordinea textului rămas este păstrată.
    """
    if not img_bytes or not placeholder:
        return 0
    count = 0
    for p in _iter_paragraphs_for_placeholder(doc):
        if placeholder not in p.text:
            continue
        # Găsim run-ul care conține placeholder-ul și îl curățăm
        for run in p.runs:
            if placeholder in run.text:
                # Împărțim textul în [before][placeholder][after]
                before, _, after = run.text.partition(placeholder)
                run.text = before
                # Inserăm imaginea ca run nou (după run-ul curent)
                img_run = p.add_run()
                try:
                    img_run.add_picture(io.BytesIO(img_bytes), width=Cm(width_cm))
                    count += 1
                except Exception:
                    # Imaginea invalidă — lasă placeholder șters fără imagine
                    pass
                if after:
                    p.add_run(after)
                break  # O singură injecție per paragraf (prima apariție)
    return count


def insert_signature_stamp(doc, signature_b64, stamp_b64):
    """Înlocuiește placeholder-urile {{SEMNATURA}} / {{STAMPILA}} cu imaginile PNG.

    Sprint 15 — Ord. MDLPA 16/2023 + uzanță juridică CPE autentificat.

    Semnătură: 5 cm lățime, ștampilă: 3 cm lățime. Dacă template-ul nu are
    placeholder, imaginile sunt adăugate la finalul documentului (ca ultimă soluție).
    """
    sig_count = 0
    stamp_count = 0

    if signature_b64:
        try:
            sig_bytes = base64.b64decode(signature_b64)
            for ph in ["{{SEMNATURA}}", "{{SIGNATURE}}", "[[SIGNATURE]]", "{SEMNATURA}"]:
                sig_count += _replace_placeholder_with_image(doc, ph, sig_bytes, width_cm=5.0)
        except Exception:
            pass

    if stamp_b64:
        try:
            stamp_bytes = base64.b64decode(stamp_b64)
            # Sprint 8 mai 2026 — Ștampila auditorului = Ø 40 mm conform Anexa 1b
            # Ord. MDLPA 348/2026 (Art. 5 alin. 5: „Este interzisă utilizarea de
            # ștampile cu alte dimensiuni decât cele precizate în anexa 1b").
            # Era 3.0 cm — corectat la 4.0 cm (40 mm) standard profesional RO.
            for ph in ["{{STAMPILA}}", "{{STAMP}}", "[[STAMP]]", "{STAMPILA}"]:
                stamp_count += _replace_placeholder_with_image(doc, ph, stamp_bytes, width_cm=4.0)
        except Exception:
            pass

    # Fallback: dacă template-ul nu are placeholder-uri pentru semnătură/ștampilă,
    # injectăm imaginile DIRECT în paragraful existent care conține textul
    # „Semnătura și ștampila auditorului" (deja prezent în template-ele MDLPA).
    # NU adăugăm paragraf nou la final — asta crea o pagină goală suplimentară
    # pentru CPE (raportat 2 mai 2026, fix imediat după dezactivare append_legal_supplement).
    if (signature_b64 or stamp_b64) and sig_count == 0 and stamp_count == 0:
        target_para = None
        # Caut prima apariție în paragrafele de top-level (corp document, NU footere)
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if "Semnătura" in t and "tampila" in t:
                target_para = p
                break
        # Și în tabele (template-ul MDLPA poate avea textul în celulă)
        if target_para is None:
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            t = (p.text or "").strip()
                            if "Semnătura" in t and "tampila" in t:
                                target_para = p
                                break
                        if target_para: break
                    if target_para: break
                if target_para: break

        try:
            if target_para is not None:
                # Păstrez textul „Semnătura și ștampila auditorului" intact (etichetă
                # conform modelului oficial MDLPA în colțul dreapta) și adaug imaginile
                # SUB text printr-un line break în același paragraf — astfel ele rămân
                # în aceeași poziție pe pagină, fără a crea o pagină nouă.
                from docx.oxml import OxmlElement as _OxEl_sig
                # Adaug w:br pentru new line în paragraf
                br = _OxEl_sig("w:br")
                last_run = target_para.runs[-1] if target_para.runs else target_para.add_run()
                last_run._r.append(br)
                # Acum adaug imaginile ca run-uri noi (vor fi pe linia următoare)
                if signature_b64:
                    sig_bytes = base64.b64decode(signature_b64)
                    target_para.add_run().add_picture(io.BytesIO(sig_bytes), width=Cm(5.0))
                    target_para.add_run("  ")
                if stamp_b64:
                    stamp_bytes = base64.b64decode(stamp_b64)
                    # Ord. MDLPA 16/2023 Art. 12 alin. (3): ștampila auditorului = Ø 40 mm fix
                    target_para.add_run().add_picture(io.BytesIO(stamp_bytes), width=Cm(4.0), height=Cm(4.0))
            else:
                # Doar dacă NU există paragraf țintă în template — fallback legacy
                p = doc.add_paragraph()
                p.add_run("Semnătură auditor / Ștampilă: ")
                if signature_b64:
                    sig_bytes = base64.b64decode(signature_b64)
                    p.add_run().add_picture(io.BytesIO(sig_bytes), width=Cm(5.0))
                    p.add_run("  ")
                if stamp_b64:
                    stamp_bytes = base64.b64decode(stamp_b64)
                    # Ștampilă Ø 40 mm fix (Ord. MDLPA 16/2023 Art. 12 alin. 3)
                    p.add_run().add_picture(io.BytesIO(stamp_bytes), width=Cm(4.0), height=Cm(4.0))
        except Exception:
            pass

    return {"signature": sig_count, "stamp": stamp_count}


def generate_qr_png(url, scale=5, border=2):
    """Generează un PNG QR scanabil pentru URL-ul dat.

    Folosește segno (pure Python, ~200 KB — NU qrcode+Pillow ~10 MB).
    Error correction M = ~15% rezistență la degradare.
    """
    if not _SEGNO_AVAILABLE or not url:
        return None
    try:
        qr = segno.make(url, error="m")
        buf = io.BytesIO()
        qr.save(buf, kind="png", scale=scale, border=border)
        return buf.getvalue()
    except Exception:
        return None


def insert_qr_code(doc, verify_url, cpe_code=""):
    """Inserează QR code scanabil pentru URL verificare.

    Sprint 15 — autentificare vizuală CPE.
    Audit 2 mai 2026 — P0.3: URL-ul: https://zephren.ro/cpe/verifica?cod={cpeCode}
    (registrul MDLPA central nu există — landing static cu form căutare manuală).

    Placeholder-uri suportate: {{QR_CODE}}, {{QR}}, [[QR_CODE]]
    Dacă template-ul nu are placeholder, QR nu se inserează (evită glitching DOCX).
    """
    if not verify_url:
        return 0
    qr_bytes = generate_qr_png(verify_url)
    if not qr_bytes:
        return 0
    count = 0
    for ph in ["{{QR_CODE}}", "{{QR}}", "[[QR_CODE]]", "{QR_CODE}"]:
        count += _replace_placeholder_with_image(doc, ph, qr_bytes, width_cm=2.5)
    return count


# ═══════════════════════════════════════════════════════
# Sprint 2 mai 2026 — BIFARE CHECKBOX-URI ÎN ANEXA 1+2
# ═══════════════════════════════════════════════════════
# Template-ele Anexa folosesc Word Form Fields legacy:
#   <w:fldChar w:fldCharType="begin">
#     <w:ffData>
#       <w:checkBox>
#         <w:default w:val="0"/>  ← 0 = unchecked, 1 = checked
#       </w:checkBox>
#     </w:ffData>
#   </w:fldChar>
# Bifarea = schimbare default w:val 0→1 + adăugare <w:checked w:val="1"/>.
# ═══════════════════════════════════════════════════════

def check_form_checkbox_in_para(para):
    """Bifează PRIMUL checkbox legacy din paragraful dat.

    Returnează True dacă a bifat, False dacă nu a găsit checkbox.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for fldChar in para._p.iter(qn("w:fldChar")):
        ffData = fldChar.find(qn("w:ffData"))
        if ffData is None:
            continue
        checkBox = ffData.find(qn("w:checkBox"))
        if checkBox is None:
            continue
        # Setez default w:val="1"
        default = checkBox.find(qn("w:default"))
        if default is not None:
            default.set(qn("w:val"), "1")
        else:
            d = OxmlElement("w:default")
            d.set(qn("w:val"), "1")
            checkBox.append(d)
        # Adaug/actualizez checked w:val="1"
        checked = checkBox.find(qn("w:checked"))
        if checked is None:
            c = OxmlElement("w:checked")
            c.set(qn("w:val"), "1")
            checkBox.append(c)
        else:
            checked.set(qn("w:val"), "1")
        return True
    return False


def check_form_checkbox_in_cell(cell):
    """Bifează PRIMUL checkbox din celula dată (toate paragrafele cellul)."""
    for p in cell.paragraphs:
        if check_form_checkbox_in_para(p):
            return True
    return False


def find_paragraph_containing(doc, text_substring, start_idx=0):
    """Returnează (idx, para) primul paragraf care conține text_substring."""
    for i, p in enumerate(doc.paragraphs[start_idx:], start=start_idx):
        if text_substring in p.text:
            return (i, p)
    return (-1, None)


def check_box_for_text(doc, text_substring, occurrence=1):
    """Bifează checkbox-ul DIN paragraful care conține text_substring.

    occurrence: 1-based index pentru ocurențe multiple.

    Returnează True dacă a bifat cu succes.
    """
    seen = 0
    for p in doc.paragraphs:
        if text_substring in p.text:
            seen += 1
            if seen == occurrence:
                return check_form_checkbox_in_para(p)
    return False


def check_box_in_table(doc, table_idx, row_idx, col_idx):
    """Bifează checkbox-ul din celula T[ti]R[ri]C[ci]."""
    if table_idx >= len(doc.tables):
        return False
    tbl = doc.tables[table_idx]
    if row_idx >= len(tbl.rows) or col_idx >= len(tbl.rows[row_idx].cells):
        return False
    return check_form_checkbox_in_cell(tbl.rows[row_idx].cells[col_idx])


def check_nth_checkbox_in_para(para, n=0):
    """Bifează al N-lea checkbox legacy din paragraf (0-based).

    Returnează True dacă a găsit și bifat checkbox-ul indexat.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    seen = -1
    for fldChar in para._p.iter(qn("w:fldChar")):
        ffData = fldChar.find(qn("w:ffData"))
        if ffData is None:
            continue
        checkBox = ffData.find(qn("w:checkBox"))
        if checkBox is None:
            continue
        seen += 1
        if seen != n:
            continue
        default = checkBox.find(qn("w:default"))
        if default is not None:
            default.set(qn("w:val"), "1")
        else:
            d = OxmlElement("w:default")
            d.set(qn("w:val"), "1")
            checkBox.append(d)
        checked = checkBox.find(qn("w:checked"))
        if checked is None:
            c = OxmlElement("w:checked")
            c.set(qn("w:val"), "1")
            checkBox.append(c)
        else:
            checked.set(qn("w:val"), "1")
        return True
    return False


def check_box_after_label(doc, label_text, option_idx=0):
    """Caută paragraful cu label_text și bifează al option_idx-lea checkbox
    din paragraful URMĂTOR (i+1). Pattern frecvent în Anexa MDLPA unde
    eticheta și opțiunile sunt pe linii consecutive.

    Audit 2 mai 2026 — pentru Conducta de recirculare a acc, Contor general
    de căldură pentru acc, Debitmetre la nivelul punctelor de consum, etc.
    """
    paragraphs = doc.paragraphs
    for i, p in enumerate(paragraphs):
        if label_text in p.text:
            # Caut în paragrafele următoare (până la 3 linii) după label
            for j in range(i + 1, min(i + 4, len(paragraphs))):
                next_p = paragraphs[j]
                # Verific dacă next_p are checkbox-uri
                from docx.oxml.ns import qn as _qn_local
                has_cb = any(
                    fc.find(_qn_local("w:ffData")) is not None and
                    fc.find(_qn_local("w:ffData")).find(_qn_local("w:checkBox")) is not None
                    for fc in next_p._p.iter(_qn_local("w:fldChar"))
                )
                if has_cb:
                    return check_nth_checkbox_in_para(next_p, option_idx)
            return False
    return False


# NOTĂ 2 mai 2026: am eliminat funcțiile merge_duplicate_year_cells +
# label_year_cells_as_built_vs_renov după ce am descoperit că template-ul
# oficial MDLPA are deja vMerge aplicat pe celula „Anul construirii/renovării
# majore" (R1C1._tc IS R2C1._tc). Eticheta apare O SINGURĂ DATĂ — nu există
# duplicate de eliminat. Populare normală a template-ului (replace AAAA →
# yearBuilt) funcționează corect fără intervenție suplimentară.


# ═══════════════════════════════════════════════════════
# Sprint 14/15/17 — PAGINĂ SUPLIMENT DATE LEGALE CPE (Etapa 1, 19 apr 2026)
# ═══════════════════════════════════════════════════════
# Template-ele oficiale MDLPA (2-12-CPE-*.docx) NU conțin placeholder-e pentru
# cod unic CPE, QR, semnătură, cadastru, valabilitate, pașaport renovare.
# Pentru a NU modifica template-ele oficiale (risc conformitate Ord. MDLPA 16/2023),
# adăugăm o pagină supliment finală cu toate metadatele legale cerute de:
#   - Ord. MDLPA 16/2023 (cod unic, semnătură, ștampilă, QR)
#   - L.238/2024 Art. 19 (înregistrare în max 30 zile + verificabilitate)
#   - EPBD 2024/1275 Art. 12 (pașaport renovare) și Art. 14 (calitate aer + EV)
#   - EPBD 2024/1275 Art. 17 (valabilitate diferențiată 5/10 ani per clasă)
# ═══════════════════════════════════════════════════════

_SUPPLEMENT_TITLE = "DATE LEGALE CONFORM ORD. MDLPA 16/2023 + L.238/2024 + EPBD 2024/1275"

_SUPPLEMENT_FOOTER_NOTE = (
    "Pagina suplimentară generată automat de Zephren conform Ord. MDLPA 16/2023, "
    "L.238/2024 Art. 19 și EPBD 2024/1275 Art. 12-14, 17. Codul unic CPE și UUID-ul "
    "pașaportului de renovare sunt deterministice (UUID v5, RFC 4122) — pot fi "
    "verificate online la URL-urile indicate."
)


def _set_cell_borders(cell, color="000000", sz="4"):
    """Aplică borduri single 0.5pt pe toate laturile celulei.

    sz=4 (1/8 pt) ≈ 0.5pt linie subțire, conform stilului tehnic MDLPA.
    """
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    # Înlocuiește orice tcBorders existent
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_pr.append(tc_borders)


def _add_label_value_row(table, label, value, bold_label=True):
    """Adaugă un rând cu etichetă (col 1) + valoare text (col 2). Valori goale → '—'."""
    row = table.add_row()
    cell_label = row.cells[0]
    cell_value = row.cells[1]
    cell_label.text = ""
    cell_value.text = ""

    p_l = cell_label.paragraphs[0]
    r_l = p_l.add_run(str(label))
    r_l.font.size = Pt(10)
    r_l.font.name = "Calibri"
    if bold_label:
        r_l.bold = True

    val_str = str(value).strip() if value is not None else ""
    p_v = cell_value.paragraphs[0]
    r_v = p_v.add_run(val_str if val_str else "—")
    r_v.font.size = Pt(10)
    r_v.font.name = "Calibri"
    if not val_str:
        r_v.italic = True

    _set_cell_borders(cell_label)
    _set_cell_borders(cell_value)
    return row


def _add_label_image_row(table, label, img_b64, width_cm=4.0):
    """Adaugă un rând cu etichetă (col 1) + imagine PNG (col 2). Lipsă → '— (lipsă)'."""
    row = table.add_row()
    cell_label = row.cells[0]
    cell_value = row.cells[1]
    cell_label.text = ""
    cell_value.text = ""

    p_l = cell_label.paragraphs[0]
    r_l = p_l.add_run(str(label))
    r_l.font.size = Pt(10)
    r_l.font.name = "Calibri"
    r_l.bold = True

    p_v = cell_value.paragraphs[0]
    if not img_b64:
        r_v = p_v.add_run("— (lipsă)")
        r_v.font.size = Pt(10)
        r_v.italic = True
    else:
        try:
            if isinstance(img_b64, (bytes, bytearray)):
                img_bytes = bytes(img_b64)
            else:
                # Strip data URL prefix dacă e prezent
                s = str(img_b64)
                if s.startswith("data:") and "," in s:
                    s = s.split(",", 1)[1]
                img_bytes = base64.b64decode(s)
            run_img = p_v.add_run()
            run_img.add_picture(io.BytesIO(img_bytes), width=Cm(width_cm))
        except Exception as e_img:
            r_v = p_v.add_run(f"[eroare imagine: {e_img}]")
            r_v.font.size = Pt(9)
            r_v.italic = True

    _set_cell_borders(cell_label)
    _set_cell_borders(cell_value)
    return row


def append_legal_supplement(doc, data):
    """Adaugă pagină supliment cu toate metadatele legale CPE.

    Apelată DOAR pentru mode == "cpe" și DOAR dacă cpe_code e prezent.
    Returnează True dacă pagina a fost adăugată, False altfel.

    Sprint 14/15/17 — Etapa 1 audit 19 apr 2026.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cpe_code = (data.get("cpe_code") or "").strip()
    if not cpe_code:
        return False

    # 1) Page break — supliment pe pagină proprie
    doc.add_page_break()

    # 2) Titlu mare centrat
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(_SUPPLEMENT_TITLE)
    title_run.bold = True
    title_run.font.size = Pt(11)
    title_run.font.name = "Calibri"

    # 3) Tabel 2 coloane
    tbl = doc.add_table(rows=0, cols=2)
    try:
        tbl.style = "Table Grid"
    except Exception:
        # Style "Table Grid" nu există în template-ele MDLPA — borduri se aplică manual via _set_cell_borders
        pass
    tbl.autofit = False
    # Lățimi: ~5.5cm etichetă, ~10.5cm valoare (total ~16cm — încape în margini A4 implicite)
    try:
        for row in [tbl.rows[i] for i in range(len(tbl.rows))]:
            row.cells[0].width = Cm(5.5)
            row.cells[1].width = Cm(10.5)
    except Exception:
        pass

    # 4) Rânduri text — Sprint 14/15/17 metadata
    rows_text = [
        ("Cod unic CPE",                cpe_code),
        ("Data emiterii",               data.get("auditor_date", "")),
        ("Data expirării",              data.get("expiry", "")),
        ("Valabilitate (ani)",          data.get("validity_years", "")),
        ("Etichetă valabilitate",       data.get("validity_label", "")),
        ("Nr. cadastral",               data.get("cadastral_number", "")),
        ("Carte funciară",              data.get("land_book", "")),
        ("Arie construită desfășurată", _format_area(data.get("area_built", ""))),
        ("Nr. apartamente",             data.get("n_apartments", "")),
        ("UUID pașaport renovare",      data.get("passport_uuid", "")),
        ("URL verificare CPE",          data.get("qr_verify_url", "")),
        ("URL pașaport renovare",       data.get("passport_url", "")),
    ]
    for label, value in rows_text:
        _add_label_value_row(tbl, label, value)

    # Aplică lățimi și pe rândurile noi
    try:
        for row in tbl.rows:
            row.cells[0].width = Cm(5.5)
            row.cells[1].width = Cm(10.5)
    except Exception:
        pass

    # 5) Rânduri imagini — semnătură + QR-uri (ștampila se aplică manual pe printout)
    _add_label_image_row(tbl, "Semnătură auditor",   data.get("signature_png_b64", ""), width_cm=4.0)

    # QR verificare CPE — generat din qr_verify_url
    qr_verify_bytes = generate_qr_png(data.get("qr_verify_url", ""), scale=4, border=2)
    _add_label_image_row(tbl, "QR verificare CPE",   qr_verify_bytes,                   width_cm=2.5)

    # QR pașaport — generat din passport_qr_url (fallback passport_url)
    passport_qr_url = data.get("passport_qr_url", "") or data.get("passport_url", "")
    qr_pass_bytes = generate_qr_png(passport_qr_url, scale=4, border=2) if passport_qr_url else None
    _add_label_image_row(tbl, "QR pașaport renovare", qr_pass_bytes,                    width_cm=2.0)

    # Re-aplică lățimile (rândurile noi se reset la default)
    try:
        for row in tbl.rows:
            row.cells[0].width = Cm(5.5)
            row.cells[1].width = Cm(10.5)
    except Exception:
        pass

    # 6) Indicatori EPBD 2024 Art. 14 — calitate aer + EV charging
    epbd_p = doc.add_paragraph()
    epbd_p.paragraph_format.space_before = Pt(6)
    epbd_run_t = epbd_p.add_run("EPBD 2024/1275 Art. 14 — Calitate aer interior + Mobilitate electrică:")
    epbd_run_t.bold = True
    epbd_run_t.font.size = Pt(10)

    co2_ppm = data.get("co2_max_ppm", "") or "—"
    pm25 = data.get("pm25_avg", "") or "—"
    ev_inst = data.get("ev_charging_points", "0") or "0"
    ev_prep = data.get("ev_charging_prepared", "0") or "0"
    epbd_p2 = doc.add_paragraph()
    epbd_run = epbd_p2.add_run(
        f"CO₂ max: {co2_ppm} ppm   |   PM2.5 mediu: {pm25} μg/m³   |   "
        f"Puncte încărcare EV: {ev_inst} instalate + {ev_prep} pregătite"
    )
    epbd_run.font.size = Pt(10)

    # 6b) BACS + SRI + n50 (Etapa 2 BUG-1, BUG-2, BUG-3 — propagate corect)
    bacs_class = (data.get("bacs_class") or "").strip()
    sri_total = (data.get("sri_total") or "").strip()
    sri_grade = (data.get("sri_grade") or "").strip()
    n50_val = (data.get("n50") or "").strip()
    if bacs_class or sri_total or n50_val:
        bs_p = doc.add_paragraph()
        bs_p.paragraph_format.space_before = Pt(4)
        bs_t = bs_p.add_run("Performanță automatizare + etanșeitate:")
        bs_t.bold = True
        bs_t.font.size = Pt(10)
        parts = []
        if bacs_class:
            parts.append(f"Clasa BACS: {bacs_class} (SR EN ISO 52120-1:2022)")
        if sri_total:
            parts.append(f"SRI: {sri_total}%" + (f" (clasa {sri_grade})" if sri_grade else ""))
        if n50_val:
            parts.append(f"n₅₀: {n50_val} h⁻¹ (etanșeitate test blower door)")
        bs_p2 = doc.add_paragraph()
        bs_p2.add_run("   |   ".join(parts)).font.size = Pt(10)

    # 6c) Penalizări Mc 001-2022 Partea III §8.10 (BUG-7 — Etapa 2)
    penalties_raw = data.get("penalties_summary") or ""
    if penalties_raw:
        try:
            import json as _json
            pen_data = _json.loads(penalties_raw)
            pen_summary = pen_data.get("summary", {})
            pen_applied = pen_data.get("applied", []) or []
            count_applied = pen_summary.get("count_applied", 0)
            total_pct = pen_summary.get("total_delta_pct", 0)
            ep_mult = pen_summary.get("ep_multiplier", 1)

            pen_p = doc.add_paragraph()
            pen_p.paragraph_format.space_before = Pt(6)
            pen_t = pen_p.add_run(
                f"Penalizări utilizare irațională energie (Mc 001-2022 Partea III §8.10):"
            )
            pen_t.bold = True
            pen_t.font.size = Pt(10)

            pen_p2 = doc.add_paragraph()
            pen_p2.add_run(
                f"Penalizări active: {count_applied}   |   "
                f"Adaos EP total: +{total_pct:.1f}%   |   "
                f"Multiplicator EP: ×{ep_mult:.3f}"
            ).font.size = Pt(10)

            # Listă penalizări aplicate
            if pen_applied:
                for p in pen_applied[:12]:  # max 12 (p0..p11)
                    pid = p.get("id", "?")
                    reason = p.get("reason", "")
                    delta = p.get("delta_EP_pct", 0)
                    li = doc.add_paragraph()
                    li.paragraph_format.left_indent = Cm(0.5)
                    li.add_run(f"• {pid.upper()} (+{delta:.0f}%): {reason}").font.size = Pt(9)
        except Exception as e_pen:
            pen_err = doc.add_paragraph()
            pen_err.add_run(f"[penalizări — eroare parse: {e_pen}]").font.size = Pt(8)

    # 7) Notă finală
    note_p = doc.add_paragraph()
    note_p.paragraph_format.space_before = Pt(8)
    note_run = note_p.add_run(_SUPPLEMENT_FOOTER_NOTE)
    note_run.font.size = Pt(8)
    note_run.italic = True

    return True


def _format_area(val):
    """Adaugă unitate ' m²' la o valoare numerică, dacă lipsește."""
    s = str(val or "").strip()
    if not s:
        return ""
    if "m²" in s or "m2" in s:
        return s
    return f"{s} m²"


# ═══════════════════════════════════════════════════════
# Etapa 4 (BUG-4) — ANEXA BLOC MULTI-APARTAMENT (19 apr 2026)
# ═══════════════════════════════════════════════════════
# Pentru blocurile de locuințe (mode == "anexa_bloc"), pe lângă procesarea
# Anexa standard (checkbox-uri, replacements, fotografii), injectăm un tabel
# detaliat cu fiecare apartament + sumar bloc + listă sisteme comune.
#
# Date de input (body['apartments']):
#   [
#     {"number": "1A", "staircase": "A", "floor": 0, "areaUseful": 58.5,
#      "orientation": ["N","E"], "occupants": 3, "corner": true,
#      "topFloor": false, "groundFloor": true,
#      "epAptM2": 192.4, "co2AptM2": 32.1, "enClass": "C", "co2Class": "D",
#      "posKey": "ground_corner", "posFactor": 1.18, "allocatedPct": 12.4},
#     ...
#   ]
#
# Sumar (body['apartmentSummary']):
#   {"totalAu": 580.0, "epAvgWeighted": 175.3, "co2AvgWeighted": 28.9,
#    "classDistribution": {"C": 5, "D": 3, "E": 1}, "count": 9}
#
# Sisteme comune (body['commonSystems']):
#   {"elevator": {"installed": true, "powerKW": 4.5, "hoursYear": 2000, "fuel": "electric"},
#    "stairsLighting": {"installed": true, "powerKW": 0.6, "hoursYear": 4380},
#    ...}
# ═══════════════════════════════════════════════════════

# Etichetele coloanelor tabelului de apartamente (Mc 001-2022 Cap. 4.7 + Anexa 7)
_APT_TABLE_HEADERS = [
    "Nr.", "Ap.", "Sc.", "Etaj", "Au [m²]", "Orient.",
    "Poziție", "× Corr.", "EP [kWh/(m²·an)]", "Clasă",
    "CO₂ [kg/(m²·an)]", "Clasă CO₂", "Comun [%]",
]

_POSITION_RO_LABELS = {
    "ground_interior": "parter int.",
    "ground_corner":   "parter colț",
    "mid_interior":    "etaj int.",
    "mid_corner":      "etaj colț",
    "top_interior":    "ultim int.",
    "top_corner":      "ultim colț",
}

_COMMON_SYSTEM_LABELS = {
    "elevator":         "Lift",
    "stairsLighting":   "Iluminat scări/holuri",
    "centralHeating":   "Centrală termică comună",
    "commonVentilation": "Ventilație comună",
    "pumpGroup":        "Grup pompe",
}


def _fmt_ro(val, decimals=1):
    """Format numeric cu virgulă românească. Acceptă None/string."""
    try:
        if val is None or val == "" or val == "—":
            return "—"
        return f"{float(val):.{decimals}f}".replace(".", ",")
    except (ValueError, TypeError):
        return "—"


def _fmt_floor(floor):
    """Formatare etaj: 0 sau 'P' → 'P'; restul → string."""
    if floor in (0, "0", "P", "p"):
        return "P"
    return str(floor) if floor not in (None, "") else "—"


def _set_table_borders_all(table, sz="4"):
    """Aplică borduri single 0.5pt pe toate celulele tabelului."""
    from docx.oxml import OxmlElement
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            existing = tc_pr.find(qn("w:tcBorders"))
            if existing is not None:
                tc_pr.remove(existing)
            tc_borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{edge}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), sz)
                b.set(qn("w:color"), "000000")
                tc_borders.append(b)
            tc_pr.append(tc_borders)


def _shade_cell(cell, fill_hex):
    """Aplică culoarea de fundal (hex fără #) pe o celulă."""
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


# Culori per clasă energetică (identic cu CLASS_COLORS frontend)
_EP_CLASS_COLORS_BG = {
    "A+": "009B00", "A": "32C831", "B": "00FF00", "C": "FFFF00",
    "D":  "F39C00", "E": "FF6400", "F": "FE4101", "G": "FE0000",
}


def insert_apartment_table(doc, apartments, summary=None):
    """Injectează un tabel cu lista apartamentelor blocului în document.

    Args:
        doc: docx.Document instanță
        apartments: list[dict] cu cheile esențiale: number, floor, areaUseful,
                    orientation (list), posKey, posFactor, epAptM2, enClass,
                    co2AptM2, co2Class, allocatedPct
        summary: dict cu totalAu, epAvgWeighted, co2AvgWeighted, classDistribution

    Returnează True dacă tabelul a fost adăugat, False dacă lista e goală.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not apartments:
        return False

    # 1) Page break — secțiunea apartamente pe pagină proprie
    doc.add_page_break()

    # 2) Titlu secțiune
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(
        f"ANEXA 2 — APARTAMENTE BLOC ({len(apartments)} apartamente)"
    )
    title_run.bold = True
    title_run.font.size = Pt(11)
    title_run.font.name = "Calibri"

    # Subtitlu metodologie
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(
        "Mc 001-2022 Cap. 4.7 + Anexa 7 — corecție poziție termică (parter/colț/ultim etaj)"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(9)

    # 3) Tabel apartamente
    n_cols = len(_APT_TABLE_HEADERS)
    table = doc.add_table(rows=1, cols=n_cols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.autofit = True

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(_APT_TABLE_HEADERS):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(8)
        _shade_cell(hdr_cells[i], "DDDDDD")

    # Data rows
    for idx, apt in enumerate(apartments):
        row = table.add_row()
        cells = row.cells
        ap_num = str(apt.get("number", "—") or "—")
        staircase = str(apt.get("staircase", "—") or "—")
        floor = _fmt_floor(apt.get("floor"))
        au_raw = apt.get("areaUseful", "")
        au = _fmt_ro(au_raw, 1)
        orient_list = apt.get("orientation") or []
        orient = " ".join(str(o) for o in orient_list) if orient_list else "—"
        pos_key = apt.get("posKey", "")
        pos_label = _POSITION_RO_LABELS.get(pos_key, pos_key or "—")
        pos_factor = apt.get("posFactor")
        pos_factor_str = _fmt_ro(pos_factor, 2) if pos_factor is not None else "—"
        ep = _fmt_ro(apt.get("epAptM2"), 1)
        en_class = str(apt.get("enClass") or "—")
        co2 = _fmt_ro(apt.get("co2AptM2"), 1)
        co2_class = str(apt.get("co2Class") or "—")
        common_pct = _fmt_ro(apt.get("allocatedPct"), 1)

        values = [
            str(idx + 1), ap_num, staircase, floor, au, orient,
            pos_label, pos_factor_str, ep, en_class, co2, co2_class, common_pct,
        ]
        for ci, val in enumerate(values):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            # Alignment per column
            if ci in (4, 7, 8, 10, 12):  # numerice → right
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif ci in (0, 9, 11):  # numerotare + clase → center
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(8)
            run.font.name = "Calibri"

        # Color clase energetice
        if en_class in _EP_CLASS_COLORS_BG:
            _shade_cell(cells[9], _EP_CLASS_COLORS_BG[en_class])
        if co2_class in _EP_CLASS_COLORS_BG:
            _shade_cell(cells[11], _EP_CLASS_COLORS_BG[co2_class])

    # Sumar row (media ponderată)
    if summary:
        sum_row = table.add_row()
        sum_cells = sum_row.cells
        # Coloana 0-3 → "MEDIE PONDERATĂ BLOC" colspan=4
        sum_cells[0].text = ""
        p_lbl = sum_cells[0].paragraphs[0]
        r_lbl = p_lbl.add_run("MEDIE PONDERATĂ BLOC")
        r_lbl.bold = True
        r_lbl.font.size = Pt(9)

        # Au total (col 4)
        sum_cells[4].text = _fmt_ro(summary.get("totalAu"), 1)
        p_au = sum_cells[4].paragraphs[0]
        p_au.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if p_au.runs:
            p_au.runs[0].bold = True
            p_au.runs[0].font.size = Pt(9)

        # EP ponderat (col 8) + clasa (col 9)
        sum_cells[8].text = _fmt_ro(summary.get("epAvgWeighted"), 1)
        sum_cells[8].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if sum_cells[8].paragraphs[0].runs:
            sum_cells[8].paragraphs[0].runs[0].bold = True
            sum_cells[8].paragraphs[0].runs[0].font.size = Pt(9)

        avg_class = str(summary.get("avgEnergyClass") or "—")
        sum_cells[9].text = avg_class
        sum_cells[9].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if sum_cells[9].paragraphs[0].runs:
            sum_cells[9].paragraphs[0].runs[0].bold = True
            sum_cells[9].paragraphs[0].runs[0].font.size = Pt(9)
        if avg_class in _EP_CLASS_COLORS_BG:
            _shade_cell(sum_cells[9], _EP_CLASS_COLORS_BG[avg_class])

        # CO₂ ponderat (col 10) + clasa (col 11)
        sum_cells[10].text = _fmt_ro(summary.get("co2AvgWeighted"), 1)
        sum_cells[10].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if sum_cells[10].paragraphs[0].runs:
            sum_cells[10].paragraphs[0].runs[0].bold = True

        avg_co2 = str(summary.get("avgCo2Class") or "—")
        sum_cells[11].text = avg_co2
        sum_cells[11].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if avg_co2 in _EP_CLASS_COLORS_BG:
            _shade_cell(sum_cells[11], _EP_CLASS_COLORS_BG[avg_co2])

        # Bold summary row + light gray background
        for c in sum_cells:
            _shade_cell(c, "F0F0F0")

    _set_table_borders_all(table)

    # 4) Distribuție clase + total (text simplu sub tabel)
    if summary and summary.get("classDistribution"):
        dist_p = doc.add_paragraph()
        dist_p.paragraph_format.space_before = Pt(6)
        r = dist_p.add_run("Distribuție clase apartamente: ")
        r.bold = True
        r.font.size = Pt(10)
        dist_items = []
        total_count = summary.get("count") or len(apartments)
        for cls, cnt in summary.get("classDistribution", {}).items():
            pct = (cnt / total_count * 100) if total_count > 0 else 0
            dist_items.append(f"{cls}={cnt} ({pct:.0f}%)")
        dist_p.add_run("   |   ".join(dist_items)).font.size = Pt(10)

    return True


def insert_common_systems_section(doc, common_systems):
    """Injectează o secțiune cu sistemele comune ale blocului (lift, iluminat, etc.).

    Args:
        doc: docx.Document
        common_systems: dict {key: {"installed": bool, "powerKW": ..., "hoursYear": ..., "fuel": ...}}

    Returnează True dacă cel puțin un sistem a fost adăugat.
    """
    if not common_systems:
        return False
    installed = [
        (key, sys_data) for key, sys_data in common_systems.items()
        if isinstance(sys_data, dict) and sys_data.get("installed")
    ]
    if not installed:
        return False

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(8)
    r = title_p.add_run("SISTEME COMUNE BLOC")
    r.bold = True
    r.font.size = Pt(10)

    table = doc.add_table(rows=1, cols=4)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    headers = ["Sistem", "Putere [kW]", "Ore/an", "Combustibil"]
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        _shade_cell(c, "DDDDDD")

    for key, sys_data in installed:
        row = table.add_row()
        cells = row.cells
        cells[0].text = _COMMON_SYSTEM_LABELS.get(key, key)
        cells[1].text = _fmt_ro(sys_data.get("powerKW"), 2)
        cells[2].text = str(sys_data.get("hoursYear") or "—")
        cells[3].text = str(sys_data.get("fuel") or "—")
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    _set_table_borders_all(table)
    return True


def replace_in_doc(doc, old_text, new_text, max_count=0):
    """Replace text across entire document (paragraphs + tables + text boxes)."""
    total = 0
    remaining = max_count
    for para in _iter_all_paragraphs(doc, include_txbx=True):
        n = replace_in_paragraph(para, old_text, new_text, remaining if max_count > 0 else 0)
        total += n
        if max_count > 0:
            remaining -= n
            if remaining <= 0:
                return total
    return total


def replace_in_txbx_only(doc, old_text, new_text):
    """Replace text ONLY in text boxes (w:txbxContent). Safe for numeric scale bars."""
    from docx.text.paragraph import Paragraph as DocxPara
    total = 0
    txbx_tag = qn("w:txbxContent")
    p_tag = qn("w:p")
    for txbx_elem in doc.element.body.iter(txbx_tag):
        for p_elem in txbx_elem.iter(p_tag):
            total += replace_in_paragraph(DocxPara(p_elem, None), old_text, new_text)
    return total


def replace_in_vml_raw(doc, old_text, new_text):
    """Replace text in VML shape text nodes (v:t) — fallback pentru barcode-uri Code 39.
    Acoperă cazul în care template-ul folosește v:textbox cu v:t în loc de w:txbxContent."""
    _VML_NS = "urn:schemas-microsoft-com:vml"
    count = 0
    for vt in doc.element.body.iter(f"{{{_VML_NS}}}t"):
        if vt.text and old_text in vt.text:
            vt.text = vt.text.replace(old_text, new_text)
            count += 1
    return count


def replace_barcode_cells(doc, code6):
    """Înlocuiește DOAR celulele 7-12 din barcode-ul CPE cu codul Zephren (6 cifre).

    Structura template MDLPA (rândul 2, tabelul de identificare):
      celule 0-5 = 'r','e','g','r','e','g'  → rămân neschimbate (MDLPA le va completa)
      celula  6  = '/'                       → separator fix
      celule 7-12= 'c','o','d','c','o','d'  → înlocuite cu cele 6 cifre Zephren
    Rezultat barcode: regreg/XXXXXX  (ex: regreg/000042)

    NOTĂ: Folosim lxml direct (row_elem.findall) în loc de row.cells din python-docx,
    deoarece row.cells expandează celulele cu gridSpan/vMerge din rândul anterior,
    deplasând indicii și cauzând înlocuirea celulelor greșite (0-5 în loc de 7-12).
    """
    if not code6:
        return False
    # Normalizare: exact 6 caractere, zero-padded dacă numeric
    if code6.isdigit():
        code6 = code6.zfill(6)[-6:]
    else:
        code6 = (code6 + "      ")[:6]

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W_TR = f"{{{W_NS}}}tr"
    W_TC = f"{{{W_NS}}}tc"
    W_T  = f"{{{W_NS}}}t"

    placeholder = "regreg/codcod"

    for table in doc.tables:
        for row_elem in table._tbl.iter(W_TR):
            # Acces direct XML — fără expansiunea merge din python-docx
            tc_elems = list(row_elem.findall(W_TC))
            if len(tc_elems) < 13:
                continue
            # Construiește textul primelor 13 celule reale
            texts = []
            for tc in tc_elems[:13]:
                t_texts = [t.text or "" for t in tc.iter(W_T)]
                texts.append("".join(t_texts))
            if "".join(texts) != placeholder:
                continue
            # Înlocuiește NUMAI celulele 7-12 (codcod → 6 cifre Zephren)
            for i in range(6):
                tc = tc_elems[7 + i]
                target = code6[i]
                t_list = list(tc.iter(W_T))
                if t_list:
                    t_list[0].text = target
                    for t in t_list[1:]:
                        t.text = ""
            return True
    return False


def set_nzeb_checkbox(doc, nzeb_ok):
    """Check (DA) or uncheck (NU) the NZEB FORMCHECKBOX in the document via XML manipulation."""
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    qW = f'{{{W}}}'

    for p in doc.element.body.iter(f'{qW}p'):
        texts = ''.join(t.text or '' for t in p.iter(f'{qW}t'))
        if 'NZEB' not in texts:
            continue
        for cb in p.iter(f'{qW}checkBox'):
            existing = cb.find(f'{qW}checked')
            if nzeb_ok:
                if existing is None:
                    etree.SubElement(cb, f'{qW}checked')
            else:
                if existing is not None:
                    cb.remove(existing)
            return


def replace_seq(doc, old_text, values):
    """Replace sequential occurrences of old_text with different values (body + tables + text boxes)."""
    idx = [0]

    def do_para(para):
        if idx[0] >= len(values):
            return
        full = para.text
        while old_text in full and idx[0] < len(values):
            replace_in_paragraph(para, old_text, values[idx[0]], count=1)
            idx[0] += 1
            full = para.text

    for para in _iter_all_paragraphs(doc, include_txbx=True):
        do_para(para)


# ═══════════════════════════════════════════════════════
# REPLACE SCALE THRESHOLDS — EP and CO2
# ═══════════════════════════════════════════════════════
def _is_generic_template(doc):
    """Detectează dacă template-ul este cel generic (3-CPE) cu placeholder-uri xxx/yyy."""
    from docx.text.paragraph import Paragraph as DocxPara
    txbx_tag = qn("w:txbxContent")
    p_tag = qn("w:p")
    t_tag = qn("w:t")
    for txbx_elem in doc.element.body.iter(txbx_tag):
        for p_elem in txbx_elem.iter(p_tag):
            texts = ''.join(t.text or '' for t in p_elem.iter(t_tag))
            if 'xxx' in texts.lower() or 'yyy' in texts.lower():
                return True
    return False


def _build_ep_intervals(thresholds):
    """Construiește lista de intervale EP din praguri: ['≤48', '48 … 68', ..., '>536']."""
    intervals = []
    intervals.append("≤" + str(thresholds[0]))
    for i in range(len(thresholds) - 1):
        intervals.append(str(thresholds[i]) + " … " + str(thresholds[i + 1]))
    intervals.append(">" + str(thresholds[-1]))
    return intervals


def _build_co2_intervals(thresholds):
    """Construiește lista de intervale CO2 din praguri: ['≤8,3', '8,3 … 11,6', ..., '> 93,1']."""
    intervals = []
    intervals.append("≤" + format_ro(thresholds[0], 1))
    for i in range(len(thresholds) - 1):
        intervals.append(format_ro(thresholds[i], 1) + " … " + format_ro(thresholds[i + 1], 1))
    intervals.append("> " + format_ro(thresholds[-1], 1))
    return intervals


def _replace_generic_placeholders(doc, new_ep_scale, new_co2_scale):
    """Înlocuiește placeholder-urile xxx/yyy din template-ul generic cu valorile reale.

    Template-ul generic (3-CPE-forma-generala-cladire.docx) conține text boxes cu:
    EP:  ≤xx, xx … xxx, xxx … xxx (×5), >xxxx
    CO2: ≤yy, yy … yyy, yyy … yyy (×5), > yyy

    Fiecare apare de 2 ori (clădire reală + clădire referință).
    Potrivirea se face pe textul COMPLET al paragrafului (nu substring).
    """
    from docx.text.paragraph import Paragraph as DocxPara
    txbx_tag = qn("w:txbxContent")
    p_tag = qn("w:p")
    t_tag = qn("w:t")

    ep_intervals = _build_ep_intervals(new_ep_scale)
    co2_intervals = _build_co2_intervals(new_co2_scale) if new_co2_scale else []

    # Contoare pentru placeholder-urile repetitive (xxx … xxx apare de 10× = 5 clase × 2 dup)
    ep_xxx_count = [0]   # câte „xxx … xxx" am înlocuit (0-9 → clasele B-F × 2)
    co2_yyy_count = [0]  # câte „yyy … yyy" am înlocuit

    def _set_para_text(para, new_text):
        """Înlocuiește tot textul din paragraf cu new_text, păstrând formatarea primului run."""
        runs = para.runs
        if not runs:
            return
        runs[0].text = new_text
        for r in runs[1:]:
            r.text = ""

    for txbx_elem in doc.element.body.iter(txbx_tag):
        for p_elem in txbx_elem.iter(p_tag):
            para = DocxPara(p_elem, None)
            full = para.text.strip()
            if not full:
                continue

            # ── EP placeholders (potrivire EXACTĂ pe text complet) ──
            if full == "\u2264xx" or full == "≤xx":
                # A+ : ≤{threshold[0]}
                _set_para_text(para, ep_intervals[0])
                continue

            if full == "xx \u2026 xxx" or full == "xx … xxx":
                # A : {threshold[0]} … {threshold[1]}
                _set_para_text(para, ep_intervals[1])
                continue

            if full == "xxx \u2026 xxx" or full == "xxx … xxx":
                # B-F : clasele 2-6, fiecare apare de 2 ori (real + ref)
                idx = ep_xxx_count[0] // 2  # 0,1→B  2,3→C  4,5→D  6,7→E  8,9→F
                interval_idx = idx + 2      # ep_intervals[2]=B, [3]=C, etc.
                if interval_idx < len(ep_intervals) - 1:  # nu depășim (ultimul e >G)
                    _set_para_text(para, ep_intervals[interval_idx])
                ep_xxx_count[0] += 1
                continue

            if full.startswith(">xxx") or full.startswith("> xxx"):
                # G : >{threshold[6]}
                _set_para_text(para, ep_intervals[7])
                continue

            # ── CO2 placeholders ──
            if not co2_intervals:
                continue

            if full == "\u2264yy" or full == "≤yy":
                _set_para_text(para, co2_intervals[0])
                continue

            if full == "yy \u2026 yyy" or full == "yy … yyy":
                _set_para_text(para, co2_intervals[1])
                continue

            if full == "yyy \u2026 yyy" or full == "yyy … yyy":
                idx = co2_yyy_count[0] // 2
                interval_idx = idx + 2
                if interval_idx < len(co2_intervals) - 1:
                    _set_para_text(para, co2_intervals[interval_idx])
                co2_yyy_count[0] += 1
                continue

            if full.startswith("> yyy") or full.startswith(">yyy"):
                _set_para_text(para, co2_intervals[7])
                continue


def replace_scales(doc, category, new_ep_scale, new_co2_scale=None):
    """Replace EP/CO2 scale values in text boxes ONLY (safe — no risk of corrupting other numbers).
    Gestionează atât template-urile specifice (cu valori numerice) cât și template-ul
    generic (3-CPE cu placeholder-uri xxx/yyy)."""

    # Detectează template generic cu placeholder-uri
    if _is_generic_template(doc):
        _replace_generic_placeholders(doc, new_ep_scale, new_co2_scale)
        return

    # Template specific — find-and-replace valori numerice existente
    old_ep = EP_TEMPLATE_SCALES.get(category, EP_TEMPLATE_SCALES["AL"])

    for i in range(7):
        if old_ep[i] != new_ep_scale[i]:
            replace_in_txbx_only(doc, str(old_ep[i]), str(new_ep_scale[i]))

    if new_co2_scale:
        old_co2 = CO2_TEMPLATE_SCALES.get(category, CO2_TEMPLATE_SCALES["AL"])
        for i in range(7):
            old_s = format_ro(old_co2[i], 1)
            new_s = format_ro(new_co2_scale[i], 1)
            if old_s != new_s:
                replace_in_txbx_only(doc, old_s, new_s)


# ═══════════════════════════════════════════════════════
# REPLACE CLASS INDICATORS — the colored arrows on scales
# Includes REPOSITIONING the shape vertically to the correct class row
# ═══════════════════════════════════════════════════════

# Vertical positions (EMU) for each class row on the scale
# Derived from template: class A posV=510540, class B posV=977265
# Spacing per class: 977265 - 510540 = 466725 EMU
_CLASS_SPACING = 466725
_CLASS_POS_V = {
    "A+": 510540 - _CLASS_SPACING,    #   43815
    "A":  510540,                      #  510540
    "B":  977265,                      #  977265
    "C":  977265 + _CLASS_SPACING,     # 1443990
    "D":  977265 + 2 * _CLASS_SPACING, # 1910715
    "E":  977265 + 3 * _CLASS_SPACING, # 2377440
    "F":  977265 + 4 * _CLASS_SPACING, # 2844165
    "G":  977265 + 5 * _CLASS_SPACING, # 3310890
}

# CO2 scale has same spacing but slightly different base
_CO2_CLASS_POS_V = {
    "A+": 963930 - _CLASS_SPACING * 2, #   30480
    "A":  963930 - _CLASS_SPACING,     #  497205
    "B":  963930,                      #  963930
    "C":  963930 + _CLASS_SPACING,     # 1430655
    "D":  963930 + 2 * _CLASS_SPACING, # 1897380
    "E":  963930 + 3 * _CLASS_SPACING, # 2364105
    "F":  963930 + 4 * _CLASS_SPACING, # 2830830
    "G":  963930 + 5 * _CLASS_SPACING, # 3297555
}


# Culorile REALE din imaginile de fundal ale template-urilor CPE MDLPA
# Extrase pixel-by-pixel din template-ul DOCX oficial (image1-16.jpeg)
_EP_CLASS_COLORS = {
    "A+": {"vml": "#009B00", "hex": "009B00"},  # verde închis
    "A":  {"vml": "#32C831", "hex": "32C831"},  # verde mediu
    "B":  {"vml": "#00FF00", "hex": "00FF00"},  # verde pur / lime
    "C":  {"vml": "#FFFF00", "hex": "FFFF00"},  # galben pur
    "D":  {"vml": "#F39C00", "hex": "F39C00"},  # chihlimbar
    "E":  {"vml": "#FF6400", "hex": "FF6400"},  # portocaliu
    "F":  {"vml": "#FE4101", "hex": "FE4101"},  # portocaliu-roșu
    "G":  {"vml": "#FE0000", "hex": "FE0000"},  # roșu pur
}

# Scala CO2 din template MDLPA — albastru (A+ → C) → gri (D → G)
_CO2_CLASS_COLORS = {
    "A+": {"vml": "#0000FE", "hex": "0000FE"},  # albastru închis
    "A":  {"vml": "#3265FF", "hex": "3265FF"},  # albastru mediu
    "B":  {"vml": "#009BFF", "hex": "009BFF"},  # albastru deschis
    "C":  {"vml": "#9CD2FF", "hex": "9CD2FF"},  # albastru foarte deschis
    "D":  {"vml": "#BEBEBE", "hex": "BEBEBE"},  # gri deschis
    "E":  {"vml": "#969696", "hex": "969696"},  # gri mediu
    "F":  {"vml": "#646464", "hex": "646464"},  # gri închis
    "G":  {"vml": "#333333", "hex": "333333"},  # gri foarte închis
}


def _text_color_for_bg(hex_color):
    """Culoarea literei din interiorul săgeții — conform specificației MDLPA:
    A+=alb, A=alb, B=negru, C=negru, D=negru, E=negru, F=negru, G=alb.
    Mapare după culoarea de fundal hex a clasei (EP și CO2)."""
    # Tabel fix hex fundal → culoare text
    _FIXED = {
        # EP scale
        "009B00": "FFFFFF",  # A+ verde închis → alb
        "32C831": "FFFFFF",  # A  verde mediu  → alb
        "00FF00": "000000",  # B  verde pur     → negru
        "FFFF00": "000000",  # C  galben        → negru
        "F39C00": "000000",  # D  portocaliu    → negru
        "FF6400": "000000",  # E  portocaliu-roș→ negru
        "FE4101": "000000",  # F  roșu-portocaliu→ negru
        "FE0000": "FFFFFF",  # G  roșu          → alb
        # CO2 scale
        "0000FE": "FFFFFF",  # A+ albastru închis → alb
        "3265FF": "FFFFFF",  # A  albastru mediu  → alb
        "009BFF": "000000",  # B  albastru deschis→ negru
        "3399CC": "000000",  # C  albastru-gri    → negru
        "808080": "000000",  # D  gri mediu       → negru
        "999999": "000000",  # E  gri deschis     → negru
        "AAAAAA": "000000",  # F  gri foarte deschis→ negru
        "333333": "FFFFFF",  # G  gri închis      → alb
    }
    key = hex_color.upper().lstrip("#")
    if key in _FIXED:
        return _FIXED[key]
    # Fallback WCAG pentru culori neprevăzute
    try:
        r = int(key[0:2], 16) / 255.0
        g = int(key[2:4], 16) / 255.0
        b = int(key[4:6], 16) / 255.0
        def lin(c):
            return c / 12.92 if c <= 0.04045 else ((c + 0.055) / 1.055) ** 2.4
        lum = 0.2126 * lin(r) + 0.7152 * lin(g) + 0.0722 * lin(b)
        return "FFFFFF" if (lum + 0.05) < 0.5 else "000000"
    except Exception:
        return "000000"


def _get_shape_fill_hex(shape_xml):
    """Extrage culoarea de umplere originală din shape (srgbClr sau VML fillcolor)."""
    m = re.search(r'<a:solidFill>\s*<a:srgbClr val="([A-Fa-f0-9]{6})"', shape_xml)
    if m:
        return m.group(1).upper()
    m = re.search(r'fillcolor="[#]?([A-Fa-f0-9]{6})"', shape_xml)
    if m:
        return m.group(1).upper()
    return None


def _is_co2_shape(shape_xml):
    """Returnează True dacă shape-ul este pe scala CO2 (culoare albastră dominantă)."""
    color = _get_shape_fill_hex(shape_xml)
    if not color:
        return False
    try:
        r = int(color[0:2], 16)
        g = int(color[2:4], 16)
        b = int(color[4:6], 16)
        # Scala CO2 din template MDLPA = albastru dominant (B > 100 și B > R * 1.5)
        return b > 100 and b > r * 1.5 and b >= g
    except Exception:
        return False


def _update_shape_color(shape_xml_str, color_info):
    """Update fill color in DrawingML (srgbClr SAU schemeClr SAU noFill) și VML (fillcolor)."""
    result = shape_xml_str
    hex_color = color_info["hex"]
    vml_color = color_info["vml"]
    explicit_fill = '<a:solidFill><a:srgbClr val="' + hex_color + '"/></a:solidFill>'

    # FIX 1: înlocuiește srgbClr explicit (culoare fixă)
    replaced = re.subn(
        r'(<a:solidFill>\s*<a:srgbClr val=")[A-Fa-f0-9]{6}(")',
        r'\g<1>' + hex_color + r'\g<2>',
        result,
        count=1
    )
    result = replaced[0]
    srgb_matched = replaced[1] > 0

    # FIX 2: dacă nu s-a găsit srgbClr, înlocuiește schemeClr (culoare din temă Word)
    # Acoperă <a:schemeClr val="..."/> (self-closing) și <a:schemeClr ...>...</a:schemeClr>
    if not srgb_matched:
        result, n1 = re.subn(
            r'<a:solidFill>\s*<a:schemeClr\b[^/]*/>\s*</a:solidFill>',
            explicit_fill,
            result,
            count=1
        )
        if n1 == 0:
            result = re.sub(
                r'<a:solidFill>\s*<a:schemeClr\b[^>]*>[\s\S]*?</a:schemeClr>\s*</a:solidFill>',
                explicit_fill,
                result,
                count=1
            )

    # FIX 3 (eliminat 19 apr 2026): textbox-urile cu litera au <a:noFill/> intenționat
    # în template-ul oficial MDLPA (transparent) — doar pentagoanele primesc culoarea.
    # Dacă înlocuim noFill cu solidFill, rezultă DUBLĂ SUPRAPUNERE vizibilă (rectangle
    # textbox peste pentagon), cu contururi imperfecte. Lăsăm noFill neatins.

    # VML fillcolor (folosit de Office pentru compatibilitate)
    # Înlocuim fillcolor existent, dar NU adăugăm fillcolor pe shape-uri care au
    # <a:noFill/> în DrawingML — acelea sunt transparente intenționat (textbox cu
    # literă peste pentagon). Suprapunerea colorabilă ar crea rectangle vizibil.
    has_no_fill_drawing = '<a:noFill' in result
    if re.search(r'fillcolor="[^"]*"', result):
        result = re.sub(
            r'fillcolor="[^"]*"',
            'fillcolor="' + vml_color + '"',
            result,
            count=1
        )
    elif not has_no_fill_drawing:
        result = re.sub(
            r'(<v:shape\b)',
            r'\1 fillcolor="' + vml_color + '"',
            result,
            count=1
        )

    return result


def _update_shape_pos_v(shape_xml_str, new_pos_v):
    """Update BOTH the modern (wp:anchor posOffset) and VML fallback (style top:) positions."""
    # 1. Update wp:anchor posOffset
    result = re.sub(
        r'(<wp:positionV[^>]*>[\s\S]*?<wp:posOffset>)-?\d+(</wp:posOffset>)',
        lambda m: m.group(1) + str(new_pos_v) + m.group(2),
        shape_xml_str
    )
    # 2. Update VML fallback style "top:XXpt" — convert EMU to pt (1pt = 12700 EMU)
    new_top_pt = round(new_pos_v / 12700, 1)
    result = re.sub(
        r'(style="[^"]*?)top:\s*-?[\d.]+pt',
        lambda m: m.group(1) + "top:" + str(new_top_pt) + "pt",
        result
    )
    return result


def replace_class_indicators(doc, ep_class_real, ep_class_ref, co2_class_real):
    """Replace class letter text AND reposition the arrow shapes on the scales."""
    import zipfile as _zf

    # Work directly on the XML string for precise shape manipulation
    xml_file = doc.part._element.getparent()
    # Access the raw XML via the document element
    body = doc.element.body
    # Serialize body to string for regex-based shape manipulation
    body_xml = etree.tostring(body, encoding="unicode")

    # Find mc:AlternateContent blocks with class indicators (H > 150000)
    alt_pattern = re.compile(r'(<mc:AlternateContent>[\s\S]*?</mc:AlternateContent>)')
    blocks = list(alt_pattern.finditer(body_xml))

    # Colectăm TOATE shape-urile indicator — atât textbox (cu litera) cât și pentagon (path)
    # Fiecare indicator are 2 shape-uri consecutive: textbox + pentagon (sau invers)
    text_indicators = []   # (match, letter, posH, posV) — textbox cu litera
    path_indicators = []   # (match, posH, posV) — pentagon/arrow fără text

    for m in blocks:
        content = m.group(1)
        # Caută litera clasei în <w:t> (Word/VML) ȘI în <a:t> (DrawingML)
        letters_wt = re.findall(r'<w:t[^>]*>([A-G]\+?)</w:t>', content)
        letters_at = re.findall(r'<a:t[^>]*>([A-G]\+?)</a:t>', content)
        letters = letters_wt if letters_wt else letters_at
        pos_h = re.search(r'positionH[^>]*>[\s\S]*?posOffset>(-?\d+)<', content)
        pos_v = re.search(r'positionV[^>]*>[\s\S]*?posOffset>(-?\d+)<', content)
        if not pos_h or not pos_v:
            continue
        h = int(pos_h.group(1))
        v = int(pos_v.group(1))
        is_path = 'coordsize' in content or '<v:path' in content

        if len(letters) > 0 and len(letters[0]) <= 2 and not is_path and v > 0:
            # Inclus indiferent de h — EP shapes au h mic (stânga), CO₂ shapes au h mare (dreapta)
            # Filtrul anterior h > 150000 excludea greșit shape-urile EP cu h mic
            text_indicators.append((m, letters[0], h, v))
        elif is_path and -100000 <= h < 700000:
            # Pentagon/arrow shapes — range extins pentru a include EP pentagon (h≈380365)
            path_indicators.append((m, h, v))

    # Clasificare prin POZIȚIE ORIZONTALĂ (h):
    # Scala EP (stânga) → h mic; Scala CO2 (dreapta) → h mare
    #
    # ATENȚIE: positionH poate fi relativ la "column", "margin" sau "page".
    # Shape-urile ancorate în coloane diferite ale tabelului (EP=stânga, CO₂=dreapta)
    # pot avea posOffset similare deoarece sunt relative la coloana lor, NU la pagină.
    # De aceea, folosim ORDINEA DIN DOCUMENT (XML order) ca fallback — în Word XML,
    # conținutul coloanei stângi apare ÎNAINTEA coloanei drepte.

    # CLASIFICARE ROBUSTĂ: detectăm CO2 indicators după POZIȚIA VERTICALĂ (V)
    # Scala EP și scala CO2 au același spacing (466725 EMU) dar offset diferit de 13335 EMU:
    #   EP:  A=510540, B=977265  |  CO2: A=497205, B=963930
    # Calculăm distanța la fiecare scală și alegem cea mai apropiată.
    # Metoda e fiabilă chiar și când shape-urile nu au fill explicit (noFill).
    def _closest_scale_co2(v):
        """Return True dacă V se potrivește mai bine scalei CO2 decât scalei EP."""
        ep_min  = min(abs(v - vv) for vv in _CLASS_POS_V.values())
        co2_min = min(abs(v - vv) for vv in _CO2_CLASS_POS_V.values())
        return co2_min < ep_min

    ep_indicators_raw = []
    co2_indicators = []
    for ind in text_indicators:
        im, letter, ih, iv = ind
        if _closest_scale_co2(iv):
            co2_indicators.append(ind)
        else:
            ep_indicators_raw.append(ind)

    # Fallback (dacă detecția V nu separă corect): ultimul indicator (cel mai jos) = CO2
    if not co2_indicators and len(ep_indicators_raw) >= 2 and co2_class_real:
        co2_indicators = [ep_indicators_raw[-1]]
        ep_indicators_raw = ep_indicators_raw[:-1]

    # Sortează indicatoarele EP după V: V mic = sus pe scală = EP_ref (clădire referință mai eficientă)
    #                                    V mare = jos pe scală = EP_real (clădire reală)
    # Această sortare e sigură pt. EP (ambele ancorate în aceeași coloană stângă)
    ep_indicators_raw.sort(key=lambda x: x[3])  # sort by v
    if len(ep_indicators_raw) >= 2:
        ep_ref_ind = ep_indicators_raw[0]   # V mic → EP_ref (mai sus = mai eficient)
        ep_real_ind = ep_indicators_raw[1]  # V mare → EP_real
        ep_indicators = [ep_real_ind, ep_ref_ind]  # [0]=real, [1]=ref
    else:
        ep_indicators = ep_indicators_raw

    # Perechi text-pentagon: pentagonul urmează imediat după textbox
    # Construim o mapare: pentru fiecare text indicator, găsim pentagonul cel mai apropiat (următor)
    def find_next_path(text_match, all_paths):
        """Find the path shape immediately AFTER a text shape in the XML."""
        text_end = text_match.end()
        best = None
        best_dist = float('inf')
        for pm, ph, pv in all_paths:
            dist = pm.start() - text_end
            if 0 < dist < best_dist:
                best = (pm, ph, pv)
                best_dist = dist
        return best

    # Replace EP real indicators (first group — higher posV = further down = typically "B")
    # Template has: ref (A, lower V) then real (B, higher V)
    # So sorted: [0]=ref(A, V=510540), [1]=real(B, V=977265) for single indicators
    # For 4 indicators: [0,1]=ref pair, [2,3]=real pair

    new_xml = body_xml

    def move_indicator(text_match, new_class, class_pos_map, old_v, is_co2=False):
        """Move both text indicator AND its paired pentagon shape, and update colors."""
        nonlocal new_xml
        new_pos = class_pos_map.get(new_class, old_v)
        delta_v = new_pos - old_v  # displacement in EMU
        color_map = _CO2_CLASS_COLORS if is_co2 else _EP_CLASS_COLORS
        color = color_map.get(new_class, color_map.get("B"))

        # 1. Update textbox (letter + position + fill color + text color)
        new_content = text_match.group(1)
        # Actualizează litera în elementele Word <w:t> (VML fallback)
        new_content = re.sub(r'(<w:t[^>]*>)[A-G]\+?(</w:t>)', r'\g<1>' + new_class + r'\g<2>', new_content)
        # Actualizează litera și în elementele DrawingML <a:t> (Choice primary)
        new_content = re.sub(r'(<a:t[^>]*>)[A-G]\+?(</a:t>)', r'\g<1>' + new_class + r'\g<2>', new_content)
        # Uniformizare font cu scala statică (Arial Bold) — template mobile folosește
        # AvantGarde Bk BT (font proprietar, substituit cu fallback neuniform de LibreOffice).
        # Păstrez size-ul original 32 (16pt) pentru vizibilitate — scala statică e 9pt dar
        # apare vizual mai mare din context (celule tabel cu auto-scale). La 16pt Arial Bold
        # mobile arată comparabil cu scala statică.
        new_content = re.sub(r'w:ascii="AvantGarde Bk BT"', 'w:ascii="Arial"', new_content)
        new_content = re.sub(r'w:hAnsi="AvantGarde Bk BT"', 'w:hAnsi="Arial"', new_content)
        new_content = re.sub(r'w:cs="AvantGarde Bk BT"', 'w:cs="Arial"', new_content)
        new_content = re.sub(r'typeface="AvantGarde Bk BT"', 'typeface="Arial"', new_content)
        # Centrare verticală în textbox (anchor="t" → "ctr")
        new_content = re.sub(r'anchor="t"', 'anchor="ctr"', new_content)
        # Centrare orizontală — inject <w:jc w:val="center"/> în <w:pPr> dacă nu există
        if '<w:jc' not in new_content:
            new_content = re.sub(r'(<w:pPr>)', r'\g<1><w:jc w:val="center"/>', new_content, count=1)
        # DrawingML paragraph alignment (a:pPr algn="ctr")
        new_content = re.sub(r'<a:pPr\b(?![^>]*algn=)', '<a:pPr algn="ctr"', new_content, count=1)
        # Reduce padding intern pentru mai mult spațiu util (16pt litera în 23pt cx)
        new_content = re.sub(r'lIns="\d+"', 'lIns="0"', new_content)
        new_content = re.sub(r'rIns="\d+"', 'rIns="0"', new_content)
        # Pentru clase 2-char ("A+"), mărește cx cu 35% ca să încapă litera completă
        # (Arial Bold 16pt "A+" ~21pt vs cx=23pt → clipping la margin).
        if len(new_class) > 1:
            new_content = re.sub(
                r'(<wp:extent cx=")(\d+)("\s+cy=")',
                lambda m: m.group(1) + str(int(int(m.group(2)) * 1.35)) + m.group(3),
                new_content
            )
            new_content = re.sub(
                r'(<a:ext cx=")(\d+)("\s+cy=")',
                lambda m: m.group(1) + str(int(int(m.group(2)) * 1.35)) + m.group(3),
                new_content
            )
            new_content = re.sub(
                r'(width:\s*)([\d.]+)(pt)',
                lambda m: m.group(1) + str(round(float(m.group(2)) * 1.35, 1)) + m.group(3),
                new_content
            )
        new_content = _update_shape_pos_v(new_content, new_pos)
        # Actualizează și culoarea de fundal a textbox-ului (solidFill) — altfel litera
        # rămâne colorată cu culoarea implicită din template (nu cu culoarea clasei reale)
        new_content = _update_shape_color(new_content, color)
        # Culoarea textului calculată din luminanța WCAG a fundalului (valabil EP și CO2)
        text_color = _text_color_for_bg(color["hex"])
        if '<w:color' not in new_content:
            # Injectează <w:color> DOAR în rPr-ul run-ului (cel imediat înainte de <w:t>)
            # Regex: </w:rPr>\s*<w:t  → rPr-ul de run, nu cel din <w:pPr> (paragraf default)
            new_content = re.sub(
                r'(<w:rPr>)([\s\S]*?)(</w:rPr>\s*<w:t)',
                r'\g<1>\g<2><w:color w:val="' + text_color + r'"/>\g<3>',
                new_content
            )
        else:
            # Înlocuiește TOATE variantele <w:color .../> (inclusiv cu w:themeColor, w:themeShade etc.)
            new_content = re.sub(r'<w:color\b[^>]*/>', '<w:color w:val="' + text_color + '"/>', new_content)

        # DrawingML text color (a:rPr solidFill) — pentru shape-uri cu text în a:t (nu w:t)
        if '<a:t>' in new_content or '<a:t ' in new_content:
            a_fill = '<a:solidFill><a:srgbClr val="' + text_color + '"/></a:solidFill>'
            # Înlocuiește solidFill din a:rPr (nu din shape fill — cel din a:spPr)
            new_content = re.sub(
                r'(<a:rPr\b[^>]*>)([\s\S]*?)(<a:solidFill>[\s\S]*?</a:solidFill>)([\s\S]*?</a:rPr>)',
                r'\g<1>\g<2>' + a_fill + r'\g<4>',
                new_content
            )
        new_xml = new_xml.replace(text_match.group(1), new_content, 1)

        # 2. Find and update paired pentagon (path shape right after textbox)
        path = find_next_path(text_match, path_indicators)
        if path:
            pm, ph, pv = path
            new_path_pos = pv + delta_v  # same displacement
            new_path_content = pm.group(1)
            new_path_content = _update_shape_pos_v(new_path_content, new_path_pos)
            new_path_content = _update_shape_color(new_path_content, color)
            new_xml = new_xml.replace(pm.group(1), new_path_content, 1)

    # EP indicators: [0]=real, [1]=ref (ordine din document)
    if len(ep_indicators) >= 1:
        m, letter, h, v = ep_indicators[0]
        move_indicator(m, ep_class_real, _CLASS_POS_V, v)

    if len(ep_indicators) >= 2:
        m, letter, h, v = ep_indicators[1]
        move_indicator(m, ep_class_ref, _CLASS_POS_V, v)

    # CO2 indicators
    for m, letter, h, v in co2_indicators:
        move_indicator(m, co2_class_real, _CO2_CLASS_POS_V, v, is_co2=True)

    # Aplică modificările IN-PLACE (nu înlocuim tot body-ul) — evită round-trip-ul XML
    # care poate schimba subtil layout-ul documentului (overflow pagina 2).
    alt_blocks_orig = [m.group(1) for m in alt_pattern.finditer(body_xml)]
    alt_blocks_new  = [m.group(1) for m in alt_pattern.finditer(new_xml)]

    MC_NS  = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    mc_tag = f"{{{MC_NS}}}AlternateContent"
    body_mc_elems = list(body.iter(mc_tag))

    # Construiește string-ul de declarații xmlns din nsmap-ul body-ului
    ns_parts = []
    for prefix, uri in body.nsmap.items():
        if prefix is None:
            ns_parts.append(f'xmlns="{uri}"')
        else:
            ns_parts.append(f'xmlns:{prefix}="{uri}"')
    ns_decls = " ".join(ns_parts)

    for i, (orig_str, new_str) in enumerate(zip(alt_blocks_orig, alt_blocks_new)):
        if orig_str != new_str and i < len(body_mc_elems):
            old_elem = body_mc_elems[i]
            parent_el = old_elem.getparent()
            idx = list(parent_el).index(old_elem)
            try:
                wrapper = etree.fromstring(
                    f'<_wr {ns_decls}>{new_str}</_wr>'.encode("utf-8")
                )
                new_mc_elem = wrapper[0]
                parent_el.remove(old_elem)
                parent_el.insert(idx, new_mc_elem)
            except Exception:
                pass  # Dacă parsing eșuează, păstrăm elementul original
    # Nu e nevoie de doc._Document__body = None — body-ul nu a fost înlocuit


# ═══════════════════════════════════════════════════════
# TOGGLE CHECKBOXES — Anexa form fields
# ═══════════════════════════════════════════════════════
def toggle_checkboxes(doc, indices):
    """Check checkboxes at given indices (0-based) in the document."""
    checkboxes = doc.element.findall(".//w:checkBox", NSMAP)
    indices_set = set(indices)
    for i, cb in enumerate(checkboxes):
        if i in indices_set:
            default_elem = cb.find("w:default", NSMAP)
            if default_elem is not None:
                default_elem.set(qn("w:val"), "1")


# ═══════════════════════════════════════════════════════
# COMPUTE CHECKBOXES — mapare date clădire → indici checkbox Anexa
# ═══════════════════════════════════════════════════════

# U referință Mc 001-2022 (nZEB)
_U_REF_RES = {"PE":0.25,"PR":0.67,"PS":0.29,"PT":0.15,"PP":0.15,"PB":0.29,"PL":0.20,"SE":0.20}
_U_REF_NRES = {"PE":0.33,"PR":0.80,"PS":0.35,"PT":0.17,"PP":0.17,"PB":0.35,"PL":0.22,"SE":0.22}


def _norm_heating_source(v):
    """Normalizează heating_source: ID-uri noi constants.js (uppercase) → legacy semantic."""
    if not v:
        return ""
    u = v.upper()
    # Termoficare
    if u.startswith("TERMO"):
        return "termoficare"
    # Pompe de căldură
    if u.startswith("PC_") or u in ("PC_AERAER",):
        return "pc_aer_apa"
    # Electric direct
    if u in ("ELEC", "ELEC_ACC", "ELEC_PANEL", "INFRARED_E"):
        return "electric_direct"
    # Sobe (foc lemn, teracotă, metalice) — tip emisie, nu cazan
    if u.startswith("SOBA_"):
        return "soba_teracota"
    # Cazane biomasă / lemn / cărbune — tip cazan cu combustibil solid
    if u.startswith("BIO_") or u in ("COCS_K", "CARBUNE", "CARBUNE_K", "LEMN"):
        return "cazan_lemn"
    # Gaz natural / GPL / motorină → CT proprie în clădire
    if any(u.startswith(p) for p in ("GAZ_", "GPL_", "MOT_", "CONV_GAZ", "INFRARED_G")):
        return "gaz_conv"
    # Cazane în cascadă cu gaz (CASC) — tratăm ca gaz_conv
    if "CASC" in u:
        return "gaz_conv"
    # Backward compat cu vechile valori lowercase
    return v  # valori vechi: "gaz_conv", "termoficare", "electric_direct", etc.


def _norm_acm_source(v):
    """Normalizează acm_source: ID-uri noi constants.js → legacy semantic."""
    if not v:
        return ""
    u = v.upper()
    if u == "TERMO_ACM":
        return "termoficare"
    if u.startswith("SOLAR"):
        return "solar_termic"
    if u.startswith("PC_ACM"):
        return "pc"
    if u in ("BOILER_E", "BOILER_E_NOAPTE", "INSTANT_E", "DESUPERHEATER"):
        return "boiler_electric"
    if u in ("BOILER_G", "BOILER_G_COND", "BOILER_GPL", "INSTANT_G",
             "CAZAN_H", "COGEN_ACM", "CENTRALIZAT_BLOC", "BOILER_BIOMASA"):
        return "ct_prop"
    return v  # backward compat


def _norm_vent_type(v):
    """Normalizează ventilation_type: ID-uri noi constants.js → legacy semantic."""
    if not v:
        return ""
    u = v.upper()
    if u == "NAT":
        return "natural_neorg"
    if u == "NAT_HIBRIDA":
        return "natural_org"
    # Mecanică cu recuperare de căldură (hasHR=true în constants.js)
    if any(u.startswith(p) for p in ("MEC_HR", "MEC_ERV")) or u in (
        "UTA", "UTA_ERV", "VRF_VENT", "VAV", "DOAS", "GEO_AER"
    ):
        return "mec_hr"
    # Mecanică simplă (hasHR=false)
    if any(u.startswith(p) for p in ("MEC_",)) or u in ("FCU", "ADIAB_VENT"):
        return "mec"
    return v  # backward compat: "natural_neorg", "natural_org", etc.


def _norm_lighting_type(v):
    """Normalizează lighting_type: ID-uri noi constants.js → legacy semantic."""
    if not v:
        return ""
    u = v.upper()
    if u.startswith("LED"):
        return "led"
    if u in ("CFL", "TUB_T8", "TUB_T8_HF", "TUB_T5", "TUB_T5_HO", "INDUCTIE"):
        return "fluorescent"
    if u in ("INCAND", "HALOGEN", "HAL_REFL", "HAL_MR16"):
        return "incandescent"
    # Metal halide, sodiu → tratăm ca "fluorescent" (eficiență medie, nu LED)
    if u in ("METAL_HAL", "METAL_HAL_HB", "SODIU_IP", "SODIU_JP"):
        return "fluorescent"
    return v  # backward compat


def _norm_lighting_control(v):
    """Normalizează lighting_control: ID-uri noi constants.js → legacy semantic."""
    if not v:
        return ""
    u = v.upper()
    if u == "MAN":
        return "manual"
    if u in ("PREZ", "PREZ_MIC"):
        return "sensor_presence"
    if u in ("PREZ_DIM", "DAYLIGHT", "PREZ_DAY", "BMS", "DALI", "AUTO_INT", "TIMER"):
        return "daylight_dimming"
    return v  # backward compat


def compute_checkboxes(data, category):
    """Compute checkbox indices to toggle based on building data.
    Returns list of 0-based indices for the 308-checkbox Anexa clădire template."""
    cbs = []
    is_res = category in ("RI", "RC", "RA")
    u_ref = _U_REF_RES if is_res else _U_REF_NRES

    # Normalizare ID-uri constants.js (uppercase) → legacy semantic
    _data = dict(data)
    _data["heating_source"]     = _norm_heating_source(data.get("heating_source", ""))
    _data["acm_source"]         = _norm_acm_source(data.get("acm_source", ""))
    _data["ventilation_type"]   = _norm_vent_type(data.get("ventilation_type", ""))
    _data["lighting_type"]      = _norm_lighting_type(data.get("lighting_type", ""))
    _data["lighting_control"]   = _norm_lighting_control(data.get("lighting_control", ""))
    data = _data

    # Parse opaque U-values
    try:
        opaque_u = json.loads(data.get("opaque_u_values", "[]"))
    except:
        opaque_u = []
    try:
        glaz_max_u = float(data.get("glazing_max_u", "0") or "0")
        if glaz_max_u != glaz_max_u or glaz_max_u < 0:  # NaN sau negativ
            glaz_max_u = 0.0
    except (ValueError, TypeError):
        glaz_max_u = 0.0

    # ══════════════════════════════
    # ANEXA 1 — RECOMANDĂRI (CB 0-47)
    # ══════════════════════════════

    # Anvelopă — bifăm dacă U calculat > U referință
    if any(e.get("type") == "PE" and float(e.get("u", 0)) > u_ref.get("PE", 0.25) for e in opaque_u):
        cbs.append(0)  # Pereți exteriori
    if any(e.get("type") == "PB" and float(e.get("u", 0)) > u_ref.get("PB", 0.29) for e in opaque_u):
        cbs.append(1)  # Planșeu subsol
    if any(e.get("type") in ("PT", "PP") and float(e.get("u", 0)) > u_ref.get(e["type"], 0.15) for e in opaque_u):
        cbs.append(2)  # Terasă/pod
    if any(e.get("type") in ("PL", "SE") and float(e.get("u", 0)) > u_ref.get(e["type"], 0.20) for e in opaque_u):
        cbs.append(3)  # Planșee contact exterior
    # CB5: tâmplărie
    u_glaz_ref = 1.30 if is_res else 1.80
    if glaz_max_u > u_glaz_ref:
        cbs.append(5)
    # CB6: grile ventilare
    vent_type = data.get("ventilation_type", "")
    if not vent_type or vent_type == "natural_neorg":
        cbs.append(6)
    # CB13: robinete termostat
    h_src = data.get("heating_source", "")
    if h_src and h_src not in ("electric_direct", "pc_aer_aer"):
        cbs.append(13)
    # CB21: automatizare
    h_ctrl = data.get("heating_control", "")
    if not h_ctrl or h_ctrl == "manual":
        cbs.append(21)
    # CB25: iluminat LED
    l_type = data.get("lighting_type", "")
    if l_type and l_type != "led":
        cbs.append(25)
    # CB26: senzori prezență
    l_ctrl = data.get("lighting_control", "")
    if not l_ctrl or l_ctrl not in ("sensor_presence", "daylight_dimming"):
        cbs.append(26)
    # CB27: regenerabile
    st_en = data.get("solar_thermal_enabled", "") == "true"
    pv_en = data.get("pv_enabled", "") == "true"
    if not st_en and not pv_en:
        cbs.append(27)
    # CB28: recuperare căldură
    if not vent_type or "hr" not in vent_type:
        cbs.append(28)

    # Cost/Savings/Payback (CB 48-64) — #7 (audit Pas 6+7 V7, 7 mai 2026)
    # Folosim valori REALE din financialAnalysis (Pas 7) primite din JS, fallback
    # la defaults dacă lipsesc. Bază: L.372/2005 R2 + Reg. UE 244/2012 (cost-optim).
    #
    # Cost total renovare (CB 48-53): <1k / 1k-10k / 10k-25k / 25k-50k / 50k-100k / >100k EUR
    try:
        total_cost_eur = float((data.get("financial_total_cost_eur", "0") or "0").replace(",", "."))
    except (ValueError, TypeError):
        total_cost_eur = 0
    if total_cost_eur > 0:
        if total_cost_eur < 1000:       cbs.append(48)
        elif total_cost_eur < 10000:    cbs.append(49)
        elif total_cost_eur < 25000:    cbs.append(50)
        elif total_cost_eur < 50000:    cbs.append(51)
        elif total_cost_eur < 100000:   cbs.append(52)
        else:                           cbs.append(53)
    else:
        cbs.append(50)  # default 10k-25k EUR

    # Savings (CB 54-59): <10% / 10-20% / 20-30% / 30-40% / 40-50% / >60%
    # M-7 din audit: template MDLPA Ord. 16/2023 NU include 50-60% (gap tipografic
    # original); pentru savings 50-59% bifăm CB59 (>60% — cea mai apropiată).
    # #12b (audit Pas 6+7 V7, 7 mai 2026) — flag pentru notă footer auditor
    # când savings cad în intervalul 50-59% (zona cu gap în template).
    try:
        savings_pct = float((data.get("financial_savings_pct", "0") or "0").replace(",", "."))
    except (ValueError, TypeError):
        savings_pct = 0
    needs_50_60_note = False
    if savings_pct > 0:
        if savings_pct < 10:        cbs.append(54)
        elif savings_pct < 20:      cbs.append(55)
        elif savings_pct < 30:      cbs.append(56)
        elif savings_pct < 40:      cbs.append(57)
        elif savings_pct < 50:      cbs.append(58)
        else:
            cbs.append(59)  # 50-60% și >60% (gap M-7)
            if 50 <= savings_pct < 60:
                needs_50_60_note = True
    else:
        cbs.append(56)  # default 20-30%
    # Salvăm flag-ul în data pentru a fi citit de funcția de injectare notă footer
    data["_needs_50_60_note"] = "1" if needs_50_60_note else ""
    data["_savings_pct_actual"] = str(savings_pct) if savings_pct > 0 else ""

    # Payback (CB 60-64): <1 / 1-3 / 3-7 / 7-10 / >10 ani
    try:
        payback_y = float((data.get("financial_payback_years", "0") or "0").replace(",", "."))
    except (ValueError, TypeError):
        payback_y = 0
    if payback_y > 0:
        if payback_y < 1:           cbs.append(60)
        elif payback_y < 3:         cbs.append(61)
        elif payback_y < 7:         cbs.append(62)
        elif payback_y < 10:        cbs.append(63)
        else:                       cbs.append(64)
    else:
        cbs.append(62)  # default 3-7 ani

    # ══════════════════════════════
    # ANEXA 2 — DATE CLĂDIRE (CB 65+)
    # ══════════════════════════════

    # Tip clădire (CB 65=existentă, 66=nouă, 67=existentă nefinalizată)
    year_b = int(data.get("year_built", "2000") or "2000")
    import datetime
    if year_b >= datetime.datetime.now().year - 1:
        cbs.append(66)
    else:
        cbs.append(65)

    # Categoria clădirii (CB 68-111)
    # Audit 2 mai 2026 — Sprint 6: BC (bloc colectiv mixt) adăugat
    cat_map = {
        "RI": [68, 69], "RC": [68, 71], "RA": [68, 71],
        "BC": [68, 71],  # Sprint 6 — fix lipsă mapping (variantă RC)
        "BI": [79, 80], "ED": [74, 76], "SA": [86, 87],
        "HC": [94, 95], "CO": [103, 104], "SP": [99, 100],
        "AL": [108, 111],
    }
    for cb in cat_map.get(category, cat_map["AL"]):
        cbs.append(cb)

    # Zone climatice (CB 112-116 = zone I-V)
    try:
        zone_num = int(data.get("climate_zone_num", "3") or "3")
    except (ValueError, TypeError):
        zone_num = 3
    zone_num = max(1, min(5, zone_num))  # clamp 1-5
    cbs.append(111 + zone_num)

    # Zone eoliene (CB 117-120)
    wind_zone = 1 if zone_num <= 2 else (2 if zone_num <= 4 else 3)
    cbs.append(116 + wind_zone)

    # Structura constructivă (CB 127-134) — matching parțial (building.structure e string lung)
    struct_map = [
        ("Zidărie portantă", 127), ("Cadre beton armat", 129),
        ("Panouri prefabricate mari", 133), ("Structură metalică", 132),
        ("Structură lemn", 131), ("Mixtă", 134),
    ]
    struct_text = data.get("structure", "")
    struct_cb = None
    for key, cb in struct_map:
        if key in struct_text:
            struct_cb = cb
            break
    if struct_cb:
        cbs.append(struct_cb)

    # ══════════════════════════════
    # ANEXA 2 — INSTALAȚII
    # ══════════════════════════════

    # ÎNCĂLZIRE (CB 135+)
    if h_src:
        cbs.append(135)  # Da, funcțională
    else:
        cbs.append(137)  # Nu

    if h_src:
        heat_src_map = {
            "gaz_conv": 144, "gaz_cond": 144, "termoficare": 146,
            "electric_direct": 139, "pc_aer_apa": 149, "pc_aer_aer": 139,
            "pc_sol_apa": 149, "pc_apa_apa": 149, "centrala_gpl": 144,
            "cazan_lemn": 138, "cazan_peleti": 138, "soba_teracota": 138,
            "pompa_caldura": 149,
        }
        h_cb = heat_src_map.get(h_src)
        if h_cb:
            cbs.append(h_cb)

    # Tip sistem încălzire
    if h_src == "soba_teracota":
        cbs.append(150)
    elif h_src == "electric_direct":
        cbs.append(154)
    elif h_src:
        cbs.append(151)  # corpuri statice

    # Distribuție
    cbs.append(160)  # inferioară (default)

    # ACM (CB 176+)
    acm_src = data.get("acm_source", "")
    if acm_src:
        cbs.append(176)  # Da
    else:
        cbs.append(178)  # Nu

    if acm_src:
        acm_map = {
            "ct_prop": 181, "boiler_electric": 180, "termoficare": 183,
            "solar_termic": 179, "pc": 186,
        }
        a_cb = acm_map.get(acm_src)
        if a_cb:
            cbs.append(a_cb)

    if acm_src == "boiler_electric":
        cbs.append(187)
    elif acm_src == "ct_prop":
        cbs.append(188)

    cbs.append(195)  # Recirculare nu există (default)

    # RĂCIRE (CB 202+)
    has_cool = data.get("cooling_has", "") == "true"
    if has_cool:
        cbs.append(202)
    else:
        cbs.append(204)

    if has_cool:
        cool_src = data.get("cooling_source", "")
        cool_map = {
            "split": 214, "chiller_aer": 205, "chiller_apa": 206,
            "pc_aer_apa": 207, "pc_apa_apa": 208, "pc_aer_aer": 209,
            "monobloc": 213,
        }
        c_cb = cool_map.get(cool_src)
        if c_cb:
            cbs.append(c_cb)
        cbs.append(229)  # Complet climatizat
        cbs.append(232)  # Fără controlul umidității

    # VENTILARE (CB 256+)
    has_vent = vent_type and vent_type != "natural_neorg"
    if has_vent:
        cbs.append(256)
    else:
        cbs.append(258)

    if vent_type == "natural_neorg":
        cbs.append(259)
    elif vent_type == "natural_org":
        cbs.append(260)
    elif has_vent:
        cbs.append(261)

    if vent_type and "hr" in vent_type:
        cbs.append(270)
    else:
        cbs.append(271)

    # ILUMINAT (CB 272+)
    if l_type:
        cbs.append(272)
    else:
        cbs.append(274)

    if l_ctrl == "manual":
        cbs.append(276)
    elif l_ctrl == "daylight_dimming":
        cbs.extend([277, 278])
    elif l_ctrl == "sensor_presence":
        cbs.extend([277, 279])
    else:
        cbs.append(275)

    if l_type == "fluorescent":
        cbs.append(281)
    elif l_type == "incandescent":
        cbs.append(282)
    elif l_type == "led":
        cbs.append(283)
    else:
        cbs.append(284)

    cbs.append(285)  # Stare bună (default)

    # REGENERABILE (CB 288+)
    if st_en:
        cbs.append(288)
    else:
        cbs.append(289)
    if pv_en:
        cbs.append(290)
    else:
        cbs.append(291)

    hp_en = data.get("heat_pump_enabled", "") == "true"
    if hp_en:
        cbs.append(292)
    else:
        cbs.append(293)

    if hp_en:
        hp_type_map = {
            "sol_apa_deschisa": 294, "sol_apa_inchisa": 295,
            "aer_apa": 296, "aer_aer": 297, "apa_aer": 298, "sol_aer": 299,
        }
        hp_cb = hp_type_map.get(data.get("heat_pump_type", ""))
        if hp_cb:
            cbs.append(hp_cb)

    bio_en = data.get("biomass_enabled", "") == "true"
    if bio_en:
        cbs.append(301)
    else:
        cbs.append(302)

    if bio_en:
        bio_type = data.get("biomass_type", "")
        if bio_type == "peleti":
            cbs.append(303)
        elif bio_type == "brichete":
            cbs.append(304)
        else:
            cbs.append(305)

    wind_en = data.get("wind_enabled", "") == "true"
    if wind_en:
        cbs.append(306)
    else:
        cbs.append(307)

    # ══════════════════════════════
    # Sprint monolith (20 apr 2026) — Extinderi CB suplimentare din date existente
    #
    # IMPORTANT: CB-urile de mai jos se activează DOAR când user-ul furnizează
    # date explicite. Evitări defaults pe indici incerți (cauzau bifare dublă
    # în testarea reală — indicii 231-236, 238-241 etc. NU corespund liniar).
    # Pentru defaults se preferă Tabel-cell-based logic (vezi Tabel 0 regim).
    # ══════════════════════════════

    # Contor căldură încălzire (default: NU BIFĂM dacă user n-a setat)
    heat_meter = data.get("heating_has_meter", "")
    heat_meter_cb = {"da": 165, "nu": 166, "nu_caz": 167}.get(heat_meter)
    if heat_meter_cb:
        cbs.append(heat_meter_cb)

    # Repartitoare costuri (fără default)
    cost_alloc = data.get("heating_cost_allocator", "")
    cost_alloc_cb = {"da": 169, "nu": 170, "nu_caz": 171}.get(cost_alloc)
    if cost_alloc_cb:
        cbs.append(cost_alloc_cb)

    # Conducta recirculare ACM (fără default — 195 era deja bifat mai sus)
    acm_recirc = data.get("acm_recirculation", "")
    acm_recirc_cb = {"functionala": 193, "nu_functioneaza": 194, "nu_exista": 195}.get(acm_recirc)
    if acm_recirc_cb:
        cbs.append(acm_recirc_cb)

    # Contor general ACM (fără default)
    acm_meter = data.get("acm_has_meter", "")
    acm_meter_cb = {"da": 196, "nu": 197, "nu_caz": 198}.get(acm_meter)
    if acm_meter_cb:
        cbs.append(acm_meter_cb)

    # Debitmetre puncte consum ACM (fără default)
    acm_flow = data.get("acm_flow_meters", "")
    acm_flow_cb = {"peste_tot": 201, "partial": 200, "nu_exista": 199}.get(acm_flow)
    if acm_flow_cb:
        cbs.append(acm_flow_cb)

    # Ventilare caracteristici (fără default)
    vent_control = data.get("ventilation_control_type", "")
    vent_ctrl_map = {
        "program": 266,
        "manual_simpla": 267,
        "temporizare": 268,
        "jaluzele_reglate": 269,
    }
    vc_cb = vent_ctrl_map.get(vent_control)
    if vc_cb:
        cbs.append(vc_cb)

    # Stare rețea iluminat (doar dacă user specifică non-default)
    light_state = data.get("lighting_network_state", "")
    if light_state == "uzata":
        cbs.append(286)
    elif light_state == "indisp":
        cbs.append(287)

    # Apartamente debranșate condominiu — BINE pentru bloc (confirmat prin screenshot)
    apt_debransate = data.get("building_has_disconnected_apartments", "")
    if category in ("RC", "RA", "BC"):
        apt_deb_cb = {"da": 156, "nu": 157}.get(apt_debransate)
        if apt_deb_cb:
            cbs.append(apt_deb_cb)

    # NOTĂ: Tratare aer (CB 234-236) și Spațiul climatizat (CB 231-233):
    # eliminate din defaults — indicii erau greșiți și cauzau bifare dublă.
    # Se vor adăuga manual dacă user le setează, NU default.
    cool_humid_ctrl = data.get("cooling_humidity_control", "")
    if cool_humid_ctrl == "cu_control" and has_cool:
        cbs.append(235)
    elif cool_humid_ctrl == "cu_partial" and has_cool:
        cbs.append(236)

    cool_space = data.get("cooling_space_scope", "")
    if cool_space == "global" and has_cool:
        cbs.append(232)
    elif cool_space == "partial" and has_cool:
        cbs.append(233)

    cool_ind_meter = data.get("cooling_individual_meter", "")
    if cool_ind_meter == "da" and has_cool:
        cbs.append(252)
    elif cool_ind_meter == "nu" and has_cool:
        cbs.append(253)

    # ══════════════════════════════
    # BACS (CB 308+) — EN 15232-1
    # Dacă template-ul are câmpuri BACS (Anexe cu CB 308-311)
    # ══════════════════════════════
    bacs_class = data.get("bacs_class", "C")
    bacs_cb_map = {"A": 308, "B": 309, "C": 310, "D": 311}
    bacs_cb = bacs_cb_map.get(bacs_class)
    if bacs_cb:
        cbs.append(bacs_cb)

    # SRI readiness (CB 312+)
    sri_total = int(data.get("sri_total", "0") or "0")
    if sri_total >= 70:
        cbs.append(312)  # SRI class A
    elif sri_total >= 50:
        cbs.append(313)  # SRI class B
    elif sri_total >= 30:
        cbs.append(314)  # SRI class C
    else:
        cbs.append(315)  # SRI class D

    return cbs


# ═══════════════════════════════════════════════════════
# Etapa 3 (19 apr 2026) — DYNAMIC CHECKBOX MAPPING (BUG-11 fix)
# ═══════════════════════════════════════════════════════
# Probleme cu maparea statică din compute_checkboxes():
#   - Indicii hardcodate (ex. CB[150] = "sobă") sunt fragili: dacă MDLPA
#     adaugă/elimină un checkbox în template, toți indicii ulteriori se mută
#   - Template clădire (308 cb) ≠ Template apartament (244 cb) → același index
#     numeric reprezintă lucruri diferite în cele două template-uri
#   - Imposibil de menținut între versiuni MDLPA
#
# Soluția — mapping semantic dinamic la runtime:
#   1. Definim CHECKBOX_KEYWORD_MAP: SEMANTIC_KEY → list_de_keywords
#   2. La generare, scanăm template-ul și pentru fiecare checkbox extragem
#      contextul textual (paragraf curent + prev + section header)
#   3. build_checkbox_index(doc) → dict {SEMANTIC_KEY: xml_index_in_doc}
#   4. compute_checkbox_keys(data, category) → list[SEMANTIC_KEY] activ
#   5. toggle_checkboxes_by_keys(doc, keys) → folosește indexul dinamic
#
# Strategie graduală: vechiul compute_checkboxes() rămâne (pentru backward compat
# pe template-ul clădire) + noul compute_checkbox_keys() se aplică ÎN PLUS cu
# semantic matching (idempotent — checkbox deja bifat nu se rebifează diferit).
# ═══════════════════════════════════════════════════════

def _normalize_text(s):
    """Normalizează text pentru match: lowercase, fără diacritice, spații colapsate."""
    import unicodedata
    if s is None:
        return ""
    s = str(s).lower().strip()
    s = unicodedata.normalize("NFD", s)
    s = "".join(c for c in s if unicodedata.category(c) != "Mn")
    s = " ".join(s.split())  # collapse whitespace
    return s


def _get_checkbox_context(cb_elem, doc=None):
    """Extrage contextul textual al unui checkbox: paragraf curent + prev + section header.

    Section header = ultimul paragraf bold/major dintr-o serie de paragrafe goale
    sau cu text care pare label de secțiune (ex. "Existența instalației de încălzire").

    Returnează un string concatenat normalizat, gata pentru substring matching.
    """
    p = cb_elem
    while p is not None and not p.tag.endswith("}p"):
        p = p.getparent()
    if p is None:
        return ""

    parts = []
    # Paragraful curent
    curr_text = "".join(t.text or "" for t in p.iter(qn("w:t")))
    parts.append(curr_text)

    # Paragraful anterior
    prev = p.getprevious()
    if prev is not None and prev.tag.endswith("}p"):
        prev_text = "".join(t.text or "" for t in prev.iter(qn("w:t")))
        parts.append(prev_text)

    # Section header — urc 5 paragrafe în sus căutând un label declarativ
    # (de obicei "Existența instalației...", "Tipul...", "Sursa...")
    header = prev
    SECTION_HEADER_HINTS = (
        "existen", "tipul", "sursa", "tip ", "starea", "structura",
        "categori", "zone climat", "zone eolien", "soluții", "solutii",
        "estimarea", "tipul cl", "anvelopa",
    )
    for _ in range(5):
        if header is None:
            break
        header = header.getprevious()
        if header is None or not header.tag.endswith("}p"):
            continue
        h_text = "".join(t.text or "" for t in header.iter(qn("w:t")))
        h_norm = _normalize_text(h_text)
        if any(h_norm.startswith(hint) for hint in SECTION_HEADER_HINTS) and len(h_norm) > 5:
            parts.append(h_text)
            break

    # Cell context — dacă e într-un tabel, ia toate paragrafele din celulă
    cell = p.getparent()
    while cell is not None and not cell.tag.endswith("}tc"):
        cell = cell.getparent()
    if cell is not None:
        cell_text = "".join(t.text or "" for t in cell.iter(qn("w:t")))
        # Adaugă doar dacă e diferit de curr_text (evită dublare)
        if _normalize_text(cell_text) != _normalize_text(curr_text):
            parts.append(cell_text)

    return _normalize_text(" | ".join(parts))


# ─── CHECKBOX_KEYWORD_MAP ───────────────────────────────────────────────────
# Cheie semantică → listă de keywords (toate trebuie să apară în context).
# Keywords-urile sunt normalizate (lowercase fără diacritice) la match.
# Pentru fiecare semantic key se selectează PRIMUL checkbox al cărui context
# match-uiește toate keywords (ordine de apariție în template).

CHECKBOX_KEYWORD_MAP = {
    # ── Anexa 1 — Recomandări anvelopă (CB ~0-8) ──
    "REC_PE_INSULATE":      ["sporirea rezisten", "pereti", "exteriori"],
    "REC_PB_INSULATE":      ["placii peste subsol"],
    "REC_PT_INSULATE":      ["sporirea rezisten", "terasei"],
    "REC_PL_INSULATE":      ["planseelor in contact"],
    "REC_SARPANTA":         ["sarpantei peste mansarda"],
    "REC_GLAZING":          ["inlocuirea tamplariei"],
    "REC_GRILES_VENT":      ["grilelor de ventilare higroreglabile"],
    "REC_SHADING":          ["dispozitive de umbrire"],

    # ── Anexa 1 — Recomandări instalații (CB ~9-30) ──
    "REC_HEAT_PIPES":       ["schimbarea conductelor", "agentului termic pentru incalzire"],
    "REC_DHW_PIPES":        ["schimbarea conductelor", "apei calde de consum"],
    "REC_HEAT_INSULATE":    ["refacerea izolatiei", "agentului termic pentru incalzire"],
    "REC_DHW_INSULATE":     ["refacerea izolatiei", "apei calde de consum"],
    "REC_THERM_VALVES":     ["robinetelor cu termostat"],
    "REC_BAL_VALVES":       ["vanelor automate de echilibare"],
    "REC_AIR_QUALITY":      ["calitatii aerului interior", "ventilare"],
    "REC_FLOW_METERS":      ["debitmetrelor"],
    "REC_HEAT_METERS":      ["contoarelor de caldura"],
    "REC_LOW_FLOW":         ["armaturilor sanitare cu consum redus"],
    "REC_DHW_RECIRC":       ["recirculare a apei calde de consum"],
    "REC_AUTOMATION":       ["sistem minim de automatizare"],
    "REC_HEAT_EQUIP":       ["echipamentelor din centrala termica"],
    "REC_VENT_EQUIP":       ["centrala de climatizare", "uzate"],
    "REC_LIGHT_LED":        ["corpurilor de iluminat", "surse economice"],
    "REC_PRESENCE_SENS":    ["senzorilor de prezenta"],
    "REC_RENEWABLES":       ["surselor regenerabile"],
    "REC_HEAT_RECOVERY":    ["recuperare a energiei termice"],

    # ── Anexa 2 — Tip clădire (CB 65-67) ──
    "BLDG_EXISTING":        ["tipul cladirii", "existenta"],
    "BLDG_NEW":             ["tipul cladirii", "noua finalizata"],
    "BLDG_UNFINISHED":      ["existenta nefinalizata"],

    # ── Anexa 2 — Zone climatice (I-V) — occurrence_idx per secțiune ──
    # Contextul fiecăruia include header "zone climatice"; distingem prin ordine.
    "CLIMA_I":   (["zone climat"], 0),
    "CLIMA_II":  (["zone climat"], 1),
    "CLIMA_III": (["zone climat"], 2),
    "CLIMA_IV":  (["zone climat"], 3),
    "CLIMA_V":   (["zone climat"], 4),

    # ── Anexa 2 — Zone eoliene (I-III) ──
    "VANT_I":    (["zone eolien"], 0),
    "VANT_II":   (["zone eolien"], 1),
    "VANT_III":  (["zone eolien"], 2),

    # ── Anexa 2 — Categoria clădirii (CB 68-111) ──
    "CAT_RES_INDIV":        ["casa individuala"],
    "CAT_RES_INSIRUITA":    ["casa insiruita"],
    "CAT_RES_BLOC":         ["bloc de locuinte"],
    "CAT_RES_CAMIN":        ["camin / internat"],
    "CAT_EDU_GRADINITA":    ["cladire de invatamant", "gradinita"],
    "CAT_EDU_SCOALA":       ["scoala /liceu/colegiu"],
    "CAT_EDU_UNIV":         ["invatamant superior"],
    "CAT_OFFICE":           ["cladire de birouri", "birouri"],
    "CAT_HOSPITAL":         ["cladire pentru sanatate", "spital"],
    "CAT_HOTEL":            ["cladire pentru turism", "hotel/motel"],
    "CAT_SPORT":            ["cladire pentru sport"],
    "CAT_COMMERCE_SMALL":   ["cladire pentru comert", "magazin comercial mic"],
    "CAT_COMMERCE_BIG":     ["magazin mare"],
    "CAT_OTHER":            ["alte tipuri de cladiri"],

    # ── Anexa 2 — Existență instalații (CB ~135, 176, 202, 256, 272) ──
    "HEAT_EXISTS_OK":       ["existenta instalatiei de incalzire", "da, functionala"],
    "HEAT_EXISTS_NONE":     ["nu", "sistem virtual de incalzire electrica"],
    "DHW_EXISTS_OK":        ["existenta instalatiei de apa calda", "da, functionala"],
    "DHW_EXISTS_NONE":      ["sistem virtual de preparare acc"],
    "COOL_EXISTS_OK":       ["existenta instalatiei de racire", "da, functionala"],
    "COOL_EXISTS_NONE":     ["se ignora consumul de energie pentru racire"],
    "VENT_EXISTS_OK":       ["existenta instalatiei de ventilare mecanica", "da, functionala"],
    "VENT_EXISTS_NONE":     ["se ignora consumul de energie electrica pentru cladiri rezidentiale"],
    "LIGHT_EXISTS_OK":      ["existenta instalatiei de iluminat", "da, functionala"],
    "LIGHT_EXISTS_NONE":    ["sistem virtual de iluminat"],

    # ── Anexa 2 — Tip ventilare (CB ~259-271) ──
    "VENT_NATURAL_NEORG":   ["exclusiv naturala neorganizata"],
    "VENT_NATURAL_ORG":     ["naturala organizata"],
    # "tipul sistemului" în context — distinge de checkbox-ul existenței (care are "existenta")
    "VENT_MECHANICAL":      ["tipul sistemului", "mecanica"],
    # VENT_HR: Da = prima apariție, Nu = a doua apariție în paragraful cu "Da Nu"
    "VENT_HR_YES":          (["recuperator de caldura", "da"], 0),
    "VENT_HR_NO":           (["recuperator de caldura"], 1),

    # ── Anexa 2 — Sursa încălzire Section B ──
    "HEAT_SRC_TERMOFICARE": ["termoficare cu racordare", "incalzire"],
    "HEAT_SRC_CT_PROP":     ["centrala termica proprie", "incalzire"],
    "HEAT_SRC_CT_EXT":      ["centrala termica in exteriorul"],
    "HEAT_SRC_ELECTRIC":    ["energie electrica", "incalzire"],
    "HEAT_SRC_PC_HEAT":     ["pompa de caldura", "incalzire"],
    "HEAT_SRC_SOBE":        ["combustibil solid", "incalzire"],
    "HEAT_TYPE_STATIC":     ["corpuri statice"],
    "HEAT_TYPE_SOBE":       ["sobe", "incalzire"],
    "HEAT_TYPE_ELECTRIC":   ["radiatoare electrice"],
    "HEAT_DIST_INF":        ["distributie inferioara"],
    "HEAT_DIST_SUP":        ["distributie superioara"],

    # ── Anexa 2 — Sursa ACM Section C ──
    "DHW_SRC_TERMOFICARE":  ["termoficare cu racordare", "apa calda"],
    "DHW_SRC_CT_PROP":      ["centrala termica", "apa calda de consum"],
    "DHW_SRC_ELECTRIC":     ["boiler", "electrica"],
    "DHW_SRC_SOLAR":        ["energie solara", "termosolare"],
    "DHW_SRC_PC":           ["pompa de caldura", "apa calda"],

    # ── Anexa 2 — Tip iluminat (CB ~281-285) ──
    "LIGHT_FLUORESCENT":    ["tipul sistemului de iluminat", "fluorescent"],
    "LIGHT_LED":            ["led mixt"],
    "LIGHT_STATE_GOOD":     ["starea retelei electrice", "buna"],

    # ── Anexa 2 — Structura constructivă (CB 127-134) ──
    # Paragrafe p163-p166 cu câte 2 checkboxes / paragraf — folosesc occurrence_idx
    # pentru a alege STÂNGA (0) vs DREAPTA (1) când contextul e identic.
    "STRUCT_ZIDARIE":      (["pereti structurali din zidarie"], 0),
    "STRUCT_BETON_PERETI": (["pereti structurali din zidarie"], 1),
    "STRUCT_BETON_CADRE":  (["cadre din beton armat"], 0),
    "STRUCT_STALPI":       (["cadre din beton armat"], 1),
    "STRUCT_LEMN":         (["structura de lemn"], 0),
    "STRUCT_METAL":        (["structura de lemn"], 1),
    "STRUCT_PANOURI":      (["structuri din panouri mari"], 0),

    # ── Anexa 2 — Regenerabile (CB 288-307) ──
    # Perechi YES/NO cu CONTEXT IDENTIC ("Există Nu există" în același paragraf):
    # folosim format tuple (keywords, occurrence_idx) — al N-lea match conform ordinii.
    "RENEW_SOLAR_TH_YES":   (["sistemul de panouri termosolare"], 0),
    "RENEW_SOLAR_TH_NO":    (["sistemul de panouri termosolare"], 1),
    "RENEW_PV_YES":         (["sistemul de panouri fotovoltaice"], 0),
    "RENEW_PV_NO":          (["sistemul de panouri fotovoltaice"], 1),
    "RENEW_HP_YES":         (["pompa de caldura"], 0),
    "RENEW_HP_NO":          (["pompa de caldura"], 1),
    "RENEW_BIOMASS_YES":    (["sistemul de utilizare a biomasei"], 0),
    "RENEW_BIOMASS_NO":     (["sistemul de utilizare a biomasei"], 1),
    "RENEW_WIND_YES":       (["centrala eoliana"], 0),
    "RENEW_WIND_NO":        (["centrala eoliana"], 1),
}


def build_checkbox_index(doc):
    """Construiește mapping dinamic SEMANTIC_KEY → xml_checkbox_index pentru un template.

    Pentru fiecare cheie din CHECKBOX_KEYWORD_MAP, găsește PRIMUL checkbox al
    cărui context (paragraf curent + prev + header secțiune + celulă tabel)
    conține TOATE keyword-urile asociate (case + diacritice insensitive).

    Returnează dict {key: idx} pentru cheile găsite. Cheile fără match nu apar
    în dict (apelantul folosește vechiul fallback hardcoded).

    Cache pe instanța doc: rezultatul e stocat în doc._zephren_cb_index.
    """
    # Cache pe document — evită re-scanare la apeluri multiple în același flow
    cached = getattr(doc, "_zephren_cb_index", None)
    if cached is not None:
        return cached

    checkboxes = doc.element.findall(".//w:checkBox", NSMAP)
    # Pre-calculează contextele
    contexts = [_get_checkbox_context(cb, doc) for cb in checkboxes]

    index = {}
    for sem_key, kw_spec in CHECKBOX_KEYWORD_MAP.items():
        # Suport pentru două formate:
        #   (a) list[str] — toate keywords trebuie să apară, primul match câștigă
        #   (b) tuple(list[str], int) — al N-lea match (ordinal) câștigă
        if isinstance(kw_spec, tuple) and len(kw_spec) == 2 and isinstance(kw_spec[1], int):
            keywords, occurrence = kw_spec
        else:
            keywords, occurrence = kw_spec, 0
        normalized_kws = [_normalize_text(k) for k in keywords]
        match_count = 0
        for idx, ctx in enumerate(contexts):
            if all(kw in ctx for kw in normalized_kws):
                if match_count == occurrence:
                    index[sem_key] = idx
                    break
                match_count += 1

    # Salvăm pe doc pentru reutilizare
    try:
        doc._zephren_cb_index = index
    except Exception:
        pass
    return index


def toggle_checkboxes_by_keys(doc, semantic_keys):
    """Bifează checkbox-uri folosind chei semantice (mapping dinamic la runtime).

    Pentru fiecare cheie din semantic_keys, găsește indexul real al checkbox-ului
    în template via build_checkbox_index(doc) și îl bifează.

    Returnează dict {found: [keys], missing: [keys]} pentru audit/debug.
    """
    if not semantic_keys:
        return {"found": [], "missing": []}
    cb_index = build_checkbox_index(doc)
    found_keys = [k for k in semantic_keys if k in cb_index]
    missing_keys = [k for k in semantic_keys if k not in cb_index]
    indices = sorted(set(cb_index[k] for k in found_keys))
    if indices:
        toggle_checkboxes(doc, indices)
    return {"found": found_keys, "missing": missing_keys}


def compute_checkbox_keys(data, category):
    """Calculează lista de chei semantice ACTIVE pentru un audit.

    Spre deosebire de compute_checkboxes() care returnează indici hardcodate,
    această funcție returnează chei care sunt rezolvate dinamic la runtime
    cu CHECKBOX_KEYWORD_MAP (independent de versiunea template).

    Returnează list[str].
    """
    keys = []
    is_res = category in ("RI", "RC", "RA")
    u_ref = _U_REF_RES if is_res else _U_REF_NRES

    # Normalizare ID-uri constants.js (uppercase) → legacy semantic
    # Aplicată o singură dată la intrarea în funcție, înainte de orice logică.
    _data = dict(data)
    _data["heating_source"]     = _norm_heating_source(data.get("heating_source", ""))
    _data["acm_source"]         = _norm_acm_source(data.get("acm_source", ""))
    _data["ventilation_type"]   = _norm_vent_type(data.get("ventilation_type", ""))
    _data["lighting_type"]      = _norm_lighting_type(data.get("lighting_type", ""))
    _data["lighting_control"]   = _norm_lighting_control(data.get("lighting_control", ""))
    data = _data

    # Parse opaque + glazing
    try:
        opaque_u = json.loads(data.get("opaque_u_values", "[]"))
    except Exception:
        opaque_u = []
    try:
        glaz_max_u = float(data.get("glazing_max_u", "0") or "0")
        if glaz_max_u != glaz_max_u or glaz_max_u < 0:
            glaz_max_u = 0.0
    except Exception:
        glaz_max_u = 0.0

    # ── Anexa 1 — recomandări anvelopă ──
    if any(e.get("type") == "PE" and float(e.get("u", 0)) > u_ref.get("PE", 0.25) for e in opaque_u):
        keys.append("REC_PE_INSULATE")
    if any(e.get("type") == "PB" and float(e.get("u", 0)) > u_ref.get("PB", 0.29) for e in opaque_u):
        keys.append("REC_PB_INSULATE")
    if any(e.get("type") in ("PT", "PP") and float(e.get("u", 0)) > u_ref.get(e["type"], 0.15) for e in opaque_u):
        keys.append("REC_PT_INSULATE")
    if any(e.get("type") in ("PL", "SE") and float(e.get("u", 0)) > u_ref.get(e["type"], 0.20) for e in opaque_u):
        keys.append("REC_PL_INSULATE")
    # Sprint 26 P1.10 — prag nZEB efectiv pentru REC_GLAZING (Mc 001-2022 Tab 2.5)
    # 1.30/1.80 era prag de PENALIZARE p1 (acum în penalties.js); aici țintă nZEB
    u_glaz_ref = 1.11 if is_res else 1.20
    if glaz_max_u > u_glaz_ref:
        keys.append("REC_GLAZING")

    # ── Anexa 1 — recomandări instalații ──
    vent_type = data.get("ventilation_type", "")
    if not vent_type or vent_type == "natural_neorg":
        keys.append("REC_GRILES_VENT")
    h_src = data.get("heating_source", "")
    if h_src and h_src not in ("electric_direct", "pc_aer_aer"):
        keys.append("REC_THERM_VALVES")
    h_ctrl = data.get("heating_control", "")
    if not h_ctrl or h_ctrl == "manual":
        keys.append("REC_AUTOMATION")
    l_type = data.get("lighting_type", "")
    if l_type and l_type != "led":
        keys.append("REC_LIGHT_LED")
    l_ctrl = data.get("lighting_control", "")
    if not l_ctrl or l_ctrl not in ("sensor_presence", "daylight_dimming"):
        keys.append("REC_PRESENCE_SENS")
    st_en = data.get("solar_thermal_enabled", "") == "true"
    pv_en = data.get("pv_enabled", "") == "true"
    if not st_en and not pv_en:
        keys.append("REC_RENEWABLES")
    if not vent_type or "hr" not in vent_type:
        keys.append("REC_HEAT_RECOVERY")

    # ════════════════════════════════════════════════════════════════════
    # Sprint 25 P0.2 — 14 chei REC_* extinse (Anexa 1 CPE completă)
    # Surse: CHECKBOX_KEYWORD_MAP linii 2142-2171 + audit S25.
    # Fiecare cheie e tratată ca opțională: dacă datele frontend lipsesc
    # (data.get → "") cheia NU se declanșează (no-op silent).
    # ════════════════════════════════════════════════════════════════════
    is_block = category in ("RC", "RA")
    # acm_src e declarat și mai jos în secțiunea Anexa 2 — definim aici pentru P0.2
    acm_src = data.get("acm_source", "")
    has_vent = bool(vent_type) and vent_type != "natural_neorg"

    # REC_SARPANTA — clădiri cu mansardă/șarpantă peste ultimul nivel
    struct_text_low = (data.get("structure", "") or "").lower()
    has_attic = ("mansard" in struct_text_low) or (data.get("attic_heated") == "true")
    if has_attic:
        keys.append("REC_SARPANTA")

    # REC_SHADING — fără rolouri/dispozitive umbrire (g_eff > 0.85 indică lipsa)
    try:
        shading_factor = float(data.get("shading_factor", "1.0") or "1.0")
    except Exception:
        shading_factor = 1.0
    if shading_factor > 0.85:
        keys.append("REC_SHADING")

    # REC_HEAT_PIPES — conducte vechi încălzire (anul instalării < 2000)
    try:
        heat_year = int(data.get("heating_year_installed", "0") or "0")
    except Exception:
        heat_year = 0
    if 0 < heat_year < 2000 and h_src and h_src not in ("electric_direct", "pc_aer_aer"):
        keys.append("REC_HEAT_PIPES")

    # REC_DHW_PIPES — conducte vechi ACM
    try:
        acm_year = int(data.get("acm_year_installed", "0") or "0")
    except Exception:
        acm_year = 0
    if 0 < acm_year < 2000 and acm_src:
        keys.append("REC_DHW_PIPES")

    # REC_HEAT_INSULATE — fără izolație conducte încălzire
    heat_pipe_insul = (data.get("heating_pipe_insulated", "") or "").lower()
    if heat_pipe_insul in ("no", "partial", "nu", "partial_izolata"):
        keys.append("REC_HEAT_INSULATE")

    # REC_DHW_INSULATE — fără izolație conducte ACM
    acm_pipe_insul = (data.get("acm_pipe_insulated", "") or "").lower()
    if acm_pipe_insul in ("no", "partial", "nu", "partial_izolata"):
        keys.append("REC_DHW_INSULATE")

    # REC_BAL_VALVES — bloc fără vane echilibrare
    has_balancing = data.get("heating_has_balancing_valves", "") == "true"
    if not has_balancing and is_block and h_src:
        keys.append("REC_BAL_VALVES")

    # REC_AIR_QUALITY — vent natural / fără filtre F7+ / CO2 > 1200 ppm
    try:
        co2_max = float(data.get("co2_max_ppm", "0") or "0")
    except Exception:
        co2_max = 0.0
    if co2_max > 1200 or vent_type == "natural_neorg":
        keys.append("REC_AIR_QUALITY")

    # REC_FLOW_METERS — apartamente fără contoare individuale ACM
    if data.get("acm_has_meter", "") == "no":
        keys.append("REC_FLOW_METERS")

    # REC_HEAT_METERS — apartamente fără contoare individuale încălzire
    if data.get("heating_has_meter", "") == "no":
        keys.append("REC_HEAT_METERS")

    # REC_LOW_FLOW — armături sanitare fără consum redus
    if data.get("acm_fixtures_low_flow", "") not in ("true", "yes", "da"):
        # declanșăm doar dacă există ACM (altfel nu are sens)
        if acm_src:
            keys.append("REC_LOW_FLOW")

    # REC_DHW_RECIRC — bloc cu ACM dar fără recirculare funcțională
    acm_recirc = (data.get("acm_recirculation", "") or "").lower()
    if is_block and acm_src and acm_recirc not in ("functioneaza", "functional", "yes", "true", "da"):
        keys.append("REC_DHW_RECIRC")

    # REC_HEAT_EQUIP — echipament CT vechi (η < 0.85) sau >15 ani
    try:
        heat_eta = float(data.get("heating_eta_gen", "1.0") or "1.0")
    except Exception:
        heat_eta = 1.0
    if 0 < heat_eta < 0.85 and h_src and h_src not in ("electric_direct", "pc_aer_aer"):
        keys.append("REC_HEAT_EQUIP")
    elif heat_year and (datetime.datetime.now().year - heat_year) > 15 and h_src not in ("electric_direct", "pc_aer_aer"):
        keys.append("REC_HEAT_EQUIP")

    # REC_VENT_EQUIP — centrală climatizare MECANICĂ uzată (>15 ani)
    # Doar dacă există ventilație mecanică (natural_neorg/natural_org NU contează)
    try:
        vent_year = int(data.get("ventilation_year_installed", "0") or "0")
    except Exception:
        vent_year = 0
    has_mechanical_vent = bool(vent_type) and vent_type not in ("natural_neorg", "natural_org", "natural")
    if has_mechanical_vent and 0 < vent_year < 2010:
        keys.append("REC_VENT_EQUIP")

    # ── Anexa 2 — tip clădire ──
    try:
        year_b = int(data.get("year_built", "2000") or "2000")
    except Exception:
        year_b = 2000
    if year_b >= datetime.datetime.now().year - 1:
        keys.append("BLDG_NEW")
    else:
        keys.append("BLDG_EXISTING")

    # ── Anexa 2 — zona climatică (I-V) și zona eoliană (I-III) ──
    # Parsăm climate_zone ("zona III" sau "zona 3") sau zona_climatica_num ("3").
    _ROMAN = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5}
    try:
        _cz_raw = (data.get("climate_zone") or data.get("zona_climatica_num") or "3").strip().upper()
        _cz_tok = _cz_raw.split()[-1]  # "zona III" → "III", "zona 3" → "3"
        _zone_num = _ROMAN.get(_cz_tok) or int(_cz_tok)
    except Exception:
        _zone_num = 3
    _zone_num = max(1, min(5, _zone_num))
    keys.append(["CLIMA_I", "CLIMA_II", "CLIMA_III", "CLIMA_IV", "CLIMA_V"][_zone_num - 1])
    # Zona eoliană: I→zone 1-2, II→zone 3-4, III→zone 5
    _wind = 1 if _zone_num <= 2 else (2 if _zone_num <= 4 else 3)
    keys.append(["VANT_I", "VANT_II", "VANT_III"][_wind - 1])

    # ── Anexa 2 — Structura constructivă (FIX Etapa 7e) ──
    # Mapare building.structure (string lung) → cheia semantică checkbox
    struct_text = (data.get("structure", "") or "").lower()
    struct_key = ""
    if "panouri" in struct_text and "mari" in struct_text:
        struct_key = "STRUCT_PANOURI"
    elif "lemn" in struct_text:
        struct_key = "STRUCT_LEMN"
    elif "metalic" in struct_text or " lsf" in struct_text or "metal" in struct_text:
        struct_key = "STRUCT_METAL"
    elif "cadre" in struct_text and "beton" in struct_text:
        struct_key = "STRUCT_BETON_CADRE"
    elif "diafragm" in struct_text or "dual" in struct_text or ("monolit" in struct_text and "beton" in struct_text):
        struct_key = "STRUCT_BETON_PERETI"
    elif "zidărie" in struct_text or "zidarie" in struct_text:
        struct_key = "STRUCT_ZIDARIE"
    if struct_key:
        keys.append(struct_key)

    # ── Anexa 2 — categoria clădirii ──
    # Audit 2 mai 2026 — Sprint 6 P0: BC (bloc colectiv mixt) era lipsă →
    # CPE generat fără bifa categorie funcțională. Adăugat mapare CAT_RES_BLOC
    # (BC e variantă a clădirii colective rezidențiale cu spații funcționale mixte).
    cat_to_keys = {
        "RI": "CAT_RES_INDIV",
        "RC": "CAT_RES_BLOC",
        "RA": "CAT_RES_BLOC",
        "BC": "CAT_RES_BLOC",  # Sprint 6 — fix lipsă mapping
        "BI": "CAT_OFFICE",
        "ED": "CAT_EDU_SCOALA",
        "SA": "CAT_HOSPITAL",
        "HC": "CAT_HOTEL",
        "CO": "CAT_COMMERCE_SMALL",
        "SP": "CAT_SPORT",
        "AL": "CAT_OTHER",
    }
    cat_key = cat_to_keys.get(category)
    if cat_key:
        keys.append(cat_key)

    # ── Anexa 2 — existență instalații ──
    if h_src:
        keys.append("HEAT_EXISTS_OK")
    else:
        keys.append("HEAT_EXISTS_NONE")
    acm_src = data.get("acm_source", "")
    if acm_src:
        keys.append("DHW_EXISTS_OK")
    else:
        keys.append("DHW_EXISTS_NONE")
    has_cool = data.get("cooling_has", "") == "true"
    if has_cool:
        keys.append("COOL_EXISTS_OK")
    else:
        keys.append("COOL_EXISTS_NONE")
    has_vent = vent_type and vent_type != "natural_neorg"
    if has_vent:
        keys.append("VENT_EXISTS_OK")
    else:
        keys.append("VENT_EXISTS_NONE")
    if l_type:
        keys.append("LIGHT_EXISTS_OK")
    else:
        keys.append("LIGHT_EXISTS_NONE")

    # ── Anexa 2 — tip ventilare ──
    if vent_type in ("natural_neorg", "natural"):
        keys.append("VENT_NATURAL_NEORG")
    elif vent_type == "natural_org":
        keys.append("VENT_NATURAL_ORG")
    elif has_vent:
        keys.append("VENT_MECHANICAL")
    if vent_type and "hr" in vent_type:
        keys.append("VENT_HR_YES")
    else:
        keys.append("VENT_HR_NO")

    # ── Anexa 2 — sursa încălzire (Section B) ──
    if h_src:
        if h_src == "termoficare":
            keys.append("HEAT_SRC_TERMOFICARE")
        elif h_src in ("gaz_conv", "gaz_cond", "centrala_gpl", "ct_prop"):
            keys.append("HEAT_SRC_CT_PROP")
        elif h_src == "ct_ext":
            keys.append("HEAT_SRC_CT_EXT")
        elif h_src in ("electric_direct",):
            keys.append("HEAT_SRC_ELECTRIC")
        elif h_src in ("pc_aer_apa", "pc_sol_apa", "pc_apa_apa", "pompa_caldura"):
            keys.append("HEAT_SRC_PC_HEAT")
        elif h_src in ("cazan_lemn", "cazan_peleti", "soba_teracota"):
            keys.append("HEAT_SRC_SOBE")
        # Tip sistem încălzire
        if h_src == "soba_teracota":
            keys.append("HEAT_TYPE_SOBE")
        elif h_src == "electric_direct":
            keys.append("HEAT_TYPE_ELECTRIC")
        else:
            keys.append("HEAT_TYPE_STATIC")
        # Distribuție (default inferioară)
        keys.append("HEAT_DIST_INF")

    # ── Anexa 2 — sursa ACM (Section C) ──
    if acm_src:
        if acm_src == "termoficare":
            keys.append("DHW_SRC_TERMOFICARE")
        elif acm_src in ("ct_prop", "boiler_gaz"):
            keys.append("DHW_SRC_CT_PROP")
        elif acm_src == "boiler_electric":
            keys.append("DHW_SRC_ELECTRIC")
        elif acm_src == "solar_termic":
            keys.append("DHW_SRC_SOLAR")
        elif acm_src == "pc":
            keys.append("DHW_SRC_PC")

    # ── Anexa 2 — iluminat ──
    if l_type == "fluorescent":
        keys.append("LIGHT_FLUORESCENT")
    elif l_type == "led":
        keys.append("LIGHT_LED")
    keys.append("LIGHT_STATE_GOOD")

    # ── Anexa 2 — regenerabile ──
    keys.append("RENEW_SOLAR_TH_YES" if st_en else "RENEW_SOLAR_TH_NO")
    keys.append("RENEW_PV_YES" if pv_en else "RENEW_PV_NO")
    hp_en = data.get("heat_pump_enabled", "") == "true"
    keys.append("RENEW_HP_YES" if hp_en else "RENEW_HP_NO")
    bio_en = data.get("biomass_enabled", "") == "true"
    keys.append("RENEW_BIOMASS_YES" if bio_en else "RENEW_BIOMASS_NO")
    wind_en = data.get("wind_enabled", "") == "true"
    keys.append("RENEW_WIND_YES" if wind_en else "RENEW_WIND_NO")

    return keys


# ═══════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════
def format_ro(val, dec=1):
    """Format number with Romanian comma decimal separator."""
    return f"{float(val):.{dec}f}".replace(".", ",")


# ═══════════════════════════════════════════════════════
# COLORARE CLASE ENERGETICE PER UTILITATE — tabelul de jos CPE
# ═══════════════════════════════════════════════════════
def _highlight_utility_class_cells(doc, data):
    """Colorează cu culoarea clasei energetice celula din tabelul de jos CPE
    care conține intervalul corespunzător consumului per utilitate.
    Culorile corespund scalei EP: A+=verde închis → G=roșu."""
    import re as _re
    from docx.oxml import OxmlElement as _OxmlElement

    # EP valori per utilitate (kWh/m²·an) — trimise din frontend
    def _parse_ro(s):
        try:
            return float((s or "0").replace(",", "."))
        except Exception:
            return 0.0

    ep_vals = {
        "incalzire": _parse_ro(data.get("ep_incalzire", "")),
        "acm":       _parse_ro(data.get("ep_acm", "")),
        "racire":    _parse_ro(data.get("ep_racire", "")),
        "ventilare": _parse_ro(data.get("ep_ventilare", "")),
        "iluminat":  _parse_ro(data.get("ep_iluminat", "")),
    }

    # CR-2 (7 mai 2026) — clase per utilitate Mc 001-2022 Tab I.1
    # Pre-calculate în JS (getServiceClass) — folosite cu prioritate față de
    # clasificarea bazată pe ep_thresholds_data (care folosea pragurile WHOLE-
    # BUILDING, ducând la clase greșite: ACM 171,8 → C în loc de G; iluminat
    # 43,2 → A+ în loc de F). Fallback la clasificare numerică dacă lipsesc.
    _CLASS_LABELS = ["A+", "A", "B", "C", "D", "E", "F", "G"]
    cls_explicit = {
        "incalzire": (data.get("cls_incalzire") or "").strip(),
        "acm":       (data.get("cls_acm") or "").strip(),
        "racire":    (data.get("cls_racire") or "").strip(),
        "ventilare": (data.get("cls_ventilare") or "").strip(),
        "iluminat":  (data.get("cls_iluminat") or "").strip(),
    }

    # Culori per clasă (0=A+, 1=A, ..., 7=G) — identice cu _EP_CLASS_COLORS
    _COL_FILLS = ["009B00", "32C831", "00FF00", "FFFF00", "F39C00", "FF6400", "FE4101", "FE0000"]

    # Mapare cuvinte-cheie rând → cheie ep
    _ROW_MAP = [
        (["ncălzire", "ncalzire", "eating"],  "incalzire"),
        (["Ap", "caldă", "calda", "ACM"],     "acm"),
        (["Răcire", "Racire", "ooling"],       "racire"),
        (["Ventilare", "entilat"],             "ventilare"),
        (["luminat", "ghting"],                "iluminat"),
    ]

    def _apply_shading(cell, fill_hex):
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = _OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        shd = tcPr.find(qn("w:shd"))
        if shd is None:
            shd = _OxmlElement("w:shd")
            tcPr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)

    def _text_color_for_bg(fill_hex):
        """Negru sau alb — alege culoarea cu contrast WCAG mai mare față de fundal."""
        def to_lin(c8):
            v = c8 / 255.0
            return v / 12.92 if v <= 0.04045 else ((v + 0.055) / 1.055) ** 2.4
        r = int(fill_hex[0:2], 16)
        g = int(fill_hex[2:4], 16)
        b = int(fill_hex[4:6], 16)
        L = 0.2126 * to_lin(r) + 0.7152 * to_lin(g) + 0.0722 * to_lin(b)
        # Pragul de echilibru: L=0.179 → contrast alb = contrast negru
        # Sub prag → alb are contrast mai mare; peste prag → negru are contrast mai mare
        return "FFFFFF" if L < 0.179 else "000000"

    def _format_ro(val):
        """Formatează număr cu 1 zecimală, virgulă românească."""
        return "{:.1f}".format(val).replace(".", ",")

    def _set_cell_text(cell, text, fill_hex, force_black=False):
        """Înlocuiește textul celulei cu valoarea reală, păstrând formatarea."""
        tc = cell._tc
        text_color = "000000" if force_black else _text_color_for_bg(fill_hex)
        # Colectăm proprietățile run-ului existent (dimensiune font etc.)
        existing_sz = None
        for p in tc.iter(qn("w:p")):
            for r in p.iter(qn("w:r")):
                rPr = r.find(qn("w:rPr"))
                if rPr is not None:
                    sz_el = rPr.find(qn("w:sz"))
                    if sz_el is not None:
                        existing_sz = sz_el.get(qn("w:val"))
                break
            if existing_sz:
                break
        # Ștergem toate paragrafele și recreem unul singur centrat
        for p in list(tc.findall(qn("w:p"))):
            tc.remove(p)
        # Paragraph nou
        p_new = _OxmlElement("w:p")
        pPr = _OxmlElement("w:pPr")
        jc = _OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)
        sp = _OxmlElement("w:spacing")
        sp.set(qn("w:before"), "0")
        sp.set(qn("w:after"), "0")
        pPr.append(sp)
        p_new.append(pPr)
        # Run nou
        r_new = _OxmlElement("w:r")
        rPr = _OxmlElement("w:rPr")
        # Bold
        b_el = _OxmlElement("w:b")
        rPr.append(b_el)
        # Culoare text
        color_el = _OxmlElement("w:color")
        color_el.set(qn("w:val"), text_color)
        rPr.append(color_el)
        # Dimensiune font (dacă am extras-o)
        if existing_sz:
            sz_new = _OxmlElement("w:sz")
            sz_new.set(qn("w:val"), existing_sz)
            rPr.append(sz_new)
            szCs_new = _OxmlElement("w:szCs")
            szCs_new.set(qn("w:val"), existing_sz)
            rPr.append(szCs_new)
        r_new.append(rPr)
        t_el = _OxmlElement("w:t")
        t_el.text = text
        r_new.append(t_el)
        p_new.append(r_new)
        tc.append(p_new)

    def _parse_range(text):
        """Returnează (lo, hi) din text ca '≤30', '30...42', '>484'."""
        text = text.replace("\xa0", " ").strip()
        le_m = _re.search(r"[≤<]\s*([\d,\.]+)", text)
        gt_m = _re.search(r"[>]\s*([\d,\.]+)", text)
        rng_m = _re.search(r"([\d,\.]+)\s*[.\-–…]+\s*([\d,\.]+)", text)
        if le_m:
            return (None, float(le_m.group(1).replace(",", ".")))
        if rng_m:
            return (float(rng_m.group(1).replace(",", ".")), float(rng_m.group(2).replace(",", ".")))
        if gt_m:
            return (float(gt_m.group(1).replace(",", ".")), None)
        return None

    def _in_range(v, rng):
        lo, hi = rng
        if lo is None: return v <= hi
        if hi is None: return v > lo
        return lo <= v <= hi

    # Găsește tabelul cu "Clasă energetică" și utility rows
    target = None
    for tbl in doc.tables:
        txt = " ".join(c.text for r in tbl.rows for c in r.cells)
        if ("lasă energetic" in txt or "lasa energetic" in txt) and \
           ("ncălzire" in txt or "ncalzire" in txt or "eating" in txt):
            target = tbl
            break
    if not target:
        return

    def _clear_placeholder_cell(cell):
        """Șterge colorarea și textul placeholder dintr-o celulă din template."""
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is not None:
            shd = tcPr.find(qn("w:shd"))
            if shd is not None:
                tcPr.remove(shd)
        for p in list(tc.findall(qn("w:p"))):
            tc.remove(p)
        p_empty = _OxmlElement("w:p")
        tc.append(p_empty)

    # ── Detectare template generic (placeholder-uri cXA, cXB în loc de valori numerice) ──
    _is_generic_util_table = False
    for row in target.rows:
        for c in row.cells:
            if "c1A" in c.text or "c2A" in c.text:
                _is_generic_util_table = True
                break
        if _is_generic_util_table:
            break

    # ── Completare placeholder-uri cXA...cXB cu intervale calculate ──
    # Template-ul generic are: ≤ c1A, c1A ... c1B, ..., >c1G
    # Le înlocuim cu praguri EP per utilitate (proporționale cu scala globală)
    if _is_generic_util_table:
        ep_thresholds_data = []
        try:
            ep_thresholds_data = [int((data.get(k, "0") or "0")) for k in
                                  ["s_ap", "s_a", "s_b", "s_c", "s_d", "s_e", "s_f"]]
        except Exception:
            ep_thresholds_data = []

        if any(t > 0 for t in ep_thresholds_data):
            # Proporții per utilitate (tipice Mc 001-2022)
            _UTILITY_FRACTIONS = {
                "1": [0.41, 0.41, 0.42, 0.43, 0.43, 0.43, 0.43],  # încălzire
                "2": [0.29, 0.29, 0.29, 0.19, 0.18, 0.17, 0.17],  # ACM
                "3": [0.18, 0.18, 0.18, 0.12, 0.12, 0.11, 0.13],  # răcire
                "4": [0.05, 0.05, 0.05, 0.04, 0.04, 0.03, 0.03],  # ventilare
                "5": [0.07, 0.07, 0.07, 0.07, 0.08, 0.08, 0.08],  # iluminat
            }
            for util_key in ["1", "2", "3", "4", "5"]:
                fracs = _UTILITY_FRACTIONS[util_key]
                util_thresholds = [max(1, round(ep_thresholds_data[i] * fracs[i])) for i in range(7)]
                # Înlocuiri: ≤ cXA → ≤ val, cXA ... cXB → val ... val, >cXG → >val
                labels = ["A", "B", "C", "D", "E", "F", "G"]
                for para in _iter_all_paragraphs(doc, include_txbx=False):
                    pass  # Skip — lucrăm direct pe celulele tabelului

                for row in target.rows:
                    for cell in row.cells:
                        ct = cell.text
                        # Procesăm orice celulă cu cX (util_key=1-5) — inclusiv cXB, cXC etc.
                        if f"c{util_key}" not in ct:
                            continue
                        # Înlocuim fiecare label cXA → prag[0], cXB → prag[1], ..., cXG → prag[6]
                        for i, lbl in enumerate(labels):
                            old_token = f"c{util_key}{lbl}"
                            if old_token in ct:
                                for para in cell.paragraphs:
                                    replace_in_paragraph(para, old_token, str(util_thresholds[i]))
                                ct = cell.text  # refresh after replacement

    for row in target.rows:
        cells = row.cells
        if len(cells) < 3:
            continue
        first_text = cells[0].text.strip()

        ep_key = None
        for keywords, key in _ROW_MAP:
            if any(kw in first_text for kw in keywords):
                ep_key = key
                break
        if ep_key is None:
            continue

        # Deduplifică celulele fuzionate (python-docx returnează duplicate la merge)
        seen_ids = set()
        unique = []
        for c in cells:
            cid = id(c._tc)
            if cid not in seen_ids:
                seen_ids.add(cid)
                unique.append(c)

        # Celule cu placeholder "Consum înc/răc/vm/il" — completăm cu intervalul
        # per utilitate (nu cel global EP). Folosim pragurile calculate per utilitate.
        _PLACEHOLDER_KEYWORDS = ["Consum ", "consum ", "CONSUM"]
        # Determinăm care utilitate este (1-5) pe baza ep_key
        _UTIL_KEY_MAP = {"incalzire": "1", "acm": "2", "racire": "3", "ventilare": "4", "iluminat": "5"}
        _util_k = _UTIL_KEY_MAP.get(ep_key, "")
        for ci, cell in enumerate(unique[1:], start=1):
            ct = cell.text.strip()
            if any(kw in ct for kw in _PLACEHOLDER_KEYWORDS):
                # Calculăm pragurile per utilitate (fracție × EP global)
                if _util_k and any(t > 0 for t in ep_thresholds_data):
                    _UTIL_FRACS = {
                        "1": [0.41, 0.41, 0.42, 0.43, 0.43, 0.43, 0.43],
                        "2": [0.29, 0.29, 0.29, 0.19, 0.18, 0.17, 0.17],
                        "3": [0.18, 0.18, 0.18, 0.12, 0.12, 0.11, 0.13],
                        "4": [0.05, 0.05, 0.05, 0.04, 0.04, 0.03, 0.03],
                        "5": [0.07, 0.07, 0.07, 0.07, 0.08, 0.08, 0.08],
                    }
                    fracs = _UTIL_FRACS.get(_util_k, [0.2]*7)
                    ut = [max(1, round(ep_thresholds_data[i] * fracs[i])) for i in range(7)]
                    if ci == 1:
                        interval_text = "\u2264 " + str(ut[0])
                    elif 2 <= ci <= 7:
                        interval_text = str(ut[ci - 2]) + "  ...  " + str(ut[ci - 1])
                    elif ci == 8:
                        interval_text = ">" + str(ut[6])
                    else:
                        interval_text = ""
                    if interval_text:
                        _set_cell_text(cell, interval_text, "FFFFFF", force_black=True)
                        _apply_shading(cell, "FFFFFF")
                    else:
                        _clear_placeholder_cell(cell)
                else:
                    _clear_placeholder_cell(cell)

        ep_val = ep_vals.get(ep_key, 0.0)
        if ep_val <= 0:
            continue
        # unique[0] = coloana etichetă; unique[1..8] = A+ → G
        # Prioritate: thresholds din date (corecte per categorie/nocool)
        # Fallback: parsare text template (compatibilitate)
        col_idx = None
        ep_thresholds_data = []
        try:
            ep_thresholds_data = [int((data.get(k, "0") or "0")) for k in
                                  ["s_ap", "s_a", "s_b", "s_c", "s_d", "s_e", "s_f"]]
        except Exception:
            ep_thresholds_data = []

        # CR-2 (7 mai 2026) — prioritate clasă explicită din JS (getServiceClass
        # cu Mc 001-2022 Tab I.1), apoi fallback la clasificarea numerică.
        explicit_cls = cls_explicit.get(ep_key, "")
        if explicit_cls in _CLASS_LABELS:
            col_idx = _CLASS_LABELS.index(explicit_cls) + 1  # unique[0]=label, unique[1]=A+
        elif any(t > 0 for t in ep_thresholds_data):
            # Clasificare directă cu pragurile reale (nu range-ul din template)
            # ATENȚIE: foloseste praguri WHOLE-BUILDING, deci poate da clase
            # greșite pentru utilități individuale. Folosit doar ca fallback.
            class_idx_direct = len(ep_thresholds_data)  # default: G
            for i, t in enumerate(ep_thresholds_data):
                if ep_val <= t:
                    class_idx_direct = i
                    break
            col_idx = class_idx_direct + 1  # unique[0]=label, unique[1]=A+ etc.
        else:
            # Fallback: parsare text template
            for i in range(1, len(unique)):
                rng = _parse_range(unique[i].text)
                if rng and _in_range(ep_val, rng):
                    col_idx = i
                    break

        if col_idx is None:
            continue

        class_idx = col_idx - 1  # 0=A+, 1=A, ..., 7=G
        if 0 <= class_idx < len(_COL_FILLS):
            fill = _COL_FILLS[class_idx]
            _apply_shading(unique[col_idx], fill)
            _set_cell_text(unique[col_idx], _format_ro(ep_val), fill, force_black=True)

    # ── Forțăm text negru pe TOATE celulele din tabelul de utilități ──
    # Template-ul generic are culoare roșie (FF0000) hardcodată pe placeholder-uri
    for row in target.rows:
        for cell in row.cells:
            tc = cell._tc
            for r_el in tc.iter(qn("w:r")):
                rPr = r_el.find(qn("w:rPr"))
                if rPr is not None:
                    color_el = rPr.find(qn("w:color"))
                    if color_el is not None:
                        color_el.set(qn("w:val"), "000000")
                    else:
                        c_new = _OxmlElement("w:color")
                        c_new.set(qn("w:val"), "000000")
                        rPr.append(c_new)


# ═══════════════════════════════════════════════════════
# MAIN HANDLER
# ═══════════════════════════════════════════════════════
class handler(BaseHTTPRequestHandler):
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    # ═══════════════════════════════════════════════════════
    # AUDIT REPORT — raport audit energetic Mc 001-2022 Partea IV
    # (Sprint 14 — schelet; Sprint 16 extinde cap. 4/7/8)
    # ═══════════════════════════════════════════════════════
    def _handle_audit_report(self, body):
        """Generează raport audit DOCX de la zero (fără template).

        Întoarce JSON {docx: base64, filename} — contract așteptat de
        AuditReport.jsx:90-97.
        """
        try:
            from docx.shared import Pt as _Pt, Cm as _Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            building = body.get("building", {}) or {}
            inst = body.get("instSummary", {}) or {}
            renew = body.get("renewSummary", {}) or {}
            auditor = body.get("auditor", {}) or {}
            en_class = body.get("energyClass", {}) or {}
            opaque = body.get("opaqueElements", []) or []
            glazing = body.get("glazingElements", []) or []
            bridges = body.get("thermalBridges", []) or []
            measured = body.get("measuredConsumption", {}) or {}

            doc = Document()
            enforce_a4_portrait(doc)

            # ── Antet ──────────────────────────────────
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = title.add_run("RAPORT AUDIT ENERGETIC")
            r.bold = True
            r.font.size = _Pt(16)

            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rs = sub.add_run("Mc 001-2022, Partea IV — Ord. MDLPA 16/2023")
            rs.font.size = _Pt(10)
            rs.italic = True

            doc.add_paragraph()

            # ── Capitol 0. Ipoteze de calcul ────
            # Sprint P2 06may2026 — listă explicită ipoteze + standarde aplicate.
            # Mc 001-2022 §1.5 cere ca raportul să declare ipotezele non-default
            # pentru reproducibilitate și verificare independentă.
            h0 = doc.add_paragraph()
            rh0 = h0.add_run("0. Ipoteze de calcul și standarde aplicate")
            rh0.bold = True
            rh0.font.size = _Pt(13)

            ipoteze_table = doc.add_table(rows=1, cols=2)
            ipoteze_table.style = "Light Grid Accent 1"
            ip_hdr = ipoteze_table.rows[0].cells
            ip_hdr[0].text = ""
            r = ip_hdr[0].paragraphs[0].add_run("Domeniu / Element")
            r.font.name = "Calibri"; r.font.size = _Pt(10); r.bold = True
            ip_hdr[1].text = ""
            r = ip_hdr[1].paragraphs[0].add_run("Standard / Ipoteză adoptată")
            r.font.name = "Calibri"; r.font.size = _Pt(10); r.bold = True

            ipoteze = [
                ("Metodologia de calcul",
                 "Mc 001-2022 (Partea I-IV) — bilanț pe utilizări UE, RER, MEPS"),
                ("Calcul transfer termic",
                 "SR EN ISO 13790:2008 (sezon încălzire) + SR EN ISO 52016-1:2017 (orar)"),
                ("Coeficienți U opaci",
                 "SR EN ISO 6946:2017 (Rsi/Rse + Σd/λ); SR EN ISO 10211:2017 (punți)"),
                ("Coeficienți U vitraje",
                 "SR EN 14351-1:2016 + SR EN 673:2011 (U_g) + SR EN 410:2011 (g_value)"),
                ("Punți termice liniare",
                 "SR EN ISO 14683:2017 + Catalog Mc 001-2022 (165 tipologii)"),
                ("Climă referință",
                 f"Zona {building.get('zonaClimatica', climate.get('zone', '—') if climate else '—')} (SR 4839/Mc 001-2022)"),
                ("Etanșeitate aer",
                 f"n50 = {building.get('n50', '—')} h⁻¹ (default Mc 001-2022 dacă neasigurat blower-door)"),
                ("Aporturi solare",
                 f"Factor umbrire mediu = {building.get('shadingFactor', '0.85')} (SR EN ISO 52010-1)"),
                ("Aporturi interne",
                 "Conform Mc 001-2022 Tab 4.10 (densitate ocupare per categorie)"),
                ("Temperatură interioară setpoint",
                 f"{(systems.get('heating', {}) or {}).get('theta_int', '20')}°C încălzire / 26°C răcire (Mc 001-2022)"),
                ("Categorie comfort IEQ",
                 "Cat. III SR EN 16798-1/NA:2019 (rezidențial) sau Cat. II nerezidențial"),
                ("Factori conversie energie",
                 "f_p_nren conform Mc 001-2022 Tab 8.3 (RO grid 2.5 / gaz 1.1 / lemn 0.2)"),
                ("Emisii CO₂",
                 "Factori SR EN 15603:2008 (cu update RO 2024 — 0.230 kgCO₂/kWh termic)"),
                ("BACS",
                 f"Clasă {(systems.get('heating', {}) or {}).get('bacsClass', 'D')} (EN 15232-1:2017 + ISO 52120-1)"),
                ("Iluminat (LENI)",
                 "SR EN 15193-1:2017/A1:2021 — densitate putere + control + ore funcționare"),
                ("Toleranțe acceptate",
                 "EP ±15% / RER ±5% / U_med ±10% / Q_inc ±15% (calibrare auditor)"),
            ]

            for k, v in ipoteze:
                row = ipoteze_table.add_row().cells
                row[0].text = ""
                p = row[0].paragraphs[0].add_run(k)
                p.font.name = "Calibri"; p.font.size = _Pt(9); p.bold = True
                row[1].text = ""
                p = row[1].paragraphs[0].add_run(v)
                p.font.name = "Calibri"; p.font.size = _Pt(9)

            doc.add_paragraph()

            # ── Capitol 1. Date identificare ───────────
            h1 = doc.add_paragraph()
            rh1 = h1.add_run("1. Date de identificare a clădirii")
            rh1.bold = True
            rh1.font.size = _Pt(13)

            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = "Light Grid Accent 1"
            def _trow(table, k, v, bold_key=True):
                """Adaugă rând cu font Calibri 10pt explicit pe ambele celule."""
                cells = table.add_row().cells
                val_str = str(v) if v not in (None, "", []) else "—"
                for cell, text, bold in ((cells[0], k, bold_key), (cells[1], val_str, False)):
                    cell.text = ""
                    r = cell.paragraphs[0].add_run(text)
                    r.font.name = "Calibri"
                    r.font.size = _Pt(10)
                    r.bold = bold
            _trow(tbl, "Adresă", building.get("address", ""))
            _trow(tbl, "Localitate", building.get("city", ""))
            _trow(tbl, "Județ", building.get("county", ""))
            _trow(tbl, "An construcție", building.get("yearBuilt", ""))
            _trow(tbl, "Categorie funcțională", building.get("category", ""))
            _trow(tbl, "Arie utilă de referință Au", f"{building.get('areaUseful', '') or '—'} m²")
            _trow(tbl, "Volum încălzit V", f"{building.get('volume', '') or '—'} m³")
            _trow(tbl, "Număr apartamente", building.get("units", ""))

            doc.add_paragraph()

            # ── Capitol 2. Auditor energetic ───────────
            h2 = doc.add_paragraph()
            rh2 = h2.add_run("2. Auditor energetic")
            rh2.bold = True
            rh2.font.size = _Pt(13)

            tbl2 = doc.add_table(rows=0, cols=2)
            tbl2.style = "Light Grid Accent 1"
            _trow(tbl2, "Nume prenume", auditor.get("name", ""))
            _trow(tbl2, "Firmă/PFA", auditor.get("company", ""))
            _trow(tbl2, "Atestat nr.", auditor.get("atestat", ""))
            _trow(tbl2, "Gradul", auditor.get("grade", ""))
            _trow(tbl2, "Telefon", auditor.get("phone", ""))
            _trow(tbl2, "Email", auditor.get("email", ""))
            _trow(tbl2, "Data", auditor.get("date", ""))
            _trow(tbl2, "Cod unic MDLPA", auditor.get("mdlpaCode", ""))

            doc.add_paragraph()

            # ── Capitol 3. Performanță energetică ──────
            h3 = doc.add_paragraph()
            rh3 = h3.add_run("3. Indicatori performanță energetică calculați")
            rh3.bold = True
            rh3.font.size = _Pt(13)

            # Sprint 06may2026 audit P0 (B7) — rotunjire la 1 zecimală
            # (era afișat 855.9586758743318 → confuz pentru raport oficial)
            def _fmt(val, decimals=1):
                if val is None or val == "" or val == "—":
                    return "—"
                try:
                    return f"{float(val):.{decimals}f}"
                except (ValueError, TypeError):
                    return str(val)

            tbl3 = doc.add_table(rows=0, cols=2)
            tbl3.style = "Light Grid Accent 1"
            _trow(tbl3, "Clasă energetică EP", en_class.get("cls", "—"))
            _trow(tbl3, "EP total [kWh/(m²·an)]",
                  _fmt(renew.get('ep_adjusted_m2', inst.get('ep_total_m2', '—')), 1))
            _trow(tbl3, "CO₂ specific [kg/(m²·an)]",
                  _fmt(renew.get('co2_adjusted_m2', inst.get('co2_total_m2', '—')), 2))
            _trow(tbl3, "RER [%]", _fmt(renew.get('rer', '—'), 1))
            _trow(tbl3, "Qf total [kWh/an]", _fmt(inst.get('qf_total', '—'), 0))
            _trow(tbl3, "LENI [kWh/(m²·an)]", _fmt(inst.get('leni', '—'), 2))

            doc.add_paragraph()

            # ── Capitol 4. Conformitate normativă (Sprint Pas 7 docs P0-4) ──
            # Implementare REALĂ a evaluării conformității contra normativelor
            # esențiale: Mc 001-2022 (EP, RER), C 107-2005 (U-uri opace),
            # SR EN 14351-1 (U vitraje), L.238/2024 Art. 6 (nZEB), Ord. 348/2026.
            h4 = doc.add_paragraph()
            rh4 = h4.add_run("4. Evaluare conformitate normativă")
            rh4.bold = True
            rh4.font.size = _Pt(13)

            # Pragurile de referință C 107-2005 / Mc 001-2022 (clădire renovată)
            U_REF = {"PE": 0.30, "PP": 0.20, "PT": 0.20, "PB": 0.35, "PL": 0.40,
                     "PSol": 0.35, "PV": 1.30, "F": 1.30}

            cap4_table = doc.add_table(rows=1, cols=4)
            cap4_table.style = "Light Grid Accent 1"
            hdr4 = cap4_table.rows[0].cells
            for i, h in enumerate(["Indicator / Element", "Valoare calculată", "Limită normativă", "Conformitate"]):
                hdr4[i].text = ""
                r = hdr4[i].paragraphs[0].add_run(h)
                r.font.name = "Calibri"; r.font.size = _Pt(10); r.bold = True

            # 4.1 EP global vs. limită Mc 001-2022 (orientativ ~150 kWh/m²·an pt RA renovat)
            ep_val = float(renew.get("ep_adjusted_m2") or inst.get("ep_total_m2") or 0)
            ep_limit_max = 150.0  # Mc 001-2022 Tab 2.4 RA renovat — orientativ
            _trow(cap4_table, "EP primar specific", f"{ep_val:.1f} kWh/(m²·an)", bold_key=False)

            # 4.2 RER vs. prag nZEB (≥ 30% pentru rezidențial)
            rer_val = float(renew.get("rer") or 0)
            _trow(cap4_table, "RER (regenerabile)", f"{rer_val:.1f} %", bold_key=False)

            # 4.3 U-uri elemente opace vs. C 107-2005
            for el in (opaque or []):
                t = el.get("type", "—")
                u_val = el.get("u")
                if u_val is None:
                    continue
                try:
                    u_num = float(u_val)
                except Exception:
                    continue
                u_lim = U_REF.get(t, 0.40)
                conform = "[OK] Conform" if u_num <= u_lim else "[X] Neconform"
                row = cap4_table.add_row().cells
                row[0].text = ""
                # Sprint 06may2026 audit P0 (B8) — eliminare slice [:30]
                # (era „Termopan PVC 4-16-4 dublu — Su" → trunchiat)
                p = row[0].paragraphs[0].add_run(f"U {t} — {el.get('name', '—')}")
                p.font.name = "Calibri"; p.font.size = _Pt(9); p.bold = True
                row[1].text = f"{u_num:.3f} W/(m²·K)"
                row[2].text = f"{u_lim:.2f} W/(m²·K)"
                row[3].text = conform
                for c in row[1:]:
                    for p_ in c.paragraphs:
                        for run in p_.runs:
                            run.font.name = "Calibri"; run.font.size = _Pt(9)

            # 4.4 U-uri vitraje
            for el in (glazing or []):
                u_val = el.get("u")
                if u_val is None:
                    continue
                try:
                    u_num = float(u_val)
                except Exception:
                    continue
                u_lim = 1.30
                conform = "[OK] Conform" if u_num <= u_lim else "[X] Neconform"
                row = cap4_table.add_row().cells
                row[0].text = ""
                # Sprint 06may2026 audit P0 (B8) — eliminare slice [:30]
                p = row[0].paragraphs[0].add_run(f"U vitraj — {el.get('name', '—')}")
                p.font.name = "Calibri"; p.font.size = _Pt(9); p.bold = True
                row[1].text = f"{u_num:.3f} W/(m²·K)"
                row[2].text = f"{u_lim:.2f} W/(m²·K)"
                row[3].text = conform
                for c in row[1:]:
                    for p_ in c.paragraphs:
                        for run in p_.runs:
                            run.font.name = "Calibri"; run.font.size = _Pt(9)

            # Conform global — EP + RER
            ep_conform = ep_val > 0 and ep_val <= ep_limit_max
            rer_conform = rer_val >= 30.0
            global_status = "[OK] Conform L.238/2024 (renovare)" if ep_conform else \
                            ("[!] Necesar reabilitare — EP > 150 kWh/(m²·an)" if ep_val > ep_limit_max else "—")
            nzeb_status = "[OK] Conform nZEB" if (ep_conform and rer_conform) else \
                          "[X] Neconform nZEB (necesar EP < 150 kWh/m² și RER ≥ 30%)"

            doc.add_paragraph()
            p4_summary = doc.add_paragraph()
            p4_summary.add_run("Conformitate globală: ").bold = True
            p4_summary.add_run(f"{global_status}\n")
            p4_summary.add_run("Conformitate nZEB (L.238/2024 Art. 6): ").bold = True
            p4_summary.add_run(f"{nzeb_status}")

            doc.add_paragraph()

            # ── Capitol 5. Anvelopa opacă + vitrată ────
            # Sprint 06may2026 audit P0 (B6) — calcul U din layers când nu e dat explicit
            # (Mc 001-2022 §3.3.1 + SR EN ISO 6946:2017: R_total = Rsi + Σ(d/λ)/1000 + Rse)
            def _calc_u_from_layers(element):
                u_explicit = element.get("u")
                if u_explicit is not None:
                    try:
                        return float(u_explicit)
                    except (ValueError, TypeError):
                        pass
                layers = element.get("layers") or []
                if not layers:
                    return None
                el_type = element.get("type", "")
                # Rezistențe superficiale SR EN ISO 6946:2017 Tab. 7
                rsi_rse = {
                    "PE": (0.13, 0.04), "PR": (0.13, 0.04),
                    "PS": (0.13, 0.04), "PA": (0.13, 0.04),
                    "PV": (0.13, 0.04),  # planșeu vitrat
                    "PI": (0.10, 0.10),  # planșeu intermediar
                    "PB": (0.17, 0.04),  # planșeu peste subsol
                    "PT": (0.10, 0.04),  # planșeu peste tavan
                    "PL": (0.10, 0.04),  # planșeu sub mansardă
                }.get(el_type, (0.13, 0.04))
                rsi, rse = rsi_rse
                r_layers = 0.0
                for ly in layers:
                    try:
                        thickness_mm = float(ly.get("thickness", 0) or 0)
                        lambda_w = float(ly.get("lambda", 0) or 0)
                        if lambda_w > 0:
                            r_layers += (thickness_mm / 1000.0) / lambda_w
                    except (ValueError, TypeError):
                        continue
                r_total = rsi + r_layers + rse
                return (1.0 / r_total) if r_total > 0 else None

            if opaque:
                h5 = doc.add_paragraph()
                rh5 = h5.add_run("5. Anvelopa termică — elemente opace")
                rh5.bold = True
                rh5.font.size = _Pt(13)

                tbl5 = doc.add_table(rows=1, cols=4)
                tbl5.style = "Light Grid Accent 1"
                hdr = tbl5.rows[0].cells
                hdr[0].text = "Element"
                hdr[1].text = "Denumire"
                hdr[2].text = "Arie [m²]"
                hdr[3].text = "U [W/(m²·K)]"
                for el in opaque:
                    cells = tbl5.add_row().cells
                    cells[0].text = str(el.get("type", ""))
                    cells[1].text = str(el.get("name", ""))
                    cells[2].text = str(el.get("area", ""))
                    u_calc = _calc_u_from_layers(el)
                    cells[3].text = f"{u_calc:.3f}" if u_calc is not None else "—"

            if glazing:
                doc.add_paragraph()
                h5b = doc.add_paragraph()
                rh5b = h5b.add_run("5.b Anvelopa termică — elemente vitrate")
                rh5b.bold = True
                rh5b.font.size = _Pt(13)

                tbl5b = doc.add_table(rows=1, cols=3)
                tbl5b.style = "Light Grid Accent 1"
                hdr = tbl5b.rows[0].cells
                hdr[0].text = "Denumire"
                hdr[1].text = "Arie [m²]"
                hdr[2].text = "U [W/(m²·K)]"
                for el in glazing:
                    cells = tbl5b.add_row().cells
                    cells[0].text = str(el.get("name", ""))
                    cells[1].text = str(el.get("area", ""))
                    cells[2].text = str(el.get("u", "") or "—")

            doc.add_paragraph()

            # ── Capitol 6. Sisteme tehnice + (opțional) consum măsurat ──
            # Sprint 06may2026 audit P0 (B5) — Cap. 6 garantat prezent
            # (anterior `if measured:` doar → numerotarea sărea 5b → 7)
            h6 = doc.add_paragraph()
            rh6 = h6.add_run("6. Sisteme tehnice (HVAC + ACM + Iluminat)")
            rh6.bold = True
            rh6.font.size = _Pt(13)

            systems_data = body.get("systems", {}) or {}
            tbl6 = doc.add_table(rows=1, cols=3)
            tbl6.style = "Light Grid Accent 1"
            hdr = tbl6.rows[0].cells
            hdr[0].text = "Sistem"
            hdr[1].text = "Caracteristică"
            hdr[2].text = "Valoare"
            for sys_key, sys_label in [
                ("heating", "Încălzire"),
                ("cooling", "Răcire"),
                ("ventilation", "Ventilare"),
                ("lighting", "Iluminat"),
                ("acm", "ACM"),
            ]:
                sys_d = systems_data.get(sys_key, {}) or {}
                if not sys_d:
                    continue
                # Selectare câmpuri reprezentative per sistem
                fields = []
                if sys_key == "heating":
                    fields = [
                        ("Sursă", sys_d.get("source", "—")),
                        ("Putere [kW]", _fmt(sys_d.get("power"), 1)),
                        ("η generare", _fmt(sys_d.get("eta_gen"), 2)),
                        ("η emisie", _fmt(sys_d.get("eta_em"), 2)),
                        ("η distribuție", _fmt(sys_d.get("eta_dist"), 2)),
                        ("η control", _fmt(sys_d.get("eta_ctrl"), 2)),
                    ]
                elif sys_key == "cooling":
                    fields = [
                        ("Sistem", sys_d.get("system", "—")),
                        ("Putere [kW]", _fmt(sys_d.get("power"), 1)),
                        ("EER", _fmt(sys_d.get("eer"), 2)),
                        ("SEER", _fmt(sys_d.get("seer"), 2)),
                    ]
                elif sys_key == "ventilation":
                    fields = [
                        ("Tip", sys_d.get("type", "—")),
                        ("Debit [m³/h]", _fmt(sys_d.get("airflow"), 0)),
                        ("η HR [%]", _fmt(sys_d.get("hrEfficiency"), 0)),
                        ("Putere ventilator [W]", _fmt(sys_d.get("fanPower"), 0)),
                    ]
                elif sys_key == "lighting":
                    fields = [
                        ("Tip", sys_d.get("type", "—")),
                        ("Densitate putere [W/m²]", _fmt(sys_d.get("pDensity"), 1)),
                        ("Control", sys_d.get("controlType", "—")),
                        ("Ore funcționare", _fmt(sys_d.get("operatingHours"), 0)),
                    ]
                elif sys_key == "acm":
                    fields = [
                        ("Sursă", sys_d.get("source", "—")),
                        ("Consumatori", _fmt(sys_d.get("consumers"), 0)),
                        ("Volum stocaj [L]", _fmt(sys_d.get("storageVolume"), 0)),
                        ("T. livrare [°C]", _fmt(sys_d.get("tSupply"), 0)),
                    ]
                first_in_group = True
                for k, v in fields:
                    cells = tbl6.add_row().cells
                    cells[0].text = sys_label if first_in_group else ""
                    cells[1].text = str(k)
                    cells[2].text = str(v)
                    first_in_group = False

            # Consum măsurat (sub-secțiune opțională Cap. 6.b)
            if measured:
                doc.add_paragraph()
                h6b = doc.add_paragraph()
                rh6b = h6b.add_run("6.b Consum măsurat vs. calculat")
                rh6b.bold = True
                rh6b.font.size = _Pt(13)
                tbl6b = doc.add_table(rows=1, cols=3)
                tbl6b.style = "Light Grid Accent 1"
                hdr = tbl6b.rows[0].cells
                hdr[0].text = "Utilitate"
                hdr[1].text = "Măsurat [kWh/an]"
                hdr[2].text = "Calculat [kWh/an]"
                for util_key, util_label in [
                    ("heating", "Încălzire"),
                    ("cooling", "Răcire"),
                    ("acm", "ACM"),
                    ("lighting", "Iluminat"),
                ]:
                    cells = tbl6b.add_row().cells
                    cells[0].text = util_label
                    cells[1].text = str(measured.get(util_key, "—"))
                    calc_key = f"qf_{util_key[0]}"
                    cells[2].text = str(inst.get(calc_key, "—"))

            doc.add_paragraph()

            # ── Capitol 7. Recomandări de reabilitare (Sprint Pas 7 docs P1-5) ──
            # Generare automată recomandări prioritizate per element/sistem cu
            # justificări tehnice și economii estimate. Înlocuiește placeholder-ul.
            h7 = doc.add_paragraph()
            rh7 = h7.add_run("7. Recomandări de reabilitare prioritizate")
            rh7.bold = True
            rh7.font.size = _Pt(13)

            # Prioritizare măsuri: anvelopă (cele mai mari pierderi) → vitraj →
            # instalații → regenerabile, în ordinea cost-eficiență.
            tbl7 = doc.add_table(rows=1, cols=4)
            tbl7.style = "Light Grid Accent 1"
            hdr7 = tbl7.rows[0].cells
            for i, h in enumerate(["Prioritate", "Măsură recomandată", "Justificare tehnică", "Reducere EP estimată"]):
                hdr7[i].text = ""
                r = hdr7[i].paragraphs[0].add_run(h)
                r.font.name = "Calibri"; r.font.size = _Pt(10); r.bold = True

            recommendations = []
            # Pereți cu U > U_REF
            for el in (opaque or []):
                u_val = el.get("u")
                if u_val is None:
                    continue
                try:
                    u_num = float(u_val)
                except Exception:
                    continue
                t = el.get("type", "")
                u_lim = U_REF.get(t, 0.40)
                # Sprint 06may2026 audit P0 (B8) — nume complet
                if u_num > u_lim * 1.5:
                    recommendations.append((1, f"Termoizolare {t} — {el.get('name', '')}",
                                            f"U={u_num:.2f} > 1.5 × U_REF ({u_lim:.2f}). Necesită ETICS 10-15 cm.",
                                            "~25-40 kWh/(m²·an)"))
                elif u_num > u_lim:
                    recommendations.append((2, f"Îmbunătățire izolație {t} — {el.get('name', '')}",
                                            f"U={u_num:.2f} > U_REF ({u_lim:.2f}).",
                                            "~10-20 kWh/(m²·an)"))

            # Vitraje cu U > 1.30
            for el in (glazing or []):
                u_val = el.get("u")
                if u_val is None:
                    continue
                try:
                    u_num = float(u_val)
                except Exception:
                    continue
                # Sprint 06may2026 audit P0 (B8) — nume complet
                if u_num > 2.0:
                    recommendations.append((1, f"Înlocuire tâmplărie — {el.get('name', '')}",
                                            f"U={u_num:.2f} > 2.0. Geam tripan Low-E.",
                                            "~15-25 kWh/(m²·an)"))
                elif u_num > 1.30:
                    recommendations.append((2, f"Înlocuire tâmplărie — {el.get('name', '')}",
                                            f"U={u_num:.2f} > U_REF (1.30). Geam dublu Low-E.",
                                            "~8-15 kWh/(m²·an)"))

            # Instalații
            systems = body.get("systems", {}) or {}
            heating_data = systems.get("heating", {}) or {}
            eta_h = float(inst.get("eta_total_h") or 0)
            if 0 < eta_h < 0.85:
                recommendations.append((2, "Modernizare sursă încălzire",
                                        f"Randament total η={eta_h*100:.0f}% sub optim. Cazan condensare sau pompă căldură.",
                                        "~30-60 kWh/(m²·an)"))

            # Regenerabile
            if rer_val < 30:
                recommendations.append((1, "Adăugare surse regenerabile (PV + pompă căldură)",
                                        f"RER actual={rer_val:.1f}% < 30% (prag nZEB).",
                                        "~40-70 kWh/(m²·an)"))

            # Ventilare cu HRV dacă lipsește
            ventilation_data = systems.get("ventilation", {}) or {}
            if not ventilation_data.get("hasHR"):
                recommendations.append((3, "Ventilare mecanică cu recuperare căldură (η ≥ 80%)",
                                        "Ventilare naturală — pierderi ~20-30% din necesarul de încălzire.",
                                        "~15-25 kWh/(m²·an)"))

            # Sortează după prioritate
            recommendations.sort(key=lambda x: x[0])
            if not recommendations:
                row = tbl7.add_row().cells
                row[0].text = "—"
                row[1].text = "Clădirea este conformă — nu sunt recomandări critice"
                row[2].text = "—"
                row[3].text = "—"
            else:
                priority_labels = {1: "[1] URGENT", 2: "[2] RECOMANDAT", 3: "[3] OPȚIONAL"}
                for prio, masura, just, red in recommendations[:10]:  # top 10
                    row = tbl7.add_row().cells
                    row[0].text = ""
                    p = row[0].paragraphs[0].add_run(priority_labels.get(prio, str(prio)))
                    p.font.name = "Calibri"; p.font.size = _Pt(9); p.bold = True
                    row[1].text = masura
                    row[2].text = just
                    row[3].text = red
                    for c in row[1:]:
                        for p_ in c.paragraphs:
                            for run in p_.runs:
                                run.font.name = "Calibri"; run.font.size = _Pt(9)

            doc.add_paragraph()

            # ── Capitol 8. Concluzii ──
            h8 = doc.add_paragraph()
            rh8 = h8.add_run("8. Concluzii")
            rh8.bold = True
            rh8.font.size = _Pt(13)
            p8 = doc.add_paragraph()
            p8.add_run(
                f"Clădirea analizată ({building.get('address', '—')}) prezintă "
                f"un consum de energie primară de {ep_val:.1f} kWh/(m²·an), "
                f"încadrată în clasa energetică {en_class.get('cls', '—')}. "
                f"Cota energiei regenerabile este de {rer_val:.1f}%. "
                f"{'Clădirea este conformă cu cerințele nZEB.' if (ep_conform and rer_conform) else 'Clădirea NU îndeplinește cerințele nZEB și necesită lucrări de reabilitare.'} "
                f"Implementarea măsurilor recomandate la Cap. 7 va asigura conformitatea "
                f"cu Mc 001-2022 și L.238/2024 Art. 6, având un orizont de recuperare "
                f"a investiției de 8-15 ani la prețurile actuale ale energiei."
            )
            p8.paragraph_format.space_before = _Pt(4)

            # ── Semnătură ──────────────────────────────
            doc.add_paragraph()
            doc.add_paragraph()
            sig = doc.add_paragraph()
            sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sig.add_run(f"Auditor energetic: {auditor.get('name', '—')}\n")
            sig.add_run(f"Atestat nr.: {auditor.get('atestat', '—')}\n")
            sig.add_run(f"Data: {auditor.get('date', '—')}")

            # Sprint 15 — embed semnătură (ștampila se aplică manual pe printout)
            sig_b64 = auditor.get("signatureDataURL", "") or ""
            if sig_b64 and "," in sig_b64:
                sig_b64 = sig_b64.split(",", 1)[1]
            if sig_b64:
                p_imgs = doc.add_paragraph()
                p_imgs.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                try:
                    p_imgs.add_run().add_picture(
                        io.BytesIO(base64.b64decode(sig_b64)), width=Cm(5.0)
                    )
                except Exception:
                    pass

            # Sprint 15 — QR code pentru verificare (dacă auditor.cpeCode există)
            # Audit 2 mai 2026 — P0.3: URL pointează la landing static cu form
            # de căutare manuală (registrul MDLPA central nu există încă).
            cpe_code_audit = auditor.get("cpeCode") or auditor.get("mdlpaCode") or ""
            if cpe_code_audit:
                from urllib.parse import quote
                verify_url_audit = f"https://zephren.ro/cpe/verifica?cod={quote(cpe_code_audit)}"
                qr_bytes = generate_qr_png(verify_url_audit)
                if qr_bytes:
                    p_qr = doc.add_paragraph()
                    p_qr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    try:
                        p_qr.add_run().add_picture(io.BytesIO(qr_bytes), width=Cm(2.5))
                        p_qr.add_run(f"\nVerificare: zephren.ro/cpe/verifica")
                    except Exception:
                        pass

            buf = io.BytesIO()
            doc.save(buf)
            docx_bytes = buf.getvalue()
            docx_b64 = base64.b64encode(docx_bytes).decode("ascii")

            addr_slug = re.sub(r"[^A-Za-z0-9]+", "_",
                               (building.get("address", "") or "cladire"))[:40]
            filename = f"raport_audit_{addr_slug or 'cladire'}.docx"

            response = json.dumps({"docx": docx_b64, "filename": filename})
            self.send_response(200)
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Content-Type", "application/json; charset=utf-8")
            self.send_header("Content-Length", str(len(response)))
            self.end_headers()
            self.wfile.write(response.encode("utf-8"))
        except Exception as e:
            import traceback
            err_body = json.dumps({
                "error": str(e),
                "trace": traceback.format_exc()[-2000:],
            })
            self.send_response(500)
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Content-Type", "application/json")
            self.end_headers()
            self.wfile.write(err_body.encode("utf-8"))

    def do_POST(self):
        try:
            content_length = int(self.headers.get("Content-Length", 0))
            raw = self.rfile.read(content_length)
            body = json.loads(raw)

            # ═══════════════════════════════════════
            # Routing pe ?type= query (Sprint 14 consolidation)
            # Fallback retro-compat: body.mode
            # ═══════════════════════════════════════
            parsed_url = urlparse(self.path)
            qs = parse_qs(parsed_url.query or "")
            doc_type = (qs.get("type", [""])[0] or "").lower()

            # ── Ramură separată: raport audit energetic ──────────────
            # Întoarce JSON {docx: base64, filename} — flux diferit de CPE,
            # generează DOCX de la zero (fără template MDLPA).
            if doc_type == "audit":
                return self._handle_audit_report(body)

            # ── Alias-uri anexa1/anexa2 pentru variantele PDF ────────
            if doc_type in ("anexa1", "cpe", ""):
                effective_mode = body.get("mode", "cpe") if doc_type == "" else "cpe"
            elif doc_type in ("anexa", "anexa2"):
                effective_mode = "anexa"
            elif doc_type == "anexa_bloc":
                # Etapa 4 (BUG-4) — Anexa 2 multi-apartament
                # Procesare identică cu "anexa" + injecție tabel apartamente la final
                effective_mode = "anexa_bloc"
            else:
                # type necunoscut → fallback la body.mode
                effective_mode = body.get("mode", "cpe")

            tpl_bytes = base64.b64decode(body["template"])
            data = body.get("data", {})
            mode = effective_mode
            category = body.get("category", "AL")

            doc = Document(io.BytesIO(tpl_bytes))

            # ═══════════════════════════════════════
            # A4 ENFORCEMENT — păstrează marginile template-ului MDLPA (5mm top/bot,
            # 17.5mm left/right) care încap CPE-ul pe 1 pagină A4. Altfel marginile
            # default Word (25mm) împing conținutul pe pagina 2.
            # ═══════════════════════════════════════
            enforce_a4_portrait(doc, preserve_margins=True, preserve_fonts=True)

            # ═══════════════════════════════════════
            # 0. SCALE EP + CO₂ — PRIMELE! (înainte de text replacements)
            # Altfel "xxxx"→volume corupe ">xxxx" din template generic
            # ═══════════════════════════════════════
            new_ep = [int(data.get(k, 0) or 0) for k in ["s_ap", "s_a", "s_b", "s_c", "s_d", "s_e", "s_f"]]
            try:
                new_co2 = [float((data.get(k, "0") or "0").replace(",", "."))
                           for k in ["co2_ap", "co2_a", "co2_b", "co2_c", "co2_d", "co2_e", "co2_f"]]
            except Exception:
                new_co2 = []
            if any(v > 0 for v in new_ep):
                replace_scales(doc, category, new_ep,
                               new_co2 if new_co2 and any(v > 0 for v in new_co2) else None)

            # CLASS INDICATORS — săgețile pe scale (tot înainte de text replacements)
            ep_cls_real = data.get("ep_class_real", "")
            ep_cls_ref = data.get("ep_class_ref", "")
            co2_cls_real = data.get("co2_class_real", "")
            if ep_cls_real:
                replace_class_indicators(doc, ep_cls_real, ep_cls_ref, co2_cls_real)

            # ═══════════════════════════════════════
            # 1. TEXT REPLACEMENTS — mapare directă
            # ═══════════════════════════════════════
            # ═══ CRITICAL: ordinea contează! ═══
            # Înlocuim pattern-urile LUNGI înainte de cele SCURTE
            # "xxxx,x" ÎNAINTE de "xxxx" (altfel "xxxx" corupe "xxxx,x")
            # "xxx,x" ÎNAINTE de "xx,x" (altfel "xx,x" corupe "xxx,x")
            # "ZZ.LL.AAAA" ÎNAINTE de "AAAA"

            # 1. Secvențiale (cele mai lungi pattern-uri primele)
            # FIX: ep_specific (kWh/m²,an) și ep_ref — nu ep_total_real (kWh/an total)
            ep_primar_vals = [data.get("ep_specific", "0,0"), data.get("ep_ref", "0,0")]
            replace_seq(doc, "xxxx,x", ep_primar_vals)

            xxx_vals = [data.get("area_ref", "0,0"), data.get("co2_val", "0,0"),
                        data.get("sre_st", "0,0"), data.get("sre_pv", "0,0"),
                        data.get("sre_pc", "0,0"), data.get("sre_bio", "0,0"),
                        data.get("sre_other", "0,0"), data.get("sre_total", "0,0")]
            replace_seq(doc, "xxx,x", xxx_vals)

            # xx,x = [qf_thermal_real, qf_electric_real, qf_thermal_ref, qf_electric_ref]
            # Ordinea corespunde celor 4 apariții din template: real(t,e) + ref(t,e)
            xx_vals = [data.get("qf_thermal", "0,0"), data.get("qf_electric", "0,0"),
                       data.get("qf_thermal_ref", "0,0"), data.get("qf_electric_ref", "0,0")]
            replace_seq(doc, "xx,x", xx_vals)

            # FIX GPS: înlocuim placeholder-ul GPS cu marker temp înainte ca nr_units să
            # corupă " x " (separator coordonate). Înlocuim [[GPS]] la final.
            gps_val = data.get("gps", "")
            replace_in_doc(doc, "II,IIII x LL,LLLL", "[[GPS]]")

            # 2. Înlocuiri simple (de la cele mai lungi la cele mai scurte)
            ordered_replacements = [
                # GPS este protejat acum cu [[GPS]] — nu mai apare " x " în coordonate
                ("ZZ.LL.AAAA", data.get("auditor_date", "")),
                ("ZZ/LL/AAAA", data.get("auditor_date", "").replace(".", "/")),
                ("XX/XXXXX", data.get("auditor_atestat", "")),
                ("zz/ll/aa", data.get("expiry", "")),
                ("GWP,G", data.get("gwp", "0,0")),
                ("RR,R", data.get("rer", "0,0")),
                ("zzz,z", data.get("area_ref", "")),
                ("yyy,y", data.get("area_gross", "")),
                ("AAAA", data.get("year", "____")),
                ("xxxx", data.get("volume", "")),
                # Eliminat: ("regim", ...) — replace prea agresiv corupea fraza
                # "în regim liber" din nota *** (ore depășire confort) → "în P+1+M liber".
                # Placeholder-ul "regim" standalone se înlocuiește mai jos cu verificare
                # strictă (paragraf cu textul EXACT "regim", după strip).
            ]

            for old, new in ordered_replacements:
                if new:
                    replace_in_doc(doc, old, new)

            # Fix dedicat pentru placeholder "regim înălțime" — text unic în template
            # MDLPA (2 cuvinte separate, apare DOAR pe rândul "Regim de înălțime:").
            # Evită coruperea frazei "în regim liber" din nota *** care folosește doar
            # cuvântul "regim" izolat, fără "înălțime" imediat după.
            regime_val = data.get("regime", "")
            if regime_val:
                replace_in_doc(doc, "regim înălțime", regime_val)

            # Adresa — completăm cele 2 rânduri din tabelul DATE CLĂDIRE:
            # R2: "Adresa clădirii: ........" → adresa completă
            # R3: ".... adresa ...." → aceeași adresă (unitatea de clădire certificată)
            address_val = data.get("address", "")
            _addr_filled = 0
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            pt = para.text
                            if ("adresa" in pt.lower() or "Adresa" in pt) and \
                               ("..." in pt or "." * 5 in pt):
                                is_label_row = "Adresa" in pt  # R2 = eticheta "Adresa clădirii:"
                                for run in para.runs:
                                    if "adresa" in run.text.lower() or "Adresa" in run.text or \
                                       "..." in run.text or "." * 5 in run.text:
                                        run.text = ""
                                if para.runs:
                                    if is_label_row:
                                        # R2: etichetă + adresă
                                        para.runs[0].text = "Adresa:  " + address_val
                                    elif _addr_filled == 1:
                                        # R3: lăsăm gol (aceeași unitate = aceeași adresă)
                                        para.runs[0].text = ""
                                _addr_filled += 1

            # Scop CPE — split text "Vânzare/Închiriere/Recepție/Inf"
            scope = data.get("scope", "")
            if scope:
                for para in doc.paragraphs:
                    full = para.text
                    if "nzare" in full or "nchir" in full or "Recep" in full:
                        # Clear all runs with scope text parts and set first one
                        scope_runs = [r for r in para.runs if any(x in r.text for x in ["Vânz", "are", "Închir", "ie", "Recep", "/Inf"])]
                        for r in scope_runs:
                            r.text = ""
                        if scope_runs:
                            scope_runs[0].text = scope
                        break
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                full = para.text
                                if "nzare" in full or "nchir" in full or "Recep" in full:
                                    scope_runs = [r for r in para.runs if any(x in r.text for x in ["Vânz", "are", "Închir", "ie", "Recep", "/Inf"])]
                                    for r in scope_runs:
                                        r.text = ""
                                    if scope_runs:
                                        scope_runs[0].text = scope

            # Program calcul — înlocuiește "................versiunea" cu "ZEPHREN v3.x"
            program_name = data.get("software", "") or "ZEPHREN"
            def fill_program_field(paragraphs):
                for para in paragraphs:
                    if "Program de calcul utilizat" in para.text or "versiunea" in para.text:
                        inserted = False
                        for run in para.runs:
                            if re.search(r'\.{2,}', run.text):
                                run.text = re.sub(r'\.{2,}', program_name if not inserted else '', run.text)
                                inserted = True
                            elif "versiunea" in run.text:
                                run.text = re.sub(r'versiunea\.?', '', run.text).strip()
                            # Curăță puncte reziduale singure (ex: ". " rămas dintr-un run split)
                            if inserted and run.text.strip() in (".", ". ", "..", "..."):
                                run.text = ""
            fill_program_field(doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        fill_program_field(cell.paragraphs)

            # Grad auditor — înlocuiește VALOAREA "I / II", NU eticheta "Gradul"
            grade_val = data.get("auditor_grade", "")
            if grade_val:
                # Spațiere vizibilă pentru grad II (altfel "II" e greu de citit)
                display_grade = "I I" if grade_val == "II" else grade_val
                replace_in_doc(doc, "I / II", display_grade)
                replace_in_doc(doc, "I/II", display_grade)
                # FIX: "gradul" este eticheta vizibilă — NU o înlocuim cu valoarea

            # Auditor name/company/phone/email
            auditor_replacements = {
                "Nume & prenume auditor energetic": data.get("auditor_name", ""),
                "Nume auditor": data.get("auditor_name", ""),
                "nume auditor": data.get("auditor_name", ""),
                "Firma/PFA": data.get("auditor_company", ""),
                "denumire firma": data.get("auditor_company", ""),
                "nr. telefon": data.get("auditor_phone", ""),
                "adresa email": data.get("auditor_email", ""),
                "cod unic": data.get("auditor_mdlpa", ""),
                "Cod unic": data.get("auditor_mdlpa", ""),
            }
            for old, new in auditor_replacements.items():
                if new:
                    replace_in_doc(doc, old, new)

            # ═══════════════════════════════════════
            # Sprint 14 — cod unic CPE (Ord. MDLPA 16/2023 + L.238/2024)
            # Placeholder-uri multiple pentru compatibilitate cu template-uri diferite.
            # Dacă template-ul nu are placeholder, codul rămâne în XML (metadata)
            # și se afișează în UI + XML export via generate-xml.js <CodUnicCPE>.
            # ═══════════════════════════════════════
            cpe_code = data.get("cpe_code", "")
            # cpe_nr = formatul barcode oficial: {nrMDLPA}/{registryIndex}
            cpe_nr = (data.get("cpe_nr", "") or "").strip()
            if cpe_code:
                for placeholder in ["[[CPE_CODE]]", "{{CPE_CODE}}", "CodUnicCPE"]:
                    replace_in_doc(doc, placeholder, cpe_code)
            # Încearcă înlocuire placeholder barcode CPE cu cpe_nr dacă există
            # (VML text box Code 39 nu poate fi înlocuit, dar orice text simplu poate)
            if cpe_nr:
                for placeholder in ["[[CPE_NR]]", "{{CPE_NR}}"]:
                    replace_in_doc(doc, placeholder, cpe_nr)
                    replace_in_vml_raw(doc, placeholder, cpe_nr)
                # Template MDLPA: barcode regreg/XXXXXX — înlocuim doar celulele 7-12
                # cpe_nr conține codul de 6 cifre (registryIndex zero-padded)
                replace_barcode_cells(doc, cpe_nr)
                # Fix Anexa 1+2: înlocuiește placeholder-ul "nr. ......" din titlul
                # "ANEXA 1/2 la Certificatul de performanță energetică nr. ......"
                # Textul e fragmentat în multe run-uri Word, deci folosim `replace_in_paragraph`
                # (joacă pe textul COMBINAT al paragrafului) cu substring exact extras prin regex.
                if mode in ("anexa", "anexa_bloc"):
                    import re as _re_cpe
                    # Regex tolerant: NBSP, spatiu normal, sau zero spatiu intre "nr." si dots
                    _nr_pattern = _re_cpe.compile(r"nr\.[\s ]*\.{3,}")
                    # Format CPE nr: {nrMDLPA}/{registryIndex} - stanga=cod MDLPA auditor,
                    # dreapta=nr secvential auto-generat Zephren (conform barcode template MDLPA).
                    # Campul cpe_nr vine direct din JS pre-format; fallback la calcul local.
                    _cpe_nr = str(data.get("cpe_nr", "") or "").strip()
                    _nr_reg_t = str(data.get("registry_index", "") or "").strip()
                    _nr_mdlpa_t = str(data.get("nr_mdlpa", "") or "").strip()
                    if _cpe_nr:
                        _title_nr = _cpe_nr
                    elif _nr_mdlpa_t and _nr_reg_t:
                        _title_nr = f"{_nr_mdlpa_t}/{_nr_reg_t}"
                    elif _nr_reg_t:
                        _title_nr = _nr_reg_t
                    elif _nr_mdlpa_t:
                        _title_nr = _nr_mdlpa_t
                    else:
                        _title_nr = cpe_code
                    for p in _iter_all_paragraphs(doc):
                        pt = p.text
                        if "Certificatul de performan" not in pt and                            "certificatul de performan" not in pt.lower():
                            continue
                        m = _nr_pattern.search(pt)
                        if not m:
                            continue
                        matched_text = pt[m.start():m.end()]
                        # `replace_in_paragraph` stie sa acopere run-uri multiple
                        replace_in_paragraph(p, matched_text, "nr. " + _title_nr, count=1)

            # ═══════════════════════════════════════
            # Sprint 15 — Semnătură + ștampilă + QR code (Ord. MDLPA 16/2023)
            # Plasat DUPĂ text replacements (pentru ca placeholder-urile să existe
            # încă în document). Imaginile înlocuiesc placeholder-ele text cu PNG.
            # ═══════════════════════════════════════
            signature_b64 = data.get("signature_png_b64", "")
            stamp_b64 = data.get("stamp_png_b64", "")
            # Pentru CPE: injectează semnătură/ștampilă în paragraful existent din template.
            # Pentru Anexa 1+2: SKIP — auditorul semnează manual pe printout (Ord. MDLPA 16/2023
            # cere doar UN exemplar semnat fizic; Anexa nu are loc dedicat pentru semnătură
            # injectată în corpul documentului — textul „Semnătura și ștampila auditorului"
            # apare DEJA în footerul oficial al fiecărei pagini).
            if signature_b64 and mode == "cpe":
                insert_signature_stamp(doc, signature_b64, "")

            # Audit 2 mai 2026 — NU MAI MODIFICĂM celula „Anul construirii/renovării majore":
            # template-ul oficial MDLPA are deja celula vertical-merged între R1+R2
            # (R1C1._tc IS R2C1._tc), deci eticheta apare O SINGURĂ DATĂ cu o singură
            # valoare AAAA care se înlocuiește în populare normală cu yearBuilt.
            # NU EXISTĂ duplicate de eliminat — versiunea anterioară (vMerge sau
            # cleanup text) STRICA template-ul oficial. Las flow-ul de populare nativ.

            qr_url = data.get("qr_verify_url", "")
            if qr_url:
                insert_qr_code(doc, qr_url, cpe_code)

            # Sprint 17 — QR code suplimentar pentru pașaport renovare (EPBD 2024/1275 Art. 12)
            # Plasat în placeholder {{QR_PASSPORT}} dedicat (dacă există în template).
            passport_qr_url = data.get("passport_qr_url", "")
            if passport_qr_url:
                try:
                    qr_pass_bytes = generate_qr_png(passport_qr_url, scale=4, border=2)
                    if qr_pass_bytes:
                        for ph in ("{{QR_PASSPORT}}", "{{QR_PASAPORT}}"):
                            _replace_placeholder_with_image(doc, ph, qr_pass_bytes, width_cm=2.2)
                except Exception as e_qrp:
                    print(f"[passport_qr] {e_qrp}", flush=True)

            # Sprint 15 — Cadastru + CF + identificare juridică (placeholder-uri opționale)
            # Sprint 17 — extensie cu pașaport renovare UUID + URL
            for ph_key, val in [
                ("{{NR_CADASTRAL}}", data.get("cadastral_number", "")),
                ("{{CARTE_FUNCIARA}}", data.get("land_book", "")),
                ("{{ARIE_CONSTRUITA}}", data.get("area_built", "")),
                ("{{NR_APARTAMENTE}}", data.get("n_apartments", "")),
                ("{{VALIDITY_YEARS}}", str(data.get("validity_years", ""))),
                ("{{VALIDITY_LABEL}}", data.get("validity_label", "")),
                ("{{PASSPORT_UUID}}", data.get("passport_uuid", "")),
                ("{{PASSPORT_URL}}", data.get("passport_url", "")),
                # Audit 2 mai 2026 — P0.4: placeholder observații auditor.
                # Daca template-ul nu contine {{AUDITOR_OBSERVATIONS}}, replace_in_doc
                # face no-op (val gol oricum sare cu `if val` de mai jos).
                ("{{AUDITOR_OBSERVATIONS}}", data.get("auditor_observations", "")),
            ]:
                if val:
                    replace_in_doc(doc, ph_key, val)

            # nZEB status — bifează checkbox-ul dacă clădirea e nZEB
            set_nzeb_checkbox(doc, data.get("nzeb", "NU") == "DA")

            # (secvențialele xxxx,x / xxx,x / xx,x au fost mutate mai sus, ordinea contează)

            # Categoria clădirii — înlocuiește placeholder-ul RA cu label-ul corect
            # FIX: "Apartament x camere, din bloc" → category_label din dropdown
            cat_label = data.get("category_label", "")
            if cat_label:
                replace_in_doc(doc, "Apartament x camere, din bloc", cat_label)
                replace_in_doc(doc, "categorie funcțională", cat_label)
                replace_in_doc(doc, "categorie functionala", cat_label)
                # Template generic: celula conține "Categoria clădirii:...categoria"
                # replace_in_paragraph e case-sensitive: "categoria" (c mic) potrivește
                # doar placeholder-ul de la poz 38, NU "Categoria" (C mare) de la poz 0
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            ct = cell.text
                            if "Categoria" in ct and "categoria" in ct:
                                for para in cell.paragraphs:
                                    if "categoria" in para.text:
                                        replace_in_paragraph(para, "categoria", cat_label, count=1)
                                        # Reducem spațiile excesive dintre etichetă și valoare
                                        for run in para.runs:
                                            if "Categoria" in run.text and "   " in run.text:
                                                run.text = run.text.rstrip() + "  "

            # Nr camere (RA) — înlocuiește " x " rămas (dacă mai există) din "Apartament x camere"
            if category == "RA":
                replace_in_doc(doc, " x ", " " + data.get("nr_units", "3") + " ")

            # FIX GPS: acum că nr_units a înlocuit " x " din template, punem GPS-ul real
            replace_in_doc(doc, "[[GPS]]", gps_val)

            # Location
            replace_in_doc(doc, "localitatea", data.get("city", ""))
            replace_in_doc(doc, "județul", data.get("county", ""))
            replace_in_doc(doc, "judetul", data.get("county", ""))
            replace_in_doc(doc, "zona climatică", data.get("climate_zone", ""))
            replace_in_doc(doc, "zona climatica", data.get("climate_zone", ""))

            # GWP lifecycle
            gwp_text = data.get("gwp", "0,0") + " kgCO2eq/m2an"
            replace_in_doc(doc, "GWP lifecycle", gwp_text)

            # Note de subsol CPE — ore supraîncălzire
            # ("regim" e deja înlocuit de ordered_replacements cu valoarea regime)
            # ".....................h" → ore supraîncălzire (0 dacă se calculează răcire)
            has_cool = data.get("cooling_has", "") == "true"
            overheat_h = "0" if has_cool else (data.get("overheating_hours", "0") or "0")
            replace_in_doc(doc, ".....................h", overheat_h + " h")

            # Aliniere stânga pentru valorile completate automat din coloana stângă a
            # tabelului "DATE PRIVIND CLĂDIREA CERTIFICATĂ" (Categoria, Adresa, Coordonate
            # GPS, Regim de înălțime). Template MDLPA folosește 8-12 spații după ":" pentru
            # pseudo-aliniere tab — după replace cu valori reale, rezultă indent vizual mare.
            # Colapsez spațiile multiple → 1 spațiu DOAR pe aceste 4 labels.
            _LEFT_ALIGN_LABELS = [
                # Coloana stângă DATE CLĂDIRE
                "Categoria clădirii:",
                "Adresa:",
                "Coordonate GPS (lat x long):",
                "Regim de înălțime:",
                # Coloana dreaptă DATE CLĂDIRE
                "Anul construirii/renovării majore:",
                "Aria de referință a pardoselii:",
                "Aria utilă / desfășurată:",
                "Volumul interior de referință:",
            ]
            from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN
            for _para in _iter_all_paragraphs(doc, include_txbx=True):
                _pt = _para.text
                # 1) Colapsează spații după ":" pentru labels din tabelul DATE CLĂDIRE
                for _lbl in _LEFT_ALIGN_LABELS:
                    if _pt.startswith(_lbl) and "  " in _pt[len(_lbl):]:
                        _new_text = _lbl + " " + _pt[len(_lbl):].lstrip()
                        if _para.runs:
                            _first = _para.runs[0]
                            _first.text = _new_text
                            for _r in _para.runs[1:]:
                                _r.text = ""
                        break
                # 2) "m²"/"m³" e într-o celulă separată de tabel (lățime 3680 DXA pentru
                #    valoare, 429 DXA pentru m²) — setarea LEFT pe celula m² nu e
                #    suficientă: rămâne spațiu de ~20mm între valoare și m² din cauza
                #    lățimii mari a celulei valorii.
                #    Fix: concatenez " m²" (arie) sau " m³" (volum) la finalul valorii
                #    și golesc celula unitate separată. Astfel unitatea apare lipită.
                #    Acoperă AMBELE template-uri: clădire (label fără "apart")
                #    și apartament (label cu "apart" / "apartamentului").
                _AREA_LABELS = (
                    # Clădire
                    "Aria de referință a pardoselii:",
                    "Aria utilă / desfășurată:",
                    # Apartament (template 4-CPE-apartament-bloc)
                    "Aria de referință a pardoselii apart",
                    "Aria utilă a apartamentului",
                )
                _VOLUME_LABELS = (
                    # Clădire
                    "Volumul interior de referință:",
                    # Apartament
                    "Volumul interior de referință al apart",
                )
                _matched_unit = None
                for _mlbl in _AREA_LABELS:
                    if _pt.startswith(_mlbl) and " m²" not in _pt and " m2" not in _pt:
                        _matched_unit = "m²"
                        break
                if _matched_unit is None:
                    for _mlbl in _VOLUME_LABELS:
                        if _pt.startswith(_mlbl) and " m³" not in _pt and " m3" not in _pt:
                            _matched_unit = "m³"
                            break
                if _matched_unit and _para.runs:
                    _para.runs[-1].text = _para.runs[-1].text.rstrip() + " " + _matched_unit
                # Golire celule unitate separate (unde textul e STRICT m2/m²/m3/m³)
                if _pt.strip() in ("m2", "m²", "m3", "m³"):
                    for _r in _para.runs:
                        _r.text = ""

            # (Scale EP/CO₂ și class indicators — mutate la secțiunea 0, înainte de text replacements)

            # ═══════════════════════════════════════
            # 4. CHECKBOXES (Anexa)
            # ═══════════════════════════════════════
            if mode in ("anexa", "anexa_bloc"):
                # Template apartament (mode=anexa, 244 cb) vs clădire (mode=anexa_bloc, 308 cb):
                # compute_checkboxes() produce DOUĂ categorii de indici:
                #   CB 0-64  = Anexa 1 (recomandări + estimare costuri/economii/recuperare)
                #              → ACEEAȘI structură în ambele template-uri → SE APLICĂ
                #   CB 65+   = Anexa 2 (tip clădire, zone, instalații)
                #              → ordine diferită pe template apartament → CAUZEAZĂ ERORI
                #              → înlocuit complet de semantic matching (mai jos)
                cb_all = compute_checkboxes(data, category)
                client_cbs = body.get("checkboxes", [])
                if client_cbs:
                    cb_all = list(set(cb_all + client_cbs))
                if mode == "anexa":
                    # Aplică NUMAI Anexa 1 (CB 0-64); Anexa 2 → semantic matching
                    cb_indices = [i for i in cb_all if i <= 64]
                else:
                    cb_indices = cb_all
                if cb_indices:
                    toggle_checkboxes(doc, cb_indices)

                # #12b (audit Pas 6+7 V7, 7 mai 2026) — Notă auditor în footer Anexa 1
                # când savings cad în intervalul 50-59% (gap tipografic în template-ul
                # oficial MDLPA Ord. 16/2023 — sare de la „40-50%" la „>60%"). Adaugă
                # un paragraf italic explicativ DUPĂ secțiunea „Estimarea economiilor"
                # fără a modifica structura template-ului oficial.
                try:
                    if data.get("_needs_50_60_note") == "1":
                        savings_actual = data.get("_savings_pct_actual", "")
                        savings_str = f"{float(savings_actual):.1f}".replace(".", ",") if savings_actual else "50-59"
                        for tbl_idx, tbl in enumerate(doc.tables):
                            txt = " ".join(c.text for r in tbl.rows for c in r.cells)
                            if "40-50%" in txt and ">60%" in txt:
                                # Găsit tabelul/secțiunea cu economii — adăugăm note
                                # paragraf imediat după acest tabel
                                last_p = doc.paragraphs[-1]
                                note_p = doc.add_paragraph()
                                note_run = note_p.add_run(
                                    f"N.B. Auditor — Economii estimate: {savings_str}%. "
                                    f"Categoria '50-60%' nu apare în template-ul oficial "
                                    f"Ord. MDLPA 16/2023 Anexa 1 (gap tipografic în original "
                                    f"care sare de la '40-50%' la '>60%'). Categoria '>60%' "
                                    f"a fost bifată ca cea mai apropiată; auditorul confirmă "
                                    f"valoarea reală în această notă conform L.372/2005 R2 "
                                    f"și Reg. UE 244/2012."
                                )
                                note_run.italic = True
                                note_run.font.size = Pt(8)
                                from docx.shared import RGBColor as _RGB
                                note_run.font.color.rgb = _RGB(0x80, 0x60, 0x00)  # amber
                                break
                except Exception as _note_err:
                    print(f"[12b notă footer 50-60%] {_note_err}", flush=True)

                # Mapping dinamic — chei semantice rezolvate la runtime
                try:
                    sem_keys = compute_checkbox_keys(data, category)
                    if sem_keys:
                        result = toggle_checkboxes_by_keys(doc, sem_keys)
                        if result.get("missing"):
                            print(
                                f"[checkbox_dynamic] {len(result['found'])} found, "
                                f"{len(result['missing'])} missing keys: "
                                f"{result['missing'][:5]}{'...' if len(result['missing']) > 5 else ''}",
                                flush=True,
                            )
                except Exception as e_dyn:
                    print(f"[checkbox_dynamic] error: {e_dyn}", flush=True)

            # ═══════════════════════════════════════
            # 4b. ANEXA 2 — TEXT REPLACEMENTS
            # ═══════════════════════════════════════
            if mode in ("anexa", "anexa_bloc"):
                # Adresa și nr certificat
                replace_in_doc(doc, "[adresa]", data.get("address", ""))
                # An construcție/renovare — FIX 20 apr 2026:
                # replace_in_doc(".................") era prea agresiv: înlocuia
                # global toate șirurile de 17 puncte (cum sunt placeholder-ele de
                # 151/150/121 puncte din template pentru "Enunțarea etapelor",
                # "Informații stimulente", "auditorul completează"), rezultând
                # "2015201520152015..." (string-ul "2015" repetat de 7-8 ori).
                # Fix chirurgical: înlocuiesc DOAR în paragraful care conține
                # "Anul construc" (eticheta semantică unică), o singură dată.
                year_full = data.get("year", "") + (" / " + data.get("year_renov", "") if data.get("year_renov") else "")
                if year_full:
                    for _p_year in doc.paragraphs:
                        if "Anul construc" in _p_year.text and "................." in _p_year.text:
                            replace_in_paragraph(_p_year, ".................", year_full, count=1)
                            break
                # ── FIX 21 apr 2026: helper pentru înlocuire COMPLETĂ paragraf ──
                # Bug raportat: replace_in_doc înlocuiește doar eticheta lăsând restul
                # ca duplicat ("Aria de referință totală: 6800 m² a pardoselii clădirii sau a unității de clădire: m2").
                # Soluție: replace COMPLET paragraf identificat prin label semantic.
                def _replace_full_para(label_match, new_full_text):
                    """Înlocuiește TEXTUL ÎNTREG al paragrafului care conține label_match."""
                    for p in doc.paragraphs:
                        if label_match in p.text:
                            # Idempotent: skip dacă noul text e deja prezent identic
                            if p.text.strip() == new_full_text.strip():
                                return True
                            # Curăț textul existent — păstrez primul run, șterg restul
                            if p.runs:
                                p.runs[0].text = new_full_text
                                for r in p.runs[1:]:
                                    r.text = ""
                            else:
                                p.add_run(new_full_text)
                            return True
                    return False

                # Arie referință totală
                au = data.get("area_ref", "")
                vol = data.get("volume", "")
                if au:
                    _replace_full_para(
                        "Aria de referință totală",
                        f"Aria de referință totală a pardoselii clădirii sau a unității de clădire: {au} m²"
                    )
                if vol:
                    _replace_full_para(
                        "Volumul interior de referință",
                        f"Volumul interior de referință V, al clădirii/unității de clădire: {vol} m³"
                    )
                # Factor formă
                try:
                    au_f = float(au.replace(",", ".")) if au else 0
                    vol_f = float(vol.replace(",", ".")) if vol else 0
                    ae = float(data.get("area_envelope", "0").replace(",", ".")) if data.get("area_envelope") else au_f * 1.3
                    se_v = ae / vol_f if vol_f > 0 else 0
                    if se_v > 0:
                        _replace_full_para(
                            "Factorul de formă",
                            f"Factorul de formă al clădirii, SE/V: {format_ro(se_v, 3)} m⁻¹"
                        )
                except Exception:
                    pass
                # Nr persoane — label diferit între template clădire și apartament:
                # clădire:   "Numărul normat de persoane din clădire/unitatea de clădire:"
                # apartament:"Numărul maxim real/normat de persoane din apartament:"
                try:
                    is_res = category in ("RI", "RC", "RA")
                    nr_pers = max(1, round(au_f / (30 if is_res else 15)))
                    nr_pers_str = str(nr_pers)
                    # Folosesc preferat valoarea trimisă din frontend, dacă există
                    nr_pers_user = (data.get("nr_persoane") or "").strip()
                    if nr_pers_user:
                        nr_pers_str = nr_pers_user
                    # Strategie robustă: caut paragraful (incl. în tabele) și înlocuiesc
                    # placeholder-ul de DOTS sau SPAȚII cu numărul de persoane.
                    # Bug anterior: `_fill_para_blank` căuta doar spații/nbsp, dar
                    # template-ul foloseste „..................." (dots) ca placeholder.
                    import re as _re_pers
                    _pers_re = _re_pers.compile(r"[\xa0\s]*\.{3,}[\xa0\s]*|[\xa0\s]{6,}")
                    filled = False
                    for _p in _iter_all_paragraphs(doc):
                        pt = _p.text
                        if "persoane din apartament" not in pt and \
                           "persoane din clădire" not in pt and \
                           "persoane din unitatea" not in pt:
                            continue
                        # Sărim dacă valoarea e deja prezentă (idempotent)
                        if f": {nr_pers_str}" in pt or f" {nr_pers_str} pers" in pt:
                            filled = True
                            break
                        m = _pers_re.search(pt)
                        if not m:
                            continue
                        matched = pt[m.start():m.end()]
                        n = replace_in_paragraph(_p, matched, " " + nr_pers_str + " ", count=1)
                        if n:
                            filled = True
                            break
                    if not filled:
                        _replace_full_para(
                            "Numărul normat de persoane",
                            f"Numărul normat de persoane din clădire/unitatea de clădire: {nr_pers} pers."
                        )
                except Exception:
                    pass
                # Detalii instalații — combustibil
                fuel_labels = {"gaz_nat": "gaz natural", "gpl": "GPL", "motorina": "motorină",
                               "lemn": "lemne", "peleti": "peleți", "carbune": "cărbune",
                               "electric": "electricitate", "biogaz": "biogaz"}
                fuel = fuel_labels.get(data.get("heating_fuel", ""), data.get("heating_fuel", ""))
                if fuel:
                    replace_in_doc(doc, "combustibil .....................", "combustibil " + fuel)
                    replace_in_doc(doc, "combustibil ...........", "combustibil " + fuel)
                # Putere nominală încălzire — _replace_full_para evită duplicat
                hp = data.get("heating_power", "0")
                if hp and hp != "0":
                    _replace_full_para(
                        "Necesarul de căldură de calcul",
                        f"Necesarul de căldură de calcul (sarcina termică necesară): {hp} kW"
                    )
                    _replace_full_para(
                        "Puterea termică instalată totală pentru încălzire",
                        f"Puterea termică instalată totală pentru încălzire: {hp} kW"
                    )

                # ── U-values per tip element ──────────────────────────────
                # Înlocuim texte tip "U pereți exteriori" / "U acoperiș" etc.
                try:
                    opaque_u = json.loads(data.get("opaque_u_values", "[]"))
                    raw_glaz = data.get("glazing_max_u", "0") or "0"
                    try:
                        glaz_u = float(raw_glaz)
                        if glaz_u != glaz_u or glaz_u < 0:
                            glaz_u = 0.0
                    except (ValueError, TypeError):
                        glaz_u = 0.0
                    glaz_g = float(data.get("glazing_g_value", "0") or "0")

                    pe_u = [e for e in opaque_u if e.get("type") == "PE"]
                    pt_u = [e for e in opaque_u if e.get("type") in ("PT","PP")]
                    pb_u = [e for e in opaque_u if e.get("type") == "PB"]
                    pl_u = [e for e in opaque_u if e.get("type") == "PL"]

                    if pe_u:
                        u_val = format_ro(pe_u[0].get("u", 0), 2)
                        replace_in_doc(doc, "U pereți exteriori", f"U pereți exteriori = {u_val} W/(m²·K)")
                        replace_in_doc(doc, "coeficientul global de transfer termic al pereților exteriori",
                                       f"U pereți ext. = {u_val} W/(m²·K)")
                    if pt_u:
                        u_val = format_ro(pt_u[0].get("u", 0), 2)
                        replace_in_doc(doc, "U terasă/acoperiș", f"U terasă = {u_val} W/(m²·K)")
                    if pb_u:
                        u_val = format_ro(pb_u[0].get("u", 0), 2)
                        replace_in_doc(doc, "U planșeu subsol", f"U planșeu subsol = {u_val} W/(m²·K)")
                    if glaz_u > 0:
                        replace_in_doc(doc, "U tâmplărie", f"U tâmplărie = {format_ro(glaz_u, 2)} W/(m²·K)")
                        replace_in_doc(doc, "coeficientul global de transfer termic al tâmplăriei",
                                       f"U tâmplărie = {format_ro(glaz_u, 2)} W/(m²·K)")
                    if glaz_g > 0:
                        replace_in_doc(doc, "factorul solar g", f"factorul solar g = {format_ro(glaz_g, 2)}")
                except Exception:
                    pass

                # ── EP breakdown pe destinații ────────────────────────────
                ep_breakdown = {
                    "ep_incalzire": data.get("ep_incalzire", ""),
                    "ep_racire": data.get("ep_racire", ""),
                    "ep_acm": data.get("ep_acm", ""),
                    "ep_ventilare": data.get("ep_ventilare", ""),
                    "ep_iluminat": data.get("ep_iluminat", ""),
                }
                ep_labels = {
                    "ep_incalzire": "EP încălzire",
                    "ep_racire": "EP răcire",
                    "ep_acm": "EP ACM",
                    "ep_ventilare": "EP ventilare",
                    "ep_iluminat": "EP iluminat",
                }
                for key, val in ep_breakdown.items():
                    if val:
                        replace_in_doc(doc, ep_labels[key], ep_labels[key] + " = " + val + " kWh/(m²·an)")

                # ── BACS clasă automatizare ───────────────────────────────
                bacs_class = data.get("bacs_class", "")
                bacs_labels = {
                    "A": "Clasa A — control predictiv (economie 25-40%)",
                    "B": "Clasa B — control avansat (economie 10-25%)",
                    "C": "Clasa C — automatizare de bază (referință)",
                    "D": "Clasa D — fără automatizare",
                }
                if bacs_class in bacs_labels:
                    replace_in_doc(doc, "Clasa BACS", "BACS: " + bacs_labels[bacs_class])
                    replace_in_doc(doc, "clasă automatizare", "Clasă automatizare BACS: " + bacs_class)

                # ── SRI — Smart Readiness Indicator ──────────────────────
                sri_val = data.get("sri_total", "")
                sri_grade = data.get("sri_grade", "")
                if sri_val:
                    replace_in_doc(doc, "SRI %", f"SRI = {sri_val}% (Clasa {sri_grade})")

                # ── Detalii ACM ───────────────────────────────────────────
                acm_vol = data.get("acm_storage_volume", "")
                acm_src_label = {
                    "ct_prop": "Centrală termică proprie",
                    "boiler_electric": "Boiler electric",
                    "termoficare": "Termoficare",
                    "solar_termic": "Solar termic",
                    "pc": "Pompă de căldură",
                }.get(data.get("acm_source", ""), "")
                if acm_vol:
                    replace_in_doc(doc, "Volumul vasului de acumulare", "Volum vas ACM: " + acm_vol + " L")
                if acm_src_label:
                    replace_in_doc(doc, "sursa de preparare ACM", "Sursă ACM: " + acm_src_label)

                # ── Infiltrații / Etanșeitate ─────────────────────────────
                n50 = data.get("n50", "")
                if n50 and float(n50 or 0) > 0:
                    replace_in_doc(doc, "n50 =", "n50 = " + format_ro(float(n50), 1) + " h⁻¹")
                    replace_in_doc(doc, "rata de infiltrații", "Rata infiltrații n50 = " + format_ro(float(n50), 1) + " h⁻¹")

                # ═══════════════════════════════════════════════════════
                # Etapa 7 (20 apr 2026) — completare Anexa 2 detaliată
                # Gap-uri identificate în audit: EER, putere frigorifică, ventilare HR,
                # putere iluminat, energie regenerabilă exportată, nr. apartamente.
                # Folosesc replacement chirurgical pe paragrafe identificate cu eticheta
                # semantică (evit bug-ul "201520152015" — fără replace_in_doc global pe puncte).
                # ═══════════════════════════════════════════════════════

                def _fill_para_blank(label_match, value, suffix=""):
                    """Inserează valoarea în paragraful identificat — INTELIGENT.

                    FIX 21 apr 2026: Vechea versiune adăuga la sfârșit ducând la
                    bug-uri "____kW 600 kW" (dublarea unității). Noua versiune:
                    1. Caut paragraful cu label_match
                    2. Idempotent: skip dacă valoarea e deja prezentă
                    3. Detectez placeholder-ul de spații/non-breaking-space DUPĂ label
                       (pattern: ":  \xa0+ \\s+ \xa0+" sau ":  \\s{3,}")
                    4. Înlocuiesc placeholder-ul cu " <value> " (păstrând unitatea
                       originală care apare DUPĂ placeholder)
                    5. Fallback: dacă nu găsesc placeholder, adaug la final cu suffix
                    """
                    import re as _re
                    if not value:
                        return False
                    val_str = str(value).strip()
                    if not val_str:
                        return False
                    for p in doc.paragraphs:
                        pt = p.text
                        if label_match not in pt:
                            continue
                        # Idempotent: skip dacă valoarea e deja
                        if val_str in pt:
                            return True
                        # Detectez placeholder de spații între label și final/unit
                        # Pattern 1: '\xa0+\s+\xa0+' (non-breaking-space block)
                        # Pattern 2: ': \\s{3,}' (3+ spații după ':')
                        replaced = False
                        for run in p.runs:
                            rt = run.text
                            if not rt:
                                continue
                            # Caut pattern de spații/nbsp consecutive (>=3 chars whitespace)
                            m = _re.search(r"[\xa0\s]{3,}", rt)
                            if m:
                                # Înlocuiesc placeholder-ul cu valoarea (păstrez 1 spațiu pre/post)
                                new_text = rt[:m.start()] + " " + val_str + " " + rt[m.end():]
                                run.text = new_text
                                replaced = True
                                break
                        if replaced:
                            return True
                        # Fallback: adaug la final dacă nu am găsit placeholder
                        if p.runs:
                            run = p.add_run(f" {val_str}{suffix}")
                            run.font.size = p.runs[0].font.size
                            run.font.name = p.runs[0].font.name
                        else:
                            p.add_run(f" {val_str}{suffix}")
                        return True
                    return False

                # ── Răcire/climatizare ─────────────────────────────────
                cooling_eer = data.get("cooling_eer", "")
                if cooling_eer:
                    _fill_para_blank("coeficientului de performanţă EER", f"EER = {cooling_eer}")
                cooling_seer = data.get("cooling_seer", "")
                if cooling_seer:
                    # Adaug SEER lângă EER în același paragraf (informativ, EN 14825)
                    for p in doc.paragraphs:
                        if "EER al sursei de răcire" in p.text and f"SEER = {cooling_seer}" not in p.text:
                            p.add_run(f"   |   SEER = {cooling_seer}")
                            break
                cooling_power = data.get("cooling_power_kw", "")
                if cooling_power:
                    # Para 336 — Necesarul de frig pentru răcire (putere frigorifică)
                    _fill_para_blank("Necesarul de frig pentru răcire", cooling_power, " kW")
                    # Para 338 — Puterea frigorifică totală instalată
                    _fill_para_blank("Puterea frigorifică totală instalată", cooling_power, " kW")
                cooled_area = data.get("cooled_area_m2", "")
                if cooled_area:
                    _fill_para_blank("Volumul de referință al zonei climatizate", cooled_area, " m²")

                # ── Ventilare ──────────────────────────────────────────
                # NU folosi _fill_para_blank pentru HR — checkbox-ul e gestionat
                # prin VENT_HR_YES/VENT_HR_NO în semantic matching.
                vent_hr_eff = data.get("ventilation_hr_efficiency_pct", "")
                if vent_hr_eff and vent_hr_eff not in ("0", "0,0", "0.0"):
                    _fill_para_blank("Eficiență declarată pe durata verii/iernii", f"{vent_hr_eff}% / {vent_hr_eff}%")
                elif not vent_hr_eff or vent_hr_eff in ("0", "0,0", "0.0"):
                    replace_in_doc(doc, "0% / 0%", "—")
                vent_label = data.get("ventilation_type_label", "")
                _vent_type_local = data.get("ventilation_type", "")
                _is_mech_vent = bool(_vent_type_local) and _vent_type_local not in (
                    "natural_neorg", "natural_org", "natural", ""
                )
                if vent_label and _is_mech_vent:
                    # Para 354 — etichetă tip ventilare mecanică (ex. "cu recuperare")
                    # Injectăm NUMAI pentru ventilare mecanică, nu pentru natural!
                    for p in doc.paragraphs:
                        if "Alt tip:" in p.text and "Cu 2 circuite" in p.text and vent_label not in p.text:
                            p.add_run(f" {vent_label}")
                            break

                # ── Iluminat ───────────────────────────────────────────
                light_power = data.get("lighting_power_kw", "")
                if light_power:
                    _fill_para_blank("Puterea electrică totală necesară a sistemului de iluminat", light_power, " kW")
                    _fill_para_blank("Puterea electrică instalată totală a sistemului de iluminat", light_power, " kW")

                # ── Regenerabile (energie exportată on-site) ───────────
                solar_th = data.get("solar_th_kwh_year", "")
                if solar_th and solar_th != "0":
                    _fill_para_blank("Energia termică exportată", solar_th, " kWh/an")
                pv_kwh = data.get("pv_kwh_year", "")
                if pv_kwh and pv_kwh != "0":
                    _fill_para_blank("Energia electrică exportată", pv_kwh, " kWh/an")
                wind_kwh = data.get("wind_kwh_year", "")
                if wind_kwh and wind_kwh != "0":
                    # Para 414 are placeholder cu puncte — chirurgical, doar în paragraful eolian
                    for p in doc.paragraphs:
                        if "Număr centrale eoliene" in p.text and "............" in p.text:
                            replace_in_paragraph(p, "............................", f"  ({wind_kwh} kWh/an produși)", count=1)
                            break

                # ═══════════════════════════════════════════════════════
                # Sprint post-deploy fix (20 apr 2026) — detalii regenerabile
                # ═══════════════════════════════════════════════════════

                # ── Regenerabile: Fill CONTEXT-AWARE (20 apr 2026 fix) ──
                # Problema: _fill_para_blank("Număr panouri", ...) înlocuiește PRIMA
                # apariție — când doar PV e activ, valoarea ajunge sub "Solar termic"
                # (și rezultă în "Nu există ⊠" cu Număr 110 — inconsistent logic).
                # Fix: fill SECVENȚIAL, secțiune cu secțiune, prin section boundaries.
                solar_th_type = data.get("solar_th_type_label", "")
                solar_th_panels = data.get("solar_th_panels", "")
                solar_th_orient = data.get("solar_th_orientation", "")
                solar_th_usage = data.get("solar_th_usage", "")
                pv_type = data.get("pv_type_label", "")
                pv_panels = data.get("pv_panels", "")
                pv_orient = data.get("pv_orientation", "")
                pv_usage = data.get("pv_usage", "")
                mount_loc = data.get("renewable_mount_location", "")

                def _fill_in_section(section_start_text, section_end_text, label_match, value):
                    """Fill placeholder în CADRUL secțiunii delimitate (între 2 headere)."""
                    if not value:
                        return False
                    in_section = False
                    for p in doc.paragraphs:
                        pt = p.text
                        if section_start_text in pt:
                            in_section = True
                            continue
                        if in_section and section_end_text in pt:
                            return False
                        if in_section and label_match in pt:
                            val_str = str(value).strip()
                            if val_str in pt:
                                return True  # idempotent
                            for run in p.runs:
                                m = re.search(r"[\xa0\s]{3,}|\.{4,}", run.text)
                                if m:
                                    run.text = run.text[:m.start()] + " " + val_str + " " + run.text[m.end():]
                                    return True
                    return False

                # Secțiunea SOLAR TERMIC: între "panouri termosolare" → "panouri fotovoltaice"
                SEC_SOLAR_START = "Sistemul de panouri termosolare"
                SEC_PV_START = "Sistemul de panouri fotovoltaice"
                SEC_PC_START = "Pompa de căldură"
                SEC_BIO_START = "Sistemul de utilizare a biomasei"
                SEC_END_G = "Energia termică exportată"

                if solar_th_type:
                    _fill_in_section(SEC_SOLAR_START, SEC_PV_START,
                                     "Tip panou (plan, cu tuburi vidate", solar_th_type)
                if solar_th_panels:
                    _fill_in_section(SEC_SOLAR_START, SEC_PV_START,
                                     "Număr panouri", solar_th_panels)
                if solar_th_type and mount_loc:
                    _fill_in_section(SEC_SOLAR_START, SEC_PV_START,
                                     "Mod montare", mount_loc)
                if solar_th_orient:
                    _fill_in_section(SEC_SOLAR_START, SEC_PV_START,
                                     "Orientare", solar_th_orient)
                if solar_th_usage:
                    _fill_in_section(SEC_SOLAR_START, SEC_PV_START,
                                     "Utilizate pentru", solar_th_usage)

                # Secțiunea PV: între "fotovoltaice" → "Pompa de căldură"
                if pv_type:
                    _fill_in_section(SEC_PV_START, SEC_PC_START,
                                     "Tip panou (monocristalin", pv_type)
                if pv_panels:
                    _fill_in_section(SEC_PV_START, SEC_PC_START,
                                     "Număr panouri", pv_panels)
                if pv_type and mount_loc:
                    _fill_in_section(SEC_PV_START, SEC_PC_START,
                                     "Mod montare", mount_loc)
                if pv_orient:
                    _fill_in_section(SEC_PV_START, SEC_PC_START,
                                     "Orientare", pv_orient)
                if pv_usage:
                    _fill_in_section(SEC_PV_START, SEC_PC_START,
                                     "Utilizate pentru", pv_usage)

                # ── Heat pump: Număr pompe ──
                hp_count = data.get("heat_pump_count", "")
                if hp_count and hp_count != "0":
                    _fill_para_blank("Număr pompe de căldură", hp_count)

                # ── Biomass: putere + alt tip precizare ──
                bio_pow = data.get("biomass_power_kw", "")
                if bio_pow:
                    _fill_para_blank("Putere nominală cazan biomasă", bio_pow, " kW")
                bio_type_label_val = data.get("biomass_type_label", "")
                # FIX 20 apr 2026: biomass "alt tip" contamina TOATE cele 7 paragrafe
                # "alt tip, precizați" din document (încălzire, răcire, iluminat etc.)
                # Fix DUBLU (server-side cleanup + frontend blank):
                # 1. Frontend NU mai trimite biomass_type_label când !enabled
                # 2. Python: skip complet dacă biomass_enabled=false SAU tip standard
                biomass_is_enabled = str(data.get("biomass_enabled", "") or "").lower() == "true"
                if not biomass_is_enabled:
                    bio_type_label_val = ""  # force empty dacă biomass not enabled
                if bio_type_label_val and bio_type_label_val.lower() not in ("peleți", "peleti", "brichete", "lemn", "lemn tocat", ""):
                    # Caut paragraful "alt tip, precizați" DUPĂ "Tip biomasă utilizată"
                    found_bio_section = False
                    for p in doc.paragraphs:
                        pt = p.text
                        if "Tip biomas" in pt or "biomas" in pt.lower()[:50]:
                            found_bio_section = True
                            continue
                        if found_bio_section and "alt tip, preciza" in pt.lower():
                            # Verific că precedent paragraf NU e heating/cooling/lighting
                            # Înlocuiesc dots/spații cu tipul biomassei
                            for run in p.runs:
                                if re.search(r"\.{4,}|[\xa0\s]{4,}", run.text):
                                    run.text = re.sub(r"\.{4,}|[\xa0\s]{4,}", " " + bio_type_label_val + " ", run.text, count=1)
                                    break
                            break

                # ══════════════════════════════════════════════════════════
                # Sprint post-deploy fix (20 apr 2026) — regim înălțime "2 (nr)" / "5 (nr)"
                # Aceste placeholder-e apar ca paragrafe separate în template,
                # NU în celule de tabel. Înlocuim text literal cu numere reale.
                # ══════════════════════════════════════════════════════════
                regime_str_rep = (data.get("regime", "") or "").upper().strip()
                import re as _re_reg
                m_sub_rep = _re_reg.search(r"(\d+)S", regime_str_rep)
                n_subsoluri = int(m_sub_rep.group(1)) if m_sub_rep else (1 if "S" in regime_str_rep else 0)
                m_et_rep = _re_reg.search(r"(\d+)E", regime_str_rep)
                n_etaje = int(m_et_rep.group(1)) if m_et_rep else (1 if "E" in regime_str_rep else 0)

                # Înlocuim paragrafele standalone "2 (nr)" și "5 (nr)"
                # (placeholder-e în template MDLPA pentru regim înălțime)
                for p in doc.paragraphs:
                    pt_stripped = p.text.strip()
                    # Match flexibil: "2 (nr)", "2(nr)", "2\n(nr)" etc.
                    txt_normalized = " ".join(pt_stripped.split())
                    if txt_normalized in ("2 (nr)", "2(nr)") and n_subsoluri > 0:
                        for r in p.runs:
                            if "2" in r.text and "(nr)" not in r.text:
                                r.text = r.text.replace("2", str(n_subsoluri), 1)
                                break
                            elif r.text.strip() == "2":
                                r.text = str(n_subsoluri)
                                break
                    elif txt_normalized in ("5 (nr)", "5(nr)") and n_etaje > 0:
                        for r in p.runs:
                            if "5" in r.text and "(nr)" not in r.text:
                                r.text = r.text.replace("5", str(n_etaje), 1)
                                break
                            elif r.text.strip() == "5":
                                r.text = str(n_etaje)
                                break

                # ══════════════════════════════════════════════════════════
                # Sprint post-deploy fix — Structura constructivă CB (fuzzy match)
                # Template are 8 CB structură; demo-urile folosesc texte descriptive
                # detaliate ("Structură prefabricată PAFP") care nu match "Zidărie portantă"
                # ══════════════════════════════════════════════════════════
                struct_text_raw = (data.get("structure", "") or "").lower()
                struct_cb_fuzzy = None
                if "pafp" in struct_text_raw or "panouri" in struct_text_raw or "prefabr" in struct_text_raw:
                    struct_cb_fuzzy = 133  # Panouri prefabricate mari
                elif "zidăr" in struct_text_raw or "zidar" in struct_text_raw or "cărăm" in struct_text_raw:
                    struct_cb_fuzzy = 127  # Zidărie
                elif "cadre" in struct_text_raw and "beton" in struct_text_raw:
                    struct_cb_fuzzy = 129  # Cadre beton armat
                elif "stâlpi" in struct_text_raw or "grinzi" in struct_text_raw:
                    struct_cb_fuzzy = 130  # Stâlpi și grinzi
                elif "lemn" in struct_text_raw:
                    struct_cb_fuzzy = 131  # Structură lemn
                elif "metalic" in struct_text_raw or "oțel" in struct_text_raw:
                    struct_cb_fuzzy = 132  # Structură metalică
                elif "mixt" in struct_text_raw:
                    struct_cb_fuzzy = 134
                if struct_cb_fuzzy is not None and mode in ("anexa", "anexa_bloc"):
                    # Bifăm CB suplimentar prin toggle (idempotent)
                    try:
                        toggle_checkboxes(doc, [struct_cb_fuzzy])
                    except Exception:
                        pass

                # ── Nr. apartamente / unități ──────────────────────────
                n_apt = data.get("n_apartments_count", "")
                if n_apt:
                    _fill_para_blank(
                        "Numărul & tipul apartamentelor/unităților de clădire/zonelor termice",
                        f" Total: {n_apt} unități"
                    )

                # ── Aria ferestre totală (info suplimentar) ────────────
                glaz_area = data.get("glazing_area_total_m2", "")
                if glaz_area:
                    # Adaug ca info suplimentar la paragraful U tâmplărie (dacă există)
                    for p in doc.paragraphs:
                        if "U tâmplărie" in p.text and "Arie totală ferestre" not in p.text:
                            p.add_run(f"   |   Arie totală ferestre: {glaz_area} m²")
                            break

                # ── Tabel apartamente/unități/zone termice (Etapa 7b — 20 apr 2026) ──
                # User raport: tabelul "Tip apart/destinație unitate/zonă" e gol în
                # Anexa generată. Trebuie completat cu denumire + arie + nr. unități.
                # Identificare: tabelul care conține header "Tip apart/" și TOTAL pe ultimul rând.
                _cat_label_table = data.get("category_label", "") or "Unitate de clădire"
                _area_ref = data.get("area_ref", "")
                _n_apt = data.get("n_apartments_count", "1") or "1"
                if _area_ref:
                    try:
                        _area_total = float(_area_ref.replace(",", "."))
                        _n_apt_int = max(1, int(float(_n_apt or "1")))
                        _area_per_unit = _area_total / _n_apt_int
                        for tbl in doc.tables:
                            if len(tbl.rows) < 2 or len(tbl.columns) < 4:
                                continue
                            header_text = tbl.rows[0].cells[0].text + " " + tbl.rows[0].cells[1].text
                            # Identific tabelul corect: header conține "Tip apart" + "Aria de referință"
                            if "Tip apart" not in header_text and "destinație" not in header_text:
                                continue
                            if "Aria de referință" not in tbl.rows[0].cells[1].text:
                                continue
                            # Tabelul corect — verific dacă deja completat (idempotent)
                            row1 = tbl.rows[1]
                            if row1.cells[1].text.strip():
                                break  # Deja completat, skip
                            # Înlocuiesc r1c0 (placeholder multi-line) cu denumirea reală
                            cell0 = row1.cells[0]
                            cell0.text = ""
                            p0 = cell0.paragraphs[0]
                            r0 = p0.add_run(_cat_label_table)
                            r0.font.size = Pt(10)
                            # r1c1 — aria per unitate
                            row1.cells[1].text = format_ro(_area_per_unit, 1)
                            # r1c2 — număr unități
                            row1.cells[2].text = str(_n_apt_int)
                            # r1c3 — aria totală pe tip
                            row1.cells[3].text = format_ro(_area_total, 1)
                            # Setez font 10pt pe celulele noi
                            for ci in (1, 2, 3):
                                for p in row1.cells[ci].paragraphs:
                                    for r in p.runs:
                                        r.font.size = Pt(10)
                            # Rândul TOTAL (ultim rând) — completez c2 + c3 cu totaluri
                            if len(tbl.rows) >= 3:
                                row_total = tbl.rows[-1]
                                if not row_total.cells[2].text.strip():
                                    row_total.cells[2].text = str(_n_apt_int)
                                if not row_total.cells[3].text.strip():
                                    row_total.cells[3].text = format_ro(_area_total, 1)
                                for ci in (2, 3):
                                    for p in row_total.cells[ci].paragraphs:
                                        for r in p.runs:
                                            r.font.size = Pt(10)
                                            r.bold = True
                            break
                    except (ValueError, TypeError) as e_apt_tbl:
                        print(f"[apt_table] eroare: {e_apt_tbl}", flush=True)

                # ── Tabel 0 — Zone climatice + eoliene + regim înălțime ──
                # FIX 21 apr 2026: BIFEZ checkbox XML existent în celulă (NU adaug text "X"
                # care ștergea form field-ul + producea X lângă bullet gol).
                # Header structure (13 cols): zona clima [I, II, II, II, II, III, III, IV, IV, IV, V, V]
                _ZONA_CLIMA_COLS = {1: 1, 2: 2, 3: 6, 4: 8, 5: 11}
                _ZONA_EOL_COLS   = {1: 1, 2: 4, 3: 7, 4: 10}
                _REGIM_COLS      = {"S": 1, "D": 3, "Mez": 5, "P": 7, "E": 9, "M": 12, "M/P": 12}

                def _check_cell_checkbox(cell):
                    """Bifează TOATE checkbox-urile XML din celulă (set w:default w:val='1')."""
                    cbs = cell._tc.findall(".//w:checkBox", NSMAP)
                    if not cbs:
                        return False
                    for cb in cbs:
                        default = cb.find("w:default", NSMAP)
                        if default is not None:
                            default.set(qn("w:val"), "1")
                        else:
                            from docx.oxml import OxmlElement
                            d = OxmlElement("w:default")
                            d.set(qn("w:val"), "1")
                            cb.append(d)
                    return True

                try:
                    zone_num = int(data.get("climate_zone_num", "3") or "3")
                    zone_num = max(1, min(5, zone_num))
                    wind_zone = 1 if zone_num <= 2 else (2 if zone_num <= 4 else 3)
                    regime_str = (data.get("regime", "") or "").upper().strip()
                    # Parse MULTI-LETTER: pentru "S+P+4E+M" → bif S, P, E, M
                    # Extrage secvențele de litere (și "MEZ" special)
                    import re as _re_reg
                    regim_keys_found = []
                    if "MEZ" in regime_str:
                        regim_keys_found.append("Mez")
                    # Caut caractere individuale relevante
                    for letter in ["S", "D", "P", "E", "M"]:
                        # "S" să nu match-uiască "MEZ" (deja tratat)
                        pattern = r"(?<![A-Z])" + letter + r"(?![A-Z])"
                        if _re_reg.search(pattern, regime_str):
                            if letter not in regim_keys_found:
                                regim_keys_found.append(letter)
                    # Parse nr subsoluri/etaje pentru înlocuirea "2 (nr)" / "5 (nr)"
                    # Formate acceptate:
                    #   "2S+P+5E+M" → 2 subsoluri, 5 etaje explicit
                    #   "P+4" → 4 etaje (fără sufix E — format scurt)
                    #   "P+4E" → 4 etaje explicit
                    #   "S+P+4" → 1 subsol (implicit), 4 etaje
                    m_sub = _re_reg.search(r"(\d+)S", regime_str)
                    if m_sub:
                        n_subsoluri = int(m_sub.group(1))
                    elif "S" in regim_keys_found:
                        n_subsoluri = 1
                    else:
                        n_subsoluri = 0
                    # Fallback: building.basement=true → min 1 subsol
                    if n_subsoluri == 0 and str(data.get("basement", "") or "").lower() in ("true", "1", "da"):
                        n_subsoluri = 1
                    # Etaje: regex cu sau fără E (P+4 = P+4E)
                    m_et = _re_reg.search(r"(\d+)E", regime_str)
                    if not m_et:
                        # Format scurt "P+X" — X = nr etaje peste parter
                        m_et_short = _re_reg.search(r"P\+(\d+)", regime_str)
                        if m_et_short:
                            m_et = m_et_short
                    n_etaje = int(m_et.group(1)) if m_et else (1 if "E" in regim_keys_found else 0)

                    for tbl in doc.tables:
                        if len(tbl.rows) < 6 or len(tbl.columns) != 13:
                            continue
                        h0 = tbl.rows[0].cells[0].text.lower()
                        if "zona climatică" not in h0 and "zona climatica" not in h0:
                            continue
                        # Bifez checkbox-uri în celulele corecte
                        col_clima = _ZONA_CLIMA_COLS.get(zone_num, 2)
                        _check_cell_checkbox(tbl.rows[1].cells[col_clima])
                        col_eol = _ZONA_EOL_COLS.get(wind_zone, 1)
                        _check_cell_checkbox(tbl.rows[3].cells[col_eol])
                        # Regim înălțime — bif MULTIPLE niveluri (nu doar primul)
                        for rk in regim_keys_found:
                            col_reg = _REGIM_COLS.get(rk)
                            if col_reg is not None and col_reg < len(tbl.rows[5].cells):
                                _check_cell_checkbox(tbl.rows[5].cells[col_reg])
                        # Înlocuire "2" și "5" literal din rândul regim cu numere reale
                        # Template are celule cu "2 (nr)" pentru S, "5 (nr)" pentru E
                        # AGRESIV: iterez TOATE paragrafele din TOATE celulele rândului 5
                        # și înlocuiesc "2" cu n_subsoluri, "5" cu n_etaje
                        regim_row_cells = tbl.rows[5].cells
                        for ci, cell in enumerate(regim_row_cells):
                            for para in cell.paragraphs:
                                para_text = para.text.strip()
                                # Cazurile posibile: "2", "2 (nr)", "2\n(nr)", "2(nr)"
                                # Normalizez whitespace
                                norm = " ".join(para_text.split())
                                if n_subsoluri > 0 and norm in ("2", "2 (nr)", "2(nr)"):
                                    # Înlocuiesc primul run care conține "2" cu n_subsoluri
                                    for r in para.runs:
                                        if r.text.strip() == "2":
                                            r.text = str(n_subsoluri)
                                            break
                                        elif "2" in r.text:
                                            r.text = r.text.replace("2", str(n_subsoluri), 1)
                                            break
                                elif n_etaje > 0 and norm in ("5", "5 (nr)", "5(nr)"):
                                    for r in para.runs:
                                        if r.text.strip() == "5":
                                            r.text = str(n_etaje)
                                            break
                                        elif "5" in r.text:
                                            r.text = r.text.replace("5", str(n_etaje), 1)
                                            break
                        break
                except Exception as e_t0:
                    print(f"[tabel_0_zone] eroare: {e_t0}", flush=True)

                # ── FALLBACK DOC-LEVEL: Regim "2 (nr)" / "5 (nr)" ──
                # Table-level fix poate eșua dacă Tabel 0 are structură diferită.
                # Fallback: iter TOATE paragrafele doc-level + paragrafele din TOATE
                # celulele tabelelor, match exact (normalized) → înlocuire.
                try:
                    # Colectare tuturor paragrafelor (doc + tabel cells)
                    all_paras = list(doc.paragraphs)
                    for tbl in doc.tables:
                        for row in tbl.rows:
                            for cell in row.cells:
                                all_paras.extend(cell.paragraphs)
                    for para in all_paras:
                        pt = para.text.strip()
                        norm = " ".join(pt.split())
                        if n_subsoluri > 0 and norm in ("2 (nr)", "2(nr)"):
                            for r in para.runs:
                                if r.text.strip() == "2":
                                    r.text = str(n_subsoluri)
                                    break
                                elif "2" in r.text:
                                    r.text = r.text.replace("2", str(n_subsoluri), 1)
                                    break
                        elif n_etaje > 0 and norm in ("5 (nr)", "5(nr)"):
                            for r in para.runs:
                                if r.text.strip() == "5":
                                    r.text = str(n_etaje)
                                    break
                                elif "5" in r.text:
                                    r.text = r.text.replace("5", str(n_etaje), 1)
                                    break
                except Exception as e_regim_fb:
                    print(f"[regim_fallback] eroare: {e_regim_fb}", flush=True)

                # ── Tabel 5 — Apartamente debranșate condominiu ──
                # Pentru clădiri RC/RA (bloc) → "Nu există" default; pentru altele skip.
                try:
                    if category in ("RC", "RA"):
                        for tbl in doc.tables:
                            if len(tbl.rows) != 2 or len(tbl.columns) != 2:
                                continue
                            h0 = tbl.rows[0].cells[0].text.lower()
                            if "apartamente debrans" in h0 or "apartamente debranș" in h0 or "debran" in h0:
                                # Default: "Nu există" — auditorul modifică dacă e cazul
                                _check_cell_checkbox(tbl.rows[1].cells[1])
                                break
                except Exception:
                    pass

                # ── Tabel 2 — Anvelopă (R-values + arii) ──
                # Template apartament: 3 rânduri, celula r1c0 are 8 paragrafe (PE 1, PE 2,
                # FE, UE, TE, Sb, CS, ...) iar c1/c2/c3 au un singur paragraf gol.
                # Fill: adăugăm paragrafe în c1/c2/c3 aliniate cu eticheta din c0.
                try:
                    opaque_u_anv = json.loads(data.get("opaque_u_values", "[]"))
                    glaz_u_anv = float((data.get("glazing_max_u", "0") or "0").replace(",", "."))
                    glaz_area_anv = float((data.get("glazing_area_total_m2", "0") or "0").replace(",", "."))
                    if opaque_u_anv or glaz_area_anv > 0:
                        R_NORMAT = {
                            "PE": 1.80, "PT": 5.00, "PP": 5.00, "PB": 4.50, "PL": 4.50,
                            "FE": 0.77, "UE": 0.77, "SE": 4.50, "CS": 1.80, "Sb": 4.50,
                        }
                        # Mapare prefix etichetă template → cod tip element
                        LABEL_PREFIX_MAP = [
                            ("PE",  ["PE"]),
                            ("FE",  ["FE"]),
                            ("UE",  ["UE"]),
                            ("TE",  ["PT", "PP"]),   # Terasă → PT sau PP
                            ("Sb",  ["Sb", "PB"]),   # Planșeu subsol → Sb sau PB
                            ("CS",  ["CS"]),
                        ]
                        def _label_to_types(label):
                            for prefix, codes in LABEL_PREFIX_MAP:
                                if label.strip().startswith(prefix):
                                    return codes
                            return []

                        for tbl in doc.tables:
                            if len(tbl.rows) != 3 or len(tbl.columns) != 4:
                                continue
                            h0 = tbl.rows[0].cells[0].text.lower()
                            if "tip element" not in h0:
                                continue
                            # Idempotență: dacă c1 are deja mai mult de 1 paragraf cu conținut, skip
                            if any(p.text.strip() for p in tbl.rows[1].cells[1].paragraphs):
                                break

                            paras_c0 = tbl.rows[1].cells[0].paragraphs
                            cells_c = [tbl.rows[1].cells[ci] for ci in range(1, 4)]

                            # Construiesc index: tip → lista de elemente (pentru multi-PE)
                            from collections import defaultdict
                            type_els = defaultdict(list)
                            for el in opaque_u_anv:
                                t = el.get("type", "")
                                area = float(el.get("area", 0) or 0)
                                u_val = float(el.get("u", 0) or 0)
                                if t and area > 0:
                                    type_els[t].append({"area": area, "u": u_val})
                            # Vitraj din glazing_area_total_m2
                            if glaz_area_anv > 0:
                                type_els["FE"].append({"area": glaz_area_anv, "u": glaz_u_anv})

                            type_seen = defaultdict(int)
                            entries = []  # (r_calc_str, r_norm_str, area_str) per paragraf c0
                            for p0 in paras_c0:
                                label = p0.text.strip()
                                if not label or label == "...":
                                    entries.append(("", "", ""))
                                    continue
                                matched_codes = _label_to_types(label)
                                found = None
                                for code in matched_codes:
                                    occ = type_seen[code]
                                    els = type_els.get(code, [])
                                    if occ < len(els):
                                        found = (code, els[occ])
                                        type_seen[code] += 1
                                        break
                                if found:
                                    code, el = found
                                    u = el["u"]
                                    area = el["area"]
                                    r_calc = 1 / u if u > 0 else 0
                                    r_norm = R_NORMAT.get(code, 0)
                                    # #9 (audit Pas 6+7 — V6, 7 mai 2026) — indicator vizual
                                    # conformitate R termică vs Mc 001-2022 Tab 2.4-2.10
                                    # (R_calc >= R_norm = conform; R_calc < R_norm = neconform).
                                    # Sufix " ✓" / " ✗" pe valoarea R_calc — ajută auditorul să
                                    # identifice rapid elementele neconforme la prima vedere.
                                    # Fallback "—" dacă lipsesc date pentru comparație.
                                    if r_calc > 0 and r_norm > 0:
                                        indicator = " ✓" if r_calc >= r_norm else " ✗"
                                        r_calc_str = format_ro(r_calc, 2) + indicator
                                    elif r_calc > 0:
                                        r_calc_str = format_ro(r_calc, 2)
                                    else:
                                        r_calc_str = "—"
                                    entries.append((
                                        r_calc_str,
                                        format_ro(r_norm, 2) if r_norm > 0 else "—",
                                        format_ro(area, 1),
                                    ))
                                else:
                                    # Audit 2 mai 2026 — element cu label în template dar
                                    # NEEXISTENT în clădirea reală (ex. UE/TE/CS pentru
                                    # apartament fără ușă exterioară directă, fără terasă,
                                    # fără casa scării). Afișez „—" pentru claritate vizuală
                                    # în loc să las celulele goale.
                                    entries.append(("—", "—", "—"))

                            # Umple c1/c2/c3: primul paragraf existent + adaug restul
                            for col_offset, val_idx in enumerate([0, 1, 2]):
                                cell = cells_c[col_offset]
                                for i, entry in enumerate(entries):
                                    val = entry[val_idx]
                                    if i == 0:
                                        para = cell.paragraphs[0]
                                        run = para.add_run(val)
                                        run.font.size = Pt(9)
                                    else:
                                        para = cell.add_paragraph()
                                        run = para.add_run(val)
                                        run.font.size = Pt(9)

                            # TOTAL arie în r2c3 (celula nemergică din rândul TOTAL)
                            total_se = sum(
                                el["area"] for els in type_els.values() for el in els
                                if el["area"] > 0
                            )
                            r2_c3 = tbl.rows[2].cells[3]
                            run_tot = r2_c3.paragraphs[0].add_run(format_ro(total_se, 1))
                            run_tot.font.size = Pt(10)
                            run_tot.bold = True
                            break
                except Exception as e_t2:
                    print(f"[tabel_2_anvelopa] eroare: {e_t2}", flush=True)

                # ── Tabel 3 — Consum specific 5 sisteme + clase + TOTAL ──
                # Template apartament: 8 rânduri, 5 coloane, header r0c2="Apartament"
                # (Vechea detecție era pentru template clădire: 8 col + "Clădirea reală" — GREȘIT)
                try:
                    sistem_data = {
                        "Încălzire":         data.get("ep_incalzire", ""),
                        "Apă caldă":         data.get("ep_acm", ""),
                        "Răcire":            data.get("ep_racire", ""),
                        "Ventilare":         data.get("ep_ventilare", ""),
                        "Iluminat":          data.get("ep_iluminat", ""),
                    }
                    ep_total_val = float((data.get("ep_specific", "0") or "0").replace(",", "."))
                    co2_total_val = float((data.get("co2_val", "0") or "0").replace(",", "."))
                    co2_ratio = co2_total_val / ep_total_val if ep_total_val > 0 else 0
                    try:
                        ep_scale = [float((data.get(k, "0") or "0").replace(",", "."))
                                    for k in ["s_ap", "s_a", "s_b", "s_c", "s_d", "s_e", "s_f"]]
                    except (ValueError, TypeError):
                        ep_scale = []
                    def _class_from_ep(ep_val):
                        if not ep_scale or all(v == 0 for v in ep_scale):
                            return ""
                        classes = ["A+", "A", "B", "C", "D", "E", "F"]
                        for i, threshold in enumerate(ep_scale):
                            if threshold > 0 and ep_val < threshold:
                                return classes[i]
                        return "G"

                    # CR-2 (7 mai 2026) — clase per utilitate explicite (Mc 001-2022 Tab I.1)
                    # primite din JS (cls_incalzire/acm/racire/ventilare/iluminat) au prioritate
                    # față de _class_from_ep care folosea WHOLE-BUILDING ep_scale.
                    _CLS_EXPLICIT_T3 = {
                        "Încălzire": (data.get("cls_incalzire") or "").strip(),
                        "Apă caldă": (data.get("cls_acm") or "").strip(),
                        "Răcire":    (data.get("cls_racire") or "").strip(),
                        "Ventilare": (data.get("cls_ventilare") or "").strip(),
                        "Iluminat":  (data.get("cls_iluminat") or "").strip(),
                    }
                    _VALID_CLASSES_T3 = {"A+", "A", "B", "C", "D", "E", "F", "G"}

                    # Suportă ambele template-uri: apartament (5 col) și clădire (8 col)
                    TABEL3_VARIANTS = [
                        # (min_cols, max_cols, h_keyword, ep_col, co2_col, cls_col, tot_row)
                        (5, 5, "Apartament",    2, 3, 4, 7),
                        (8, 8, "Clădirea reală", 2, 3, 4, 7),
                    ]
                    sistem_to_row = {
                        "Încălzire": 2, "Apă caldă": 3, "Răcire": 4,
                        "Ventilare": 5, "Iluminat": 6,
                    }
                    for tbl in doc.tables:
                        ncols = len(tbl.columns)
                        nrows = len(tbl.rows)
                        if nrows < 8:
                            continue
                        matched_variant = None
                        for (mn, mx, kw, ec, cc, kc, tr) in TABEL3_VARIANTS:
                            if mn <= ncols <= mx and kw in tbl.rows[0].cells[2].text:
                                matched_variant = (ec, cc, kc, tr)
                                break
                        if not matched_variant:
                            continue
                        ep_col, co2_col, cls_col, tot_row = matched_variant
                        # Idempotency
                        if tbl.rows[2].cells[ep_col].text.strip():
                            break
                        for nume, row_idx in sistem_to_row.items():
                            ep_str = sistem_data.get(nume, "")
                            if not ep_str:
                                continue
                            try:
                                ep_val = float(ep_str.replace(",", "."))
                            except (ValueError, TypeError):
                                continue
                            row = tbl.rows[row_idx]
                            row.cells[ep_col].text  = format_ro(ep_val, 1)
                            row.cells[co2_col].text = format_ro(ep_val * co2_ratio, 1) if co2_ratio > 0 else "—"
                            # CR-2 — prioritate clasă explicită din JS, fallback la _class_from_ep
                            explicit_cls_t3 = _CLS_EXPLICIT_T3.get(nume, "")
                            cls = explicit_cls_t3 if explicit_cls_t3 in _VALID_CLASSES_T3 else _class_from_ep(ep_val)
                            if cls:
                                row.cells[cls_col].text = cls
                            for ci in (ep_col, co2_col, cls_col):
                                for p in row.cells[ci].paragraphs:
                                    for r in p.runs:
                                        r.font.size = Pt(10)
                        # TOTAL
                        if ep_total_val > 0:
                            tbl.rows[tot_row].cells[ep_col].text  = format_ro(ep_total_val, 1)
                            tbl.rows[tot_row].cells[co2_col].text = format_ro(co2_total_val, 1)
                            tbl.rows[tot_row].cells[cls_col].text = data.get("energy_class", "") or ""
                            for ci in (ep_col, co2_col, cls_col):
                                for p in tbl.rows[tot_row].cells[ci].paragraphs:
                                    for r in p.runs:
                                        r.font.size = Pt(10)
                                        r.bold = True
                        break
                except Exception as e_t3:
                    print(f"[tabel_3_sisteme] eroare: {e_t3}", flush=True)

                # ── Layout: elimină paragrafe goale excesive înainte de titluri de secțiune ──
                # Template-ul Anexa apartament are 4+ paragrafe goale înainte de "DATE TEHNICE",
                # creând spațiu alb mare. Reducem la maxim 1 paragraf gol consecutiv.
                # IMPORTANT: iterăm NUMAI copiii DIRECȚI ai body (nu interiorul tabelelor)
                # — altfel ștergem rânduri din tabele.
                try:
                    WNS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
                    _body = doc.element.body
                    # Colectăm numai <w:p> copii direcți ai body (nu din <w:tbl>)
                    _body_paras = [ch for ch in _body if ch.tag == f"{WNS}p"]
                    i = 0
                    while i < len(_body_paras):
                        p_el = _body_paras[i]
                        txt = "".join(t.text or "" for t in p_el.iter(f"{WNS}t")).strip()
                        if txt:
                            i += 1
                            continue
                        # Paragraf gol — colectăm consecutivele
                        seq_empty = [p_el]
                        j = i + 1
                        while j < len(_body_paras):
                            nxt = _body_paras[j]
                            nxt_txt = "".join(t.text or "" for t in nxt.iter(f"{WNS}t")).strip()
                            if nxt_txt:
                                break
                            seq_empty.append(nxt)
                            j += 1
                        # Păstrăm maxim 1, ștergem restul
                        for extra in seq_empty[1:]:
                            _body.remove(extra)
                        i = j
                except Exception as e_layout:
                    print(f"[layout_empty_paras] eroare: {e_layout}", flush=True)

                # ── Cleanup: șterge punctele rămase după auto-fill în câmpuri "precizați" ──
                # Regulă: dacă un paragraf cu "precizați" conține text real (valoare) ȘI
                # mai are și "....." → scoatem punctele. Dacă e necompletat → le păstrăm.
                try:
                    for para in _iter_all_paragraphs(doc, include_txbx=True):
                        low = para.text.lower()
                        if "preciza" not in low:
                            continue
                        # Verifică dacă există text real după "precizați" (nu doar dots/spații)
                        after_match = re.split(r"preciza[tț][ie]\s*", para.text, flags=re.IGNORECASE)
                        if len(after_match) < 2:
                            continue
                        after_label = after_match[-1]
                        real_text = re.sub(r"[\.\s\xa0 ]", "", after_label)
                        if not real_text:
                            continue  # Câmp necompletat — păstrăm dots
                        # Are dots rămase → le ștergem
                        for run in para.runs:
                            if re.search(r"\.{4,}", run.text):
                                run.text = re.sub(r"\.{4,}", "", run.text).rstrip()
                except Exception as e_dots:
                    print(f"[cleanup_precizati_dots] eroare: {e_dots}", flush=True)

                # ══════════════════════════════════════════════════════════
                # Sprint monolith (20 apr 2026) — Populare TABELE RĂMASE
                # T5 Corpuri statice (din heating_radiators JSON)
                # T7 Spații neîncălzite (din building_unheated_spaces JSON)
                # T8 Gradul ocupare încălzire (defaults rezidențial/nerezidențial)
                # T11 Obiecte sanitare (din acm_fixtures JSON, fallback defaults)
                # T12 Gradul ocupare răcire (defaults)
                # ══════════════════════════════════════════════════════════

                # ── Tabel 5 — Corpuri statice (Tip | Număr | Putere termică) ──
                # Populare AGRESIVĂ: force defaults pentru orice clădire cu încălzire funcțională
                try:
                    radiators_json = data.get("heating_radiators", "[]")
                    radiators = json.loads(radiators_json) if radiators_json else []
                    # Fallback FORCE: dacă lista e goală dar avem încălzire → 1 rând default
                    if not radiators:
                        try:
                            pw_str = str(data.get("heating_power", "") or "0").replace(",", ".").strip()
                            pw = float(pw_str) if pw_str else 0.0
                            # Chiar dacă pw=0, tot populăm cu placeholder pentru a nu lăsa tabelul gol
                            radiator_type_guess = data.get("heating_radiator_type", "") or "Radiator oțel"
                            # Heuristic: 1.5 kW/radiator; fallback 1 radiator când pw=0
                            n_rads = int(max(1, pw / 1.5)) if pw > 0 else 1
                            radiators = [{
                                "type": radiator_type_guess,
                                "count_private": n_rads,
                                "count_common": 0,
                                "power_kw": format_ro(pw, 1) if pw > 0 else "—",
                            }]
                        except (ValueError, TypeError):
                            radiators = [{
                                "type": "Radiator oțel",
                                "count_private": 1,
                                "count_common": 0,
                                "power_kw": "—",
                            }]
                    if radiators:
                        for tbl in doc.tables:
                            # Header complex: 4 rânduri, uneori 9 coloane (subheadere merged)
                            # Detectare pe label "Tip corp static" oriunde în primele 2 rânduri
                            if len(tbl.rows) < 3:
                                continue
                            h0 = " ".join(tbl.rows[0].cells[0].text.split())
                            h1 = " ".join(tbl.rows[1].cells[0].text.split()) if len(tbl.rows) > 1 else ""
                            if "Tip corp static" not in h0 and "Tip corp static" not in h1:
                                continue
                            # Identificare rând data + rând TOTAL
                            data_row_idx = None
                            total_row_idx = None
                            for ri, row in enumerate(tbl.rows):
                                txt0 = row.cells[0].text.strip()
                                if txt0 == "..." or txt0 == "…":
                                    data_row_idx = ri
                                elif "TOTAL" in txt0.upper():
                                    total_row_idx = ri
                            if data_row_idx is None:
                                # Fallback: asumăm r2 = data, r-1 = TOTAL
                                data_row_idx = 2
                                total_row_idx = len(tbl.rows) - 1
                            # Idempotent: skip dacă deja populat cu altceva decât "..."
                            data_row = tbl.rows[data_row_idx]
                            if data_row.cells[0].text.strip() and data_row.cells[0].text.strip() not in ("...", "…"):
                                break
                            # Identificare coloane cheie:
                            # [0]=Tip corp, [1]=Zona, [2]=spațiu locuit, [3]=spațiu comun, [4+]=Putere
                            n_cells = len(data_row.cells)
                            if n_cells < 4:
                                break
                            rad = radiators[0]
                            # Col 0 = tip (clear "..." first)
                            data_row.cells[0].text = str(rad.get("type", "Radiator"))
                            # Col 1 = zona (opțional; "Zona 1" default)
                            if n_cells >= 5:
                                data_row.cells[1].text = "Zona 1"
                            # Col 2 = spațiu locuit
                            cp = int(rad.get("count_private", 0) or 0)
                            cc = int(rad.get("count_common", 0) or 0)
                            locuit_col = 2 if n_cells >= 5 else 1
                            common_col = 3 if n_cells >= 5 else 2
                            power_col = n_cells - 1  # ultima coloană = putere
                            data_row.cells[locuit_col].text = str(cp) if cp else "—"
                            data_row.cells[common_col].text = str(cc) if cc else "—"
                            data_row.cells[power_col].text = str(rad.get("power_kw", "—"))
                            for ci in range(n_cells):
                                for p in data_row.cells[ci].paragraphs:
                                    for r in p.runs:
                                        r.font.size = Pt(9)
                            # Rând TOTAL
                            if total_row_idx is not None:
                                trow = tbl.rows[total_row_idx]
                                if locuit_col < len(trow.cells):
                                    trow.cells[locuit_col].text = str(cp) if cp else "—"
                                if common_col < len(trow.cells):
                                    trow.cells[common_col].text = str(cc) if cc else "—"
                                # Putere totală în ultima coloană
                                try:
                                    pw_val = float(str(rad.get("power_kw", "0") or "0").replace(",", "."))
                                    if pw_val > 0 and power_col < len(trow.cells):
                                        trow.cells[power_col].text = format_ro(pw_val, 1)
                                except (ValueError, TypeError):
                                    pass
                                for ci in range(len(trow.cells)):
                                    for p in trow.cells[ci].paragraphs:
                                        for r in p.runs:
                                            r.font.size = Pt(10)
                                            r.bold = True
                            break
                except Exception as e_t5:
                    print(f"[tabel_5_corpuri_statice] eroare: {e_t5}", flush=True)

                # ── Tabel 7 — Spații neîncălzite (Cod | Diametru tronson | Lungime tronson) ──
                try:
                    unheated_json = data.get("building_unheated_spaces", "[]")
                    unheated = json.loads(unheated_json) if unheated_json else []
                    if unheated:
                        for tbl in doc.tables:
                            if len(tbl.rows) != 3:
                                continue
                            # Header r0c0 conține "Codul spațiului neîncălzit"
                            h0 = tbl.rows[0].cells[0].text
                            if "Codul spațiului" not in h0 and "Codul spa" not in h0:
                                continue
                            # Idempotent: skip dacă deja populat (r0c1 ≠ "ZU1")
                            if tbl.rows[0].cells[1].text.strip() and tbl.rows[0].cells[1].text.strip() not in ("ZU1", "ZU2", "..."):
                                # deja populat cu date reale
                                pass
                            n_cols = len(tbl.rows[0].cells) - 1  # exclud coloana label
                            for i, sp in enumerate(unheated[:n_cols]):
                                col_idx = 1 + i
                                tbl.rows[0].cells[col_idx].text = str(sp.get("code", f"ZU{i+1}"))
                                tbl.rows[1].cells[col_idx].text = str(sp.get("diameter_mm", "—"))
                                tbl.rows[2].cells[col_idx].text = str(sp.get("length_m", "—"))
                            for ri in range(3):
                                for ci in range(1, 1 + n_cols):
                                    for p in tbl.rows[ri].cells[ci].paragraphs:
                                        for r in p.runs:
                                            r.font.size = Pt(9)
                            break
                except Exception as e_t7:
                    print(f"[tabel_7_spatii_neincalzite] eroare: {e_t7}", flush=True)

                # ── Tabel 8 + T12 — Grad ocupare (încălzire + răcire) ──
                # Defaults rezidențial / nerezidențial
                try:
                    is_res_t8 = category in ("RI", "RC", "RA", "BC")
                    # Defaults EN 16798-1 + Mc 001-2022
                    if is_res_t8:
                        sched_defaults = {
                            "Programul (h)":         {"zi_lucru": "16", "noaptea": "24", "weekend": "24"},
                            "Programul [h]":         {"zi_lucru": "16", "noaptea": "24", "weekend": "24"},
                            "Temperatura interioară (grdC)": {"zi_lucru": "20", "noaptea": "18", "weekend": "20"},
                            "Temperatura interioară [grdC]": {"zi_lucru": "20", "noaptea": "18", "weekend": "20"},
                            "Grad de ocupare zilnic/săptămânal/lunar [m²/pers]": {"zi_lucru": "30", "noaptea": "30", "weekend": "30"},
                        }
                    else:
                        sched_defaults = {
                            "Programul (h)":         {"zi_lucru": "10", "noaptea": "0", "weekend": "0"},
                            "Programul [h]":         {"zi_lucru": "10", "noaptea": "0", "weekend": "0"},
                            "Temperatura interioară (grdC)": {"zi_lucru": "20", "noaptea": "15", "weekend": "15"},
                            "Temperatura interioară [grdC]": {"zi_lucru": "26", "noaptea": "28", "weekend": "28"},
                            "Grad de ocupare zilnic/săptămânal/lunar [m²/pers]": {"zi_lucru": "15", "noaptea": "15", "weekend": "15"},
                        }
                    for tbl in doc.tables:
                        if len(tbl.rows) < 2 or len(tbl.columns) < 4:
                            continue
                        h0 = tbl.rows[0].cells[0].text.strip().lower()
                        if h0 != "zona":
                            continue
                        # Check header contains "Zi de lucru" and "Noaptea"
                        header_concat = " ".join(c.text for c in tbl.rows[0].cells).lower()
                        if "zi de lucru" not in header_concat or "noaptea" not in header_concat:
                            continue
                        # Idempotent: skip dacă celula (1,1) e deja populată
                        if tbl.rows[1].cells[1].text.strip():
                            continue
                        # Pentru fiecare rând de date (după header r0)
                        for ri in range(1, len(tbl.rows)):
                            row = tbl.rows[ri]
                            label = " ".join(row.cells[0].text.split()).strip()
                            # Match label la defaults
                            sd = None
                            for k, v in sched_defaults.items():
                                if k in label or label in k:
                                    sd = v
                                    break
                            if not sd:
                                continue
                            # c1=Zi de lucru, c2=Noaptea, c3=Zi de weekend (și c4= "...." ignore)
                            if len(row.cells) > 1 and not row.cells[1].text.strip():
                                row.cells[1].text = sd["zi_lucru"]
                            if len(row.cells) > 2 and not row.cells[2].text.strip():
                                row.cells[2].text = sd["noaptea"]
                            if len(row.cells) > 3 and not row.cells[3].text.strip():
                                row.cells[3].text = sd["weekend"]
                            for ci in (1, 2, 3):
                                if ci < len(row.cells):
                                    for p in row.cells[ci].paragraphs:
                                        for r in p.runs:
                                            r.font.size = Pt(9)
                except Exception as e_t8:
                    print(f"[tabel_8_12_grad_ocupare] eroare: {e_t8}", flush=True)

                # ── Tabel 11 — Obiecte sanitare ACM (Lavoare, Cadă, Spălătoare, etc.) ──
                # Populare AGRESIVĂ pentru rezidențial: chiar dacă fixtures[] are toate "0",
                # tot aplicăm defaults bazate pe n_apartments
                try:
                    fixtures_json = data.get("acm_fixtures", "{}")
                    try:
                        fixtures = json.loads(fixtures_json) if fixtures_json else {}
                    except (ValueError, TypeError):
                        fixtures = {}
                    # Detectare "efectiv gol" — toate valorile goale/0
                    fixtures_empty = not fixtures or all(
                        not v or str(v).strip() in ("", "0") for v in fixtures.values()
                    )
                    is_res_t11 = category in ("RI", "RC", "RA", "BC")
                    if fixtures_empty and is_res_t11:
                        n_apt_int = max(1, int(str(data.get("n_apartments_count", "1") or "1").split(".")[0]))
                        fixtures = {
                            "lavoare": str(n_apt_int),
                            "cada_baie": str(n_apt_int),
                            "spalatoare": str(n_apt_int),
                            "rezervor_wc": str(n_apt_int),
                            "bideuri": "0",
                            "pisoare": "0",
                            "dus": str(n_apt_int),
                            "masina_spalat_vase": "0",
                            "masina_spalat_rufe": str(n_apt_int),
                        }
                    elif fixtures_empty and not is_res_t11:
                        # Nerezidențial minimal: 1 lavoar + 1 WC (evită tabel complet gol)
                        fixtures = {
                            "lavoare": "1",
                            "cada_baie": "0",
                            "spalatoare": "0",
                            "rezervor_wc": "1",
                            "bideuri": "0",
                            "pisoare": "0",
                            "dus": "0",
                            "masina_spalat_vase": "0",
                            "masina_spalat_rufe": "0",
                        }
                    if fixtures:
                        # Map între cheia semantică și label-ul din template
                        # Fix 20 apr 2026: normalizare diacritice pentru match robust
                        def _norm_label(s):
                            """Normalizare pentru match: jos + fara diacritice + strip"""
                            return (s or "").lower().replace("ș", "s").replace("ş", "s") \
                                .replace("ț", "t").replace("ţ", "t") \
                                .replace("ă", "a").replace("â", "a").replace("î", "i") \
                                .strip()
                        fix_key_map_norm = {
                            "lavoare": "lavoare",
                            "cada de baie": "cada_baie",
                            "spalatoare": "spalatoare",
                            "rezervor wc": "rezervor_wc",
                            "bideuri": "bideuri",
                            "masina de spalat vase": "masina_spalat_vase",
                            "masini spalat vase": "masina_spalat_vase",
                            "pisoare": "pisoare",
                            "masina de spalat rufe": "masina_spalat_rufe",
                            "masini spalat rufe": "masina_spalat_rufe",
                            "dus": "dus",
                        }
                        for tbl in doc.tables:
                            # Header check: r0c0 = "Lavoare" și r0c1 = "[nr.]"
                            if len(tbl.rows) < 1 or len(tbl.rows[0].cells) < 2:
                                continue
                            r0c0 = tbl.rows[0].cells[0].text.strip()
                            r0c1 = tbl.rows[0].cells[1].text.strip()
                            if "Lavoare" not in r0c0 or "[nr.]" not in r0c1:
                                continue
                            # FIX 20 apr 2026: folosesc matching NORMALIZAT (diacritice-insensitive)
                            # pentru a match "Duș"/"Dus"/"Spălătoare"/"Spalatoare"/"Mașina"/"Masina"
                            for ri in range(len(tbl.rows)):
                                for ci in range(0, len(tbl.rows[ri].cells), 2):
                                    if ci + 1 >= len(tbl.rows[ri].cells):
                                        continue
                                    label_cell = tbl.rows[ri].cells[ci]
                                    value_cell = tbl.rows[ri].cells[ci + 1]
                                    label_txt = label_cell.text.strip()
                                    label_norm = _norm_label(label_txt)
                                    if "[nr.]" in value_cell.text:
                                        key = fix_key_map_norm.get(label_norm)
                                        if key:
                                            count = fixtures.get(key, "0")
                                            # Înlocuiesc [nr.] cu valoarea numerică
                                            for p in value_cell.paragraphs:
                                                for r in p.runs:
                                                    if "[nr.]" in r.text:
                                                        r.text = r.text.replace("[nr.]", str(count))
                            break
                except Exception as e_t11:
                    print(f"[tabel_11_sanitare] eroare: {e_t11}", flush=True)

                # ── Etapa 7d (20 apr 2026) — gap-uri suplimentare paragrafe ──
                # După audit exhaustiv, identificate aceste gap-uri P0 (date exist):
                #   p208 — Necesarul de căldură (kW)
                #   p266, p267 — Putere termică ACM (kW)
                #   p349 — Debit aer proaspăt (m³/h)
                #   p424, p425 — Energie regenerabilă (duplicate p422/p423)
                #   p426 — EPP (kWh/m²·an)
                #   p427 — RERP (%)
                #   p428 — Emisii CO2 (kg/m²·an)
                #   p429 — SRI (smart readiness indicator)

                # p208 — Necesarul de căldură de calcul
                heat_power = data.get("heating_power", "")
                if heat_power and heat_power not in ("", "0"):
                    _fill_para_blank("Necesarul de căldură de calcul", heat_power, " kW")

                # p266, p267 — Putere termică ACM (folosesc acm_power dacă există, fallback heating)
                acm_power = data.get("acm_power", "") or data.get("heating_power", "")
                if acm_power and acm_power not in ("", "0"):
                    _fill_para_blank("Puterea termică necesară pentru prepararea acc", acm_power, " kW")
                    _fill_para_blank("Puterea termică maximă instalată pentru prepararea acc", acm_power, " kW")

                # p349 — Debit aer proaspăt minim (supply / exhaust)
                # Template format: "din apartament:    /     m³/h" (două câmpuri separate de "/")
                # _fill_para_blank umple PRIMUL câmp; al doilea rămas gol → spațiu mare.
                # Fix: înlocuim TOATE secvențele de spații din toate run-urile paragrafului.
                vent_flow = data.get("ventilation_flow_m3h", "")
                if vent_flow and vent_flow != "0":
                    import re as _re_vf
                    _vf_filled = False
                    for _vfp in doc.paragraphs:
                        if "Debitul minim de aer proaspăt" in _vfp.text:
                            for _vfr in _vfp.runs:
                                if _re_vf.search(r"[\xa0\s]{3,}", _vfr.text):
                                    _vfr.text = _re_vf.sub(r"[\xa0\s]{3,}", f" {vent_flow} ", _vfr.text)
                                    _vf_filled = True
                            break
                    if not _vf_filled:
                        _fill_para_blank("Debitul minim de aer proaspăt", vent_flow, " m³/h")

                # p424, p425 — duplicate cu p422/p423 (acelaș text dar contains "din surse regenerabile")
                if data.get("solar_th_kwh_year") and data["solar_th_kwh_year"] != "0":
                    _fill_para_blank("Energia termică exportată din surse regenerabile",
                                     data["solar_th_kwh_year"], " kWh/an")
                if data.get("pv_kwh_year") and data["pv_kwh_year"] != "0":
                    _fill_para_blank("Energia electrică exportată din surse regenerabile",
                                     data["pv_kwh_year"], " kWh/an")

                # p426 — EPP (Indicatorul energiei primare)
                # FIX 20 apr 2026 (Sprint monolith): elimin skip pe "0,0" — completez întotdeauna
                # Placeholder rămas gol e mai rău decât valoarea 0 (vizual pare "necompletat").
                epp = data.get("ep_specific", "")
                if epp:
                    _fill_para_blank("Indicatorul energiei primare EPP", epp, " kWh/(m²·an)")
                    _fill_para_blank("Indicatorul energiei primare EP", epp, " kWh/(m²·an)")

                # p427 — RERP % (renewable energy ratio prime)
                rerp = data.get("rer", "")
                if rerp:
                    _fill_para_blank("Indicele RERP", rerp, "%")
                    _fill_para_blank("Indicele RER", rerp, "%")

                # p428 — Indicator emisii CO2
                co2_ind = data.get("co2_val", "")
                if co2_ind:
                    _fill_para_blank("Indicatorul emisiilor de CO", co2_ind, " kgCO₂/(m²·an)")

                # p429 — SRI complet (label + valoare + clasă)
                sri_v = data.get("sri_total", "")
                sri_g = data.get("sri_grade", "")
                if sri_v:
                    _fill_para_blank("Indicele SRI (smart readiness indicator)",
                                     f"{sri_v}% (Clasa {sri_g})" if sri_g else f"{sri_v}%")

                # ═══════════════════════════════════════════════════════
                # Sprint monolith (20 apr 2026) — completări suplimentare Anexa 2
                # Umplere câmpuri care au date în payload dar nu erau procesate:
                # - Heat pump usage (Utilizată/e pentru) + SCOP/SEER
                # - Biomass alt tip precizare
                # - Wind centrals count + putere + caracteristici
                # - Energie termică/electrică exportată (non-regenerabil — duplicat)
                # - Echipamente ACM (Boiler acumulare număr/volum, instant putere)
                # - Necesar umidificare
                # - Diametru nominal + presiune racord centralizat
                # - Contor căldură încălzire + răcire + ACM
                # - Elemente reglaj termic + hidraulic
                # - Spații climatizate speciale + complet/global/parțial
                # - Ventilare: caracteristici + recuperator tip
                # - Iluminat alt tip precizare
                # ═══════════════════════════════════════════════════════

                # ── Heat pump: utilizare + SCOP/SEER ──────────────────
                hp_covers = data.get("heat_pump_covers", "")
                hp_covers_label = {
                    "heating_only": "încălzire",
                    "cooling_only": "răcire",
                    "heating_cooling": "încălzire și răcire",
                    "heating_acm": "încălzire și preparare acc",
                    "all": "încălzire, răcire și preparare acc",
                }.get(hp_covers, hp_covers) if hp_covers else ""
                if hp_covers_label:
                    _fill_para_blank("Utilizată/e pentru", hp_covers_label)
                hp_scop = data.get("heat_pump_scop_heating", "")
                hp_seer = data.get("heat_pump_scop_cooling", "")
                if hp_scop or hp_seer:
                    scop_text = f"SCOP = {hp_scop}" if hp_scop else ""
                    seer_text = f"SEER = {hp_seer}" if hp_seer else ""
                    value_str = " / ".join([v for v in (scop_text, seer_text) if v])
                    _fill_para_blank("Valoarea medie SCOP/SEER", value_str)
                    _fill_para_blank("Valoarea medie SCOP", value_str)

                # ── Biomass: alt tip precizare + putere ───────────────
                bio_type = data.get("biomass_type", "")
                # Case-insensitive: "PELETI"/"BRICHETE" (uppercase din Step3) → skip,
                # altfel contamina prima apariție "alt tip, precizați" din building-type section.
                if bio_type and bio_type.lower() not in ("peleti", "brichete", "lemn", "lemn tocat", ""):
                    _fill_para_blank("alt tip, precizați", bio_type)

                # ── Wind centrals detail ──────────────────────────────
                wind_count = data.get("wind_centrals_count", "")
                wind_power = data.get("wind_power_kw", "")
                if wind_count:
                    _fill_para_blank("Număr centrale eoliene", wind_count)
                if wind_power:
                    _fill_para_blank("Putere nominală [kW]", wind_power)

                # ── Energie termică exportată (non-regenerabil duplicat) ──
                if data.get("solar_th_kwh_year") and data["solar_th_kwh_year"] != "0":
                    _fill_para_blank("Energia termică exportată:",
                                     data["solar_th_kwh_year"], " kWh/an")
                if data.get("pv_kwh_year") and data["pv_kwh_year"] != "0":
                    _fill_para_blank("Energia electrică exportată:",
                                     data["pv_kwh_year"], " kWh/an")

                # ── ACM echipamente ───────────────────────────────────
                acm_storage = data.get("acm_storage_volume", "")
                if acm_storage:
                    # "Boiler cu acumulare (număr/volum)"
                    _fill_para_blank("Boiler cu acumulare", f"1 / {acm_storage} L")
                acm_inst_power = data.get("acm_instant_power_kw", "")
                if acm_inst_power:
                    _fill_para_blank("Preparare locală cu aparate de tip instant", f"1 / {acm_inst_power} kW")

                # ── Necesar umidificare (opțional, gol dacă nu se calculează) ──
                humid_power = data.get("humidification_power_kw", "")
                if humid_power:
                    _fill_para_blank("Necesarul de energie pentru umidificare", humid_power, " kW")

                # ── Racord centralizat de căldură: diametru + presiune ──
                pipe_diam = data.get("heating_pipe_diameter_mm", "")
                if pipe_diam:
                    _fill_para_blank("diametru nominal:", pipe_diam, " mm")
                pipe_press = data.get("heating_pipe_pressure_mca", "")
                if pipe_press:
                    _fill_para_blank("disponibil de presiune (nominal):", pipe_press, " mCA")

                # ── Ventilare caracteristici + recuperator tip ────────
                vent_fan_count = data.get("ventilation_fan_count", "")
                if vent_fan_count:
                    for p in doc.paragraphs:
                        if "Numărul total de ventilatoare" in p.text and "......" in p.text:
                            replace_in_paragraph(p, "......................", vent_fan_count, count=1)
                            break
                vent_hr_type = data.get("ventilation_hr_type", "")
                if vent_hr_type:
                    _fill_para_blank("Tip:", vent_hr_type)

                # ── Iluminat alt tip precizare ────────────────────────
                light_other = data.get("lighting_other_type", "")
                if light_other:
                    _fill_para_blank("Mixt (precizați)", light_other)

                # ── Număr total puncte consum ACM ─────────────────────
                acm_consume_points = data.get("acm_consume_points_count", "")
                if acm_consume_points:
                    _fill_para_blank("Număr total de puncte de consum acc", acm_consume_points)

                # ── Cooling: agent frigorific + dezumidificare ────────
                cool_refrigerant = data.get("cooling_refrigerant", "")
                if cool_refrigerant:
                    _fill_para_blank("Tip agent frigorific utilizat", cool_refrigerant)
                cool_dehum = data.get("cooling_dehum_power_kw", "")
                if cool_dehum:
                    _fill_para_blank("Necesarul de frig pentru dezumidificare", cool_dehum, " kW")

                # ── Cooling: număr unități interior/exterior (split) ──
                cool_int_units = data.get("cooling_indoor_units", "")
                cool_ext_units = data.get("cooling_outdoor_units", "")
                if cool_int_units:
                    _fill_para_blank("Număr de unități interioare", cool_int_units)
                if cool_ext_units:
                    _fill_para_blank("Număr de unități exterioare", cool_ext_units)

                # ── Cooling: diametru + presiune racord ───────────────
                cool_diam = data.get("cooling_pipe_diameter_mm", "")
                if cool_diam:
                    _fill_para_blank("diametru nominal:", cool_diam, " mm")

                # ── ACM: diametru + presiune racord ───────────────────
                # FIX 20 apr: racord la sursa centralizată ACM are 3 câmpuri goale:
                # - diametru nominal mm, - necesar de presiune mCA, multiplu ___ puncte
                # Labels: "- diametru nominal:", "- necesar de presiune (nominal):", "puncte"
                acm_diam = data.get("acm_pipe_diameter_mm", "")
                acm_pts = data.get("acm_consume_points_count", "")
                # Căutăm după "Tipul echipamentelor" sau "Boiler cu acumulare" pentru
                # a detecta secțiunea ACM, apoi completăm următoarele paragrafe
                acm_section_found = False
                for p in doc.paragraphs:
                    pt = p.text
                    if "Tipul echipamentelor de preparare" in pt or "Boiler cu acumulare" in pt:
                        acm_section_found = True
                        continue
                    if not acm_section_found:
                        continue
                    # Ieșire din secțiune când ajungem la răcire
                    if "INFORMAȚII PRIVIND INSTALAȚIA DE RĂCIRE" in pt.upper():
                        break
                    # Fill diameter nominal (ACM)
                    if acm_diam and "diametru nominal" in pt.lower() and "mm" in pt:
                        for run in p.runs:
                            m = re.search(r"[\xa0\s]{3,}|\.{4,}", run.text)
                            if m:
                                run.text = run.text[:m.start()] + f" {acm_diam} " + run.text[m.end():]
                                break
                    # Fill presiune (ACM)
                    pipe_press = data.get("heating_pipe_pressure_mca", "")
                    if pipe_press and "presiune" in pt.lower() and "mCA" in pt:
                        for run in p.runs:
                            m = re.search(r"[\xa0\s]{3,}|\.{4,}", run.text)
                            if m:
                                run.text = run.text[:m.start()] + f" {pipe_press} " + run.text[m.end():]
                                break
                    # Fill multiplu puncte
                    if acm_pts and "multiplu" in pt.lower() and "puncte" in pt.lower():
                        for run in p.runs:
                            m = re.search(r"[\xa0\s]{3,}|\.{4,}", run.text)
                            if m:
                                run.text = run.text[:m.start()] + f" {acm_pts} " + run.text[m.end():]
                                break

                # ── Footer: "Numărul certificatului în registrul auditorului: X / Y" ──
                # Structura template: "Numărul certificatului în registrul auditorului............"
                # Înlocuim NUMAI dots-urile cu "registryIndex / nrMDLPA" (textul dinaintea dots rămâne intact).
                nr_registru = str(data.get("registry_index", "") or "").strip()
                nr_mdlpa_val = str(data.get("nr_mdlpa", "") or "").strip()
                # Format footer: "Numarul certificatului in registrul auditorului: X"
                # X = nr secvential (registryIndex) — codul propriu al auditorului.
                # Formatul complet CPE nr ({nrMDLPA}/{registryIndex}) apare in barcode si titlu.
                if nr_registru or nr_mdlpa_val:
                    if nr_registru:
                        nr_footer = nr_registru
                    else:
                        nr_footer = nr_mdlpa_val
                    try:
                        for section in doc.sections:
                            for footer_type in ("first_page_footer", "even_page_footer", "footer"):
                                footer = getattr(section, footer_type, None)
                                if footer is None:
                                    continue
                                for p in footer.paragraphs:
                                    if "registrul auditorului" not in p.text:
                                        continue
                                    # Înlocuiește run-ul cu dots (placeholder) cu nr_footer.
                                    # Textul "Numărul certificatului în registrul auditorului" rămâne intact.
                                    replaced = False
                                    # Înlocuim dots cu ": NR" (cu colon + spațiu prefix, ca în template oficial MDLPA)
                                    for run in p.runs:
                                        if re.search(r"\.{4,}", run.text):
                                            run.text = re.sub(r"\.{4,}", ": " + nr_footer, run.text, count=1)
                                            replaced = True
                                            break
                                    if not replaced:
                                        # Fallback: dacă dots sunt în același run cu textul
                                        for run in p.runs:
                                            if re.search(r"[ \s]{4,}", run.text) and "auditorului" in run.text:
                                                run.text = run.text.rstrip() + ": " + nr_footer
                                                replaced = True
                                                break
                    except Exception:
                        pass

                # ── Număr sobe ─────────────────────────────────────────
                stove_count = data.get("stove_count", "")
                if stove_count:
                    _fill_para_blank("Numărul sobelor / combustibilul utilizat", stove_count)

                # ── CT proprie/exterior combustibil + termoficare ─────
                heat_location = data.get("heating_gen_location", "")
                fuel_heating = fuel_labels.get(data.get("heating_fuel", ""), data.get("heating_fuel", ""))
                if heat_location == "CT_PROP" and fuel_heating:
                    _fill_para_blank("Centrală termică proprie în clădire, cu combustibil", fuel_heating)
                elif heat_location == "CT_EXT" and fuel_heating:
                    _fill_para_blank("Centrală termică în exteriorul clădirii, cu combustibil", fuel_heating)
                # Alt tip sursă
                other_heat_source = data.get("heating_other_source_text", "")
                if other_heat_source:
                    _fill_para_blank("Altă sursă sau sursă mixtă (precizați)", other_heat_source)
                # Tip corp static (generic pentru încălzire cu corpuri statice)
                radiator_type = data.get("heating_radiator_type", "")
                if radiator_type:
                    _fill_para_blank("Încălzire cu corpuri statice", f" — {radiator_type}")

                # ══════════════════════════════════════════════════════════
                # FIX 20 apr 2026: marcaj "—" pentru secțiuni EXOTICE neaplicabile
                # (planșeu/plafon/perete încălzitor, tuburi radiante, aer cald)
                # Aceste secțiuni au ~8 placeholder-e ".........." care rămân goale
                # pentru majoritatea clădirilor (unde nu există aceste instalații).
                # Marchează cu "—" pentru a arăta că software-ul a procesat câmpul,
                # dar nu există date aplicabile (distinct de "nefilled").
                # ══════════════════════════════════════════════════════════
                heating_src_key = (data.get("heating_source", "") or "").lower()
                has_radiant = any(k in heating_src_key for k in ["radiant", "planseu", "plafon", "perete", "pardoseala"])
                has_radiant_tube = "tub" in heating_src_key or "radiant_tube" in heating_src_key
                has_hot_air = "aer_cald" in heating_src_key or "generator_aer" in heating_src_key
                has_electric_radiant = "electric_radiant" in heating_src_key or "cable" in heating_src_key

                # Pentru secțiuni NEAPLICABILE, marchez cu "—" în placeholder.
                # Fix 2 mai 2026: normalizare diacritice (ș/ş, ț/ţ) + iterație ÎN TABELE
                # (template-ul oficial pune multe placeholder-uri în celule de tabel).
                import unicodedata as _ud
                def _norm_diacritics(s):
                    """Normalizează ş→ș, ţ→ț pentru match tolerant."""
                    return s.replace("ş", "ș").replace("Ş", "Ș") \
                            .replace("ţ", "ț").replace("Ţ", "Ț")
                def _fill_na_placeholder(label, marker="—"):
                    """Înlocuiește placeholder "..." sau spații cu marker "—" pentru
                    secțiuni neaplicabile. Caută în corp + tabele, tolerant la diacritice.
                    """
                    label_n = _norm_diacritics(label)
                    for p in _iter_all_paragraphs(doc):
                        pt_n = _norm_diacritics(p.text)
                        if label_n not in pt_n:
                            continue
                        for run in p.runs:
                            rt = run.text
                            if not rt:
                                continue
                            if re.search(r"[\xa0\s]{3,}|\.{4,}", rt):
                                run.text = re.sub(
                                    r"[\xa0\s]{3,}|\.{4,}",
                                    f" {marker} ", rt, count=1)
                                return True
                    return False

                # Planșeu/plafon/perete încălzitor — labels potrivite cu textul EXACT
                # din template (verificat 2 mai 2026 prin docx inspection)
                if not has_radiant:
                    _fill_na_placeholder("Aria planşeelor/plafoanelor/pereților")
                # Cabluri electrice încălzitoare — dacă nu e activ
                if not has_electric_radiant:
                    _fill_na_placeholder("Lungimea şi tipul cablurilor electrice")
                # Tuburi radiante
                if not has_radiant_tube:
                    _fill_na_placeholder("Tip/putere tub radiant")
                    _fill_na_placeholder("Număr/lungime tuburi radiante")
                # Generator aer cald
                if not has_hot_air:
                    _fill_na_placeholder("Tip/putere generator aer cald")
                    _fill_na_placeholder("Număr/debit aer")

                # Debit nominal agent termic (pt. încălzire) — placeholder înainte de l/h
                heating_flow = data.get("heating_flow_lh", "") or data.get("heating_flow_rate", "")
                if heating_flow:
                    _fill_para_blank("Debitul nominal de agent termic", str(heating_flow), " l/h")
                else:
                    _fill_na_placeholder("Debitul nominal de agent termic")

                # Similar pentru necesar umidificare (rar aplicabil)
                humid_power_val = data.get("humidification_power_kw", "")
                if not humid_power_val:
                    _fill_na_placeholder("Necesarul de energie pentru umidificare", marker="0")

            # ═══════════════════════════════════════
            # 5. FOTO CLĂDIRE — inserare imagine în text box
            # ═══════════════════════════════════════
            photo_b64 = body.get("photo")
            if photo_b64 and photo_b64.startswith("data:image"):
                # Suport png/jpeg/jpg/webp/gif/bmp
                match = re.match(r"data:image/(png|jpeg|jpg|webp|gif|bmp);base64,(.+)", photo_b64, re.DOTALL)
                if match:
                    img_data = base64.b64decode(match.group(2))
                    from docx.oxml.ns import qn as docx_qn
                    from docx.shared import Emu

                    # Dimensiunile frame-ului FOTO din template (EMU): 647700 x 584200
                    FOTO_W_EMU = 620000   # puțin sub lățimea frame-ului (~1,7 cm)
                    FOTO_H_EMU = 560000   # puțin sub înălțimea frame-ului (~1,55 cm)

                    # Calculează dimensiunile imaginii păstrând aspect ratio
                    try:
                        import struct
                        raw = img_data
                        img_w_px, img_h_px = None, None
                        # PNG: lățime/înălțime la bytes 16-24
                        if raw[:4] == b'\x89PNG':
                            img_w_px, img_h_px = struct.unpack('>II', raw[16:24])
                        # JPEG: cauta SOF marker
                        elif raw[:2] == b'\xff\xd8':
                            i = 2
                            while i < len(raw) - 8:
                                if raw[i] == 0xff and raw[i+1] in (0xC0, 0xC2):
                                    img_h_px = struct.unpack('>H', raw[i+5:i+7])[0]
                                    img_w_px = struct.unpack('>H', raw[i+7:i+9])[0]
                                    break
                                seg_len = struct.unpack('>H', raw[i+2:i+4])[0]
                                i += 2 + seg_len
                    except Exception:
                        img_w_px, img_h_px = None, None

                    # Alege dimensiunea finală respectând frame-ul și aspect ratio
                    if img_w_px and img_h_px and img_w_px > 0 and img_h_px > 0:
                        ratio = img_w_px / img_h_px
                        if ratio >= 1:  # landscape sau pătrat → constrâns de lățime
                            pic_w = FOTO_W_EMU
                            pic_h = int(FOTO_W_EMU / ratio)
                            if pic_h > FOTO_H_EMU:  # totuși prea înalt
                                pic_h = FOTO_H_EMU
                                pic_w = int(FOTO_H_EMU * ratio)
                        else:  # portrait → constrâns de înălțime
                            pic_h = FOTO_H_EMU
                            pic_w = int(FOTO_H_EMU * ratio)
                            if pic_w > FOTO_W_EMU:
                                pic_w = FOTO_W_EMU
                                pic_h = int(FOTO_W_EMU / ratio)
                    else:
                        # Fallback: folosim lățimea frame-ului
                        pic_w = FOTO_W_EMU
                        pic_h = None

                    # Găsește TOATE text box-urile care conțin "FOTO" și înlocuiește-le
                    txbx_tag = qn("w:txbxContent")
                    p_tag = qn("w:p")
                    from docx.text.paragraph import Paragraph as DocxPara

                    foto_txbx_list = []
                    for txbx_elem in doc.element.body.iter(txbx_tag):
                        for p_elem in txbx_elem.iter(p_tag):
                            if "FOTO" in (DocxPara(p_elem, None).text or ""):
                                foto_txbx_list.append(txbx_elem)
                                break  # un singur text box per găsire

                    for foto_txbx in foto_txbx_list:
                        # Șterge TOATE paragrafele existente din text box
                        for p_elem in list(foto_txbx.findall(p_tag)):
                            foto_txbx.remove(p_elem)

                        # Creează un singur paragraf centrat cu imaginea
                        new_p = etree.SubElement(foto_txbx, p_tag)
                        pPr = etree.SubElement(new_p, qn("w:pPr"))
                        sp = etree.SubElement(pPr, qn("w:spacing"))
                        sp.set(qn("w:after"), "0")
                        jc = etree.SubElement(pPr, qn("w:jc"))
                        jc.set(qn("w:val"), "center")
                        new_r = etree.SubElement(new_p, qn("w:r"))

                        # Adaugă imaginea via python-docx (calea standard)
                        img_stream = io.BytesIO(img_data)
                        img_stream.seek(0)
                        temp_para = doc.add_paragraph()
                        temp_run = temp_para.add_run()
                        if pic_h:
                            temp_run.add_picture(img_stream, width=Emu(pic_w), height=Emu(pic_h))
                        else:
                            temp_run.add_picture(img_stream, width=Emu(pic_w))
                        drawing = temp_run._element.find(docx_qn("w:drawing"))
                        if drawing is not None:
                            new_r.append(drawing)
                        temp_para._element.getparent().remove(temp_para._element)

            # ═══════════════════════════════════════
            # 5b. ETAPE IMPLEMENTARE + STIMULENTE FINANCIARE (Anexa 1+2, secțiunea I)
            # ═══════════════════════════════════════
            # Înlocuiește textul placeholder al auditorului cu textul furnizat de
            # auditor din formularul „Completare automată detaliată" din Pas 6.
            # Dacă câmpul e gol, textul original din template rămâne neschimbat.
            if mode in ("anexa", "anexa_bloc"):
                _etape = (data.get("etape_implementare") or "").strip()
                _stimulente = (data.get("stimulente_financiare") or "").strip()
                _solutii_anvelopa = (data.get("solutii_anvelopa") or "").strip()
                _solutii_instalatii = (data.get("solutii_instalatii") or "").strip()
                _masuri_organizare = (data.get("masuri_organizare") or "").strip()
                _masuri_locale = (data.get("masuri_locale") or "").strip()
                _regenerabile_custom = (data.get("regenerabile_custom") or "").strip()

                def _replace_placeholder_para(marker_text, new_text, occurrence=1):
                    """Înlocuiește conținutul paragrafului care conține marker_text.

                    Textul nou poate conține newline-uri — fiecare linie devine
                    o pereche w:t/w:br în același run, pentru a păstra formatarea
                    originală a paragrafului (dimensiune font, stil).

                    occurrence: 1-based index — care apariție a marker_text să fie
                    înlocuită (util când același text apare la mai multe secțiuni).
                    """
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    seen = 0
                    for p in doc.paragraphs:
                        if marker_text in p.text:
                            seen += 1
                            if seen != occurrence:
                                continue
                            # Salvăm rPr-ul primului run pentru a păstra formatarea
                            saved_rPr = None
                            if p.runs:
                                rPr = p.runs[0]._r.find(qn("w:rPr"))
                                if rPr is not None:
                                    saved_rPr = copy.deepcopy(rPr)
                            # Ștergem toate run-urile existente
                            for run in list(p.runs):
                                run._r.getparent().remove(run._r)
                            # Adăugăm un run nou cu rPr păstrat + w:t/w:br alternând
                            new_run = p.add_run("")
                            # Eliminăm w:t implicit creat de add_run
                            for child in list(new_run._r):
                                new_run._r.remove(child)
                            # Atașăm rPr salvat (dacă există)
                            if saved_rPr is not None:
                                new_run._r.append(saved_rPr)
                            # Adăugăm fiecare linie ca w:t, separate prin w:br
                            lines = new_text.split("\n")
                            for i, line in enumerate(lines):
                                if i > 0:
                                    new_run._r.append(OxmlElement("w:br"))
                                t = OxmlElement("w:t")
                                t.text = line
                                t.set(qn("xml:space"), "preserve")
                                new_run._r.append(t)
                            break

                if _etape:
                    _replace_placeholder_para(
                        "auditorul energetic va completa mai departe lista cu etapele adaptate",
                        _etape,
                    )
                if _stimulente:
                    _replace_placeholder_para(
                        "auditorul energetic va completa mai departe lista cu stimulentele financiare",
                        _stimulente,
                    )
                # Marker "soluții adaptate obiectivului certificat" apare de 3 ori:
                # ocurența 1 = anvelopă (para [13]), 2 = instalații (para [38]), 3 = 3.A organizare (para [50])
                MARKER_OBIECTIV = "soluții adaptate obiectivului certificat"
                if _solutii_anvelopa:
                    _replace_placeholder_para(MARKER_OBIECTIV, _solutii_anvelopa, occurrence=1)
                if _solutii_instalatii:
                    _replace_placeholder_para(MARKER_OBIECTIV, _solutii_instalatii, occurrence=2)
                if _masuri_organizare:
                    _replace_placeholder_para(MARKER_OBIECTIV, _masuri_organizare, occurrence=3)
                # Marker pentru 3.B măsuri locale — text diferit ("clădirii certificate" + fără "obiectivului")
                if _masuri_locale:
                    _replace_placeholder_para(
                        "soluții adaptate clădirii certificate",
                        _masuri_locale,
                    )
                # Marker surse regenerabile — para [339/420]
                if _regenerabile_custom:
                    _replace_placeholder_para(
                        "alte echipamente care utilizează sursele regenerabile",
                        _regenerabile_custom,
                    )

                # ─── Nr. persoane apartament/clădire (para [120]) ───
                _nr_persoane = (data.get("nr_persoane") or "").strip()
                if _nr_persoane:
                    _replace_placeholder_para(
                        "Numărul maxim real/normat de persoane",
                        f"Numărul maxim real/normat de persoane din apartament:             {_nr_persoane}       pers.",
                    )

                # ═══════════════════════════════════════
                # 5c. BIFARE CHECKBOX-URI ÎN ANEXA (Sprint 2 mai 2026)
                # ═══════════════════════════════════════
                # Maparea code → text-anchor din template:
                _structure = data.get("structure_code", "")
                _structure_anchors = {
                    "zidarie":        "pereţi structurali din zidărie",
                    "beton_armat":    "pereţi structurali din beton armat",
                    "cadre_ba":       "cadre din beton armat",
                    "stalpi_grinzi":  "stâlpi şi grinzi",
                    "lemn":           "structura de lemn",
                    "metalica":       "structură metalică",
                    "alt_tip":        None,  # nu bifez nimic
                }
                if _structure and _structure in _structure_anchors and _structure_anchors[_structure]:
                    check_box_for_text(doc, _structure_anchors[_structure])

                # Tip sistem încălzire
                _heating = data.get("heating_source_type", "")
                _heating_anchors = {
                    "centrala_proprie":     "Centrală termică proprie în clădire",
                    "centrala_exterioara":  "Centrală termică în exteriorul clădirii",
                    "aparate_independente": "Încălzire cu alte aparate independente",
                    "alt_tip":              None,
                }
                if _heating and _heating in _heating_anchors and _heating_anchors[_heating]:
                    check_box_for_text(doc, _heating_anchors[_heating])

                # Contor căldură + Repartitoare costuri (paragrafele NEXT după label)
                # Contor de căldură: există(0) / există fără viză(1) / nu există(2)
                _contor_caldura = data.get("contor_caldura", "")
                if _contor_caldura == "exista":
                    check_box_after_label(doc, "Contor de căldură", option_idx=0)
                elif _contor_caldura == "nu_exista":
                    check_box_after_label(doc, "Contor de căldură", option_idx=2)

                # Repartitoare de costuri: există(0) / nu există(1)
                _repartitoare = data.get("repartitoare_costuri", "")
                if _repartitoare == "exista":
                    check_box_after_label(doc, "Repartitoare de costuri", option_idx=0)
                elif _repartitoare == "nu_exista":
                    check_box_after_label(doc, "Repartitoare de costuri", option_idx=1)

                # ACM Conducta recirculare: funcțională(0) / există dar nu funcționează(1) / nu există(2)
                _acm_rec = data.get("acm_recirculare", "")
                if _acm_rec == "functionala":
                    check_box_after_label(doc, "Conducta de recirculare", option_idx=0)
                elif _acm_rec == "exista_nefunctional":
                    check_box_after_label(doc, "Conducta de recirculare", option_idx=1)
                elif _acm_rec == "nu_exista":
                    check_box_after_label(doc, "Conducta de recirculare", option_idx=2)

                # Contor general căldură pentru ACM: există(0) / nu există(1) / nu este cazul(2)
                _contor_acm = data.get("contor_acm", "")
                if _contor_acm == "exista":
                    check_box_after_label(doc, "Contor general de căldură pentru acc", option_idx=0)
                elif _contor_acm == "nu_exista":
                    check_box_after_label(doc, "Contor general de căldură pentru acc", option_idx=1)
                elif _contor_acm == "nu_este_cazul":
                    check_box_after_label(doc, "Contor general de căldură pentru acc", option_idx=2)

                # Debitmetre puncte consum: nu există(0) / parțial(1) / peste tot(2)
                _debitmetre = data.get("debitmetre", "")
                if _debitmetre == "nu_exista":
                    check_box_after_label(doc, "Debitmetre la nivelul punctelor de consum", option_idx=0)
                elif _debitmetre == "partial":
                    check_box_after_label(doc, "Debitmetre la nivelul punctelor de consum", option_idx=1)
                elif _debitmetre == "peste_tot":
                    check_box_after_label(doc, "Debitmetre la nivelul punctelor de consum", option_idx=2)

                # Control iluminat
                _ilum = data.get("iluminat_control", "")
                if _ilum == "fara_reglare":
                    check_box_for_text(doc, "Fără reglare (on/off)")
                elif _ilum == "manuala":
                    check_box_for_text(doc, "Reglare manuală")
                elif _ilum == "automat":
                    check_box_for_text(doc, "Automat funcție de")

                # Bifare zone climatică/eoliană/regim înălțime în Tabel #0 al Anexa
                # Tabelul are structura: 6 rânduri × 13 coloane (col 0 = label, col 1-12 = opțiuni)
                # R0 = zone climatice (I/II/II/II/II/III/III/IV/IV/IV/V/V), R1 = celule bifare
                # R2 = zone eoliene (I/I/I/II/II/II/III/III/III/IV/IV/IV), R3 = celule bifare
                # R4 = regim înălțime (S/S/D/D/Mez/Mez/P/P/E/E/E/M/P), R5 = celule bifare
                _climate = str(data.get("zona_climatica", "")).strip().upper()  # I, II, III, IV, V
                _wind = str(data.get("zona_eoliana", "")).strip().upper()
                _regim = str(data.get("regim_inaltime", "")).strip()
                # Caut tabelul cu „Zona climatică în care este amplasată"
                _zone_table = None
                for tbl in doc.tables:
                    if any("Zona climatică" in c.text for r in tbl.rows for c in r.cells):
                        _zone_table = tbl
                        break
                if _zone_table is not None and len(_zone_table.rows) >= 6:
                    rows = _zone_table.rows
                    # R0: zone climatice headers; R1: bifare. Caut col cu zona corectă în R0
                    if _climate:
                        for ci, cell in enumerate(rows[0].cells):
                            if cell.text.strip() == _climate and ci > 0:
                                check_form_checkbox_in_cell(rows[1].cells[ci])
                                break
                    if _wind:
                        for ci, cell in enumerate(rows[2].cells):
                            if cell.text.strip() == _wind and ci > 0:
                                check_form_checkbox_in_cell(rows[3].cells[ci])
                                break
                    # Regim înălțime: bifează TOATE tipurile aplicabile
                    # P+4E → bifează P și E; S → bifează S; completează (nr) cu nr. etaje
                    if _regim:
                        regim_parts = set()
                        nr_etaje_val = None
                        if "S" in _regim and "+" not in _regim:
                            regim_parts.add("S")
                        elif "Mez" in _regim or "mezanin" in _regim.lower():
                            regim_parts.add("Mez")
                        elif _regim.upper().startswith("D"):
                            regim_parts.add("D")
                        elif "P+" in _regim:
                            regim_parts.add("P")
                            try:
                                nr_etaje_val = int(
                                    _regim.split("P+")[1].split("E")[0].strip() or "0"
                                )
                                if nr_etaje_val >= 1:
                                    regim_parts.add("E")
                                if nr_etaje_val >= 6:
                                    regim_parts.add("M/P")
                            except Exception:
                                regim_parts.add("E")
                        else:
                            regim_parts.add("P")
                        # Bifează toate coloanele care corespund tipurilor detectate
                        for ci, cell in enumerate(rows[4].cells):
                            if cell.text.strip() in regim_parts and ci > 0:
                                check_form_checkbox_in_cell(rows[5].cells[ci])
                        # Completează (nr) cu numărul de etaje — căutare GLOBALĂ
                        # peste TOATE paragrafele documentului (placeholder-ul poate fi
                        # în paragrafe top-level sau alte tabele, nu doar în rows[4]/[5])
                        if nr_etaje_val is not None and nr_etaje_val > 0:
                            replaced_nr = 0
                            for _p_global in _iter_all_paragraphs(doc):
                                if "(nr)" in _p_global.text:
                                    n = replace_in_paragraph(
                                        _p_global, "(nr)", str(nr_etaje_val)
                                    )
                                    replaced_nr += n
                            if replaced_nr:
                                print(f"[regim] (nr) -> {nr_etaje_val} ({replaced_nr} loc)",
                                      flush=True)

            # ═══════════════════════════════════════
            # 7b. PAGINĂ SUPLIMENT — dezactivat 2 mai 2026
            # ═══════════════════════════════════════
            # if mode == "cpe": append_legal_supplement(doc, data)

            # ═══════════════════════════════════════
            # 7c. ANEXA BLOC — tabel apartamente + sisteme comune
            # ═══════════════════════════════════════
            # Injectat ÎNAINTE de secțiunea H pentru ca H să fie ULTIMA secțiune.
            if mode == "anexa_bloc":
                apartments = body.get("apartments", []) or []
                apartment_summary = body.get("apartmentSummary") or {}
                common_systems = body.get("commonSystems") or {}
                try:
                    if apartments:
                        added_apts = insert_apartment_table(doc, apartments, apartment_summary)
                        if added_apts:
                            print(
                                f"[anexa_bloc] tabel apartamente injectat: {len(apartments)} ap.",
                                flush=True,
                            )
                    if common_systems:
                        insert_common_systems_section(doc, common_systems)
                except Exception as e_bloc:
                    print(f"[anexa_bloc] eroare injecție: {e_bloc}", flush=True)

            # ═══════════════════════════════════════
            # 6. ANEXĂ FOTOGRAFII CLĂDIRE — Secțiunea H (ULTIMA secțiune)
            # ═══════════════════════════════════════
            # Mutat după 7c (apt. bloc) pentru ca H să fie mereu ultima pagină.
            building_photos = body.get("buildingPhotos", [])
            if mode in ("anexa", "anexa_bloc"):
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml.ns import qn as docx_qn
                from docx.oxml import OxmlElement

                # Separator pagină nouă
                doc.add_page_break()

                # Titlu secțiune H (conform structurii oficiale Anexa CPE)
                title_p = doc.add_paragraph()
                title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_p.add_run("H. DOCUMENTARE FOTOGRAFICĂ A CLĂDIRII")
                title_run.bold = True
                title_run.font.size = Pt(11)

                if not building_photos:
                    # Placeholder când auditorul nu a încărcat poze
                    sub_p = doc.add_paragraph()
                    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    sub_run = sub_p.add_run(
                        "Auditorul energetic va atașa aici fotografii reprezentative ale clădirii\n"
                        "(exterior, interior, instalații, eventuale defecte/degradări, termoviziune IR)."
                    )
                    sub_run.italic = True
                    sub_run.font.size = Pt(9)
                    sub_run.font.color.rgb = None
                    # 6 casete rezervate cu chenar punctat (3 rânduri × 2 coloane)
                    placeholder_p = doc.add_paragraph()
                    placeholder_p.add_run("\n").font.size = Pt(8)
                    for _ in range(3):
                        ph_tbl = doc.add_table(rows=1, cols=2)
                        tbl_pr = ph_tbl._tbl.tblPr
                        if tbl_pr is None:
                            tbl_pr = OxmlElement("w:tblPr")
                            ph_tbl._tbl.insert(0, tbl_pr)
                        tbl_borders = OxmlElement("w:tblBorders")
                        for bn in ("top", "left", "bottom", "right", "insideH", "insideV"):
                            be = OxmlElement(f"w:{bn}")
                            be.set(docx_qn("w:val"), "dotted")
                            be.set(docx_qn("w:sz"), "8")
                            be.set(docx_qn("w:color"), "999999")
                            tbl_borders.append(be)
                        tbl_pr.append(tbl_borders)
                        for cell in ph_tbl.rows[0].cells:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            placeholder_run = cell.paragraphs[0].add_run("[Spațiu rezervat foto]")
                            placeholder_run.font.size = Pt(8)
                            placeholder_run.italic = True
                            # 3 paragrafe goale per celulă (~3 cm înălțime, evită pagini goale)
                            for _ in range(3):
                                cell.add_paragraph()
                    building_photos = []

                cat_labels = {
                    "exterior": "Exterior",
                    "interior": "Interior",
                    "ir": "Termoviziune IR",
                    "instalatii": "Instalații",
                    "defecte": "Defecte/Degradări",
                    "altele": "Altele",
                }

                # Grupare pe categorii
                grouped = {}
                for ph in building_photos:
                    cat = ph.get("zone", "altele")
                    grouped.setdefault(cat, []).append(ph)

                for cat, photos in grouped.items():
                    # Subtitlu categorie
                    cat_p = doc.add_paragraph()
                    cat_run = cat_p.add_run(f"{cat_labels.get(cat, cat)}  ({len(photos)})")
                    cat_run.bold = True
                    cat_run.font.size = Pt(9)

                    # Tabel 2 coloane pentru poze
                    i = 0
                    while i < len(photos):
                        row_photos = photos[i:i+2]
                        tbl = doc.add_table(rows=1, cols=2)
                        # Elimină borduri tabel
                        tbl_pr = tbl._tbl.tblPr
                        if tbl_pr is None:
                            tbl_pr = OxmlElement("w:tblPr")
                            tbl._tbl.insert(0, tbl_pr)
                        tbl_borders = OxmlElement("w:tblBorders")
                        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                            border_el = OxmlElement(f"w:{border_name}")
                            border_el.set(docx_qn("w:val"), "none")
                            tbl_borders.append(border_el)
                        tbl_pr.append(tbl_borders)

                        for j in range(2):
                            cell = tbl.rows[0].cells[j]
                            if j < len(row_photos):
                                ph = row_photos[j]
                                ph_url = ph.get("url", "")
                                if ph_url and "," in ph_url:
                                    try:
                                        img_data = base64.b64decode(ph_url.split(",")[1])
                                        img_stream = io.BytesIO(img_data)
                                        # Detectează dimensiuni JPEG/PNG fără Pillow
                                        import struct as _struct
                                        raw = img_data
                                        img_w_px, img_h_px = None, None
                                        if raw[:4] == b'\x89PNG':
                                            img_w_px, img_h_px = _struct.unpack('>II', raw[16:24])
                                        elif raw[:2] == b'\xff\xd8':
                                            ii = 2
                                            while ii < len(raw) - 8:
                                                if raw[ii] == 0xff and raw[ii+1] in (0xC0, 0xC2):
                                                    img_h_px = _struct.unpack('>H', raw[ii+5:ii+7])[0]
                                                    img_w_px = _struct.unpack('>H', raw[ii+7:ii+9])[0]
                                                    break
                                                seg_len = _struct.unpack('>H', raw[ii+2:ii+4])[0]
                                                ii += 2 + seg_len
                                        # Calculează height păstrând aspect ratio
                                        pic_w = Inches(2.8)
                                        if img_w_px and img_h_px and img_w_px > 0:
                                            ratio = img_h_px / img_w_px
                                            pic_h = Inches(2.8 * ratio)
                                        else:
                                            pic_h = Inches(2.1)  # fallback 4:3
                                        img_p = cell.paragraphs[0]
                                        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        img_run = img_p.add_run()
                                        img_run.add_picture(img_stream, width=pic_w, height=pic_h)
                                    except Exception as e:
                                        import sys
                                        print(f"[FOTO ERR] {e}", file=sys.stderr)
                                label = ph.get("label", "")
                                note = ph.get("note", "")
                                if label:
                                    lp = cell.add_paragraph(label)
                                    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    if lp.runs:
                                        lp.runs[0].font.size = Pt(7)
                                if note:
                                    np_ = cell.add_paragraph(note)
                                    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    if np_.runs:
                                        np_.runs[0].font.size = Pt(6)
                        i += 2

            # ═══════════════════════════════════════
            # 6b. COLORARE CELULE CLASE ENERGETICE PER UTILITATE
            # ═══════════════════════════════════════
            # Evidențiază cu culoarea clasei celula din tabelul de jos care conține
            # intervalul în care se încadrează consumul specific al fiecărei utilități
            if mode == "cpe":
                _highlight_utility_class_cells(doc, data)

            # ═══════════════════════════════════════
            # 6b. TEXT ROȘU → NEGRU pe tot documentul
            # Template-ul generic are placeholder-uri cu FF0000 (roșu).
            # După înlocuire, textul trebuie să fie negru.
            # ═══════════════════════════════════════
            _w_color_tag = qn("w:color")
            _w_rPr_tag = qn("w:rPr")
            _w_val = qn("w:val")
            for r_el in doc.element.body.iter(qn("w:r")):
                rPr = r_el.find(_w_rPr_tag)
                if rPr is not None:
                    color_el = rPr.find(_w_color_tag)
                    if color_el is not None and color_el.get(_w_val) == "FF0000":
                        color_el.set(_w_val, "000000")

            # ═══════════════════════════════════════
            # 7. STRIP TRAILING EMPTY PARAGRAPHS
            # ═══════════════════════════════════════
            # Elimină paragrafele goale de la finalul documentului (după tabelul de bare cod)
            # Tabelul "COD UNIC DE BARE" se păstrează — face parte din modelul oficial CPE
            body_el = doc.element.body
            sect_pr = body_el.find(qn("w:sectPr"))
            if sect_pr is not None:
                def _get_texts(el):
                    return [t.text for t in el.iter(qn("w:t")) if t.text and t.text.strip()]

                # Elimină paragrafele goale de la final (după conținut/tabel bare cod)
                children = list(body_el)
                sect_idx = children.index(sect_pr)
                for child in reversed(children[:sect_idx]):
                    if child.tag != qn("w:p"):
                        break
                    if _get_texts(child):
                        break
                    body_el.remove(child)

                # ─── CPE: elimină spacing body paragraphs pentru a preveni overflow pagina 2 ──
                # Înlocuirile de text (firma, adresă, etc.) pot mări celule → tabelele cresc
                # → conținut depășește pagina 1. Reducem spacing-ul paragrafelor de la nivel body.
                # Economie: ~11mm (suficient pentru orice date reale de audit)
                if mode == "cpe":
                    for child in list(body_el):
                        if child.tag == qn("w:p"):
                            pPr = child.find(qn("w:pPr"))
                            if pPr is None:
                                continue
                            sp = pPr.find(qn("w:spacing"))
                            if sp is None:
                                continue
                            for attr in (qn("w:after"), qn("w:before")):
                                val = sp.get(attr)
                                if val and int(val) > 0:
                                    sp.set(attr, "0")
                        elif child.tag == qn("w:tbl"):
                            # Reducere spacing în tabelul de bare cod
                            all_t = " ".join(t.text for t in child.iter(qn("w:t")) if t.text)
                            if "COD" in all_t.upper() and "BARE" in all_t.upper():
                                for p in child.iter(qn("w:p")):
                                    pPr2 = p.find(qn("w:pPr"))
                                    if pPr2 is None:
                                        continue
                                    sp2 = pPr2.find(qn("w:spacing"))
                                    if sp2 is None:
                                        continue
                                    for attr in (qn("w:after"), qn("w:before")):
                                        val = sp2.get(attr)
                                        if val and int(val) > 0:
                                            sp2.set(attr, "0")

            # ═══════════════════════════════════════
            # 8. SAVE & RETURN
            # ═══════════════════════════════════════
            buf = io.BytesIO()
            doc.save(buf)
            result = buf.getvalue()

            # Filename per mode (Etapa 4: anexa_bloc primește nume distinct)
            _filename_map = {
                "anexa": "Anexa-1-2-CPE.docx",
                "anexa_bloc": "Anexa-Bloc-Multi-Apartament.docx",
            }
            _filename = _filename_map.get(mode, "CPE.docx")

            self.send_response(200)
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            self.send_header("Content-Disposition", f"attachment; filename={_filename}")
            self.send_header("Content-Length", str(len(result)))
            self.end_headers()
            self.wfile.write(result)

        except Exception as e:
            self.send_response(500)
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Content-Type", "application/json")
            self.end_headers()
            self.wfile.write(json.dumps({"error": str(e)}).encode())
